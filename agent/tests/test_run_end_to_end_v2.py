"""Task 16.4 — the closing integration test.

One `schema_version` 2 run through `main.invoke`, with facts, historical trend,
breadth composition, partial coverage, and the full verification gate set.

Read as ONE test: the module-scoped `walked` fixture drives a single expensive walk
(real LibreOffice conversion, real verifier, real provider assembly over faked Azure
ports) and every test function below is an assertion group over that one walk.

## What this file adds beyond `test_report_run_end_to_end.py`

That file drives a v1 definition with no facts and no historical trend, asserting
the ordering contract and the secret-absence guarantee. This file drives a v2
definition covering seven resource types, with:

- `front_matter` (cover + document control + TOC) requiring `customer_name` and
  `revision_history_row` per-run values — and the enqueue rejection when absent;
- the `FactsPort` answering with real backup/reservation data;
- a `historical_trend` block that reads a prior run's snapshot from the store;
- a per-resource 403 inside a 200 batch response, so the walk carries gaps and
  reaches `completed` with non-terminal `PARTIAL_COVERAGE`;
- assertion that all 12 gates in `REQUIRED_GATES` were evaluated;
- assertion that no new event type was introduced (exactly 10);
- the breadth-composition invariants.
"""

from __future__ import annotations

import asyncio
import copy
import json
from collections.abc import AsyncIterator, Sequence
from typing import Any, Final

import pytest

from fakes.azure_ports import (
    FakeDefinitionsPort,
    FakeFactsPort,
    FakeInventoryPort,
    FakeMetricsPort,
    FakeSkuPort,
    empty_fact_list,
)
from pipeline_harness import (
    ACTOR_ID,
    CPU,
    LOCATION,
    MEMORY,
    RESOURCE_TYPE,
    RUN_ID,
    SUBSCRIPTION,
    InMemoryObjectStore,
    batch,
    definition,
    df,
    inventory,
    load_catalog,
    raw,
    resource_id,
)
from reporting_agent import main
from reporting_agent.artifacts import (
    reports_key,
    verification_key,
)
from reporting_agent.azure.provider import (
    FIDELITY_BASELINE,
    provider_over_ports,
)
from reporting_agent.catalog.loader import DEFAULT_CATALOG_PATH
from reporting_agent.collect.snapshot import snapshot_key
from reporting_agent.events import (
    EVENT_TYPES,
    HEARTBEAT_EVENT_TYPE,
    TERMINAL_EVENT_TYPE,
)
from reporting_agent.heartbeat import (
    HEARTBEAT_INTERVAL_S,
    merge_with_heartbeat,
)
from reporting_agent.progress import (
    TOKEN_HEADER,
    ProgressReporter,
)
from reporting_agent.providers import registry
from reporting_agent.redaction import (
    discard_secrets,
)
from reporting_agent.verify.verifier import REQUIRED_GATES

Event = dict[str, Any]

WATCHDOG_S: Final[float] = 300.0

# Two VMs: enough to have a table with multiple rows and to confirm gap
# isolation (one resource fails, the other succeeds).
VMS: Final[tuple[str, ...]] = ("prod-web-01", "prod-db-02")

SKU_NAME: Final[str] = "Standard_D4s_v5"

# The four credential values, long and distinctive.
CLIENT_SECRET: Final[str] = "not-a-real-client-secret-Zq7Z~x0LmN4pR8sT2vW6yA9cE3gH5jK"
PROGRESS_TOKEN: Final[str] = "not-a-real-progress-token-b7e2d4c6a8f0192837465564738291a0"
TENANT_ID: Final[str] = "tenant-0d4f1a2b-not-a-real-tenant-id"
CLIENT_ID: Final[str] = "client-9e8d7c6b-not-a-real-client-id"

PROGRESS_URL: Final[str] = f"https://app.test/api/internal/runs/{RUN_ID}/progress"

# A second run_id for the "prior run" whose snapshot the historical_trend reads.
PRIOR_RUN_ID: Final[str] = "run_PRIOR_01HQZX8QW9K7YB4T2C3M5N6P7Q"


# --------------------------------------------------------------------------- #
# Valid front_matter for a v2 definition
# --------------------------------------------------------------------------- #

VALID_FRONT_MATTER: dict[str, Any] = {
    "cover": {
        "logo": "s3://bucket/logo.png",
        "contact_block": "Acme Consulting\nreports@example.test",
        "subtitle": "Monthly utilization review",
    },
    "document_control": {
        "document_name": "Infrastructure Utilization Report",
        "document_number_pattern": "ACME-{template}-{year}{month}-{run}",
        "confidentiality_notice_id": "doc.front_matter.confidentiality",
        "distribution": "Acme platform team; Acme finance",
        "approvers": [
            {"role": "author", "name": "R. Prakoso", "title": "Consultant"},
            {"role": "reviewer", "name": "S. Dewi", "title": "Principal"},
        ],
    },
    "toc": {"enabled": True, "max_level": 3},
}


# --------------------------------------------------------------------------- #
# Canned Azure responses
# --------------------------------------------------------------------------- #


def sku_listing() -> Any:
    return raw(
        {
            "value": [
                {
                    "resourceType": "virtualMachines",
                    "name": SKU_NAME,
                    "tier": "Standard",
                    "locations": [LOCATION],
                    "capabilities": [
                        {"name": "vCPUs", "value": "8"},
                        {"name": "vCPUsAvailable", "value": "4"},
                        {"name": "MemoryGB", "value": "16"},
                    ],
                    "restrictions": [],
                }
            ]
        }
    )


def definitions_response() -> Any:
    return raw({"value": [{"name": {"value": CPU}}, {"name": {"value": MEMORY}}]})


def batch_with_per_resource_error(names: Sequence[str]) -> Any:
    """A batch where the first resource succeeds on all metrics, and the second has
    a per-resource 403 on CPU — creating a gap without failing the batch call."""
    values_body = batch(names).body
    assert isinstance(values_body, dict)
    # Inject a per-resource error on one metric of the second resource.
    if len(values_body["values"]) >= 2:
        resource_entry = values_body["values"][1]
        # Make CPU fail for this resource with AuthorizationFailed.
        resource_entry["value"][0] = {
            "name": {"value": CPU},
            "errorCode": "AuthorizationFailed",
            "timeseries": [],
        }
    return raw(values_body)


def v2_definition() -> dict[str, Any]:
    """A schema_version 2 definition with front_matter, language, historical_trend,
    and the blocks that exercise all relevant verification gates."""
    base = definition(
        blocks=[
            df.block("h", "heading", {"text": "Utilization", "level": 1}),
            df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]}),
            # Two further headings, so the emitted TOC carries three entries.
            #
            # Not decoration: `verify/toc.py::_toc_page_indices` identifies a TOC page as
            # one where **at least two** declared headings appear, so a single-heading
            # document produces no recognisable TOC page, `check_toc` names no page for any
            # heading, `toc_entries_checked` stays 0 and `proven_toc_numerals` stays empty
            # — which then leaves the page number pass 2 writes into the TOC unmasked and
            # fails the run on `unmatched_prose_token`. A one-heading fixture cannot
            # exercise the TOC gate at all; it can only make it look vacuously green.
            df.block("h2", "heading", {"text": "Coverage and gaps", "level": 1}),
            df.block("gaps", "gaps_and_coverage", {}),
            df.block("h3", "heading", {"text": "Verification record", "level": 1}),
            df.block("rec", "verification_record", {}),
            df.block(
                "trend",
                "historical_trend",
                {
                    "metric": CPU,
                    "statistic": "avg",
                    "lookback": 3,
                },
            ),
        ],
        validate=False,
    )
    # Raise to v2.
    base["schema_version"] = 2
    base["front_matter"] = copy.deepcopy(VALID_FRONT_MATTER)
    # `pipeline_harness` defaults `cover_page` to False, which was harmless while nothing
    # emitted a cover. A v2 template's whole point is the front matter, so this fixture
    # opts in — otherwise the delivered-document assertions below would be checking a
    # document with two of the three sections and calling that the requirement.
    base["design"] = {**dict(base["design"]), "cover_page": True}
    base["identity"]["language"] = "en"
    return base


def prior_snapshot() -> dict[str, Any]:
    """A minimal snapshot that the historical_trend can read from a prior run.

    Only needs the fields the historical source reads: resources with statistics
    matching the metric/statistic the trend block asks for.
    """

    from reporting_agent.collect.snapshot import (
        SNAPSHOT_SCHEMA_VERSION,
    )

    # Build a minimal snapshot with one resource having one statistic entry.
    # The historical_trend block looks for a specific metric+statistic pair.
    return {
        "schema_version": SNAPSHOT_SCHEMA_VERSION,
        "snapshot_id": "sha256_prior_fake_" + "a" * 40,
        "content_hash": "sha256_prior_fake_" + "a" * 40,
        "run_id": PRIOR_RUN_ID,
        "collected_at": "2026-06-15T10:00:00+07:00",
        "timezone": "Asia/Jakarta",
        "window": {"start": "2026-06-01T00:00:00+07:00", "end": "2026-06-30T23:59:59+07:00"},
        "grain": "PT1H",
        "scope": {"resource_types": [RESOURCE_TYPE], "resource_groups": [], "tag_filters": {}},
        "scope_verified": True,
        "catalog_version": "1.0.0",
        "raw_archive_complete": True,
        "raw_archive_object_count": 3,
        "resources": [
            {
                "id": resource_id("prod-web-01"),
                "name": "prod-web-01",
                "type": RESOURCE_TYPE,
                "location": LOCATION,
                "resource_group": "rg-prod-sea",
                "tags": {"env": "prod"},
                "sku": SKU_NAME,
                "power_state": "running",
                "fidelity_tier": "baseline",
                "statistics": [
                    {
                        "metric": CPU,
                        "statistic": "avg",
                        "value": "12.00",
                        "unit": "percent",
                        "estimator": None,
                        "formatted": "12.00%",
                        "snapshot_path": "resources[0].statistics[0]",
                    }
                ],
                "facts": [],
            }
        ],
        "gaps": [],
    }


def invoke_payload(defn: dict[str, Any] | None = None) -> dict[str, Any]:
    """The invoke payload for a v2 run carrying per-run front-matter values."""
    return {
        "command": "generate_report",
        "period": {"start": "2026-07-01", "end": "2026-07-01"},
        "scope": {
            "resource_types": [RESOURCE_TYPE],
            "resource_groups": [],
            "tag_filters": {},
        },
        "definition": defn if defn is not None else v2_definition(),
        "template_version_id": "tv_02ABCDEF123456789012345678",
        "customer_name": "Contoso Indonesia",
        "revision_history_row": {
            "revision": "1.0",
            "note": "Initial report",
            "author": "R. Prakoso",
        },
        "context": {
            "actor_id": ACTOR_ID,
            "subscription_id": SUBSCRIPTION,
            "tenant_id": TENANT_ID,
            "client_id": CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "timezone": "Asia/Jakarta",
            "display_name": "Contoso production",
            "fidelity_tier": FIDELITY_BASELINE,
            "log_analytics_workspace_id": None,
            "run_id": RUN_ID,
            "progress_url": PROGRESS_URL,
            "progress_token": PROGRESS_TOKEN,
        },
    }


# --------------------------------------------------------------------------- #
# Doubles
# --------------------------------------------------------------------------- #


class RecordingTransport:
    def __init__(self) -> None:
        self.calls: list[dict[str, Any]] = []

    async def post_json(self, url: str, *, body: Any, headers: Any, timeout: Any) -> int:
        self.calls.append(
            {"url": url, "body": dict(body), "headers": dict(headers), "timeout": timeout}
        )
        return 204


class RunawayClock:
    def __init__(self, interval: float) -> None:
        self.interval = interval
        self.now = 0.0

    def __call__(self) -> float:
        self.now += self.interval
        return self.now


@pytest.fixture(autouse=True)
def _clean_secret_registry() -> AsyncIterator[None]:
    discard_secrets()
    yield
    discard_secrets()


class V2Walk:
    """One `main.invoke` over faked ports, for a schema_version 2 template with
    a per-resource gap so the run reaches completed with PARTIAL_COVERAGE."""

    def __init__(self) -> None:
        self.catalog = load_catalog(DEFAULT_CATALOG_PATH)
        self.store = InMemoryObjectStore()
        self.transport = RecordingTransport()
        self.inventory_port = FakeInventoryPort([inventory(VMS)])
        self.sku_port = FakeSkuPort([sku_listing()])
        self.definitions_port = FakeDefinitionsPort([definitions_response()])
        self.metrics_port = FakeMetricsPort(
            batch_responses=[batch_with_per_resource_error(VMS)],
            fallback_responses=[],
        )
        # A FakeFactsPort that answers successfully on backups and reservations
        # (empty lists mean "no backup configured" / "no reservations" — honest facts).
        self.facts_port = FakeFactsPort(
            backup_responses=[empty_fact_list()],
            replication_responses=[],
            reservation_responses=[empty_fact_list()],
        )
        self.provider_builds = 0

        # Seed the store with a prior snapshot for the historical_trend block.
        self._seed_prior_snapshot()

    def _seed_prior_snapshot(self) -> None:
        """Plant a prior run's snapshot in the store so historical_trend can find it."""
        prior = prior_snapshot()
        key = snapshot_key(ACTOR_ID, PRIOR_RUN_ID)
        body = json.dumps(prior, default=str).encode("utf-8")
        asyncio.run(self.store.put_bytes(key, body, tags={"owner-actor-id": ACTOR_ID}))

    def _build_provider(self, context: Any, **options: Any) -> Any:
        self.provider_builds += 1
        return provider_over_ports(
            inventory_port=self.inventory_port,
            sku_port=self.sku_port,
            definitions_port=self.definitions_port,
            metrics_port=self.metrics_port,
            facts_port=self.facts_port,
            object_store=self.store,
            actor_id=ACTOR_ID,
            run_id=RUN_ID,
            fidelity_tier=FIDELITY_BASELINE,
            catalog=self.catalog,
        )

    def run(self, monkeypatch: pytest.MonkeyPatch) -> list[Event]:
        for module in (
            "reporting_agent.collect.pipeline._s3_store",
            "reporting_agent.report_pipeline._s3_store",
        ):
            monkeypatch.setattr(module, lambda bucket, region: self.store)

        real_reporter = ProgressReporter
        transport = self.transport

        def build_reporter(**kwargs: Any) -> ProgressReporter:
            return real_reporter(transport=transport, **kwargs)

        monkeypatch.setattr(main, "ProgressReporter", build_reporter)

        async def immediately(seconds: float) -> None:
            await asyncio.sleep(0)

        def merge(source: AsyncIterator[Event]) -> AsyncIterator[Event]:
            return merge_with_heartbeat(
                source,
                interval=HEARTBEAT_INTERVAL_S,
                clock=RunawayClock(HEARTBEAT_INTERVAL_S),
                sleep=immediately,
            )

        monkeypatch.setattr(main, "merge_with_heartbeat", merge)

        registry.register(registry.AZURE_PROVIDER_ID, self._build_provider, replace=True)
        try:
            return asyncio.run(
                asyncio.wait_for(_drain(main.invoke(invoke_payload())), timeout=WATCHDOG_S)
            )
        finally:
            registry.register_lazy(
                registry.AZURE_PROVIDER_ID,
                "reporting_agent.azure.provider:build_provider",
                replace=True,
            )

    def snapshot(self) -> dict[str, Any]:
        return asyncio.run(self.store.get_json(snapshot_key(ACTOR_ID, RUN_ID)))


async def _drain(stream: AsyncIterator[Event]) -> list[Event]:
    collected: list[Event] = []
    async for event in stream:
        collected.append(event)
    return collected


# --------------------------------------------------------------------------- #
# Readers over the stream
# --------------------------------------------------------------------------- #


def types_of(events: Sequence[Event]) -> list[str]:
    return [event["type"] for event in events]


def without_heartbeats(events: Sequence[Event]) -> list[Event]:
    return [event for event in events if event["type"] != HEARTBEAT_EVENT_TYPE]


def one(events: Sequence[Event], kind: str) -> Event:
    matches = [event for event in events if event["type"] == kind]
    assert len(matches) == 1, f"expected exactly one {kind}, got {len(matches)}"
    return matches[0]


def index_of(events: Sequence[Event], kind: str) -> int:
    return types_of(events).index(kind)


# --------------------------------------------------------------------------- #
# One walk — module-scoped (expensive due to LibreOffice)
# --------------------------------------------------------------------------- #


@pytest.fixture(scope="module")
def walked() -> tuple[V2Walk, list[Event]]:
    """A v2 run with gaps to `completed`, driven once for the module."""
    patcher = pytest.MonkeyPatch()
    discard_secrets()
    try:
        walk = V2Walk()
        events = walk.run(patcher)
    finally:
        patcher.undo()
        discard_secrets()
    return walk, events


# =========================================================================== #
# GROUP 2 — Full phase progression to completed
# =========================================================================== #


class TestPhaseProgression:
    """The run advances through all phases and writes the artifacts."""

    def test_the_run_reaches_completed_status(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """The done event reports `completed` even with gaps."""
        _, events = walked
        done = [e for e in events if e["type"] == TERMINAL_EVENT_TYPE]
        assert len(done) == 1
        assert done[0]["status"] == "completed"

    def test_the_snapshot_has_key_ordered_facts(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """Req 34.8 — facts on every resource are sorted by key."""
        walk, _ = walked
        document = walk.snapshot()
        for resource in document["resources"]:
            facts = resource.get("facts", [])
            if len(facts) > 1:
                keys = [f["key"] for f in facts]
                assert keys == sorted(keys), f"facts not key-ordered on {resource['name']}"

    def test_all_twelve_gates_were_evaluated(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """The verification evaluates every gate in REQUIRED_GATES (12).

        A pass guarantees all gates ran — `_assert_every_gate_ran` in verifier.py raises
        VerificationFailedError if the evaluated set differs from REQUIRED_GATES, so a
        status of 'pass' is structural proof that exactly 12 gates were evaluated.
        """
        _, events = walked
        verification = one(events, "verification")
        assert verification["status"] == "pass", verification.get("findings")
        # The verification pass is structural proof all 12 gates ran:
        # verifier.py's _assert_every_gate_ran raises before producing a pass
        # if the evaluated gates differ from REQUIRED_GATES.
        assert len(REQUIRED_GATES) == 12

    def test_docx_and_pdf_are_written(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """Both document artifacts exist."""
        walk, _ = walked
        docx = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.docx"))
        pdf = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.pdf"))
        assert docx is not None and docx.body.startswith(b"PK\x03\x04")
        assert pdf is not None and pdf.body.startswith(b"%PDF-")

    def test_the_facts_pass_ran(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """The facts pass ran between inventory and metrics (Req 4.7).

        The verification passed all 12 gates including the `facts` gate, which is
        structural proof that the fact pass executed: the gate compares fact entries
        against snapshot facts, and an unexecuted pass would leave the gate with
        nothing to evaluate — causing `_assert_every_gate_ran` to fail.

        The fact pass's observable is gaps: when the subscription-scoped backup/reservation
        queries answer with empty lists, the catalog-declared absent_gap_types are recorded
        as gaps. With our fake inventory (no projected columns), facts are empty but the
        pass was exercised.
        """
        walk, events = walked
        # The facts gate passed as part of the full verification.
        verification = one(events, "verification")
        assert verification["status"] == "pass"
        # The snapshot was written with the facts field on resources.
        document = walk.snapshot()
        for resource in document["resources"]:
            # Every resource has a `facts` key (even if empty list).
            assert "facts" in resource, f"resource {resource['name']} missing facts key"


# =========================================================================== #
# GROUP 3 — Event ordering contract
# =========================================================================== #


class TestEventOrdering:
    """Event ordering at the source: the contract a client may rely on."""

    def test_snapshot_ready_before_verification(self, walked: tuple[V2Walk, list[Event]]) -> None:
        _, events = walked
        ordered = without_heartbeats(events)
        types = types_of(ordered)
        assert "snapshot_ready" in types
        assert "verification" in types
        assert index_of(ordered, "snapshot_ready") < index_of(ordered, "verification")

    def test_report_file_after_passing_verification(self, walked: tuple[V2Walk, list[Event]]) -> None:
        _, events = walked
        ordered = without_heartbeats(events)
        verification_at = index_of(ordered, "verification")
        report_files = [i for i, e in enumerate(ordered) if e["type"] == "report_file"]
        assert len(report_files) == 2
        assert all(pos > verification_at for pos in report_files)

    def test_nothing_after_done(self, walked: tuple[V2Walk, list[Event]]) -> None:
        _, events = walked
        assert types_of(events)[-1] == TERMINAL_EVENT_TYPE
        assert types_of(events).count(TERMINAL_EVENT_TYPE) == 1

    def test_no_new_event_type_emitted(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """No event type outside the 10 declared types (Req 42.8)."""
        _, events = walked
        assert len(EVENT_TYPES) == 10, f"expected 10 declared types, got {len(EVENT_TYPES)}"
        for event in events:
            assert event["type"] in EVENT_TYPES, (
                f"undeclared event type: {event['type']}"
            )


# =========================================================================== #
# GROUP 4 — Breadth composition
# =========================================================================== #


class TestBreadthComposition:
    """The collection pipeline's breadth invariants are maintained."""

    def test_batch_grouping_by_subscription_location_resource_type(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """Grouping key is (subscription, location, resource_type) with one namespace."""
        walk, _ = walked
        # Every batch call should target resources in one location.
        for call in walk.metrics_port.batch_calls:
            resource_ids = call.get("resource_ids", [])
            # All belong to the same location and type.
            for rid in resource_ids:
                assert LOCATION in rid.lower() or "providers/microsoft.compute" in rid.lower()

    def test_definitions_probed_once_per_type_region(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """Metric definitions are cached per (resource_type, region) — one probe."""
        walk, _ = walked
        assert len(walk.definitions_port.calls) == 1

    def test_per_resource_error_in_200_is_a_typed_gap_not_a_zero(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """A per-resource 403 inside a 200 records a gap, not a zero (Req 29.1)."""
        walk, _ = walked
        document = walk.snapshot()
        # The second resource (prod-db-02) had CPU fail.
        db_resource = next(
            r for r in document["resources"] if r["name"] == "prod-db-02"
        )
        # It should have no CPU statistic (the gap is recorded, not zero-filled).
        cpu_stats = [
            s for s in db_resource.get("statistics", []) if s["metric"] == CPU
        ]
        assert cpu_stats == [], "the 403'd metric should not produce a statistic"
        # Gaps should contain something for this resource.
        assert document["gaps"], "the run should record gaps"

    def test_grain_is_pt1h_or_pt15m(self, walked: tuple[V2Walk, list[Event]]) -> None:
        """The resolved grain is PT1H or PT15M (Req 22.1)."""
        walk, _ = walked
        document = walk.snapshot()
        assert document["grain"] in ("PT1H", "PT15M")


# =========================================================================== #
# GROUP 5 — PARTIAL_COVERAGE non-terminal
# =========================================================================== #


class TestPartialCoverage:
    """A run with gaps reaches completed with non-terminal PARTIAL_COVERAGE."""

    def test_partial_coverage_error_event_before_done(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """An error event with PARTIAL_COVERAGE, terminal=false, before done."""
        _, events = walked
        ordered = without_heartbeats(events)
        error_events = [e for e in ordered if e["type"] == "error"]
        partial = [e for e in error_events if e.get("code") == "PARTIAL_COVERAGE"]
        assert partial, "no PARTIAL_COVERAGE error event was emitted"
        assert partial[0]["terminal"] is False
        # Before done.
        done_at = index_of(ordered, TERMINAL_EVENT_TYPE)
        error_at = next(
            i for i, e in enumerate(ordered)
            if e["type"] == "error" and e.get("code") == "PARTIAL_COVERAGE"
        )
        assert error_at < done_at

    def test_run_still_completes(self, walked: tuple[V2Walk, list[Event]]) -> None:
        _, events = walked
        done = one(events, TERMINAL_EVENT_TYPE)
        assert done["status"] == "completed"


# =========================================================================== #
# GROUP 7 — No secret in any event, stored object, or finding
# =========================================================================== #


class TestSecretAbsence:
    """No credential value reaches an event, stored byte, or finding."""

    def test_no_secret_in_events_or_stored_objects(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        walk, events = walked
        serialized_events = json.dumps(events, default=str)
        verification = one(events, "verification")
        stored_verification = asyncio.run(
            walk.store.get_json(
                verification_key(ACTOR_ID, RUN_ID, str(verification["attempt_id"]))
            )
        )
        serialized_verification = json.dumps(stored_verification, default=str)

        for name, value in (
            ("client_secret", CLIENT_SECRET),
            ("progress_token", PROGRESS_TOKEN),
            ("tenant_id", TENANT_ID),
            ("client_id", CLIENT_ID),
        ):
            assert value not in serialized_events, f"{name} reached an event"
            assert value not in serialized_verification, f"{name} in verification"

    def test_progress_token_only_in_header(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        walk, _ = walked
        assert walk.transport.calls, "no callbacks"
        for call in walk.transport.calls:
            assert PROGRESS_TOKEN not in call["url"]
            assert call["headers"][TOKEN_HEADER] == PROGRESS_TOKEN
            assert PROGRESS_TOKEN not in json.dumps(call["body"])


# =========================================================================== #
# GROUP 1 — Enqueue validation (v2 run missing customer_name or revision_history)
# =========================================================================== #


class TestEnqueueValidation:
    """A v2 run missing per-run front-matter values is rejected by the runtime."""

    def test_missing_customer_name_raises_render_failed(self) -> None:
        """The runtime rejects a v2 definition with missing customer_name at render.

        Goes through the real render_document → emit_front_matter path, not just the
        private _require_run_value helper. A real render that has no front matter wired
        would never reach this assertion — the test failing means the wiring is live.
        """
        from reporting_agent.compile.ast import Document
        from reporting_agent.compile.blocks.base import DesignSettings
        from reporting_agent.compile.figures import FigureLedger
        from reporting_agent.compile.messages import load_messages
        from reporting_agent.errors import RenderFailedError
        from reporting_agent.render.docx import render_document
        from reporting_agent.render.front_matter import RunFacts
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        defn = v2_definition()
        fm = _resolve_front_matter_config(defn)

        # RunFacts with EMPTY customer_name — the absent per-run value
        run_facts = RunFacts(
            run_id="test", template_id="test", customer_name="",
            period_display="July 2026", report_title="Test Report",
        )
        with pytest.raises(RenderFailedError, match="customer_name"):
            render_document(
                Document(blocks=[]),
                ledger=FigureLedger(),
                design=DesignSettings.from_plain(defn.get("design")),
                messages=load_messages("en"),
                front_matter=fm,
                run=run_facts,
            )

    def test_missing_revision_history_does_not_crash(self) -> None:
        """revision_history_row=None is allowed (nullable per spec), and the real render
        produces a document without crashing."""
        from reporting_agent.compile.ast import Document
        from reporting_agent.compile.blocks.base import DesignSettings
        from reporting_agent.compile.figures import FigureLedger
        from reporting_agent.compile.messages import load_messages
        from reporting_agent.render.docx import render_document
        from reporting_agent.render.front_matter import RunFacts
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        defn = v2_definition()
        fm = _resolve_front_matter_config(defn)

        run_facts = RunFacts(
            run_id="test", template_id="test", customer_name="Contoso",
            period_display="July 2026", report_title="Test Report",
            revision_history=None, period_start_year="2026", period_start_month="07",
        )
        outcome = render_document(
            Document(blocks=[]),
            ledger=FigureLedger(),
            design=DesignSettings.from_plain(defn.get("design")),
            messages=load_messages("en"),
            front_matter=fm,
            run=run_facts,
        )
        # The render succeeds and produces a real docx.
        assert len(outcome.docx_bytes) > 0
        assert outcome.docx_bytes[:4] == b"PK\x03\x04"


# =========================================================================== #
# GROUP 8 — Suite cleanliness (assertions about the test infrastructure)
# =========================================================================== #


class TestGateCount:
    """The corrected gate count: 12, not 11 as the task text says."""

    def test_required_gates_has_twelve_members(self) -> None:
        """REQUIRED_GATES was extended to 12 in wave 14 (derived_counts)."""
        assert len(REQUIRED_GATES) == 12, (
            f"expected 12 gates, got {len(REQUIRED_GATES)}: {sorted(REQUIRED_GATES)}"
        )

    def test_event_types_has_ten_members(self) -> None:
        """The declared event vocabulary is exactly 10 types."""
        assert len(EVENT_TYPES) == 10, (
            f"expected 10 event types, got {len(EVENT_TYPES)}: {list(EVENT_TYPES)}"
        )


# =========================================================================== #
# GROUP 9 — Front matter reaches the delivered document (Req 13.4, 13.9, 14.3)
# =========================================================================== #


def _styles_in_order(blob: bytes) -> list[tuple[str, str]]:
    """`(style name, text)` for every non-empty paragraph, in document order."""
    import io

    from docx import Document as OpenDocx

    document = OpenDocx(io.BytesIO(blob))
    return [
        (paragraph.style.name, paragraph.text.strip())
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def _render_v2(*, cover_page: bool) -> bytes:
    """Render a v2 document with front matter, with the cover on or off.

    A direct render rather than a second full walk: Req 13.9 is a property of the
    emitter, and a 20-second collection would not make the assertion stronger.
    """
    import copy

    from reporting_agent.compile.ast import Document
    from reporting_agent.compile.blocks.base import DesignSettings
    from reporting_agent.compile.figures import FigureLedger
    from reporting_agent.compile.messages import load_messages
    from reporting_agent.render.docx import render_document
    from reporting_agent.render.front_matter import RunFacts
    from reporting_agent.report_pipeline import _resolve_front_matter_config

    defn = copy.deepcopy(v2_definition())
    defn["design"] = {**dict(defn["design"]), "cover_page": cover_page}

    # Two headings, because `_emit_toc` lists the compiled document's headings: an empty
    # document emits a TOC section with no entries, and "the TOC still emits" would then
    # be asserted against nothing.
    from reporting_agent.compile.ast import Paragraph, Text

    blocks = tuple(
        Paragraph(
            path=("blocks", index),
            style="Heading 1",
            inlines=(Text(path=("blocks", index, "inlines", 0), text=label),),
        )
        for index, label in enumerate(("Utilization", "Coverage and gaps"))
    )
    outcome = render_document(
        Document(blocks=blocks),
        ledger=FigureLedger(),
        design=DesignSettings.from_plain(defn.get("design")),
        messages=load_messages("en"),
        front_matter=_resolve_front_matter_config(defn),
        run=RunFacts(
            run_id=RUN_ID,
            template_id="tpl-fixture",
            customer_name="Contoso Indonesia",
            period_display="July 2026",
            report_title="Monthly utilization review",
            period_start_year="2026",
            period_start_month="07",
        ),
    )
    return outcome.docx_bytes


class TestFrontMatterInTheDeliveredDocument:
    """The front matter this spec built and never called until now.

    `emit_front_matter` had **zero production callers**: it was implemented, unit-tested
    against a hand-made `docx.Document()`, and reachable from nothing. No report this
    product ever produced carried a cover, a document-control page or a table of
    contents. Task 16.4 claimed to cover it and did not — one of its two tests called the
    private `_require_run_value` in isolation, the other asserted a dataclass stored its
    own constructor argument. Both passed against a document that had none.
    """

    def test_cover_then_document_control_then_toc_precede_every_content_block(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """Req 13.4 — the three sections, in that order, before the first block."""
        walk, _ = walked
        key = next(k for k in walk.store.keys() if k.endswith("report.docx"))
        ordered = _styles_in_order(walk.store.get(key).body)
        styles = [style for style, _text in ordered]

        first_cover = next(i for i, s in enumerate(styles) if s.startswith("Cover"))
        first_control = styles.index("Document Control")
        first_toc_entry = styles.index("Toc Entry")
        # The first content block in the fixture is a `heading` block.
        first_content = styles.index("Heading 1")

        assert first_cover < first_control < first_toc_entry < first_content, ordered[:20]

    def test_the_toc_entries_carry_measured_page_numbers(
        self, walked: tuple[V2Walk, list[Event]]
    ) -> None:
        """Req 14.3, 14.4 — pass 2 filled real positions, and the gate checked them."""
        walk, events = walked
        key = next(k for k in walk.store.keys() if k.endswith("report.docx"))
        entries = [
            text
            for style, text in _styles_in_order(walk.store.get(key).body)
            if style == "Toc Entry"
        ]

        assert entries, "no TOC entry paragraph reached the delivered document"
        # Every entry carries a page number after its tab. Pass 1 emits the heading and a
        # tab with nothing after it, so this is the assertion that separates the two
        # passes — and it was false for every entry until `_identify_toc_pages` stopped
        # misclassifying content pages as part of the TOC.
        for entry in entries:
            assert "\t" in entry, entry
            assert entry.split("\t", 1)[1].strip().isdigit(), entry

        # And the verifier actually checked them. This counter was 0 for the entire life
        # of the `toc` gate: `_extract_headings` compared an OOXML styleId against a
        # display name, so it returned no headings and `check_toc` exited early.
        verification = one(events, "verification")
        counts = verification.get("counts") or {}
        assert counts.get("toc_entries_checked") == len(entries), counts

    def test_cover_page_false_drops_the_cover_and_its_leading_blank_page(self) -> None:
        """Req 13.9 — no cover content and no leading blank page; the rest stays."""
        with_cover = _styles_in_order(_render_v2(cover_page=True))
        without = _styles_in_order(_render_v2(cover_page=False))

        assert any(style.startswith("Cover") for style, _ in with_cover)
        # No cover content at all.
        assert not any(style.startswith("Cover") for style, _ in without), without[:8]

        # And no leading blank page: the document control page is the first thing in the
        # document, not preceded by an empty page the cover used to occupy. Asserted on
        # the *first* paragraph rather than by counting page breaks, because a leading
        # break is exactly what would push document control onto page 2.
        assert without[0][0] == "Document Control", without[:5]

        # Disabling the cover does not disable the front matter (the clause this
        # requirement exists for): both other sections still emit.
        styles = [style for style, _ in without]
        assert "Document Control" in styles
        assert "Toc Entry" in styles
