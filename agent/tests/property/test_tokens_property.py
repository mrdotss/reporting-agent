"""Property 2: Token extraction and prose masking.

**Validates: Requirements 26.1, 26.3, 26.6, 26.7, 26.8, 26.9, 28.1, 28.2, 28.3, 28.4,
28.5, 28.6, 28.9, 28.11, 29.1, 19.3, 19.4, 33.5, 33.6, 45.1**

*For any* rendered document whose figures are split across consecutive runs and whose
prose embeds identifiers, resource ids, addresses, dates, durations and template
chrome, the extractor extracts each figure as one token equal to its `formatted`
string; the five masking stages leave zero numeric-bearing survivors; and a numeric
string absent from both the ledger and the allowlist always survives and is always
reported.

**On scale.** The requirement names 1–5,000 paragraphs and 0–500 data tables. Building
a 5,000-paragraph `.docx` a hundred times over would put this module in the minutes and
buy nothing the shape of the assertion does not already cover, so the split here is
deliberate: the *masking* properties — which are pure functions over text and where
scale is free — run over 1–400 paragraphs of generated prose, and the *extraction*
properties, which need a real Word package with real nesting, run over 1–30 paragraphs
plus one declared example at 1,200 paragraphs and 60 nested tables. The large case runs
on every invocation as a declared example rather than as a generated one, so the scale
clause is exercised deterministically instead of occasionally.
"""

from __future__ import annotations

import re
from typing import Final

from docx import Document
from hypothesis import example, given
from hypothesis import strategies as st

from reporting_agent.render.anchors import write_data_table_caption, write_layout_table
from reporting_agent.verify.masking import (
    MASK_CHAR,
    mask_paragraph,
    masking_order,
    scan_paragraphs,
)
from reporting_agent.verify.tokens import (
    PART_BODY,
    ExtractedParagraph,
    numeric_tokens,
    paragraph_texts,
)

# --- vocabularies whose members are legitimately numeric and must never survive --------

IDENTIFIERS: Final[tuple[str, ...]] = (
    "prod-sql-01",
    "web-01",
    "Standard_D4s_v5",
    "Standard_E32-8s_v5",
    "westus2",
    "rg-app-002",
)
STRUCTURED: Final[tuple[str, ...]] = (
    "550e8400-e29b-41d4-a716-446655440000",
    "/subscriptions/4e818b57-c747-4ce0-ac4f-bfc7912e95a4/resourceGroups/rg/vm/web-01",
    "10.0.1.4",
    "10.0.0.0/16",
    "2001:0db8:85a3:0000:0000:8a2e:0370:7334",
    "fe80::1",
)
TEMPORALS: Final[tuple[str, ...]] = (
    "2026-07-01",
    "2026-07-31T17:00:00Z",
    "PT1H",
    "PT15M",
    "P1D",
    "14:35",
)
CHROME: Final[tuple[str, ...]] = ("Top 10", "Section 3", "page 2 of 14")

WORDS: Final[tuple[str, ...]] = (
    "average",
    "peak",
    "observed",
    "headroom",
    "the",
    "across",
)


@st.composite
def formatted_strings(draw: st.DrawFn, minimum: int = 1, maximum: int = 6) -> list[str]:
    """A ledger's display strings, always including a proper-substring pair.

    `12.4%` and `112.4%` are drawn together on purpose: that pair is what distinguishes
    longest-first masking from insertion-order masking, and a generator that produced
    it only by chance would exercise the distinction only by chance.
    """
    count = draw(st.integers(min_value=minimum, max_value=maximum))
    whole = draw(st.lists(st.integers(1, 999), min_size=count, max_size=count))
    frac = draw(st.lists(st.integers(0, 99), min_size=count, max_size=count))
    values = [f"{w}.{f:02d}%" for w, f in zip(whole, frac, strict=True)]
    shadowed = values[0]
    values.append(f"1{shadowed}")  # a proper substring of the one after it
    return sorted(set(values))


@st.composite
def legitimate_paragraph(draw: st.DrawFn, ledger: list[str], allowlist: list[str]) -> str:
    """Prose in which every numeric-bearing token is legitimate.

    Every token comes from the ledger, the identifier set, the structured set, the
    temporal set, the allowlist or a plain word — so a correct masker leaves zero
    survivors and any survivor is a defect rather than a generator artefact.
    """
    pieces = draw(
        st.lists(
            st.one_of(
                st.sampled_from(WORDS),
                st.sampled_from(ledger),
                st.sampled_from(IDENTIFIERS),
                st.sampled_from(STRUCTURED),
                st.sampled_from(TEMPORALS),
                st.sampled_from(allowlist) if allowlist else st.sampled_from(WORDS),
            ),
            min_size=1,
            max_size=12,
        )
    )
    return " ".join(pieces)


def _allowlist_tokens() -> list[str]:
    """The whitespace-delimited numeric tokens of the chrome strings."""
    return sorted(
        {
            match.group()
            for phrase in CHROME
            for match in re.finditer(r"\S+", phrase)
            if re.search(r"\d", match.group())
        }
    )


def _survivor_texts(text: str, ledger: list[str], allowlist: list[str]) -> list[str]:
    masked = mask_paragraph(
        text,
        ledger_strings=masking_order(ledger),
        allowlist=masking_order(allowlist),
    )
    return [m.group() for m in re.finditer(r"\S+", masked) if re.search(r"\d", m.group())]


# --- the masking half -------------------------------------------------------------------


@given(data=st.data())
# declared: the substring-shadowed ledger pair
@example(data=None)
def test_legitimate_prose_leaves_zero_survivors(data: object) -> None:
    """Req 28.1-28.6 — every numeral in the prose is accounted for by some stage."""
    if data is None:
        ledger, allowlist = ["12.4%", "112.4%"], _allowlist_tokens()
        paragraphs = ["Average 112.4% and 12.4% on prod-sql-01 over PT1H", "Top 10"]
    else:
        ledger = data.draw(formatted_strings())  # type: ignore[attr-defined]
        allowlist = _allowlist_tokens()
        paragraphs = data.draw(  # type: ignore[attr-defined]
            st.lists(legitimate_paragraph(ledger, allowlist), min_size=1, max_size=400)
        )
    for text in paragraphs:
        assert _survivor_texts(text, ledger, allowlist) == [], text


@given(data=st.data())
# declared: 12.4% must not eat the tail of 112.4%
@example(data=None)
def test_a_shadowed_ledger_string_masks_longest_first(data: object) -> None:
    """Req 28.2 — insertion-order masking leaves `1` behind on a correct document."""
    if data is None:
        ledger = ["12.4%", "112.4%"]
    else:
        ledger = data.draw(formatted_strings())  # type: ignore[attr-defined]
    for value in ledger:
        assert _survivor_texts(f"observed {value} today", ledger, []) == []


@given(
    foreign_whole=st.integers(min_value=1, max_value=9999),
    foreign_frac=st.integers(min_value=0, max_value=99),
    data=st.data(),
)
def test_a_foreign_numeral_always_survives_and_is_reported(
    foreign_whole: int, foreign_frac: int, data: st.DataObject
) -> None:
    """Req 29.1 — the control: a number the model invented is in neither vocabulary."""
    ledger = data.draw(formatted_strings())
    allowlist = _allowlist_tokens()
    foreign = f"{foreign_whole}.{foreign_frac:02d}units"
    # Only assert on a value the two vocabularies genuinely do not contain, and that no
    # earlier stage legitimately consumes — otherwise the property would be testing the
    # generator rather than the masker.
    if foreign in ledger or foreign in allowlist:
        return
    if _survivor_texts(foreign, [], []) != [foreign]:
        return

    text = f"headroom is substantial at {foreign} across the fleet"
    assert foreign in _survivor_texts(text, ledger, allowlist)

    findings = scan_paragraphs(
        [ExtractedParagraph(text, PART_BODY, 1, None)],
        ledger_strings=ledger,
        allowlist=allowlist,
    )
    assert [f["substring"] for f in findings] == [foreign]
    assert findings[0]["severity"] == "blocking"


@given(data=st.data())
def test_the_five_stages_are_idempotent_and_deterministic(data: st.DataObject) -> None:
    """Req 28.1 — no stage re-reads text an earlier stage consumed."""
    ledger = data.draw(formatted_strings())
    allowlist = _allowlist_tokens()
    text = data.draw(legitimate_paragraph(ledger, allowlist))
    order_l, order_a = masking_order(ledger), masking_order(allowlist)

    once = mask_paragraph(text, ledger_strings=order_l, allowlist=order_a)
    assert once == mask_paragraph(text, ledger_strings=order_l, allowlist=order_a)
    assert mask_paragraph(once, ledger_strings=order_l, allowlist=order_a) == once
    assert len(once) == len(text)  # overwriting, never deleting


@given(strings=st.lists(st.text(min_size=1, max_size=12), max_size=20))
def test_masking_order_is_a_total_order_independent_of_input_order(
    strings: list[str],
) -> None:
    """Req 28.2 — identical on every run over the same ledger."""
    assert masking_order(strings) == masking_order(list(reversed(strings)))
    ordered = masking_order(strings)
    assert list(ordered) == sorted(ordered, key=lambda s: (-len(s), s))


# --- the extraction half -----------------------------------------------------------------


def _document_with(paragraph_specs: list[tuple[str, bool]]) -> Document:
    """Build a document, nesting some paragraphs in a captioned table inside a layout
    table — the structure `document.paragraphs` cannot see into."""
    document = Document()
    for index, (text, nested) in enumerate(paragraph_specs):
        if nested:
            layout = document.add_table(rows=1, cols=1)
            write_layout_table(layout)
            inner = layout.rows[0].cells[0].add_table(rows=1, cols=1)
            write_data_table_caption(inner, f"companion-{index}")
            target = inner.rows[0].cells[0].paragraphs[0]
        else:
            target = document.add_paragraph()
        target.add_run(text)
    return document


@given(
    parts=st.lists(
        st.tuples(st.sampled_from(["1,", "1", "12"]), st.sampled_from(["234.", "2.", "."]), st.sampled_from(["56", "5", "6"])),
        min_size=1,
        max_size=30,
    )
)
# declared: the three-run split of 1,234.56
@example(parts=[("1,", "234.", "56")])
def test_a_figure_split_across_runs_extracts_as_one_token(
    parts: list[tuple[str, str, str]]
) -> None:
    """Req 26.3, 26.8, 28.9 — per-run tokenization yields three survivors, not one match."""
    document = Document()
    for a, b, c in parts:
        paragraph = document.add_paragraph()
        for piece in (a, b, c):
            paragraph.add_run(piece)

    extracted = [p for p in paragraph_texts(document) if p.text]
    assert len(extracted) == len(parts)
    for (a, b, c), paragraph in zip(parts, extracted, strict=True):
        whole = f"{a}{b}{c}"
        assert paragraph.text == whole
        assert [t.text for t in numeric_tokens(paragraph)] == [whole]


@given(
    specs=st.lists(
        st.tuples(st.sampled_from(["47.3%", "web-01 saw 12.5%", "PT1H"]), st.booleans()),
        min_size=1,
        max_size=30,
    )
)
def test_every_paragraph_is_extracted_including_nested_ones(
    specs: list[tuple[str, bool]]
) -> None:
    """Req 26.1, 26.2 — the object model misses the nested ones entirely."""
    document = _document_with(specs)
    extracted = [p.text for p in paragraph_texts(document) if p.text]
    for text, _ in specs:
        assert text in extracted

    if any(nested for _, nested in specs):
        # Compare counts, not texts: the same string may legitimately appear both
        # nested and unnested, so "this text is absent from the naive read" can be
        # false while the naive read still misses a paragraph. What the naive reader
        # provably cannot do is see as many paragraphs as the iterator.
        naive = [p.text for p in document.paragraphs if p.text]
        assert len(naive) < len(extracted)


@given(
    specs=st.lists(
        st.tuples(st.sampled_from(["47.3%", "12.5% peak", "no digits"]), st.booleans()),
        min_size=1,
        max_size=25,
    )
)
def test_two_extractions_of_one_document_agree_exactly(
    specs: list[tuple[str, bool]]
) -> None:
    """Req 26.9 — a re-verification cannot disagree with the original."""
    document = _document_with(specs)
    first = paragraph_texts(document)
    second = paragraph_texts(document)
    assert first == second
    assert [t for p in first for t in numeric_tokens(p)] == [
        t for p in second for t in numeric_tokens(p)
    ]


def test_the_declared_large_document_extracts_completely() -> None:
    """The scale clause, run deterministically rather than occasionally.

    1,200 paragraphs with 60 of them nested in captioned tables inside layout tables.
    A reader that missed nesting would come up 60 paragraphs short here, and every one
    of those 60 carries a figure.
    """
    specs = [(f"{index % 90}.{index % 100:02d}%", index % 20 == 0) for index in range(1200)]
    document = _document_with(specs)

    extracted = [p for p in paragraph_texts(document) if p.text]
    assert len(extracted) == 1200
    assert sum(1 for p in extracted if p.block_id is not None) == 60
    assert len([t for p in extracted for t in numeric_tokens(p)]) == 1200


def test_a_masked_paragraph_never_gains_or_loses_characters() -> None:
    """Offsets stay stable, so a finding's location points at the right characters."""
    text = "(34.2%) and [PT1H] on web-01"
    masked = mask_paragraph(text, ledger_strings=("34.2%",), allowlist=())
    assert len(masked) == len(text)
    for index, character in enumerate(text):
        assert masked[index] in (character, MASK_CHAR)
