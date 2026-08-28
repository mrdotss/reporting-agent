"""Tests for verify/toc.py — the table-of-contents verification gate."""

from __future__ import annotations

import io
import tempfile

import pytest

from reporting_agent.render.toc import (
    ADOPTED_APPROACH,
    TOC_APPROACH_NONE,
    TOC_APPROACH_TWO_PASS,
    should_emit_toc,
    toc_entries_from_document,
)
from reporting_agent.verify.toc import (
    TocPass,
    _build_proven_toc_numerals,
    _extract_headings,
    _first_character_page,
    _named_pages_from_pdf,
    _observed_pages,
    _toc_page_indices,
    check_toc,
)
from reporting_agent.verify.tokens import ExtractedParagraph


# ---------------------------------------------------------------------------
# Unit tests for internal helpers
# ---------------------------------------------------------------------------


class TestTocPageIndices:
    """Test that we identify TOC pages correctly.

    Every fixture here carries the **leader run** a rendered contents entry actually
    has — `Heading......4` — because that run is what distinguishes a contents entry
    from a page that merely mentions a heading beside a number. These fixtures used to
    be space-separated (`"Executive Overview 2"`), a shape no converter emits for a tab
    stop and one that made the tests agree with a rule loose enough to classify an
    ordinary table row as a table of contents.
    """

    def test_identifies_page_with_multiple_heading_entries(self):
        pages = (
            "Executive Overview.........2 Memory And Capacity.........4 Network.........6",
            "This is body content page with no headings",
            "More body content",
        )
        headings = ("Executive Overview", "Memory And Capacity", "Network")
        indices = _toc_page_indices(pages, headings)
        assert 0 in indices
        assert 1 not in indices
        assert 2 not in indices

    def test_page_with_only_one_heading_is_not_toc(self):
        pages = (
            "Executive Overview.........2",
            "Memory And Capacity section content",
        )
        headings = ("Executive Overview", "Memory And Capacity")
        indices = _toc_page_indices(pages, headings)
        assert len(indices) == 0

    def test_a_content_page_listing_headings_beside_numbers_is_not_toc(self):
        """The regression this rule exists for.

        A per-resource heading is named for its resource, and a resource's tables repeat
        that name beside its values — `prod-web-01  12.00`. Under the old rule (any
        whitespace, then digits) two such rows made the page a table of contents, which
        excluded it from the content search, and the heading that really lived on it was
        then reported as "not found on any content page". Two blocking findings on a
        correct document.
        """
        pages = (
            "Virtual Machine Utilization.........4 prod-db-02.........4 prod-web-01.........4",
            "PERIOD TIME VALUE prod-web-01 12.00 2026-07-01 prod-db-02 15.00 2026-07-01",
        )
        headings = ("Virtual Machine Utilization", "prod-db-02", "prod-web-01")
        indices = _toc_page_indices(pages, headings)
        assert indices == frozenset({0}), (
            "the contents page is page 0; page 1 is a table that happens to repeat two "
            "of the headings beside numbers"
        )


class TestNamedPagesFromPdf:
    def test_extracts_page_numbers_from_toc_entries(self):
        pages = (
            "Executive Overview.........2 Memory And Capacity.........4",
            "Body page content",
        )
        headings = ("Executive Overview", "Memory And Capacity")
        toc_indices = frozenset({0})
        named = _named_pages_from_pdf(pages, headings, toc_indices)
        assert named == {"Executive Overview": 2, "Memory And Capacity": 4}

    def test_a_number_merely_following_a_heading_is_not_read_as_its_page(self):
        """`_named_pages_from_pdf` reads the same shape `_toc_page_indices` selects by.

        If it did not, a page classified by the strict rule could still have its numbers
        read by a loose one, and the gate would compare a real page against a value it
        picked up from a table cell.
        """
        pages = ("prod-web-01 12 prod-db-02 15",)
        headings = ("prod-web-01", "prod-db-02")
        named = _named_pages_from_pdf(pages, headings, frozenset({0}))
        assert named == {}

    def test_returns_empty_when_no_toc_pages(self):
        pages = ("Body content only",)
        headings = ("A Heading",)
        named = _named_pages_from_pdf(pages, headings, frozenset())
        assert named == {}


class TestObservedPages:
    def test_finds_headings_on_non_toc_pages(self):
        pages = (
            "Executive Overview 2 Memory 4",  # TOC page
            "Executive Overview section body",  # Content page 2
            "Memory section body",  # Content page 3
        )
        headings = ("Executive Overview", "Memory")
        toc_indices = frozenset({0})
        observed = _observed_pages(pages, headings, toc_indices)
        assert observed == {"Executive Overview": 2, "Memory": 3}


class TestFirstCharacterPage:
    def test_skips_toc_pages(self):
        pages = (
            "Heading A 3",  # TOC page (index 0)
            "Heading A body text",  # Content page 2
        )
        result = _first_character_page(pages, "Heading A", frozenset({0}))
        assert result == 2

    def test_returns_none_when_not_found(self):
        pages = ("Some other text",)
        result = _first_character_page(pages, "Missing Heading", frozenset())
        assert result is None


class TestBuildProvenTocNumerals:
    def test_builds_mapping_for_correct_entries(self):
        paragraphs = [
            ExtractedParagraph(
                text="Executive Overview 2",
                part="body",
                ordinal=1,
                block_id=None,
            ),
            ExtractedParagraph(
                text="Memory 4",
                part="body",
                ordinal=2,
                block_id=None,
            ),
            ExtractedParagraph(
                text="Some other text",
                part="body",
                ordinal=3,
                block_id=None,
            ),
        ]
        named = {"Executive Overview": 2, "Memory": 4}
        observed = {"Executive Overview": 2, "Memory": 4}

        proven = _build_proven_toc_numerals(paragraphs, named, observed)
        assert 1 in proven
        assert "2" in proven[1]
        assert 2 in proven
        assert "4" in proven[2]
        assert 3 not in proven

    def test_excludes_mismatched_entries(self):
        paragraphs = [
            ExtractedParagraph(
                text="Executive Overview 2",
                part="body",
                ordinal=1,
                block_id=None,
            ),
        ]
        named = {"Executive Overview": 2}
        observed = {"Executive Overview": 5}  # Mismatch!

        proven = _build_proven_toc_numerals(paragraphs, named, observed)
        assert proven == {}

    def test_empty_when_no_headings(self):
        proven = _build_proven_toc_numerals([], {}, {})
        assert proven == {}


# ---------------------------------------------------------------------------
# Integration: check_toc with synthetic PDF
# ---------------------------------------------------------------------------


class TestCheckTocWithUnreadablePdf:
    """When PDF is not parseable, the gate returns zero findings gracefully."""

    def test_returns_empty_on_unparseable_pdf(self):
        fake_pdf = b"%PDF-1.7 not a real pdf"
        paragraphs: list[ExtractedParagraph] = []

        # Create a minimal docx with a heading for the function to read
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Test Heading", style="Heading 1")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        opened = DocxDocument(buf)

        result = check_toc(fake_pdf, paragraphs=paragraphs, document=opened)
        assert result.entries_checked == 0
        assert result.findings == ()
        assert result.proven_toc_numerals == {}


# ---------------------------------------------------------------------------
# Adopted approach
# ---------------------------------------------------------------------------


class TestAdoptedApproach:
    def test_adopted_approach_is_two_pass(self):
        assert ADOPTED_APPROACH == TOC_APPROACH_TWO_PASS

    def test_should_emit_toc_is_true(self):
        assert should_emit_toc() is True


# ---------------------------------------------------------------------------
# HTML TOC emission
# ---------------------------------------------------------------------------


class TestEmitTocHtml:
    def test_emits_heading_list_with_no_page_numbers(self):
        from reporting_agent.render.html import emit_toc_html
        from reporting_agent.compile.ast import Document, Paragraph, Text, FigurePath

        doc = Document(
            blocks=(
                Paragraph(
                    path=FigurePath("h1"),
                    style="Heading 1",
                    inlines=(Text(path=FigurePath("h1/t0"), text="Executive Overview"),),
                ),
                Paragraph(
                    path=FigurePath("h2"),
                    style="Heading 2",
                    inlines=(Text(path=FigurePath("h2/t0"), text="Sub Section"),),
                ),
                Paragraph(
                    path=FigurePath("p1"),
                    style="Body Text",
                    inlines=(Text(path=FigurePath("p1/t0"), text="Body paragraph"),),
                ),
            )
        )
        result = emit_toc_html(doc)
        assert "Executive Overview" in result
        assert "Sub Section" in result
        assert "Body paragraph" not in result
        # No page numbers
        assert "data-page" not in result
        # Has the nav structure
        assert '<nav class="rpt-toc"' in result
        assert '<ol class="rpt-toc-list">' in result

    def test_returns_empty_for_no_headings(self):
        from reporting_agent.render.html import emit_toc_html
        from reporting_agent.compile.ast import Document, Paragraph, Text, FigurePath

        doc = Document(
            blocks=(
                Paragraph(
                    path=FigurePath("p1"),
                    style="Body Text",
                    inlines=(Text(path=FigurePath("p1/t0"), text="Only body text"),),
                ),
            )
        )
        result = emit_toc_html(doc)
        assert result == "" or "rpt-toc" not in result

    def test_excludes_heading_4_and_deeper(self):
        from reporting_agent.render.html import emit_toc_html
        from reporting_agent.compile.ast import Document, Paragraph, Text, FigurePath

        doc = Document(
            blocks=(
                Paragraph(
                    path=FigurePath("h1"),
                    style="Heading 1",
                    inlines=(Text(path=FigurePath("h1/t0"), text="Level 1"),),
                ),
                Paragraph(
                    path=FigurePath("h4"),
                    style="Heading 4",
                    inlines=(Text(path=FigurePath("h4/t0"), text="Level 4 Deep"),),
                ),
            )
        )
        result = emit_toc_html(doc)
        assert "Level 1" in result
        assert "Level 4 Deep" not in result


# ---------------------------------------------------------------------------
# HTML notice resolves through messages (not a raw catalog id)
# ---------------------------------------------------------------------------


class TestHtmlNoticeResolvesText:
    """The empty-scope notice in a chart emits resolved text, not the catalog id."""

    def test_notice_is_resolved_text_not_id(self):
        from reporting_agent.render.html import emit_html, NOTICE_ROW_CLASS
        from reporting_agent.compile.ast import (
            Chart, Document, FigurePath, Series,
        )
        from reporting_agent.compile.blocks.base import EMPTY_SCOPE_TEXT

        # Build a Messages instance that resolves the id.
        from reporting_agent.compile.messages import Messages
        messages = Messages(language="en", _table={EMPTY_SCOPE_TEXT: "No resources in scope"})

        doc = Document(
            blocks=(
                Chart(
                    path=FigurePath("chart:0"),
                    chart_type="line",
                    title="Test Chart",
                    unit="%",
                    encoding="sequential",
                    series=(Series(path=FigurePath("chart:0.0"), key="s1", label="Series 1", points=()),),
                ),
            )
        )
        outcome = emit_html(doc, messages=messages)
        # The resolved text appears, not the catalog id.
        assert "No resources in scope" in outcome.html
        assert EMPTY_SCOPE_TEXT not in outcome.html  # The raw id must NOT appear.

    def test_period_label_emitted_when_present(self):
        from reporting_agent.render.html import emit_html
        from reporting_agent.compile.ast import (
            Chart, Document, FigurePath, Series,
        )
        from reporting_agent.compile.messages import Messages
        from reporting_agent.compile.blocks.base import EMPTY_SCOPE_TEXT

        # Chart with period_label and no points — exercises the period_label path
        # without needing a valid Figure (which requires a compiling_against context).
        messages = Messages(
            language="en",
            _table={EMPTY_SCOPE_TEXT: "No resources in scope"},
        )
        doc = Document(
            blocks=(
                Chart(
                    path=FigurePath("chart:0"),
                    chart_type="line",
                    title="CPU",
                    unit="%",
                    encoding="sequential",
                    period_label="Jul 2026 – Aug 2026 (UTC+7)",
                    series=(Series(path=FigurePath("chart:0.0"), key="s1", label="CPU", points=()),),
                ),
            )
        )
        outcome = emit_html(doc, messages=messages)
        assert "Jul 2026" in outcome.html
        assert "rpt-chart-period" in outcome.html
        # Period label is before the indication / series
        idx_period = outcome.html.index("rpt-chart-period")
        idx_series = outcome.html.index("rpt-series-set")
        assert idx_period < idx_series


# ---------------------------------------------------------------------------
# Masking integration: proven_toc_numerals
# ---------------------------------------------------------------------------


class TestProvenTocNumeralsMasking:
    """Test that proven_toc_numerals suppresses page numbers in masking."""

    def test_admitted_numeral_is_not_a_finding(self):
        from reporting_agent.verify.masking import scan_paragraphs

        paragraphs = [
            ExtractedParagraph(
                text="Executive Overview 7",
                part="body",
                ordinal=5,
                block_id=None,
            )
        ]
        # Without proven_toc_numerals, "7" survives
        findings_without = scan_paragraphs(
            paragraphs, ledger_strings=[], allowlist=[]
        )
        assert len(findings_without) == 1

        # With proven_toc_numerals admitting "7" in paragraph 5
        findings_with = scan_paragraphs(
            paragraphs,
            ledger_strings=[],
            allowlist=[],
            proven_toc_numerals={5: frozenset({"7"})},
        )
        assert len(findings_with) == 0

    def test_numeral_not_proven_for_different_paragraph_is_still_finding(self):
        from reporting_agent.verify.masking import scan_paragraphs

        paragraphs = [
            ExtractedParagraph(
                text="Some text with 7 in it",
                part="body",
                ordinal=10,
                block_id=None,
            )
        ]
        # Proven for paragraph 5, but this is paragraph 10
        findings = scan_paragraphs(
            paragraphs,
            ledger_strings=[],
            allowlist=[],
            proven_toc_numerals={5: frozenset({"7"})},
        )
        assert len(findings) == 1

    def test_uncompared_numeral_stays_unmatched(self):
        """A numeral in a page-number position the verifier compared to nothing stays
        unmatched_prose_token (Req 14.12)."""
        from reporting_agent.verify.masking import scan_paragraphs

        paragraphs = [
            ExtractedParagraph(
                text="Executive Overview 7",
                part="body",
                ordinal=5,
                block_id=None,
            )
        ]
        # proven_toc_numerals admits "4" but not "7" for paragraph 5
        findings = scan_paragraphs(
            paragraphs,
            ledger_strings=[],
            allowlist=[],
            proven_toc_numerals={5: frozenset({"4"})},
        )
        assert len(findings) == 1
