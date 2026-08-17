"""Anchored cell equality (Req 27).

The load-bearing test here is
:func:`test_a_transposition_fails_the_anchored_pass_and_passes_a_containment_check`. It does
not merely assert that a transposition is caught — it asserts, in the same test and over the
same document, that the check this pass replaced records **nothing**. Without that second half
the first is a test that would still pass against a verifier quietly doing containment; with
it, the test is an executable statement of why this module exists.
"""

from __future__ import annotations

import io
from collections.abc import Sequence
from typing import Final

from docx import Document as open_docx
from docx.oxml.ns import qn

import definition_factory as df
import snapshot_factory as sf
from reporting_agent.compile.ast import compiling_against
from reporting_agent.compile.blocks import compile_document
from reporting_agent.compile.blocks.base import (
    EMPTY_SCOPE_TEXT,
    NOTICE_COLUMN_HEADER,
    DesignSettings,
)
from reporting_agent.compile.figures import (
    ANCHOR_TABLE,
    BlockCursor,
    FigureLedger,
    TableAnchor,
)
from reporting_agent.compile.snapshot_view import build_snapshot_view
from reporting_agent.render.anchors import write_data_table_caption, write_layout_table
from reporting_agent.render.docx import render_document
from reporting_agent.verify.anchors import (
    AnchorPass,
    TableGrid,
    check_tables,
    containment_discrepancies,
    read_grids,
)
from reporting_agent.verify.findings import (
    FINDING_DUPLICATE_TABLE_ANCHOR,
    FINDING_TABLE_ANCHOR_MISSING,
    FINDING_TABLE_ANCHOR_UNEXPECTED,
    FINDING_TABLE_CELL_MISMATCH,
    FINDING_TABLE_COLUMN_UNRESOLVED,
    FINDING_TABLE_ROW_UNRESOLVED,
    FINDING_TABLE_ROWS_ABSENT,
    SEVERITY_BLOCKING,
)

_W_T: Final[str] = qn("w:t")
_W_TBL: Final[str] = qn("w:tbl")
_W_TBLPR: Final[str] = qn("w:tblPr")
_W_TBLCAPTION: Final[str] = qn("w:tblCaption")
_W_VAL: Final[str] = qn("w:val")

DEFAULT_DESIGN: Final[dict[str, object]] = {
    "preset": "editorial",
    "accent_color": "#1f6f78",
    "density": "normal",
    "table_style": "hairline",
    "number_format": {"decimal_places": 2, "group_thousands": True},
    "cover_page": False,
    "logo": None,
    "page_size": "A4",
}


# --------------------------------------------------------------------------- #
# Rendering a real document, then reading it back
# --------------------------------------------------------------------------- #


def render(blocks: list[dict], **kwargs: object):
    """Compile and render, returning `(compiled, docx bytes)`."""
    view = build_snapshot_view(sf.two_vm_snapshot())
    design = {**DEFAULT_DESIGN, **(kwargs.pop("design", None) or {})}
    compiled = compile_document(
        df.definition(blocks, design=design, **kwargs), view=view  # type: ignore[arg-type]
    )
    outcome = render_document(
        compiled.document,
        ledger=compiled.ledger,
        design=DesignSettings.from_plain(design),
    )
    return compiled, outcome.docx_bytes


def grids_of(payload: bytes) -> tuple[TableGrid, ...]:
    return read_grids(open_docx(io.BytesIO(payload)))


def resource_table_blocks() -> list[dict]:
    return [df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]})]


def types_of(outcome: AnchorPass) -> list[str]:
    return [str(finding["type"]) for finding in outcome.findings]


# --------------------------------------------------------------------------- #
# The reader
# --------------------------------------------------------------------------- #


def test_a_grid_separates_the_header_row_from_the_data_rows() -> None:
    """Req 27.10's "data row" is every row other than the header row."""
    _, payload = render(resource_table_blocks())
    grids = grids_of(payload)

    assert grids
    grid = grids[0]
    assert grid.headers
    assert grid.rows
    # Two VMs in the fixture snapshot, so two data rows and one header above them.
    assert len(grid.rows) == 2
    assert all(len(row) == len(grid.headers) for row in grid.rows)


def test_a_grid_reads_direct_child_rows_only_so_a_nested_table_is_not_spliced_in() -> None:
    """The one place this module deliberately does *not* descend.

    A chart's companion table nested inside a `row` block's layout cell is a table in its own
    right. An `iter()`-based row walk would splice its rows into the outer table's row list,
    and every anchor of the outer table would then resolve against a grid belonging to two
    tables.
    """
    document = open_docx()
    outer = document.add_table(rows=2, cols=1)
    write_data_table_caption(outer, "tbl:outer:0")
    outer.rows[0].cells[0].text = "Resource"
    outer.rows[1].cells[0].text = "web-01"

    nested = outer.rows[1].cells[0].add_table(rows=2, cols=1)
    write_data_table_caption(nested, "tbl:nested:0")
    nested.rows[0].cells[0].text = "Point"
    nested.rows[1].cells[0].text = "2026-07-01"

    grids = {grid.identity: grid for grid in read_grids(document)}

    assert set(grids) == {"tbl:outer:0", "tbl:nested:0"}
    assert grids["tbl:outer:0"].row_keys == ("web-01",)
    assert grids["tbl:nested:0"].row_keys == ("2026-07-01",)


def test_a_layout_table_is_absent_from_the_grids_by_construction() -> None:
    """Req 26.5 — no caption, no grid, whatever the table holds."""
    document = open_docx()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    layout.rows[0].cells[0].text = "99.9%"

    assert read_grids(document) == ()


def test_a_cell_split_across_runs_reads_back_as_one_string() -> None:
    """Req 26.4 — the concatenation, at cell level. Word splits `1,234.56` routinely."""
    document = open_docx()
    table = document.add_table(rows=2, cols=1)
    write_data_table_caption(table, "tbl:x:0")
    table.rows[0].cells[0].text = "Value"
    paragraph = table.rows[1].cells[0].paragraphs[0]
    for piece in ("1,", "234", ".56"):
        paragraph.add_run(piece)

    assert read_grids(document)[0].rows == (("1,234.56",),)


# --------------------------------------------------------------------------- #
# The clean case
# --------------------------------------------------------------------------- #


def test_an_unmutated_render_records_no_finding_and_matches_every_anchor() -> None:
    """Property 3.1, as a unit test over a real compiled document."""
    compiled, payload = render(resource_table_blocks())
    outcome = check_tables(compiled.ledger, grids_of(payload))

    assert outcome.findings == ()
    assert outcome.anchors_checked > 0
    assert outcome.tables_resolved > 0
    assert len(outcome.matched) == outcome.anchors_checked
    assert outcome.faulted == frozenset()
    assert outcome.blocking_identities == frozenset()


def test_the_counts_distinguish_a_pass_that_checked_nothing() -> None:
    """Req 27.13 — a pass over zero anchors is a pass, and must not read like a full one."""
    outcome = check_tables(FigureLedger(), ())

    assert outcome.findings == ()
    assert outcome.anchors_checked == 0
    assert outcome.tables_resolved == 0


def test_every_anchor_the_renderer_records_carries_a_resolvable_column_key() -> None:
    """The defect this pass exposed: an anchor's column key is the **header text**.

    A chart's companion table declares `Column(key="value", header="Value")`, so recording
    `Column.key` would record a string that appears nowhere in the emitted grid — and every
    chart figure in an otherwise correct document would fail as `table_column_unresolved`.
    """
    compiled, payload = render(
        [df.block("ts", "timeseries_chart", {"metrics": [df.CPU_AVG]})]
    )
    grids = {grid.identity: grid for grid in grids_of(payload)}

    anchors = compiled.ledger.anchors()
    assert anchors
    for anchor in anchors.values():
        assert anchor.column_key in grids[anchor.anchor_id].headers

    assert check_tables(compiled.ledger, grids.values()).findings == ()


# --------------------------------------------------------------------------- #
# The argument for the whole module
# --------------------------------------------------------------------------- #


def _transpose_two_columns(payload: bytes, identity: str, left: int, right: int) -> bytes:
    """Swap two columns' **values** across every data row, leaving the headers in place.

    Every `formatted` string is still present in the document afterwards — attached to the
    wrong things — which is precisely the state a containment check cannot see.
    """
    document = open_docx(io.BytesIO(payload))
    for table in document.element.body.iter(_W_TBL):
        properties = table.find(_W_TBLPR)
        caption = properties.find(_W_TBLCAPTION) if properties is not None else None
        if caption is None or caption.get(_W_VAL) != identity:
            continue
        rows = table.findall(qn("w:tr"))
        for row in rows[1:]:
            cells = row.findall(qn("w:tc"))
            a = [node.text for node in cells[left].iter(_W_T)]
            b = [node.text for node in cells[right].iter(_W_T)]
            _write_cell(cells[left], "".join(piece or "" for piece in b))
            _write_cell(cells[right], "".join(piece or "" for piece in a))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _write_cell(cell_element: object, text: str) -> None:
    """Overwrite a cell's text, keeping its first `w:t` and blanking the rest."""
    nodes = list(cell_element.iter(_W_T))  # type: ignore[attr-defined]
    if not nodes:
        return
    nodes[0].text = text
    for node in nodes[1:]:
        node.text = ""


def test_a_transposition_fails_the_anchored_pass_and_passes_a_containment_check() -> None:
    """Req 27.3, Property 3.2, negative test 44.3 — the whole argument, executable.

    Both halves are required. The first asserts this pass catches a document in which every
    VM's average and peak are swapped. The second asserts that the check this pass replaced
    records **zero** discrepancies over the very same document — so this test fails against a
    verifier that quietly checks containment, which is the only way the assertion is worth
    anything.
    """
    compiled, payload = render(resource_table_blocks())
    identity = next(iter(compiled.ledger.table_identities()))
    clean = grids_of(payload)
    columns = clean[0].headers
    left, right = len(columns) - 2, len(columns) - 1

    mutated = grids_of(_transpose_two_columns(payload, identity, left, right))
    outcome = check_tables(compiled.ledger, mutated)

    mismatches = [f for f in outcome.findings if f["type"] == FINDING_TABLE_CELL_MISMATCH]
    assert mismatches, "a transposition must record at least one table_cell_mismatch"
    assert all(f["severity"] == SEVERITY_BLOCKING for f in mismatches)

    assert containment_discrepancies(compiled.ledger, mutated) == ()


def test_a_column_reordering_that_carries_its_header_verifies_cleanly() -> None:
    """Property 3.3 — the case a positional implementation gets backwards.

    Moving a column *with* its header changes nothing about which cell holds which figure, so
    the document is correct and must verify. A verifier resolving by ordinal fails here and
    passes the transposition above; this pass does the opposite, which is the point.
    """
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    permuted = tuple(
        TableGrid(
            identity=g.identity,
            ordinal=g.ordinal,
            headers=_rotate(g.headers),
            rows=tuple(_rotate(row) for row in g.rows),
        )
        for g in grids
    )

    assert check_tables(compiled.ledger, permuted).findings == ()


def _rotate(values: Sequence[str]) -> tuple[str, ...]:
    """Move every column after the key column one place left, wrapping.

    The key column stays put: it is the row key, and moving it would be a *row* permutation
    rather than a column one.
    """
    if len(values) < 3:
        return tuple(values)
    head, tail = values[0], list(values[1:])
    return (head, *tail[1:], tail[0])


def test_a_row_reordering_that_carries_its_key_verifies_cleanly() -> None:
    """Property 3.3's other half. Rows resolve by key, so their order is not information."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    reversed_rows = tuple(
        TableGrid(
            identity=g.identity, ordinal=g.ordinal, headers=g.headers, rows=g.rows[::-1]
        )
        for g in grids
    )

    assert check_tables(compiled.ledger, reversed_rows).findings == ()


# --------------------------------------------------------------------------- #
# Each resolution failure is its own finding
# --------------------------------------------------------------------------- #


def test_a_mutated_cell_names_the_table_the_row_and_the_column() -> None:
    """Req 27.8, Property 3.4 — a finding a reader can act on without opening the document."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    mutated = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=target.headers,
            rows=((*target.rows[0][:-1], "999.9%"), *target.rows[1:]),
        ),
        *grids[1:],
    )

    findings = [
        f
        for f in check_tables(compiled.ledger, mutated).findings
        if f["type"] == FINDING_TABLE_CELL_MISMATCH
    ]

    assert len(findings) == 1
    finding = findings[0]
    assert finding["table_id"] == target.identity
    assert finding["row_key"] == target.rows[0][0]
    assert finding["column_key"] == target.headers[-1]
    assert finding["observed"] == "999.9%"
    assert finding["expected"] == target.rows[0][-1]
    assert finding["ast_path"]


def test_a_removed_caption_is_table_anchor_missing() -> None:
    """Property 3.6 — the identity is the only thing tying an anchor to a table."""
    compiled, _ = render(resource_table_blocks())
    outcome = check_tables(compiled.ledger, ())

    assert types_of(outcome) == [FINDING_TABLE_ANCHOR_MISSING] * outcome.anchors_checked
    assert outcome.matched == frozenset()
    assert len(outcome.faulted) == outcome.anchors_checked
    assert compiled.ledger.table_identities()


def test_a_column_key_matching_two_columns_is_unresolved_and_names_the_count() -> None:
    """Req 27.6 — two matches have no single cell to compare, so neither is chosen."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    duplicated = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=(*target.headers, target.headers[-1]),
            rows=tuple((*row, row[-1]) for row in target.rows),
        ),
        *grids[1:],
    )

    findings = [
        f
        for f in check_tables(compiled.ledger, duplicated).findings
        if f["type"] == FINDING_TABLE_COLUMN_UNRESOLVED
    ]

    assert findings
    assert all(f["match_count"] == 2 for f in findings)
    assert all(f["column_key"] == target.headers[-1] for f in findings)


def test_a_missing_column_is_unresolved_with_a_zero_match_count() -> None:
    """Req 27.6's other half — absent and ambiguous are one finding type, two counts."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    dropped = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=(*target.headers[:-1], "Something Else"),
            rows=target.rows,
        ),
        *grids[1:],
    )

    findings = [
        f
        for f in check_tables(compiled.ledger, dropped).findings
        if f["type"] == FINDING_TABLE_COLUMN_UNRESOLVED
    ]

    assert findings
    assert all(f["match_count"] == 0 for f in findings)


def test_a_row_key_matching_two_rows_is_unresolved_and_names_the_count() -> None:
    """Req 27.7 — the row-side mirror of the column rule."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    duplicated = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=target.headers,
            rows=(*target.rows, target.rows[0]),
        ),
        *grids[1:],
    )

    findings = [
        f
        for f in check_tables(compiled.ledger, duplicated).findings
        if f["type"] == FINDING_TABLE_ROW_UNRESOLVED
    ]

    assert findings
    assert all(f["match_count"] == 2 for f in findings)
    assert all(f["row_key"] == target.rows[0][0] for f in findings)


def test_a_missing_row_is_unresolved_rather_than_a_mismatch() -> None:
    """A dropped row is not a wrong number; conflating the two would mislabel the defect."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    dropped = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=target.headers,
            rows=target.rows[1:],
        ),
        *grids[1:],
    )

    outcome = check_tables(compiled.ledger, dropped)

    assert FINDING_TABLE_ROW_UNRESOLVED in types_of(outcome)
    assert FINDING_TABLE_CELL_MISMATCH not in types_of(outcome)


def test_two_tables_sharing_an_identity_is_one_duplicate_finding() -> None:
    """Req 27.5. One collision is one defect, whatever it does to the anchors under it."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    doubled = (*grids, TableGrid(identity=grids[0].identity, ordinal=99, headers=(), rows=()))

    outcome = check_tables(compiled.ledger, doubled)

    duplicates = [f for f in outcome.findings if f["type"] == FINDING_DUPLICATE_TABLE_ANCHOR]
    assert len(duplicates) == 1
    assert duplicates[0]["table_id"] == grids[0].identity
    # And every anchor under it reports that it could not be resolved.
    assert FINDING_TABLE_ANCHOR_MISSING in types_of(outcome)


def test_a_captioned_table_the_ledger_never_registered_is_unexpected() -> None:
    """Req 27.5 — a caption the renderer did not write means the document is not the one
    this ledger describes."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    intruder = (
        *grids,
        TableGrid(identity="tbl:smuggled:0", ordinal=99, headers=("A",), rows=(("1",),)),
    )

    outcome = check_tables(compiled.ledger, intruder)

    unexpected = [
        f for f in outcome.findings if f["type"] == FINDING_TABLE_ANCHOR_UNEXPECTED
    ]
    assert [f["table_id"] for f in unexpected] == ["tbl:smuggled:0"]


# --------------------------------------------------------------------------- #
# Req 27.10 / 27.11 — the distinction the two negative tests exist to keep apart
# --------------------------------------------------------------------------- #


def test_a_table_that_rendered_no_rows_over_a_non_empty_scope_fails() -> None:
    """Req 27.10, negative test 44.4 — the block that silently dropped its rows."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    emptied = (
        TableGrid(
            identity=target.identity, ordinal=target.ordinal, headers=target.headers, rows=()
        ),
        *grids[1:],
    )

    outcome = check_tables(
        compiled.ledger, emptied, scope_counts={target.identity: 2}
    )

    absent = [f for f in outcome.findings if f["type"] == FINDING_TABLE_ROWS_ABSENT]
    assert len(absent) == 1
    assert "2 resource(s)" in absent[0]["message"]
    assert "0 data rows" in absent[0]["message"]


def test_an_empty_scope_notice_row_with_no_anchors_records_nothing() -> None:
    """Req 27.11, negative test 44.5 — the block whose scope legitimately matched nothing.

    Both halves of the exemption are exercised: the notice is the table's only data row, and
    the table carries no anchors. A verifier that checked only the row count would fail this
    correct document, and one that checked only the anchors would wave through a table that
    rendered a notice *and* figures.
    """
    ledger = FigureLedger()
    ledger.register_table("empty:0", "tbl:empty:0")  # type: ignore[arg-type]
    notice = TableGrid(
        identity="tbl:empty:0",
        ordinal=1,
        headers=(NOTICE_COLUMN_HEADER,),
        rows=((EMPTY_SCOPE_TEXT,),),
    )

    outcome = check_tables(ledger, (notice,), scope_counts={"tbl:empty:0": 7})

    assert outcome.findings == ()


def test_a_table_with_no_recorded_scope_count_records_no_rows_absent() -> None:
    """An unknown scope count is not a finding about the document."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    emptied = (
        TableGrid(
            identity=target.identity, ordinal=target.ordinal, headers=target.headers, rows=()
        ),
        *grids[1:],
    )

    outcome = check_tables(compiled.ledger, emptied)

    assert FINDING_TABLE_ROWS_ABSENT not in types_of(outcome)


def test_a_zero_scope_count_records_no_rows_absent() -> None:
    """Zero rows over zero resources is the correct document, not a dropped section."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    emptied = (
        TableGrid(
            identity=target.identity, ordinal=target.ordinal, headers=target.headers, rows=()
        ),
        *grids[1:],
    )

    outcome = check_tables(compiled.ledger, emptied, scope_counts={target.identity: 0})

    assert FINDING_TABLE_ROWS_ABSENT not in types_of(outcome)


# --------------------------------------------------------------------------- #
# Determinism and the reported sets
# --------------------------------------------------------------------------- #


def test_findings_are_ordered_by_table_then_row_then_column() -> None:
    """Req 27.14 — two verifications of one document produce identical results."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    wrecked = tuple(
        TableGrid(
            identity=g.identity,
            ordinal=g.ordinal,
            headers=g.headers,
            rows=tuple((row[0], *("0" for _ in row[1:])) for row in g.rows),
        )
        for g in grids
    )

    first = check_tables(compiled.ledger, wrecked).findings
    second = check_tables(compiled.ledger, tuple(reversed(wrecked))).findings

    assert first == second
    keys = [(f.get("table_id"), f.get("row_key"), f.get("column_key")) for f in first]
    assert keys == sorted(keys, key=lambda k: tuple(str(part or "") for part in k))


def test_the_faulted_set_holds_exactly_the_paths_that_recorded_a_finding() -> None:
    """Req 29.8 reads this: an entry unrendered *because* its anchor failed records no
    second finding, so one rendering defect stays one finding."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    mutated = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=target.headers,
            rows=((*target.rows[0][:-1], "0.0%"), *target.rows[1:]),
        ),
        *grids[1:],
    )

    outcome = check_tables(compiled.ledger, mutated)

    assert len(outcome.faulted) == 1
    assert outcome.matched.isdisjoint(outcome.faulted)
    assert len(outcome.matched) + len(outcome.faulted) == outcome.anchors_checked
    assert outcome.blocking_identities == frozenset({target.identity})


def test_an_anchor_carrying_no_row_or_column_is_not_counted_as_checked() -> None:
    """`BlockCursor.anchor_table` records a bare identity; the emitter completes the triple
    for every cell it writes. An anchor still bare is a figure in no cell, and this pass has
    nothing to resolve for it — backward completeness picks it up through the prose path."""
    ledger = FigureLedger()
    view = build_snapshot_view(sf.two_vm_snapshot())
    with compiling_against(view):
        cursor = BlockCursor(block_id="kpi", ledger=ledger).child("nodes", 0)
        figure = cursor.figure(view.values()[0])
    ledger.register_table(figure.path, "tbl:kpi:0")
    ledger.record_anchor(figure.path, TableAnchor(kind=ANCHOR_TABLE, anchor_id="tbl:kpi:0"))

    outcome = check_tables(ledger, ())

    assert ledger.anchors()
    assert outcome.anchors_checked == 0
    assert outcome.findings == ()


def test_containment_reports_a_string_that_is_genuinely_absent() -> None:
    """The containment helper is not a no-op: it finds a missing string, which is what makes
    its silence on a transposition meaningful."""
    compiled, payload = render(resource_table_blocks())
    grids = grids_of(payload)
    target = grids[0]
    wiped = (
        TableGrid(
            identity=target.identity,
            ordinal=target.ordinal,
            headers=target.headers,
            rows=tuple((row[0], *("" for _ in row[1:])) for row in target.rows),
        ),
        *grids[1:],
    )

    assert containment_discrepancies(compiled.ledger, wiped)
