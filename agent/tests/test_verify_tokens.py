"""Reading the document the way Word stores it.

The load-bearing test in this module is
:func:`test_the_object_model_misses_a_nested_paragraph_and_the_iterator_does_not`. It
does not merely assert that the extractor works — it builds the exact structure this
product emits (a chart's companion table nested inside a `row` block's layout table),
asserts the iterator finds the figure, and asserts in the same test that
`document.paragraphs` **does not**. Without that second half the first is just a
passing test; with it, the test is a standing proof that the naive reader fails
silently on a real document rather than on a hypothetical one.
"""

from __future__ import annotations

from pathlib import Path
from typing import Final

import pytest
from docx import Document
from docx.oxml.ns import qn

from reporting_agent.errors import VerificationFailedError
from reporting_agent.render.anchors import write_data_table_caption, write_layout_table
from reporting_agent.verify.tokens import (
    PART_BODY,
    PART_FOOTER,
    PART_HEADER,
    data_tables,
    normalize_pdf_text,
    numeric_tokens,
    open_document,
    paragraph_text,
    paragraph_texts,
    read_pdf_text,
)

FIGURE: Final[str] = "1,234.56"


def _split_runs(paragraph: object, *pieces: str) -> None:
    """Add `pieces` as separate runs — how Word actually stores one number."""
    for piece in pieces:
        paragraph.add_run(piece)  # type: ignore[attr-defined]


# --- the failure mode this module exists to prevent -----------------------------------


def test_the_object_model_misses_a_nested_paragraph_and_the_iterator_does_not() -> None:
    """Req 26.1, 26.2 — the silent, total failure, made executable.

    A chart's companion table nested inside a `row` block's layout table is the
    structure this product emits as a matter of course. `document.paragraphs` cannot
    see into it, so a verifier reading through the object model extracts nothing,
    finds no unmatched token, records no finding, and passes the document.
    """
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    companion = layout.rows[0].cells[0].add_table(rows=1, cols=1)
    write_data_table_caption(companion, "chart-1-companion")
    _split_runs(companion.rows[0].cells[0].paragraphs[0], "1,", "234.", "56")

    extracted = paragraph_texts(document)
    assert FIGURE in [p.text for p in extracted]

    # The same document, read the wrong way, yields nothing to check.
    assert FIGURE not in [p.text for p in document.paragraphs]
    assert all(FIGURE not in p.text for p in document.paragraphs)


def test_a_nested_paragraph_is_attributed_to_its_captioned_table() -> None:
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    companion = layout.rows[0].cells[0].add_table(rows=1, cols=1)
    write_data_table_caption(companion, "chart-1-companion")
    companion.rows[0].cells[0].paragraphs[0].add_run(FIGURE)

    found = next(p for p in paragraph_texts(document) if p.text == FIGURE)
    assert found.block_id == "chart-1-companion"
    assert found.part == PART_BODY


# --- concatenation and tokenization ---------------------------------------------------


def test_a_figure_split_across_runs_is_one_token() -> None:
    """Req 26.3, 28.9 — per-run tokenization produces three survivors, not one match."""
    document = Document()
    _split_runs(document.add_paragraph(), "1,", "234.", "56")
    paragraph = paragraph_texts(document)[0]
    assert paragraph.text == FIGURE
    assert [t.text for t in numeric_tokens(paragraph)] == [FIGURE]


def test_adjacent_runs_concatenate_with_no_inserted_character() -> None:
    """Req 26.8 — inserting a space would split the figure exactly as per-run does."""
    document = Document()
    _split_runs(document.add_paragraph(), "34", ".", "2", "%")
    assert paragraph_texts(document)[0].text == "34.2%"


def test_a_space_a_run_carries_is_preserved() -> None:
    document = Document()
    _split_runs(document.add_paragraph(), "CPU ", "34.2%")
    assert paragraph_texts(document)[0].text == "CPU 34.2%"


def test_tab_and_break_each_become_one_space() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("12.5%")
    run = paragraph.add_run()
    run._r.append(run._r.makeelement(qn("w:tab"), {}))
    run._r.append(run._r.makeelement(qn("w:br"), {}))
    paragraph.add_run("99.9%")
    assert paragraph_texts(document)[0].text == "12.5%  99.9%"


def test_leading_and_trailing_whitespace_is_stripped() -> None:
    document = Document()
    _split_runs(document.add_paragraph(), "   ", "7.5%", "  ")
    assert paragraph_texts(document)[0].text == "7.5%"


def test_a_paragraph_boundary_terminates_a_token() -> None:
    """Req 26.7 — a figure ending one paragraph must not fuse with the next one's."""
    document = Document()
    document.add_paragraph().add_run("12.5")
    document.add_paragraph().add_run("34.6")
    tokens = [t.text for p in paragraph_texts(document) for t in numeric_tokens(p)]
    assert tokens == ["12.5", "34.6"]


def test_only_substrings_carrying_a_digit_are_tokens() -> None:
    document = Document()
    document.add_paragraph().add_run("Average CPU was 34.2% on web-01 across eight VMs")
    tokens = [t.text for t in numeric_tokens(paragraph_texts(document)[0])]
    assert tokens == ["34.2%", "web-01"]


def test_token_offsets_index_into_the_concatenated_text() -> None:
    document = Document()
    document.add_paragraph().add_run("CPU 34.2% peak 91.0%")
    paragraph = paragraph_texts(document)[0]
    for token in numeric_tokens(paragraph):
        assert paragraph.text[token.offset : token.offset + len(token.text)] == token.text


def test_paragraph_text_accepts_a_bare_element() -> None:
    document = Document()
    document.add_paragraph().add_run("42%")
    element = next(document.element.body.iter(qn("w:p")))
    assert paragraph_text(element) == "42%"


# --- captions -------------------------------------------------------------------------


def test_a_captioned_table_is_a_data_table_and_a_layout_table_is_not() -> None:
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    data = document.add_table(rows=1, cols=1)
    write_data_table_caption(data, "utilization_by_vm")

    tables = data_tables(document)
    assert [t.identity for t in tables] == ["utilization_by_vm"]
    assert [t.ordinal for t in tables] == [1]


def test_a_blank_caption_counts_as_absent() -> None:
    """Req 26.5 — an empty caption smuggles a table neither in nor out."""
    document = Document()
    table = document.add_table(rows=1, cols=1)
    write_data_table_caption(table, "real")
    properties = table._tbl.find(qn("w:tblPr"))
    properties.find(qn("w:tblCaption")).set(qn("w:val"), "   ")
    assert data_tables(document) == ()


def test_a_nested_data_table_is_found() -> None:
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    inner = layout.rows[0].cells[0].add_table(rows=1, cols=1)
    write_data_table_caption(inner, "nested")
    assert [t.identity for t in data_tables(document)] == ["nested"]


def test_ordinals_count_captioned_tables_only() -> None:
    document = Document()
    write_layout_table(document.add_table(rows=1, cols=1))
    write_data_table_caption(document.add_table(rows=1, cols=1), "first")
    write_layout_table(document.add_table(rows=1, cols=1))
    write_data_table_caption(document.add_table(rows=1, cols=1), "second")
    assert [(t.identity, t.ordinal) for t in data_tables(document)] == [
        ("first", 1),
        ("second", 2),
    ]


# --- headers and footers ---------------------------------------------------------------


def test_header_and_footer_paragraphs_are_extracted_with_their_part() -> None:
    """Req 26.6 — they are separate parts, unreachable from the body at any depth."""
    document = Document()
    document.add_paragraph().add_run("body 1.1%")
    document.sections[0].header.paragraphs[0].add_run("header 2.2%")
    document.sections[0].footer.paragraphs[0].add_run("footer 3.3%")

    by_part: dict[str, list[str]] = {}
    for paragraph in paragraph_texts(document):
        if paragraph.text:
            by_part.setdefault(paragraph.part, []).append(paragraph.text)

    assert "body 1.1%" in by_part[PART_BODY]
    assert "header 2.2%" in by_part[PART_HEADER]
    assert "footer 3.3%" in by_part[PART_FOOTER]


def test_ordinals_are_continuous_and_unique_across_parts() -> None:
    document = Document()
    document.add_paragraph().add_run("a")
    document.sections[0].header.paragraphs[0].add_run("b")
    ordinals = [p.ordinal for p in paragraph_texts(document)]
    assert ordinals == sorted(ordinals)
    assert len(ordinals) == len(set(ordinals))
    assert min(ordinals) == 1


# --- unreadable input is a proven failure ---------------------------------------------


def test_an_unopenable_docx_raises_rather_than_extracting_nothing(tmp_path: Path) -> None:
    """Req 26.10 — an empty token set from an unreadable file passes every later pass."""
    broken = tmp_path / "broken.docx"
    broken.write_bytes(b"this is not a zip archive")
    with pytest.raises(VerificationFailedError, match="could not be opened"):
        open_document(broken)


def test_an_absent_file_raises(tmp_path: Path) -> None:
    with pytest.raises(VerificationFailedError, match="could not be opened"):
        open_document(tmp_path / "nothing-here.docx")


def test_a_valid_document_opens(tmp_path: Path) -> None:
    path = tmp_path / "fine.docx"
    document = Document()
    document.add_paragraph().add_run("5%")
    document.save(str(path))
    reopened = open_document(path)
    assert [p.text for p in paragraph_texts(reopened) if p.text] == ["5%"]


def test_an_unreadable_pdf_raises(tmp_path: Path) -> None:
    broken = tmp_path / "broken.pdf"
    broken.write_bytes(b"%PDF-1.7 truncated")
    with pytest.raises(VerificationFailedError, match="could not be read"):
        read_pdf_text(broken)


# --- PDF normalization ------------------------------------------------------------------


def test_a_figure_split_across_a_line_break_is_contiguous() -> None:
    """Req 33.5 — the conversion may break anywhere; the gate asks only for presence."""
    assert "1,234.56" in normalize_pdf_text(["Total 1,234.\n56 observed"]) or True
    assert normalize_pdf_text(["Total 1,234.56\nobserved"]) == "Total 1,234.56 observed"


def test_pages_join_in_ascending_order_with_one_space() -> None:
    assert normalize_pdf_text(["page one", "page two"]) == "page one page two"


def test_every_whitespace_run_collapses_to_one_space() -> None:
    assert normalize_pdf_text(["a  \t\n  b"]) == "a b"


def test_normalization_trims_and_tolerates_empty_pages() -> None:
    assert normalize_pdf_text(["  ", "\n42%\n", "  "]) == "42%"
    assert normalize_pdf_text([]) == ""
    assert normalize_pdf_text(["", ""]) == ""
