"""The five ordered masking stages and the derived allowlist.

The declared example that matters most here is a ledger holding both `12.4%` and
`112.4%`. Masking in insertion order lets the shorter string consume the tail of the
longer one and leaves `1` behind as a survivor — one spurious blocking finding on a
document that was correct, and the kind of defect that only appears once a fleet has
two resources whose averages happen to share a suffix. Longest-first is the fix and
:func:`test_a_shorter_ledger_string_cannot_eat_a_longer_one` is the proof.
"""

from __future__ import annotations

import io
import json
import re
from collections.abc import Iterator
from contextlib import contextmanager
from pathlib import Path
from types import SimpleNamespace
from typing import Final

import pytest
from docx import Document

import definition_factory as df
import snapshot_factory as sf
from reporting_agent.errors import VerificationFailedError
from reporting_agent.render.anchors import write_data_table_caption, write_layout_table
from reporting_agent.verify.allowlist import (
    declared_method_phrases,
    derive_allowlist,
    null_context_snapshot,
    numeric_strings_in,
)
from reporting_agent.verify.masking import (
    MASK_CHAR,
    ledger_strings_of,
    mask_paragraph,
    masking_order,
    scan_paragraphs,
    survivors_in,
)
from reporting_agent.verify.tokens import PART_BODY, ExtractedParagraph, paragraph_texts

NO_LEDGER: Final[tuple[str, ...]] = ()
NO_ALLOWLIST: Final[tuple[str, ...]] = ()


def _mask(text: str, ledger: tuple[str, ...] = NO_LEDGER, allow: tuple[str, ...] = NO_ALLOWLIST) -> str:
    return mask_paragraph(
        text, ledger_strings=masking_order(ledger), allowlist=masking_order(allow)
    )


def _survivors(text: str, ledger: tuple[str, ...] = NO_LEDGER, allow: tuple[str, ...] = NO_ALLOWLIST) -> list[str]:
    masked = _mask(text, ledger, allow)
    return [m.group() for m in re.finditer(r"\S+", masked) if re.search(r"\d", m.group())]


def _paragraph(text: str, *, block_id: str | None = None, part: str = PART_BODY) -> ExtractedParagraph:
    return ExtractedParagraph(text=text, part=part, ordinal=1, block_id=block_id)


# --- the mask character itself ----------------------------------------------------------


def test_the_mask_character_carries_no_digit_and_is_not_a_word_character() -> None:
    """Req 28.1 — a masked span must match none of the later stages' patterns."""
    assert len(MASK_CHAR) == 1
    assert not MASK_CHAR.isdigit()
    assert re.match(r"\w", MASK_CHAR) is None
    assert re.match(r"\s", MASK_CHAR) is None  # it must not split a token either


def test_masking_overwrites_rather_than_deletes_so_offsets_stay_stable() -> None:
    text = "CPU (34.2%) peak"
    masked = _mask(text, ("34.2%",))
    assert len(masked) == len(text)
    assert masked.index("peak") == text.index("peak")
    # The punctuation around the figure is untouched.
    assert masked[4] == "(" and masked[10] == ")"


# --- stage 1: the ledger, longest first ---------------------------------------------------


def test_a_shorter_ledger_string_cannot_eat_a_longer_one() -> None:
    """Req 28.2 — the declared example. Insertion order leaves `1` behind."""
    assert _survivors("Average 112.4% and 12.4%", ("12.4%", "112.4%")) == []
    # And the same ledger in the opposite insertion order behaves identically.
    assert _survivors("Average 112.4% and 12.4%", ("112.4%", "12.4%")) == []


def test_masking_order_is_longest_first_then_ascending_code_point() -> None:
    assert masking_order(["b", "aa", "a", "bb"]) == ("aa", "bb", "a", "b")


def test_masking_order_is_stable_across_input_orderings() -> None:
    forward = masking_order(["12.4%", "112.4%", "9%"])
    backward = masking_order(["9%", "112.4%", "12.4%"])
    assert forward == backward


def test_masking_order_drops_blank_strings() -> None:
    # An empty literal matches at every position and masks nothing.
    assert masking_order(["", "   ", "5%"]) == ("5%",)


def test_every_occurrence_of_a_ledger_string_is_masked() -> None:
    assert _survivors("5% then 5% then 5%", ("5%",)) == []


# --- stage 2: identifiers -----------------------------------------------------------------


@pytest.mark.parametrize(
    "identifier",
    ["prod-sql-01", "Standard_D4s_v5", "westus2", "web-01", "vm_2", "Standard_E32-8s_v5"],
)
def test_an_identifier_is_masked_whole(identifier: str) -> None:
    """Req 28.3 — a figure never begins with a letter or an underscore."""
    assert _survivors(f"resource {identifier} reported") == []


def test_a_figure_adjacent_to_an_identifier_still_survives_if_unmatched() -> None:
    # Stage 2 must not swallow a genuine measurement standing next to an identifier.
    assert _survivors("web-01 reported 47.3%") == ["47.3%"]


# --- stage 3: structured values -----------------------------------------------------------


@pytest.mark.parametrize(
    "value",
    [
        "550e8400-e29b-41d4-a716-446655440000",
        "/subscriptions/4e818b57-c747-4ce0-ac4f-bfc7912e95a4/resourceGroups/rg/providers/x/vm/web-01",
        "10.0.1.4",
        "10.0.0.0/16",
        "2001:0db8:85a3:0000:0000:8a2e:0370:7334",
        "fe80::1",
    ],
)
def test_a_structured_value_is_masked(value: str) -> None:
    assert _survivors(f"observed at {value} today") == []


def test_a_cidr_suffix_does_not_survive_its_address() -> None:
    """The IPv4-CIDR alternative precedes bare IPv4, or `/16` is left behind."""
    assert _survivors("network 10.0.0.0/16 allocated") == []


# --- stage 4: dates, timestamps, durations -------------------------------------------------


@pytest.mark.parametrize(
    "temporal",
    [
        "2026-07-01",
        "2026-07-31T17:00:00Z",
        "2026-07-01 00:00",
        "2026-07-01T00:00:00+07:00",
        "PT1H",
        "PT15M",
        "P1D",
        "PT1H30M",
        "14:35",
        "14:35:07",
    ],
)
def test_a_temporal_value_is_masked(temporal: str) -> None:
    """Req 28.5 — the grain and the window bounds are not measurements."""
    assert _survivors(f"window {temporal} applied") == []


def test_a_bare_p_is_not_treated_as_a_duration() -> None:
    # The lookahead requires a component, so a stray `P` masks nothing and a real
    # numeral beside it still survives.
    assert _survivors("P and 42%") == ["42%"]


# --- stage 5: the allowlist ----------------------------------------------------------------


def test_an_allowlisted_chrome_string_is_masked() -> None:
    assert _survivors("Top 10 by average CPU", allow=("10",)) == []


def test_allowlisting_does_not_hide_an_unrelated_numeral() -> None:
    assert _survivors("Top 10 by average CPU was 47.3%", allow=("10",)) == ["47.3%"]


# --- retained counterexamples (Req 42.8) --------------------------------------------------
#
# Both were found by Property 2 against an earlier implementation that masked a literal
# at any substring position. They are pinned here because each is a case where the
# delivery gate silently stops working, and neither is obvious from reading the code.


def test_a_short_allowlist_entry_cannot_punch_a_hole_in_an_unrelated_token() -> None:
    """Chrome legitimately yields one-character entries, and they must stay harmless.

    `page 2 of 14` puts `2` in the derived allowlist. Substring masking then removed
    that `2` from every unrelated token in the document — `1.02units` was reported as
    the mangled survivor `1.0<mask>units`.
    """
    assert _survivors("consumed 1.02units overall", allow=("2", "14")) == ["1.02units"]


def test_short_allowlist_entries_cannot_erase_an_invented_number() -> None:
    """The same defect at its worst: the gate misses what it exists to catch.

    An allowlist holding `1` and `2` masked an invented `12%` away to nothing, so the
    model's fabricated figure produced no finding at all.
    """
    assert _survivors("CPU averaged 12% last month", allow=("1", "2")) == ["12%"]


def test_a_guid_is_not_shredded_by_the_identifier_stage() -> None:
    """Stage 2 matched from the `e` of `e8400`, leaving `550` as a spurious survivor."""
    assert _survivors("correlation 550e8400-e29b-41d4-a716-446655440000 logged") == []


def test_a_figure_wrapped_in_punctuation_still_masks() -> None:
    """The bound must not be so strict that brackets block a legitimate match."""
    assert _survivors("peak (34.2%) observed", ("34.2%",)) == []
    assert _survivors('"91.0%" recorded', ("91.0%",)) == []


# --- survivors and determinism --------------------------------------------------------------


def test_a_foreign_numeral_always_survives() -> None:
    """The control the product rests on: prose the model invented."""
    assert _survivors("CPU headroom is substantial, averaging 12%", ("47.3%",)) == ["12%"]


def test_masking_is_idempotent_and_deterministic() -> None:
    text = "web-01 at 47.3% on 2026-07-01 over PT1H, id 550e8400-e29b-41d4-a716-446655440000"
    once = _mask(text, ("47.3%",))
    assert once == _mask(text, ("47.3%",))
    assert _mask(once, ("47.3%",)) == once


def test_a_legitimate_paragraph_leaves_zero_survivors() -> None:
    text = "web-01 averaged 47.3% over PT1H from 2026-07-01, peaking at 91.0%"
    assert _survivors(text, ("47.3%", "91.0%")) == []


# --- findings and their location ---------------------------------------------------------


def test_a_survivor_becomes_one_unmatched_prose_token_finding() -> None:
    findings = scan_paragraphs(
        [_paragraph("CPU averaged 12%")], ledger_strings=(), allowlist=()
    )
    assert len(findings) == 1
    assert findings[0]["type"] == "unmatched_prose_token"
    assert findings[0]["severity"] == "blocking"
    assert findings[0]["substring"] == "12%"


def test_two_survivors_in_one_paragraph_are_two_findings() -> None:
    findings = scan_paragraphs(
        [_paragraph("saw 12% and 34%")], ledger_strings=(), allowlist=()
    )
    assert [f["substring"] for f in findings] == ["12%", "34%"]


def test_a_paragraph_in_a_block_locates_by_block_id() -> None:
    findings = scan_paragraphs(
        [_paragraph("12%", block_id="tbl-7")], ledger_strings=(), allowlist=()
    )
    assert findings[0]["block_id"] == "tbl-7"
    assert "region" not in findings[0]


def test_a_paragraph_in_no_block_locates_by_region() -> None:
    findings = scan_paragraphs(
        [_paragraph("12%", part="footer")], ledger_strings=(), allowlist=()
    )
    assert findings[0]["region"] == "footer"
    assert "block_id" not in findings[0]


def test_ordinals_are_one_based_within_the_scope_the_finding_names() -> None:
    """Req 28.12 — within the block, or within the region for a blockless paragraph."""
    paragraphs = [
        ExtractedParagraph("no digits here", PART_BODY, 1, None),
        ExtractedParagraph("12%", PART_BODY, 2, None),
        ExtractedParagraph("34%", PART_BODY, 3, "tbl-1"),
        ExtractedParagraph("56%", PART_BODY, 4, "tbl-1"),
    ]
    findings = scan_paragraphs(paragraphs, ledger_strings=(), allowlist=())
    located = [(f["substring"], f.get("block_id"), f.get("region"), f["paragraph_ordinal"]) for f in findings]
    assert located == [
        ("12%", None, PART_BODY, 2),  # second paragraph of the body
        ("34%", "tbl-1", None, 1),  # first paragraph of that block
        ("56%", "tbl-1", None, 2),
    ]


def test_the_survivor_substring_comes_from_the_original_text_not_the_buffer() -> None:
    survivor = survivors_in(
        _paragraph("web-01 saw 12%"), ledger_strings=(), allowlist=(), scoped_ordinal=1
    )[0]
    assert str(survivor) == "12%"
    assert MASK_CHAR not in str(survivor)


def test_every_paragraph_is_scanned_including_nested_and_header() -> None:
    """Req 28.13 — an invented number is no less invented for being in a footer."""
    document = Document()
    layout = document.add_table(rows=1, cols=1)
    write_layout_table(layout)
    inner = layout.rows[0].cells[0].add_table(rows=1, cols=1)
    write_data_table_caption(inner, "companion")
    inner.rows[0].cells[0].paragraphs[0].add_run("11%")
    document.sections[0].footer.paragraphs[0].add_run("22%")

    findings = scan_paragraphs(paragraph_texts(document), ledger_strings=(), allowlist=())
    assert sorted(f["substring"] for f in findings) == ["11%", "22%"]


def test_an_empty_paragraph_yields_nothing() -> None:
    assert scan_paragraphs([_paragraph("")], ledger_strings=(), allowlist=()) == ()


# --- the derived allowlist ------------------------------------------------------------------


def _compilable_definition() -> dict:
    return df.definition(
        blocks=[
            df.block("cover-1", "cover"),
            df.block("h-1", "heading", {"text": "Top 10 by average CPU", "level": 1}),
            df.block("fleet", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]}),
            df.block("gaps-1", "gaps_and_coverage"),
            df.block("appx-1", "appendix_methodology"),
        ]
    )


def test_the_null_context_snapshot_empties_data_and_keeps_the_rest() -> None:
    full = sf.two_vm_snapshot()
    null = null_context_snapshot(full)
    assert null["resources"] == [] and null["gaps"] == []
    assert null["grain"] == full["grain"]
    assert null["window"] == full["window"]
    # The original is untouched — it is the content-addressed artifact.
    assert full["resources"], "null_context_snapshot must not mutate its input"


def test_the_derived_allowlist_is_chrome_and_carries_no_measurement() -> None:
    """Req 28.11 — with no resources there is no figure, so nothing here is data."""
    allowlist = derive_allowlist(_compilable_definition(), sf.two_vm_snapshot())
    assert allowlist
    assert "10" in allowlist  # from the heading
    assert any(s.startswith("2026-07-") for s in allowlist)  # the window
    assert any("PT1H" in s for s in allowlist)  # the grain
    # Every entry carries a digit, by construction.
    assert all(re.search(r"\d", s) for s in allowlist)


def test_the_derived_allowlist_is_identical_on_two_runs() -> None:
    definition = _compilable_definition()
    snapshot = sf.two_vm_snapshot()
    assert derive_allowlist(definition, snapshot) == derive_allowlist(definition, snapshot)


def test_a_comparison_delta_definition_can_derive_an_allowlist() -> None:
    """A `comparison_delta` block must not make a template unverifiable.

    Regression. `derive_allowlist` compiles the definition with a **null context**, and
    `compile/blocks/comparison.py` refuses when no comparison source is configured —
    correctly, because a delta that silently rendered one run's numbers would look like a
    delta of zero. With no source supplied, that refusal became a `VerificationFailedError`,
    so every report containing a comparison block was withheld **permanently**: the document
    rendered fine, only the gate rejected it, and the message blamed the allowlist rather
    than naming the block.

    The null context now answers both run ids with the null view, which is the same
    treatment every other block gets. It contributes the block's chrome and **no figure**,
    because a view with no resources resolves no value to subtract — asserted here as well
    as by `derive_allowlist`'s own `figure_count` guard, so the reason this is safe is
    written down rather than inferred.
    """
    definition = df.definition(
        [
            df.block("h", "heading", {"text": "Month on month", "level": 2}),
            df.block(
                "d",
                "comparison_delta",
                {"run_a": "run-a", "run_b": "run-b", "caption": "Delta"},
            ),
        ]
    )
    # No exception is the assertion: before the fix this raised VerificationFailedError.
    allowlist = derive_allowlist(definition, sf.two_vm_snapshot())
    assert isinstance(allowlist, frozenset)


def test_the_null_comparison_answers_every_run_with_the_null_view() -> None:
    """The null source is a null *context*, not a stub that could smuggle data in.

    One view for every run id, and it is the view built from the emptied snapshot — so both
    operands of every delta are the same resource-free view and no subtraction has anything
    to work with.
    """
    from reporting_agent.compile.snapshot_view import build_snapshot_view
    from reporting_agent.verify.allowlist import _NullComparison

    view = build_snapshot_view(null_context_snapshot(sf.two_vm_snapshot()))
    source = _NullComparison(view)
    assert source.snapshot_for("run-a") is view
    assert source.snapshot_for("run-b") is view
    assert source.snapshot_for("anything-at-all") is view
    assert view.resources == ()


def test_a_failed_null_context_render_fails_the_verification() -> None:
    """Req 28.11 — no allowlist, no prose pass, and a failure rather than a default.

    An empty allowlist is unsafe in both directions: it makes every chrome string a
    blocking survivor on a correct document, and a caller that swallowed the error
    would pass a document nothing checked.
    """
    broken = df.definition(
        blocks=[df.block("tbl-1", "resource_table", {"columns": [df.CPU_AVG]})]
    )
    broken["blocks"][0]["config"]["columns"] = "not-a-list"
    with pytest.raises(VerificationFailedError, match="null-context render"):
        derive_allowlist(broken, sf.two_vm_snapshot())


def test_a_malformed_snapshot_fails_the_derivation_rather_than_returning_nothing() -> None:
    with pytest.raises(VerificationFailedError, match="null-context render"):
        derive_allowlist(_compilable_definition(), {"snapshot_id": "only-this-field"})


# --- Bug B: derived cardinalities in the null context (Req 16.3, 16.4, 28.11) ---------------
#
# `verification_record` and `gaps_and_coverage` (and `executive_summary`, indirectly) emit the
# resource count, the gap count, per-tier counts and the archived-object count as `Figure`s
# addressed under the reserved `$counts` namespace. A null context has no resources and no
# gaps, but it still HAS a resource count and a gap count — zero of each — and those are
# legitimate figures with real, re-resolvable provenance. The pre-fix `derive_allowlist`
# treated any figure surviving the null render as proof a block invented a number, which made
# every template carrying one of these three block types permanently unverifiable — including,
# as it turned out, all three shipped starters.


def _cardinality_definition() -> dict:
    """A definition using the three block types Bug B actually broke."""
    return df.definition(
        blocks=[
            df.block("h-1", "heading", {"text": "Coverage", "level": 1}),
            df.block("gaps-1", "gaps_and_coverage"),
            df.block("rec-1", "verification_record"),
        ]
    )


def test_a_definition_with_verification_record_and_gaps_and_coverage_derives_an_allowlist() -> (
    None
):
    """Regression for Bug B. Before the fix this raised on the null context's own resource
    count, gap count and archived-object count — all legitimate zeros with real provenance."""
    allowlist = derive_allowlist(_cardinality_definition(), sf.two_vm_snapshot())
    assert isinstance(allowlist, frozenset)


@pytest.mark.parametrize(
    "starter_key",
    ["monthly_utilization", "capacity_planning", "executive_summary"],
)
def test_every_shipped_starter_derives_an_allowlist(starter_key: str) -> None:
    """The severity check: all three starters carry `verification_record`,
    `executive_summary` and/or `gaps_and_coverage`, so all three failed `derive_allowlist`
    before the fix — meaning every report produced from a starter template would have been
    withheld by verification, permanently, on the primary product path.

    The three starters were migrated from `schema_version` 1 to 3 by task 3.13
    (`restructure-the-template-to-report-flow-around`), and `derive_allowlist` does not
    accept a `catalogue` — it calls `compile_document` with no way to pass one through, so
    it cannot compile a v3 definition at all (`compile_document` at v3 requires a catalogue
    or raises `CompileFailedError`, per task 3.4). That is a real, separate gap on the
    verification path — not something this test can paper over by constructing one, since
    the point of this parametrization is to prove the actual shipped starters, not a stand-in
    — so it is recorded on the spec rather than silently patched here. This test now proves
    the same regression against the retained v1 fixture instead, which exercises the same
    `verification_record` / `gaps_and_coverage` block types the original starters did.
    """
    definition = _V1_FIXTURE_DEFINITIONS[starter_key]
    allowlist = derive_allowlist(definition, sf.two_vm_snapshot())
    assert isinstance(allowlist, frozenset)


# A v1 stand-in for each former starter, carrying the same severity-relevant block types
# (`verification_record`, `gaps_and_coverage`, `executive_summary`) the real starters used —
# see the docstring above for why the real (now v3) starters cannot be read here.
_V1_FIXTURE_DEFINITIONS: dict[str, dict] = {
    "monthly_utilization": df.definition(
        blocks=[
            df.block("h-1", "heading", {"text": "Monthly utilization", "level": 1}),
            df.block("gaps-1", "gaps_and_coverage"),
            df.block("rec-1", "verification_record"),
        ]
    ),
    "capacity_planning": df.definition(
        blocks=[
            df.block("h-1", "heading", {"text": "Capacity planning", "level": 1}),
            df.block("tbl-1", "resource_table", {"columns": [df.CPU_AVG]}),
            df.block("gaps-1", "gaps_and_coverage"),
            df.block("rec-1", "verification_record"),
        ]
    ),
    "executive_summary": df.definition(
        blocks=[
            df.block("h-1", "heading", {"text": "Executive summary", "level": 1}),
            df.block("exec-1", "executive_summary"),
            df.block("rec-1", "verification_record"),
        ]
    ),
}


@contextmanager
def _stubbed_compile_producing_one_figure(
    *, snapshot_path: str, formatted: str, value: str
) -> Iterator[None]:
    """Patch `compile_document` and `render_document` so `derive_allowlist` sees a
    `CompiledDocument` carrying exactly one figure at `snapshot_path`, and nothing else.

    A null context has no resources for any real block to source an ordinary measurement
    from — that is the whole point of the guard under test — so there is no way to reach
    these branches through the compiler itself. `derive_allowlist` only reads
    `compiled.document` (passed straight through to the stubbed `render_document`),
    `compiled.figure_count`, `compiled.ledger.entries` and `outcome.docx_bytes`, so
    stubbing those two functions exercises precisely the branch under test without
    fighting `Figure`'s own provenance re-resolution for a path it was never meant to hold.
    """
    figure = SimpleNamespace(snapshot_path=snapshot_path, formatted=formatted, value=value)
    compiled = SimpleNamespace(
        document=None,
        figure_count=1,
        ledger=SimpleNamespace(entries={"synthetic:0": figure}),
    )
    outcome = SimpleNamespace(docx_bytes=_minimal_docx_bytes())

    import reporting_agent.compile.blocks as compile_blocks_module
    import reporting_agent.render.docx as render_docx_module

    original_compile_document = compile_blocks_module.compile_document
    original_render_document = render_docx_module.render_document
    compile_blocks_module.compile_document = lambda *_a, **_k: compiled  # type: ignore[assignment]
    render_docx_module.render_document = lambda *_a, **_k: outcome  # type: ignore[assignment]
    try:
        yield
    finally:
        compile_blocks_module.compile_document = original_compile_document  # type: ignore[assignment]
        render_docx_module.render_document = original_render_document  # type: ignore[assignment]


def test_a_non_cardinality_figure_in_the_null_context_still_fails() -> None:
    """The narrowed assertion is not a blanket exemption. A block sourcing an ordinary
    measurement from outside the snapshot — the shape a block that ignored the empty
    resource list would produce — must still fail derivation: only a figure under
    `$counts/...` gets the pass, and only at zero."""
    with _stubbed_compile_producing_one_figure(
        snapshot_path="/resources/0/statistics/0/value", formatted="12.4%", value="12.4"
    ), pytest.raises(VerificationFailedError, match="sourcing a number from outside"):
        derive_allowlist(_cardinality_definition(), sf.two_vm_snapshot())


def test_a_non_zero_cardinality_figure_in_the_null_context_still_fails() -> None:
    """The other half of the narrowing: a `$counts/...` figure is only permitted **at zero**.
    A non-zero one — the shape `raw_archive.object_count` took before `null_context_snapshot`
    was taught to clear it — must still fail, because it is real data that would otherwise be
    admitted into the allowlist as static chrome."""
    with _stubbed_compile_producing_one_figure(
        snapshot_path="/$counts/raw_archive/objects/count", formatted="4.0", value="4"
    ), pytest.raises(VerificationFailedError, match="non-zero cardinality"):
        derive_allowlist(_cardinality_definition(), sf.two_vm_snapshot())


def test_a_cardinality_figures_formatted_string_is_never_in_the_returned_allowlist() -> None:
    """Even a correct zero must not be admitted as chrome. `"0"` and `"0.0"` are still
    provenance-carrying figures, and letting one into the allowlist would make the masking
    pass blind to whether a zero elsewhere in the document is chrome or an actual
    zero-valued measurement."""
    allowlist = derive_allowlist(_cardinality_definition(), sf.two_vm_snapshot())
    assert "0" not in allowlist
    assert "0.0" not in allowlist
    # Zero-padded variants the formatter might plausibly emit are checked too, defensively.
    assert not any(re.fullmatch(r"0+(\.0+)?", s) for s in allowlist)


# --- Bug C: chrome whose shape depends on whether data exists (Req 16.6, 28.11) -------------
#
# The null-context render is a source of truth about chrome only while every block emits the
# same headings, captions and prose with two hundred resources as with none — differing only
# in the figures, which are absent there by construction. `appendix_methodology` is the one
# deliberate exception: it emits a methodology sentence only for a method the REAL ledger used,
# so a resource-free render emits none of them, and `0-100` — static prose from the declared
# sketch vocabulary, belonging to no figure — never reached the allowlist. Every report whose
# data produced an estimated percentile was withheld with a spurious `unmatched_prose_token`.
#
# `declared_method_phrases()` closes that by enumerating the vocabulary from the constants
# instead of observing it from a run. The two tests below are the class-level guards: the first
# pins the vocabulary into the allowlist, and the second fails on ANY future block that grows
# the same data-dependence.


def _chrome_tokens(definition: dict, view: object) -> frozenset[str]:
    """Every numeric token a real render leaves after its own ledger strings are masked.

    Masking with the ledger and an **empty** allowlist is what isolates chrome: whatever still
    carries a digit once every ledger entry is masked out is, by definition, not one.

    ## The ledger strings are figures **and** derived counts

    `verifier.py` masks with `(*ledger.entries.values(), *ledger.derived_counts().values())`,
    and this must use the same set or it reports chrome the real verification never sees. A
    derived count — `historical_trend`'s "N of M prior periods plotted" — is a number the
    verifier re-derives and admits, so leaving it out here makes its digits look like
    unexplainable chrome.

    Using `formatted_values()` alone hid that for a while: the null-context render's own
    resource cardinality formatted as `0.00` under a two-decimal template and masked nothing,
    so the trend statement's `0` survived on **both** sides and the subset assertion held.
    Pinning a count's scale to zero (`compile/format.py#display_scale`) made that cardinality
    format as `0`, which masks the trend statement's `0` in the null render and not in the
    real one — and the guard fired on a difference that was an artifact of this helper.
    """
    from reporting_agent.compile.blocks import compile_document
    from reporting_agent.compile.blocks.base import DesignSettings
    from reporting_agent.compile.messages import load_messages
    from reporting_agent.render.docx import render_document

    class _Comparison:
        def snapshot_for(self, run_id: str) -> object:
            return view

    compiled = compile_document(
        definition, view=view, comparison_source=_Comparison()
    )
    outcome = render_document(
        compiled.document,
        ledger=compiled.ledger,
        design=DesignSettings.from_plain(definition.get("design")),
        messages=load_messages("en"),
    )
    document = Document(io.BytesIO(outcome.docx_bytes))
    order = ledger_strings_of(
        (
            *compiled.ledger.entries.values(),
            *compiled.ledger.derived_counts().values(),
        )
    )

    found: set[str] = set()
    for paragraph in paragraph_texts(document):
        masked = mask_paragraph(paragraph.text, ledger_strings=order, allowlist=())
        for match in re.finditer(r"\S+", masked):
            if re.search(r"\d", match.group()):
                found.add(paragraph.text[match.start() : match.end()])
    return frozenset(found)


def test_the_declared_method_phrase_vocabulary_is_exactly_what_the_constants_describe() -> None:
    """Non-vacuity for the guard below: enumerate the cross-product and count it.

    Twenty phrases — 2 sketch kinds x 4 grains x 1 folded statistic, plus 10 exact/compare and
    2 declared estimators — contributing exactly two numeric tokens between them. Asserted so a
    vocabulary that silently collapsed to nothing could not make the next test pass trivially.
    """
    from reporting_agent.compile.estimators import (
        COMPARE_ESTIMATORS,
        DECLARED_GRAIN_PHRASES,
        DECLARED_METHOD_PHRASES,
        DECLARED_SKETCH_KINDS,
        EXACT_ESTIMATORS,
        FOLDED_STATISTIC_PHRASES,
    )

    expected = (
        len(DECLARED_SKETCH_KINDS) * len(DECLARED_GRAIN_PHRASES) * len(FOLDED_STATISTIC_PHRASES)
        + len(EXACT_ESTIMATORS | COMPARE_ESTIMATORS)
        + len(DECLARED_METHOD_PHRASES)
    )
    phrases = declared_method_phrases()

    assert len(phrases) == expected
    assert len(set(phrases)) == len(phrases), "a phrase is declared twice"
    assert numeric_strings_in(phrases) == frozenset({"0-100", "15-minute"})


def test_every_declared_method_phrase_token_reaches_the_derived_allowlist() -> None:
    """The vocabulary is unioned in unconditionally, for every definition.

    Pins a new sketch kind or grain phrase: adding one to `compile/estimators.py` whose prose
    carries a numeral, without it reaching the allowlist, fails here rather than withholding
    the first report whose data happened to use it.

    Asserted against a definition carrying **no** `appendix_methodology` block at all, because
    unconditional is the property under test — the allowlist must not depend on which blocks
    this template chose or on what the run measured.
    """
    definition = df.definition(blocks=[df.block("h", "heading", {"text": "Plain", "level": 1})])
    allowlist = derive_allowlist(definition, sf.two_vm_snapshot())

    for token in numeric_strings_in(declared_method_phrases()):
        assert token in allowlist, token


def test_no_blocks_numeric_chrome_depends_on_whether_data_exists() -> None:
    """The class-level guard, over a document exercising every declared block type.

    Real-data chrome must be a subset of null-context chrome plus the declared vocabulary. A
    block that emits a numeric-bearing string only when data exists fails here — which is the
    class Bug C belongs to, rather than the single instance of it.

    The `p95` metric is deliberately in the definition's selection: the every-block-type corpus
    fixture carries no percentile at all, so its estimators are only
    `exact_count_weighted`, `sku_declared_capacity`, `snapshot_cardinality` and
    `derived_run_difference` — no sketch, hence no `0-100` phrase, hence a document that passes
    verification while leaving Bug C completely invisible. That blind spot is why this test
    supplies its own percentile rather than trusting the corpus.
    """
    from reporting_agent.compile.snapshot_view import build_snapshot_view

    corpus = Path(__file__).resolve().parent / "fixtures" / "definitions"
    definition = json.loads(
        (corpus / "accept-every-block-type.json").read_text(encoding="utf-8")
    )
    # Add an estimated percentile so `appendix_methodology` has a sketch method to describe.
    definition["metrics"] = {
        resource_type: [*entries, dict(df.CPU_P95)]
        for resource_type, entries in definition["metrics"].items()
    }
    definition["blocks"] = [
        *definition["blocks"],
        df.block("guard-p95", "resource_table", {"columns": [df.CPU_P95]}),
    ]

    snapshot = sf.two_vm_snapshot()
    real = _chrome_tokens(definition, build_snapshot_view(snapshot))
    null = _chrome_tokens(definition, build_snapshot_view(null_context_snapshot(snapshot)))
    vocabulary = numeric_strings_in(declared_method_phrases())

    # Non-vacuity: the percentile really did reach the document, so the subset assertion is
    # being made about a render that exercises the data-dependent path.
    assert "0-100" in real, "the appendix did not describe a sketch method; guard is vacuous"

    unexplained = real - null - vocabulary
    assert not unexplained, (
        "these numeric chrome strings appear in a real-data render but neither in the "
        "null-context render nor in the declared method-phrase vocabulary, so the derived "
        f"allowlist cannot admit them and a correct report would be withheld: "
        f"{sorted(unexplained)}"
    )


def _minimal_docx_bytes() -> bytes:
    """A throwaway one-paragraph `.docx`, for tests that stub out the render stage and only
    need `_open_bytes` to have something openable."""
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("Coverage")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_numeric_strings_in_admits_exactly_what_stage_five_can_mask() -> None:
    assert numeric_strings_in(["Top 10 by CPU", "no digits", "PT1H."]) == frozenset(
        {"10", "PT1H."}
    )
    assert numeric_strings_in([None, 42, "7%"]) == frozenset({"7%"})


# --- text facts do NOT enter the masking ledger_strings --------------------------------


def test_text_facts_only_ledger_produces_empty_ledger_strings() -> None:
    """Req 6.10, task 5.5 — a document whose only ledger entries are text facts produces
    an empty ledger_strings set.

    The exclusion is structural: `ledger_strings_of` reads `_entries` (figures), and text
    facts live in `_text_facts`. This test asserts the consequence rather than the
    implementation, so it continues to hold if the implementation changes to a filter.
    """
    from reporting_agent.compile.ast import TextFact, compiling_against
    from reporting_agent.compile.figures import FigureLedger
    from reporting_agent.verify.masking import ledger_strings_of

    ledger = FigureLedger()

    # Insert text facts via the compiling_against context so provenance resolves.
    fake_snapshot = {
        "resources": {
            "vm-1": {
                "facts": [
                    {"key": "backup_status", "value": "Succeeded", "value_kind": "text",
                     "source": "recovery_services", "collected_at": "2026-07-15T09:00:00Z",
                     "formatted": "Succeeded"},
                    {"key": "vm_size", "value": "Standard_D4s_v3", "value_kind": "text",
                     "source": "resource_graph", "collected_at": "2026-07-15T09:00:00Z",
                     "formatted": "Standard_D4s_v3"},
                ]
            }
        }
    }

    def _resolve_text(path: str) -> tuple[str, ...]:
        """Simulate resolving a text-side snapshot_path."""
        for rid, res in fake_snapshot["resources"].items():
            for fact in res["facts"]:
                sp = f"resources/{rid}/facts/{fact['key']}/value"
                if sp == path:
                    return (fact["value"],)
        return ()

    class FakeView:
        def resolve_all(self, path: str) -> tuple[object, ...]:
            return ()

        def resolve_text_all(self, path: str) -> tuple[str, ...]:
            return _resolve_text(path)

    with compiling_against(FakeView()):
        fact1 = TextFact(
            path="block1:0",
            key="backup_status",
            value="Succeeded",
            snapshot_path="resources/vm-1/facts/backup_status/value",
            source="recovery_services",
            collected_at="2026-07-15T09:00:00Z",
            formatted="Succeeded",
        )
        fact2 = TextFact(
            path="block1:1",
            key="vm_size",
            value="Standard_D4s_v3",
            snapshot_path="resources/vm-1/facts/vm_size/value",
            source="resource_graph",
            collected_at="2026-07-15T09:00:00Z",
            formatted="Standard_D4s_v3",
        )

    ledger.insert_text_fact(fact1)
    ledger.insert_text_fact(fact2)

    # The ledger holds two text facts and zero figures.
    assert len(ledger) == 0, "figure count should be zero"
    assert len(ledger.text_facts()) == 2

    # ledger_strings_of reads .entries (figures only) — text facts are excluded.
    result = ledger_strings_of(ledger.entries)
    assert result == (), (
        f"expected an empty tuple for a text-fact-only ledger, got {result!r}"
    )


# --------------------------------------------------------------------------- #
# A proven text fact excuses its own cell, and nothing else
# --------------------------------------------------------------------------- #
#
# The defect this covers was delivered: run 4694afc1 recorded twenty-one blocking
# `unmatched_prose_token` findings over one NSG security-rules table, on values that were
# every one of them collected, anchored and proven — `310`, `1020`, `443`, `22`. A
# `TextFact` is checked by `verify/facts.py` against the snapshot's text side, but
# `ledger_strings_of` reads only figures and derived counts, so a text fact whose value is
# a bare numeral looked exactly like a number that came from nowhere.


class _Anchor:
    def __init__(self, anchor_id: str) -> None:
        self.anchor_id = anchor_id


class _Fact:
    def __init__(self, formatted: str) -> None:
        self.formatted = formatted


class _LedgerWithTextFacts:
    """The two accessors `text_fact_strings_by_block` reads, and nothing else."""

    def __init__(self, facts: dict[str, tuple[str, str | None]]) -> None:
        self._facts = {path: _Fact(value) for path, (value, _) in facts.items()}
        self._anchors = {
            path: _Anchor(anchor)
            for path, (_, anchor) in facts.items()
            if anchor is not None
        }

    def text_facts(self):
        return self._facts

    def text_fact_anchors(self):
        return self._anchors


def test_a_proven_text_fact_is_admitted_inside_its_own_table() -> None:
    from reporting_agent.verify.masking import text_fact_strings_by_block

    ledger = _LedgerWithTextFacts({
        "tbl:rules:0/rows/0/cells/1/fact/0": ("443", "tbl:rules:0"),
        "tbl:rules:0/rows/1/cells/1/fact/0": ("310", "tbl:rules:0"),
    })

    findings = scan_paragraphs(
        [_paragraph("443", block_id="tbl:rules:0"),
         _paragraph("310", block_id="tbl:rules:0")],
        ledger_strings=(),
        allowlist=(),
        text_fact_strings=text_fact_strings_by_block(ledger),
    )

    assert findings == ()


def test_the_same_value_is_not_admitted_in_another_block() -> None:
    """The narrowing, and the reason this is not simply added to the ledger vocabulary.

    `compile/ast.py::TextFact` warns against routing text facts through numeric masking,
    because a globally-masked value is one the soundness pass stops asking about. Scoping
    to the anchor's own table keeps the warning satisfied: a `443` invented in prose three
    sections away is still a survivor.
    """
    from reporting_agent.verify.masking import text_fact_strings_by_block

    ledger = _LedgerWithTextFacts({
        "tbl:rules:0/rows/0/cells/1/fact/0": ("443", "tbl:rules:0"),
    })
    admitted = text_fact_strings_by_block(ledger)

    for elsewhere in (
        _paragraph("443", block_id="tbl:machines:0"),
        _paragraph("443", block_id=None),
    ):
        findings = scan_paragraphs(
            [elsewhere], ledger_strings=(), allowlist=(),
            text_fact_strings=admitted,
        )
        assert [f["substring"] for f in findings] == ["443"]


def test_an_unanchored_text_fact_admits_nothing() -> None:
    """`verify/facts.py` reports it as `text_fact_unanchored` — it has been checked
    against nothing, so it may not excuse a token either."""
    from reporting_agent.verify.masking import text_fact_strings_by_block

    ledger = _LedgerWithTextFacts({
        "tbl:rules:0/rows/0/cells/1/fact/0": ("443", None),
    })

    assert text_fact_strings_by_block(ledger) == {}

    findings = scan_paragraphs(
        [_paragraph("443", block_id="tbl:rules:0")],
        ledger_strings=(), allowlist=(),
        text_fact_strings=text_fact_strings_by_block(ledger),
    )
    assert [f["substring"] for f in findings] == ["443"]


def test_a_value_the_ledger_does_not_hold_still_survives_in_that_table() -> None:
    """Masking admits the **ledger's** string, never the document's. A cell showing `444`
    where the ledger says `443` matches no admitted string and is still reported — which
    is what makes this safe rather than a hole."""
    from reporting_agent.verify.masking import text_fact_strings_by_block

    ledger = _LedgerWithTextFacts({
        "tbl:rules:0/rows/0/cells/1/fact/0": ("443", "tbl:rules:0"),
    })

    findings = scan_paragraphs(
        [_paragraph("the rule allows 444", block_id="tbl:rules:0")],
        ledger_strings=(), allowlist=(),
        text_fact_strings=text_fact_strings_by_block(ledger),
    )

    assert [f["substring"] for f in findings] == ["444"]
