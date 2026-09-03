"""Tests for `render/front_matter.py` — proving real emission into python-docx documents.

Tests cover:
- emit_front_matter emits actual paragraphs and tables into a real Document
- clause (a): cover disabled => no cover content, no leading blank page, document control
  unchanged
- clause (b): no signature image => EMPTY ruled signature box (never typed name)
- clause (c): absent per-run value => RENDER_FAILED naming that value
- message resolution for document-control labels in both languages
"""

from __future__ import annotations

import asyncio
import base64

import pytest
from docx.oxml.ns import qn

from reporting_agent.compile.messages import load_messages
from reporting_agent.errors import RenderFailedError
from reporting_agent.render.front_matter import (
    SIGNATURE_BOX_HEIGHT_TWIPS,
    ApproverEntry,
    CoverConfig,
    DocumentControlConfig,
    FrontMatterConfig,
    FrontMatterLogo,
    RevisionHistoryRow,
    RunFacts,
    emit_front_matter,
    front_matter_sections,
)
from reporting_agent.render.themes import load_theme

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

_ONE_PIXEL_PNG: bytes = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUAAdafFs0AAAAASUVORK5CYII="
)
"""The smallest valid PNG: a single transparent pixel. Real magic bytes and
a real, decodable image — not an arbitrary byte string — so a test using it
exercises the actual `python-docx` `add_picture` path rather than assuming
it would work."""

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _messages():
    """This file constructs messages per test rather than holding a module constant;
    this is that call in one place for the classes below."""
    return load_messages("en")


def _run(
    *,
    run_id: str = "run-001",
    template_id: str = "tmpl-abc",
    year: str = "2026",
    month: str = "07",
    customer: str = "Acme Corp",
) -> RunFacts:
    return RunFacts(
        run_id=run_id,
        template_id=template_id,
        customer_name=customer,
        period_display="July 2026",
        report_title="Infrastructure Report",
        period_start_year=year,
        period_start_month=month,
    )

def _paragraphs_text(doc: object) -> list[str]:
    """Extract all paragraph texts from a Document."""
    return [p.text for p in doc.paragraphs]


def _tables(doc: object) -> list[object]:
    """Extract all tables from a Document."""
    return list(doc.tables)


def _document_text(doc: object) -> list[str]:
    """Every paragraph in the document, including the ones inside table cells.

    The cover and the document control page carry their values in labelled
    ``Layout Table`` blocks rather than in `f"{label}: {value}"` paragraphs, so "the
    document states X" is a question about the whole body and not about
    `doc.paragraphs` alone.
    """
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts


def _approvers_table(doc: object) -> object:
    """The approvers table, found by its own style rather than by position.

    It was `doc.tables[0]`, which silently meant "whichever table the emitter happens to
    write first" — so adding the document control page's naming block above it pointed
    every signature assertion at the wrong table. `Table Signature` is the approvers
    table's identity (`render/themes.py`), and it is the only table carrying it.
    """
    from reporting_agent.render.themes import SIGNATURE_TABLE_STYLE

    matches = [t for t in doc.tables if t.style is not None and t.style.name == SIGNATURE_TABLE_STYLE]
    assert len(matches) == 1, (
        f"expected exactly one {SIGNATURE_TABLE_STYLE!r} table, found {len(matches)}"
    )
    return matches[0]


# ---------------------------------------------------------------------------
# Basic emission — cover enabled, document control emitted
# ---------------------------------------------------------------------------


class TestEmitCoverEnabled:
    """When cover is enabled, the function emits cover content followed by doc control."""

    def test_cover_title_is_emitted(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=True))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _paragraphs_text(doc)
        # The first paragraph is the report title (Cover Title style)
        assert texts[0] == "Infrastructure Report"

    def test_cover_customer_and_period_emitted(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=True))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _document_text(doc)
        assert "Acme Corp" in texts
        assert "July 2026" in texts

    def test_document_control_heading_emitted_after_cover(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=True))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _paragraphs_text(doc)
        # Document control heading appears (it's the resolved message id)
        assert "Document control" in texts

    def test_document_number_on_cover_and_control(self) -> None:
        """Document number appears identically on cover AND document control page."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=True),
            document_control=DocumentControlConfig(
                document_number_pattern="RPT-{template}-{year}{month}-{run}",
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _document_text(doc)
        expected_number = "RPT-tmpl-abc-202607-run-001"
        # Once on the cover and once on the document control page — the same string
        # both times, which is the whole of Req 13.8.
        assert texts.count(expected_number) == 2
        # And the document control page labels it.
        assert "Document number" in texts


# ---------------------------------------------------------------------------
# Clause (a) — cover disabled
# ---------------------------------------------------------------------------


class TestCoverDisabled:
    """Where cover-page flag is FALSE: no cover content, no leading blank page,
    but document control is emitted unchanged."""

    def test_no_cover_title_when_disabled(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _paragraphs_text(doc)
        # The report title should NOT appear as a cover paragraph
        # (it would be in Cover Title style which only the cover uses)
        # First substantive paragraph should be the document control heading
        non_empty = [t for t in texts if t.strip()]
        assert non_empty[0] == "Document control"

    def test_no_page_break_before_document_control(self) -> None:
        """No leading blank page — no page break emitted before document control."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        # The first paragraph should be document control heading (not a break)
        first_para = doc.paragraphs[0]
        assert first_para.text == "Document control"

    def test_document_control_still_emitted_when_cover_disabled(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                document_name="Monthly Infra Report",
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _document_text(doc)
        assert any("Monthly Infra Report" in t for t in texts)

    def test_approvers_table_still_emitted_when_cover_disabled(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        assert len(tables) >= 1  # approvers table


# ---------------------------------------------------------------------------
# Clause (b) — empty ruled signature box, NOT typed name
# ---------------------------------------------------------------------------


class TestSignatureBox:
    """A role with no supplied signature image gets an EMPTY RULED SIGNATURE BOX
    at the theme's declared height, and emphatically NOT that role's typed name."""

    def test_signature_cell_is_empty_not_typed_name(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(role="author", name="Alice Smith", title="Engineer"),
                    ApproverEntry(role="reviewer", name="Bob Jones", title="Manager"),
                ),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        assert len(tables) >= 1
        approvers_table = _approvers_table(doc)

        # Check every signature cell (column 3) in data rows
        for row_idx in range(1, len(approvers_table.rows)):
            sig_cell = approvers_table.rows[row_idx].cells[3]
            # Signature cell MUST be empty — never the typed name
            assert sig_cell.text == "", (
                f"Row {row_idx} signature cell contains {sig_cell.text!r} — "
                f"must be empty (an empty ruled box, not a typed name)"
            )

    def test_name_appears_in_name_column_not_signature(self) -> None:
        """The approver's name goes in column 2 (Name), never column 3 (Signature)."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(role="author", name="Alice Smith", title="Lead"),
                ),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        # Row 1 is the "author" row (row 0 is header)
        author_row = approvers_table.rows[1]
        assert author_row.cells[2].text == "Alice Smith"
        assert author_row.cells[3].text == ""  # empty ruled box

    def test_signature_row_height_is_declared(self) -> None:
        """Each approver row has w:trHeight set to the theme's declared height."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(role="author", name="A", title="T"),
                ),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        # Check row 1 (first data row) has the height set
        tr = approvers_table.rows[1]._tr
        tr_height_els = tr.findall(f".//{qn('w:trHeight')}")
        assert len(tr_height_els) >= 1
        height_val = tr_height_els[0].get(qn("w:val"))
        assert height_val == str(SIGNATURE_BOX_HEIGHT_TWIPS)

    def test_signed_cell_contains_the_image_and_still_no_text(self) -> None:
        """A row carrying a signature image places that image in the signature
        cell, and the cell's text remains empty — the image is what fills the
        box, never a typed name alongside or instead of it (Req 13.3, 13.4)."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(
                        role="author",
                        name="Alice Smith",
                        title="Engineer",
                        signature_image=_ONE_PIXEL_PNG,
                    ),
                    ApproverEntry(role="reviewer", name="Bob Jones", title="Manager"),
                ),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        author_row = approvers_table.rows[1]
        reviewer_row = approvers_table.rows[2]

        # The signed row's signature cell has no text...
        assert author_row.cells[3].text == ""
        # ...and DOES contain a placed picture.
        signed_pictures = author_row.cells[3]._tc.findall(
            f".//{qn('w:drawing')}"
        )
        assert len(signed_pictures) == 1

        # The unsigned row's signature cell has no text and no picture.
        assert reviewer_row.cells[3].text == ""
        unsigned_pictures = reviewer_row.cells[3]._tc.findall(
            f".//{qn('w:drawing')}"
        )
        assert len(unsigned_pictures) == 0

        # The name still prints in the name column for the signed row.
        assert author_row.cells[2].text == "Alice Smith"

    def test_a_signed_rows_height_equals_an_unsigned_rows(self) -> None:
        """Req 13.3 — a signed row and an unsigned row occupy the same space,
        so the document's pagination does not depend on who signed."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(
                        role="author",
                        name="Alice Smith",
                        signature_image=_ONE_PIXEL_PNG,
                    ),
                    ApproverEntry(role="reviewer", name="Bob Jones"),
                ),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)

        def row_height(row_idx: int) -> str | None:
            tr = approvers_table.rows[row_idx]._tr
            els = tr.findall(f".//{qn('w:trHeight')}")
            return els[0].get(qn("w:val")) if els else None

        signed_height = row_height(1)
        unsigned_height = row_height(2)
        assert signed_height == unsigned_height == str(SIGNATURE_BOX_HEIGHT_TWIPS)

    def test_the_role_column_shows_the_positional_label_not_the_stored_id(
        self,
    ) -> None:
        """Req 12.3 — the role column reads `Author`, not the stored id
        `author`. The label is resolved through the message catalog, and the
        stored role id never reaches the rendered document."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=False),
            document_control=DocumentControlConfig(
                approvers=(ApproverEntry(role="author", name="Alice"),),
            ),
        )
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        role_cells = [
            approvers_table.rows[i].cells[0].text
            for i in range(1, len(approvers_table.rows))
        ]
        assert "author" not in role_cells
        assert "Author" in role_cells


# ---------------------------------------------------------------------------
# Clause (c) — absent per-run values raise RENDER_FAILED
# ---------------------------------------------------------------------------


class TestAbsentPerRunValues:
    """A per-run value that is absent raises RENDER_FAILED naming that value."""

    def test_absent_customer_name(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = RunFacts(
            run_id="r1", template_id="t1", customer_name="",
            period_display="July 2026", report_title="Report",
        )
        msgs = load_messages("en")
        with pytest.raises(RenderFailedError, match="customer_name"):
            emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

    def test_absent_period_display(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = RunFacts(
            run_id="r1", template_id="t1", customer_name="C",
            period_display="", report_title="Report",
        )
        msgs = load_messages("en")
        with pytest.raises(RenderFailedError, match="period_display"):
            emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

    def test_absent_report_title(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = RunFacts(
            run_id="r1", template_id="t1", customer_name="C",
            period_display="P", report_title="   ",
        )
        msgs = load_messages("en")
        with pytest.raises(RenderFailedError, match="report_title"):
            emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

    def test_absent_run_id(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = RunFacts(
            run_id="", template_id="t1", customer_name="C",
            period_display="P", report_title="R",
        )
        msgs = load_messages("en")
        with pytest.raises(RenderFailedError, match="run_id"):
            emit_front_matter(doc, front_matter=config, run=run, messages=msgs)


# ---------------------------------------------------------------------------
# Message resolution — both languages
# ---------------------------------------------------------------------------


class TestMessageResolution:
    """Every fixed string is resolved by id through messages (Req 15.3)."""

    def test_document_control_heading_in_english(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _paragraphs_text(doc)
        assert "Document control" in texts

    def test_document_control_heading_in_indonesian(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("id")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _paragraphs_text(doc)
        assert "Kendali dokumen" in texts

    def test_approver_headers_resolved_in_english(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        header_row = approvers_table.rows[0]
        # Blank, as `ReportA.dc.html` leaves it: the column holds Author, Quality
        # Control, Reviewed By and Customer, and a reader does not need those captioned.
        assert header_row.cells[0].text == ""
        assert header_row.cells[1].text == "Company"
        assert header_row.cells[2].text == "Name"
        assert header_row.cells[3].text == "Signature"

    def test_approver_headers_resolved_in_indonesian(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(cover=CoverConfig(enabled=False))
        run = _run()
        msgs = load_messages("id")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        tables = _tables(doc)
        approvers_table = _approvers_table(doc)
        header_row = approvers_table.rows[0]
        # Blank, as `ReportA.dc.html` leaves it: the column holds Author, Quality
        # Control, Reviewed By and Customer, and a reader does not need those captioned.
        assert header_row.cells[0].text == ""
        assert header_row.cells[1].text == "Perusahaan"
        assert header_row.cells[2].text == "Nama"
        assert header_row.cells[3].text == "Tanda tangan"


# ---------------------------------------------------------------------------
# Integration — full config
# ---------------------------------------------------------------------------


class TestFullIntegration:
    """End-to-end with realistic config."""

    def test_full_config_produces_paragraphs_and_table(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            cover=CoverConfig(enabled=True, subtitle="Monthly Review"),
            document_control=DocumentControlConfig(
                document_name="Infra Report",
                document_number_pattern="RPT-{template}-{year}{month}-{run}",
                distribution="Internal",
                approvers=(
                    ApproverEntry(role="author", name="A", title="T"),
                ),
            ),
        )
        run = RunFacts(
            run_id="r42", template_id="INFRA", customer_name="BigCo",
            period_display="Aug 2026", report_title="Cloud Report",
            revision_history=RevisionHistoryRow(
                revision="1.0", note="Initial", author="Alice",
            ),
            period_start_year="2026", period_start_month="08",
        )
        msgs = load_messages("en")

        emit_front_matter(doc, front_matter=config, run=run, messages=msgs)

        texts = _document_text(doc)
        # Cover content
        assert "Cloud Report" in texts
        assert "BigCo" in texts
        assert "Monthly Review" in texts
        # Document number on cover
        assert "RPT-INFRA-202608-r42" in texts
        # Document control
        assert "Document control" in texts
        assert any("Infra Report" in t for t in texts)
        assert any("RPT-INFRA-202608-r42" in t for t in texts)
        # Revision history — the revision and its note are the two cells of one row.
        assert "Revision history" in texts
        assert "1.0" in texts
        assert any("Initial" in t for t in texts)
        # Distribution
        assert any("Internal" in t for t in texts)
        # Table exists
        assert len(_tables(doc)) >= 1


# ---------------------------------------------------------------------------
# The two values the app collected and the document never showed
# ---------------------------------------------------------------------------


class TestApproverCompanyColumn:
    """The COMPANY column shows the approver's company.

    `ApproverEntry.company` was declared, collected by the wizard's approver rows, and
    dropped by `report_pipeline._front_matter_from` — so the column headed "Company"
    rendered `title`, the person's job title, instead.
    """

    def test_company_is_shown_when_supplied(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(
                        role="author",
                        name="Mayer Reflino Sitorus",
                        title="Cloud Engineer",
                        company="PT Helios Informatika Nusantara",
                    ),
                )
            )
        )
        emit_front_matter(
            doc, front_matter=config, run=_run(), messages=load_messages("en")
        )

        row = _approvers_table(doc).rows[1]
        assert row.cells[1].text == "PT Helios Informatika Nusantara"
        assert row.cells[2].text == "Mayer Reflino Sitorus"

    def test_title_remains_the_fallback(self) -> None:
        """A profile authored before this fix put its company in `title`, because that
        is the field the column rendered. Those keep rendering exactly as they did."""
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            document_control=DocumentControlConfig(
                approvers=(
                    ApproverEntry(role="author", name="A. Person", title="Acme Ltd"),
                )
            )
        )
        emit_front_matter(
            doc, front_matter=config, run=_run(), messages=load_messages("en")
        )

        assert _approvers_table(doc).rows[1].cells[1].text == "Acme Ltd"


class TestDistributionRows:
    """Req 12.6's v3 rows reach the document as a table, not as a Python repr."""

    def test_v3_rows_render_as_their_own_values(self) -> None:
        from reporting_agent.render.front_matter import DistributionRow

        doc = load_theme("editorial")
        config = FrontMatterConfig(
            document_control=DocumentControlConfig(
                distribution_rows=(
                    DistributionRow(
                        recipient="Doni Prasetyo",
                        company="PT. Herlina Indah",
                        note="Softcopy",
                    ),
                )
            )
        )
        emit_front_matter(
            doc, front_matter=config, run=_run(), messages=load_messages("en")
        )

        texts = _document_text(doc)
        assert "Doni Prasetyo" in texts
        assert "PT. Herlina Indah" in texts
        assert "Softcopy" in texts
        # The defect: `str()` over the row list put its repr in a signed document.
        assert not any("DistributionRow(" in t or "{'recipient'" in t for t in texts)

    def test_the_v1_v2_free_text_block_still_renders(self) -> None:
        doc = load_theme("editorial")
        config = FrontMatterConfig(
            document_control=DocumentControlConfig(distribution="Internal only")
        )
        emit_front_matter(
            doc, front_matter=config, run=_run(), messages=load_messages("en")
        )

        assert "Internal only" in _document_text(doc)


class TestResolveFrontMatterConfigCarriesEveryField:
    """`report_pipeline._resolve_front_matter_config` is the wire between the definition
    the app writes and the renderer that draws it.

    Both bugs above were **here**, not in the renderer: `company` was not read, and
    `distribution` was coerced with `str()` whatever shape it held. These assert the
    wire carries what the app puts on it.
    """

    def test_the_approver_company_survives_the_wire(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config({
            "front_matter": {
                "document_control": {
                    "approvers": [
                        {
                            "role": "author",
                            "name": "Mayer Reflino Sitorus",
                            "title": "Cloud Engineer",
                            "company": "PT Helios Informatika Nusantara",
                        }
                    ]
                }
            }
        })

        approver = config.document_control.approvers[0]
        assert approver.company == "PT Helios Informatika Nusantara"
        assert approver.title == "Cloud Engineer"

    def test_v3_distribution_rows_arrive_as_rows(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config({
            "front_matter": {
                "document_control": {
                    "distribution": [
                        {
                            "recipient": "Doni Prasetyo",
                            "company": "PT. Herlina Indah",
                            "note": "Softcopy",
                        }
                    ]
                }
            }
        })

        control = config.document_control
        assert control.distribution is None
        assert len(control.distribution_rows) == 1
        assert control.distribution_rows[0].recipient == "Doni Prasetyo"

    def test_the_v1_v2_string_still_arrives_as_a_string(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config({
            "front_matter": {"document_control": {"distribution": "Internal only"}}
        })

        control = config.document_control
        assert control.distribution == "Internal only"
        assert control.distribution_rows == ()


class TestFrontMatterImages:
    """The bytes behind `signature_key` and `logo_key`.

    `ApproverEntry.signature_image` was declared when the front matter was written and
    `_place_signature_image` has always known how to draw one, but nothing ever loaded
    the bytes — so an approver who uploaded a signature got the identical empty ruled box
    as one who had not. `report_pipeline._load_front_matter_images` is the missing middle,
    and these are the properties that make it safe to have: it reads only this run's own
    objects, and it never fails a report.
    """

    # Real 1×1 images, not a magic number followed by zeros: the docx emitter hands
    # these to python-docx, which parses the header to size the shape and refuses a
    # truncated file.
    PNG = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
        "IQAAAABJRU5ErkJggg=="
    )
    JPEG = base64.b64decode(
        "/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsLDBkSEw8UHRofHh0a"
        "HBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/wAALCAABAAEBAREA/8QAFAABAAAAAAAA"
        "AAAAAAAAAAAACf/EABQQAQAAAAAAAAAAAAAAAAAAAAD/2gAIAQEAAD8AVN//2Q=="
    )

    class _Store:
        """The two `ObjectStore` behaviours this loader depends on, and nothing else."""

        def __init__(self, objects: dict[str, bytes]) -> None:
            self.objects = objects
            self.reads: list[str] = []

        async def get_bytes(self, key: str) -> bytes:
            self.reads.append(key)
            if key not in self.objects:
                raise KeyError(key)
            return self.objects[key]

    def _definition(self, *, logo_key: str | None, signature_key: str | None) -> dict:
        cover: dict[str, object] = {"logo": "https://cdn.example/logo.png"}
        if logo_key is not None:
            cover["logo_key"] = logo_key
        approver: dict[str, object] = {"role": "author", "name": "R. Prakoso"}
        if signature_key is not None:
            approver["signature_key"] = signature_key
        return {
            "front_matter": {
                "cover": cover,
                "document_control": {"approvers": [approver]},
            }
        }

    def _load(self, definition: dict, store: object, actor: str = "alice") -> dict:
        # `asyncio.run` rather than a plugin marker, matching `test_collect_pipeline.py`
        # and the rest of the suite: the loader is the only awaitable here.
        from reporting_agent.report_pipeline import _load_front_matter_images

        return asyncio.run(
            _load_front_matter_images(definition, store=store, actor_id=actor)
        )

    def test_both_kinds_of_image_are_read(self) -> None:
        store = self._Store({
            "alice/logos/l.png": self.PNG,
            "alice/signatures/s.png": self.JPEG,
        })

        images = self._load(
            self._definition(
                logo_key="alice/logos/l.png",
                signature_key="alice/signatures/s.png",
            ),
            store,
        )

        assert images == {
            "alice/logos/l.png": self.PNG,
            "alice/signatures/s.png": self.JPEG,
        }

    def test_another_tenants_key_is_never_read(self) -> None:
        """Ownership is checked *before* the read, not after: a definition naming
        another tenant's object must not produce a GET at all. The prefix check is on
        the whole first segment, so `alice-evil/...` is not `alice`'s."""
        store = self._Store({
            "bob/logos/l.png": self.PNG,
            "alice-evil/signatures/s.png": self.PNG,
        })

        images = self._load(
            self._definition(
                logo_key="bob/logos/l.png",
                signature_key="alice-evil/signatures/s.png",
            ),
            store,
        )

        assert images == {}
        assert store.reads == []

    def test_a_key_that_cannot_be_read_is_not_fatal(self) -> None:
        """Withholding a delivered report because a logo would not load is the wrong
        trade by a wide margin — and the resulting document is the one criterion 13.6
        clause (b) already requires for an approver who supplied no signature."""
        store = self._Store({"alice/logos/l.png": self.PNG})

        images = self._load(
            self._definition(
                logo_key="alice/logos/l.png",
                signature_key="alice/signatures/missing.png",
            ),
            store,
        )

        assert images == {"alice/logos/l.png": self.PNG}

    def test_a_definition_naming_no_images_reads_nothing(self) -> None:
        store = self._Store({"alice/logos/l.png": self.PNG})

        images = self._load(
            self._definition(logo_key=None, signature_key=None), store
        )

        assert images == {}
        assert store.reads == []

    def test_the_loaded_bytes_reach_both_config_fields(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config(
            self._definition(
                logo_key="alice/logos/l.png",
                signature_key="alice/signatures/s.png",
            ),
            {"alice/logos/l.png": self.PNG, "alice/signatures/s.png": self.JPEG},
        )

        assert config.cover.logo_key == "alice/logos/l.png"
        assert config.cover.logo_image == self.PNG
        assert config.document_control.approvers[0].signature_image == self.JPEG

    def test_an_unreadable_key_leaves_the_fields_none(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config(
            self._definition(
                logo_key="alice/logos/l.png",
                signature_key="alice/signatures/s.png",
            ),
            {},
        )

        assert config.cover.logo_key == "alice/logos/l.png"
        assert config.cover.logo_image is None
        assert config.document_control.approvers[0].signature_image is None


class TestCoverLogoIsDrawn:
    """A configured logo reaches both emitters, and an unfetchable one reserves the same
    space rather than shifting every element on the cover."""

    # Real 1×1 images, not a magic number followed by zeros: the docx emitter hands
    # these to python-docx, which parses the header to size the shape and refuses a
    # truncated file.
    PNG = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
        "IQAAAABJRU5ErkJggg=="
    )
    JPEG = base64.b64decode(
        "/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsLDBkSEw8UHRofHh0a"
        "HBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/wAALCAABAAEBAREA/8QAFAABAAAAAAAA"
        "AAAAAAAAAAAACf/EABQQAQAAAAAAAAAAAAAAAAAAAAD/2gAIAQEAAD8AVN//2Q=="
    )

    def _logo(self, image: bytes | None) -> object:
        return FrontMatterLogo(height_pt=34.0, width_pt=104.0, image=image)

    def test_the_html_emitter_inlines_the_bytes_as_a_data_uri(self) -> None:
        from reporting_agent.render.html import emit_front_matter_html

        markup = emit_front_matter_html([self._logo(self.PNG)])

        expected = base64.b64encode(self.PNG).decode("ascii")
        assert f'src="data:image/png;base64,{expected}"' in markup

    def test_the_media_type_comes_from_the_bytes_not_the_key(self) -> None:
        from reporting_agent.render.html import emit_front_matter_html

        markup = emit_front_matter_html([self._logo(self.JPEG)])

        assert "data:image/jpeg;base64," in markup
        assert "image/png" not in markup

    def test_the_reading_copy_makes_no_network_request_for_the_logo(self) -> None:
        """The bytes are already in hand, and the process that rasterises this markup
        has no access to the artifact bucket and no business reaching the internet."""
        from reporting_agent.render.html import emit_front_matter_html

        markup = emit_front_matter_html([self._logo(self.PNG)])

        assert "http://" not in markup
        assert "https://" not in markup

    def test_an_absent_logo_still_reserves_its_block(self) -> None:
        from reporting_agent.render.html import emit_front_matter_html

        markup = emit_front_matter_html([self._logo(None)])

        assert "height:34pt" in markup
        assert "<img" not in markup

    def test_the_docx_emitter_places_the_picture(self) -> None:
        from reporting_agent.render.front_matter import _emit_section

        document = load_theme("editorial")
        before = len(document.inline_shapes)
        _emit_section(document, self._logo(self.PNG))

        assert len(document.inline_shapes) == before + 1

    def test_the_configs_bytes_reach_the_section(self) -> None:
        """The join between the two halves: `_load_front_matter_images` puts the bytes on
        `CoverConfig`, and `front_matter_sections` is what carries them to the emitters.
        Without this, both halves can be right and the cover still draw nothing."""
        from reporting_agent.render.front_matter import front_matter_sections

        sections = front_matter_sections(
            front_matter=FrontMatterConfig(
                cover=CoverConfig(
                    enabled=True,
                    logo="https://cdn.example/logo.png",
                    logo_key="alice/logos/l.png",
                    logo_image=self.PNG,
                ),
                document_control=DocumentControlConfig(document_name="Report"),
            ),
            run=_run(),
            messages=load_messages("en"),
        )

        logo = next(s for s in sections if isinstance(s, FrontMatterLogo))
        assert logo.image == self.PNG

    def test_a_cover_with_no_stored_bytes_still_reserves_the_block(self) -> None:
        from reporting_agent.render.front_matter import front_matter_sections

        sections = front_matter_sections(
            front_matter=FrontMatterConfig(
                cover=CoverConfig(enabled=True, logo="https://cdn.example/logo.png"),
                document_control=DocumentControlConfig(document_name="Report"),
            ),
            run=_run(),
            messages=load_messages("en"),
        )

        logo = next(s for s in sections if isinstance(s, FrontMatterLogo))
        assert logo.image is None
        assert logo.height_pt > 0

    def test_the_docx_emitter_places_nothing_when_there_are_no_bytes(self) -> None:
        from reporting_agent.render.front_matter import _emit_section

        document = load_theme("editorial")
        _emit_section(document, self._logo(None))

        assert len(document.inline_shapes) == 0


class TestConfidentialityNotice:
    """Requirement 12.7 — the notice is Brand-owned, printed under its own heading.

    It never appeared in a delivered document, and not for one reason: the Brand editor
    exposed no control, nothing resolved the Brand's value into the definition, and the
    wizard deliberately wrote nothing, citing a Brand that never supplied it. This covers
    the renderer's half — that a notice which *is* configured reaches both emitters under a
    heading, and that one which is not prints nothing rather than an empty heading.
    """

    NOTICE = (
        "The information in this document is confidential and is owned by "
        "PT. Helios Informatika Nusantara."
    )

    def _sections(self, **control) -> tuple:
        return front_matter_sections(
            front_matter=FrontMatterConfig(
                cover=CoverConfig(enabled=False),
                document_control=DocumentControlConfig(
                    document_name="Report", **control
                ),
            ),
            run=_run(),
            messages=_messages(),
        )

    def _heading_texts(self, sections) -> list[str]:
        from reporting_agent.render.front_matter import FrontMatterHeading

        return [s.text for s in sections if isinstance(s, FrontMatterHeading)]

    def _note_texts(self, sections) -> list[str]:
        from reporting_agent.render.front_matter import FrontMatterNote

        return [s.text for s in sections if isinstance(s, FrontMatterNote)]

    def test_the_prose_is_printed_under_its_own_heading(self) -> None:
        sections = self._sections(confidentiality_notice=self.NOTICE)

        assert (
            _messages().text("doc.front_matter.confidentiality")
            in self._heading_texts(sections)
        )
        assert self.NOTICE in self._note_texts(sections)

    def test_the_heading_comes_before_the_notice(self) -> None:
        from reporting_agent.render.front_matter import (
            FrontMatterHeading,
            FrontMatterNote,
        )

        sections = self._sections(confidentiality_notice=self.NOTICE)
        heading = _messages().text("doc.front_matter.confidentiality")

        index_of_heading = next(
            i
            for i, s in enumerate(sections)
            if isinstance(s, FrontMatterHeading) and s.text == heading
        )
        index_of_notice = next(
            i
            for i, s in enumerate(sections)
            if isinstance(s, FrontMatterNote) and s.text == self.NOTICE
        )
        assert index_of_heading < index_of_notice

    def test_no_notice_prints_no_heading(self) -> None:
        """An empty heading over nothing is worse than no heading."""
        sections = self._sections()

        assert (
            _messages().text("doc.front_matter.confidentiality")
            not in self._heading_texts(sections)
        )

    def test_the_prose_wins_over_a_catalogue_id(self) -> None:
        """A person wrote one of these for this account; the other names copy every
        tenant shares."""
        sections = self._sections(
            confidentiality_notice=self.NOTICE,
            confidentiality_notice_id="doc.front_matter.confidentiality",
        )

        assert self.NOTICE in self._note_texts(sections)

    def test_the_catalogue_id_still_works_on_its_own(self) -> None:
        """v1 and v2 definitions carry an id and no prose, and must render as they did."""
        sections = self._sections(
            confidentiality_notice_id="doc.front_matter.confidentiality"
        )

        assert (
            _messages().text("doc.front_matter.confidentiality")
            in self._note_texts(sections)
        )

    def test_both_emitters_print_it(self) -> None:
        from reporting_agent.render.html import emit_front_matter_html

        sections = self._sections(confidentiality_notice=self.NOTICE)

        markup = emit_front_matter_html(list(sections))
        assert self.NOTICE in markup

        document = load_theme("editorial")
        for section in sections:
            from reporting_agent.render.front_matter import _emit_section

            _emit_section(document, section)
        assert self.NOTICE in _document_text(document)

    def test_the_definition_carries_it_to_the_config(self) -> None:
        from reporting_agent.report_pipeline import _resolve_front_matter_config

        config = _resolve_front_matter_config({
            "front_matter": {"document_control": {"confidentiality_notice": self.NOTICE}}
        })

        assert config.document_control.confidentiality_notice == self.NOTICE
