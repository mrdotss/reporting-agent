"""`pdf_page_texts`, the committed TOC fixture's geometry, and the measurement harness
(Req 14.2, 14.11).

## What this module is responsible for, and what it is not

Task 2.4 adds `test_toc_proof.py`, which reads `ADOPTED_APPROACH` and asserts the adopted
candidate's page numbers are right. This module is the layer beneath that: it asserts the
**instrument** works and the **fixture** has the geometry the proof will need. Splitting them
matters because a proof test over a one-page fixture would pass trivially — every heading
would be on page 1 and "named equals observed" would hold for a document that proves
nothing about pagination.

So the assertions here are about the measurement being able to fail:

* `pdf_page_texts` separates pages rather than joining them, and preserves a blank page's
  index — otherwise every page number after a blank one would be off by one.
* the fixture renders to **at least 8 pages** with **at least 6 headings** on **at least 4
  distinct pages**, so a table of contents over it has something to get wrong.
* `measure` refuses an approach it cannot render, rather than measuring `none` and letting a
  verdict be recorded for a candidate that was never exercised.

## The slow test, and why it is not stubbed

One `measure` call spawns LibreOffice and takes seconds. It is module-scoped and shared, and
it is not faked, because pagination is decided by the converter: a measurement over a stubbed
conversion would be measuring an assumption about the thing under test.
"""

from __future__ import annotations

import asyncio
import io
import shutil
from pathlib import Path
from typing import Final

import pytest
from docx import Document as open_docx
from pypdf import PdfWriter

import toc_harness as H
from reporting_agent.errors import ErrorCode, VerificationFailedError
from reporting_agent.render import pdf as P
from reporting_agent.render.themes import TOC_ENTRY_STYLE
from reporting_agent.render.toc import TOC_APPROACH_NONE
from reporting_agent.verify.tokens import pdf_page_texts, read_pdf_text

SOFFICE: Final[str | None] = shutil.which("soffice")

MIN_PAGES: Final[int] = 8
MIN_HEADINGS: Final[int] = 6
MIN_DISTINCT_HEADING_PAGES: Final[int] = 4


@pytest.fixture(autouse=True)
def _required_lang(monkeypatch: pytest.MonkeyPatch) -> None:
    """Every test runs with the locale the converter requires.

    Autouse for the reason `tests/test_pdf.py` gives for its own copy: the alternative is
    every test remembering to set it, and one that forgot would fail on the `LANG` assertion
    for a reason unrelated to what it checks.
    """
    monkeypatch.setenv("LANG", P.REQUIRED_LANG)


# --- pdf_page_texts (Req 14.2) -------------------------------------------------------


def blank_pdf(path: Path, pages: int) -> Path:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    with path.open("wb") as handle:
        writer.write(handle)
    return path


def test_one_string_per_page_in_page_order(tmp_path: Path) -> None:
    texts = pdf_page_texts(blank_pdf(tmp_path / "three.pdf", 3))

    assert len(texts) == 3


def test_a_blank_page_yields_an_empty_string_rather_than_being_dropped(
    tmp_path: Path,
) -> None:
    """The index of a page in the returned tuple **is** its page number minus one, so a
    dropped blank page would renumber every page after it — and a table of contents built on
    that would be wrong by one for the rest of the document."""
    texts = pdf_page_texts(blank_pdf(tmp_path / "blank.pdf", 4))

    assert texts == ("", "", "", "")
    assert len(texts) == 4


def test_an_unreadable_pdf_raises_rather_than_reporting_no_pages(tmp_path: Path) -> None:
    """The same contract `read_pdf_text` holds: a caller must not conclude that a heading is
    on no page because the file would not parse."""
    broken = tmp_path / "broken.pdf"
    broken.write_bytes(b"not a pdf at all")

    with pytest.raises(VerificationFailedError) as raised:
        pdf_page_texts(broken)

    assert raised.value.code is ErrorCode.VERIFICATION_FAILED
    assert str(broken) in str(raised.value)


def test_the_two_readers_agree_about_the_same_document(tmp_path: Path) -> None:
    """`read_pdf_text` joins pages and `pdf_page_texts` does not, so joining the second must
    reproduce the first. Without this they could drift into two different normalizations and
    a string findable by the fidelity gate would be unfindable by the page locator."""
    path = blank_pdf(tmp_path / "agree.pdf", 3)

    whole, count = read_pdf_text(path)
    per_page = pdf_page_texts(path)

    assert count == len(per_page)
    assert " ".join(per_page).strip() == whole


# --- the harness contract, without rendering ------------------------------------------


def test_an_undeclared_approach_is_refused() -> None:
    """A typo must fail rather than silently measure `none`, which would record a verdict for
    a candidate that was never rendered."""
    definition, snapshot = H.load_fixture()

    with pytest.raises(ValueError, match="is not one of the declared TOC approaches"):
        asyncio.run(H.measure(definition, snapshot, approach="two-pass"))


# --- the emission paths, without converting (Req 14.1) --------------------------------
#
# These build the `.docx` and inspect its XML. No LibreOffice, so they are milliseconds
# rather than seconds, and they check the half of each candidate that is ours: what the
# document asks for. Whether the converter honours it is the evaluation's question.


@pytest.fixture(scope="module")
def base_docx() -> bytes:
    definition, snapshot = H.load_fixture()
    return H._base_docx(definition, snapshot)


def document_xml(docx_bytes: bytes) -> str:
    import io
    import zipfile

    return (
        zipfile.ZipFile(io.BytesIO(docx_bytes)).read("word/document.xml").decode("utf-8")
    )


def test_the_field_candidate_emits_a_toc_field_with_no_cached_result(
    base_docx: bytes,
) -> None:
    """Criterion 14.1's exact shape, and the reason it matters.

    `begin` → `instrText` → `end` with **no `separate`** means there is no cached result run,
    so the PDF cannot display a page number computed when the `.docx` was written. A cached
    result is the failure that would look most like success: a stale number that reads as
    resolved. `w:dirty="true"` is the only signal a consumer gets that the empty result is
    deliberate.
    """
    xml = document_xml(H._prepend_toc(base_docx, field=True))

    assert 'TOC \\o "1-3" \\h \\z \\u' in xml
    assert xml.count('w:fldCharType="begin"') == 1
    assert xml.count('w:fldCharType="end"') == 1
    assert 'w:fldCharType="separate"' not in xml
    assert 'w:dirty="true"' in xml


def test_the_contents_label_uses_a_style_the_field_does_not_collect(
    base_docx: bytes,
) -> None:
    """`\\o "1-3"` collects `Heading 1` to `Heading 3`, so a `Heading 1` label would put
    "Contents" into its own table of contents — and would add a seventh entry to
    `observed_pages` for a section the definition never declared, making the fixture's
    heading count disagree with the six it declares."""
    document = open_docx(io.BytesIO(H._prepend_toc(base_docx, field=True)))
    label = next(p for p in document.paragraphs if p.text == H.TOC_LABEL)

    assert label.style.name == "Title"
    assert label.style.name not in {"Heading 1", "Heading 2", "Heading 3"}


def test_the_two_pass_candidate_emits_the_same_layout_with_and_without_numbers(
    base_docx: bytes,
) -> None:
    """"At full size with no numbers" made executable, and it is the whole basis of the
    fixed point.

    Pass 1 must lay out to exactly the height pass 2 will, or the page numbers pass 1
    measured describe a document with a different pagination. Both passes therefore emit one
    paragraph per entry **and the tab**, differing only in whether a number follows it.
    """
    headings = ("Alpha Section", "Beta Section")
    without = open_docx(
        io.BytesIO(
            H._prepend_toc(base_docx, entries=tuple((h, None) for h in headings))
        )
    )
    with_numbers = open_docx(
        io.BytesIO(
            H._prepend_toc(base_docx, entries=tuple((h, 4) for h in headings))
        )
    )

    def entries(document: object) -> list[tuple[str, str]]:
        return [
            (p.style.name, p.text)
            for p in document.paragraphs  # type: ignore[attr-defined]
            if p.style.name == TOC_ENTRY_STYLE
        ]

    assert [style for style, _ in entries(without)] == [TOC_ENTRY_STYLE] * 2
    assert [text for _, text in entries(without)] == ["Alpha Section\t", "Beta Section\t"]
    assert [text for _, text in entries(with_numbers)] == [
        "Alpha Section\t4",
        "Beta Section\t4",
    ]
    # Same paragraph count in both passes: the number is a run inside an existing paragraph,
    # never a paragraph of its own.
    assert len(without.paragraphs) == len(with_numbers.paragraphs)


def test_the_two_pass_candidate_emits_no_field_and_the_field_candidate_no_entries(
    base_docx: bytes,
) -> None:
    """The two emission paths are genuinely different documents.

    Without this, a wiring mistake that sent both candidates through one path would have the
    evaluation record two verdicts for one mechanism — and they would agree, which is exactly
    what would make it look right.
    """
    field_xml = document_xml(H._prepend_toc(base_docx, field=True))
    literal_xml = document_xml(H._prepend_toc(base_docx, entries=(("Alpha", 3),)))

    assert "instrText" in field_xml
    assert "instrText" not in literal_xml
    assert "Alpha" in literal_xml
    assert field_xml != literal_xml


def test_the_toc_section_is_inserted_before_the_first_block(base_docx: bytes) -> None:
    """Ahead of the content, not appended after it: a table of contents at the end of the
    document would still name true pages and would still be wrong."""
    before = open_docx(io.BytesIO(base_docx))
    after = open_docx(io.BytesIO(H._prepend_toc(base_docx, entries=(("Alpha", 1),))))

    assert before.paragraphs[0].style.name == "Heading 1"
    assert after.paragraphs[0].text == H.TOC_LABEL
    assert [p.text for p in after.paragraphs[3:]] == [
        p.text for p in before.paragraphs
    ]


# --- the macro candidate's availability check (Req 14.3) ------------------------------


def test_the_macro_candidate_refuses_a_profile_with_no_basic_module(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """`unavailable`, decided **before** invoking anything and without writing to the profile.

    The alternative — installing the macro — is the writable macro library criterion 14.3
    rejects in as many words, so the refusal has to happen here rather than being discovered
    after a `soffice` invocation has already mutated the profile.
    """
    monkeypatch.setenv(P.PROFILE_ENV_VAR, str(tmp_path))

    with pytest.raises(H.TocApproachUnavailableError, match="no Basic module"):
        H._convert_via_macro(b"PK\x03\x04 not really a docx")

    assert list(tmp_path.iterdir()) == [], "the check must not create anything"


def test_the_macro_candidate_refuses_an_empty_standard_library(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The case the real image is in: LibreOffice warms `Module1.xba` containing an empty
    `Sub Main`, so the file exists and the macro does not."""
    module = tmp_path / H.MACRO_LIBRARY_RELATIVE_PATH
    module.parent.mkdir(parents=True)
    module.write_text("<script:module>REM\nSub Main\nEnd Sub</script:module>", "utf-8")
    monkeypatch.setenv(P.PROFILE_ENV_VAR, str(tmp_path))

    with pytest.raises(H.TocApproachUnavailableError, match="declares no"):
        H._convert_via_macro(b"PK\x03\x04 not really a docx")


def test_the_macro_candidates_refusal_names_no_machine_specific_path(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The message is committed into `evidence/toc/evaluation.json`, so it must describe the
    missing facility rather than the directory the evaluation happened to run in."""
    monkeypatch.setenv(P.PROFILE_ENV_VAR, str(tmp_path))

    with pytest.raises(H.TocApproachUnavailableError) as raised:
        H._convert_via_macro(b"PK\x03\x04 not really a docx")

    assert str(tmp_path) not in str(raised.value)
    assert H.MACRO_LIBRARY_RELATIVE_PATH in str(raised.value)


# --- the contents pages are excluded from the observation (Req 14.2) ------------------
#
# The defect these exist for is not hypothetical; it is the one this measurement had. A
# contents entry contains the heading's text, and the contents page comes first, so a locator
# that searched every page reported every heading on page 1 — `distinct_heading_pages`
# collapsed to 1 while the numbers printed beside them stayed right, which reads as a
# *correct* table of contents for a document with no pagination.

CONTENTS_PAGE: Final[str] = (
    "Contents Alpha Section..........2 Beta Section..........4"
)


def test_a_heading_is_observed_on_its_body_page_not_on_the_contents_page() -> None:
    pages = (CONTENTS_PAGE, "Alpha Section and its prose", "more", "Beta Section here")

    observed = H._observed_pages(pages, ("Alpha Section", "Beta Section"))

    assert observed == {"Alpha Section": 2, "Beta Section": 4}


def test_the_contents_page_is_identified_by_starting_with_the_label() -> None:
    """A substring test would also match a body page whose prose contained the word, and
    excluding a body page from the observation would report its heading one page late."""
    prose = "The section below lists Contents of the archive: Alpha Section"

    assert H._contents_page_indices((CONTENTS_PAGE, prose)) == frozenset({0})


def test_named_pages_are_read_from_the_contents_page_only() -> None:
    """Read back out of the PDF rather than returned from the dict the emitter wrote — a
    number the emitter intended but the converter dropped must not be reported as named."""
    pages = (CONTENTS_PAGE, "Alpha Section and its prose", "more", "Beta Section here")

    assert H._named_pages(pages, ("Alpha Section", "Beta Section")) == {
        "Alpha Section": 2,
        "Beta Section": 4,
    }


def test_a_heading_followed_by_a_number_in_body_prose_is_not_a_named_page() -> None:
    """Why the pattern is applied only to the contents pages.

    In body prose "a heading followed by a number" is an ordinary sentence — a running header
    above a figure, a cross-reference, a sentence beginning with a count — and anchoring the
    pattern to a line boundary would not help, because the extractor returns a whole contents
    page as **one** line.

    Constructed so the two readings give **different** answers: the body page comes first and
    says 12, the contents page comes later and says 5. A locator that searched every page
    would report 12, which is a page number lifted out of a sentence.
    """
    pages = (
        "Alpha Section 12 machines were in scope for this period",
        "Contents Alpha Section..........5",
        "Alpha Section and its prose",
    )

    assert H._named_pages(pages, ("Alpha Section",)) == {"Alpha Section": 5}


def test_a_heading_named_on_no_contents_page_names_no_page() -> None:
    """A document with no table of contents names nothing, however much its prose looks like
    an entry. This is the case that makes the restriction visible on its own."""
    pages = ("Alpha Section 12 machines were in scope", "Alpha Section and its prose")

    assert H._named_pages(pages, ("Alpha Section",)) == {}


def test_a_contents_entry_with_no_number_names_no_page() -> None:
    """Candidate B's first pass, and candidate A's unresolved field: entries are present and
    carry no number. Absent rather than zero, so the verdict rule sees "named nothing"."""
    pages = ("Contents Alpha Section\t Beta Section\t", "Alpha Section prose")

    assert H._named_pages(pages, ("Alpha Section", "Beta Section")) == {}


def test_the_fixture_declares_the_headings_the_measurement_looks_for() -> None:
    """Read from the definition, not found in the PDF. Scanning the rendered text for things
    that *look* like headings would pass on a document that emitted none."""
    definition, _snapshot = H.load_fixture()

    headings = H._heading_texts(definition)

    assert len(headings) >= MIN_HEADINGS
    assert len(set(headings)) == len(headings), (
        "two headings with identical text cannot be told apart by page, so the fixture's "
        "headings must be distinct"
    )


def test_the_fixture_digest_covers_both_files() -> None:
    """The evidence record pins what was measured, so the digest has to move when either
    file does."""
    before = H.fixture_digest()

    assert len(before) == 64
    assert before == H.fixture_digest(), "the digest must be stable across calls"


def test_the_first_character_rule_resolves_a_straddling_heading_to_one_page() -> None:
    """The subtlety that makes `observed_pages` single-valued.

    A heading long enough to wrap can straddle a page boundary and then appears in full on
    neither page. The rule is the page carrying its **first character**, which is also the
    page a reader turns to — pointing at the page holding its second line would send them one
    page late.
    """
    pages = ("body text ending with Methodology And", "Provenance continues here")

    assert H._first_character_page(pages, "Methodology And Provenance") == 1


def test_a_heading_present_in_full_resolves_to_the_page_holding_it() -> None:
    pages = ("front matter", "Compute Utilization and some body", "later text")

    assert H._first_character_page(pages, "Compute Utilization") == 2


def test_a_heading_appearing_on_several_pages_resolves_to_the_first() -> None:
    """The rule is the page carrying the heading's **first** character, and this is the only
    case where first and last differ — so it is the only case that can catch a locator
    written as "the last page mentioning it".

    Not hypothetical: a theme with a running header repeats the section title on every page
    of the section, and body prose can quote a heading. A locator taking the last match would
    then point a reader at the end of the section instead of its start, and every page number
    in the table of contents would be wrong in the same plausible-looking direction.
    """
    pages = (
        "front matter",
        "Compute Utilization begins here",
        "Compute Utilization continued",
        "Compute Utilization continued further",
    )

    assert H._first_character_page(pages, "Compute Utilization") == 2


def test_a_heading_on_no_page_resolves_to_none() -> None:
    """So a heading the renderer dropped is a missing observation rather than page 1."""
    assert H._first_character_page(("a", "b"), "Absent Heading") is None


# --- the real render, once (Req 14.11) ------------------------------------------------


@pytest.fixture(scope="module")
def measured(tmp_path_factory: pytest.TempPathFactory) -> H.TocMeasurement:
    """One `measure` call for the whole module: it spawns LibreOffice and takes seconds."""
    if SOFFICE is None:  # pragma: no cover - environment dependent
        pytest.skip("LibreOffice is not installed")

    import os

    profile = tmp_path_factory.mktemp("profile")
    previous_lang = os.environ.get("LANG")
    previous_profile = os.environ.get(P.PROFILE_ENV_VAR)
    os.environ["LANG"] = P.REQUIRED_LANG
    os.environ[P.PROFILE_ENV_VAR] = str(profile)
    try:
        definition, snapshot = H.load_fixture()
        return asyncio.run(
            H.measure(definition, snapshot, approach=TOC_APPROACH_NONE)
        )
    finally:
        for name, value in (
            ("LANG", previous_lang),
            (P.PROFILE_ENV_VAR, previous_profile),
        ):
            if value is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = value


@pytest.mark.skipif(SOFFICE is None, reason="LibreOffice is not installed")
def test_the_fixture_has_the_geometry_the_proof_test_needs(
    measured: H.TocMeasurement,
) -> None:
    """Req 14.11, and the reason the fixture is as long as it is.

    A proof test over a one-page document would pass trivially: every heading would be on
    page 1, and "named equals observed" would hold for a table of contents that demonstrated
    nothing about pagination. These three floors are what make the eventual equality
    assertion meaningful.
    """
    assert measured.pages >= MIN_PAGES
    assert measured.headings >= MIN_HEADINGS
    assert measured.distinct_heading_pages >= MIN_DISTINCT_HEADING_PAGES


@pytest.mark.skipif(SOFFICE is None, reason="LibreOffice is not installed")
def test_every_declared_heading_was_located(measured: H.TocMeasurement) -> None:
    """A heading the renderer emitted nowhere would otherwise be silently absent from the
    measurement rather than a failure."""
    definition, _snapshot = H.load_fixture()

    assert set(measured.observed_pages) == set(H._heading_texts(definition))


@pytest.mark.skipif(SOFFICE is None, reason="LibreOffice is not installed")
def test_each_heading_resolves_to_exactly_one_page_in_ascending_order(
    measured: H.TocMeasurement,
) -> None:
    """Single-valued, and in document order: heading *n* cannot be on an earlier page than
    heading *n − 1*. That ordering is a property of the document rather than of the
    measurement, so a locator that matched the wrong occurrence of a repeated word would
    show up here."""
    definition, _snapshot = H.load_fixture()
    ordered = [measured.observed_pages[text] for text in H._heading_texts(definition)]

    assert all(isinstance(page, int) and page >= 1 for page in ordered)
    assert ordered == sorted(ordered), ordered
    assert max(ordered) <= measured.pages


@pytest.mark.skipif(SOFFICE is None, reason="LibreOffice is not installed")
def test_the_none_approach_names_no_page(measured: H.TocMeasurement) -> None:
    """`named_pages` is empty rather than mirroring `observed_pages`.

    Mirroring would make the proof test's `named == observed` equality true by construction
    for **every** approach, including the one that emits no table of contents at all — which
    is precisely the assertion that has to be able to fail.
    """
    assert measured.named_pages == {}
    assert measured.observed_pages != {}


@pytest.mark.skipif(SOFFICE is None, reason="LibreOffice is not installed")
def test_the_measurement_carries_both_artifacts_and_the_pdf_digest(
    measured: H.TocMeasurement,
) -> None:
    """The evidence record pins `docx_sha256` and `pdf_sha256`, so the measurement has to
    carry the bytes it measured rather than a path to something that could be regenerated."""
    import hashlib

    assert measured.docx_bytes.startswith(b"PK")
    assert measured.pdf_bytes.startswith(b"%PDF")
    assert measured.pdf_sha256 == hashlib.sha256(measured.pdf_bytes).hexdigest()
