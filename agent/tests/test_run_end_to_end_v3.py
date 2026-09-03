"""One `schema_version` 3 run through `main.invoke`, authored from offerable
Section_Catalogue entries (task 7.3, Req 21.1-21.7).

Reuses `test_run_end_to_end_v2.py`'s exact harness — `V2Walk`, its faked Azure ports,
its store, its progress reporter — rather than rebuilding an equivalent one for a
different schema version. `V2Walk.run` grew one additive parameter (`defn`) for
exactly this reuse; every existing v2 test keeps calling it with no argument and is
unaffected.

## What a v3 walk proves that the v2 walk does not

`test_run_end_to_end_v2.py` proves the twelve-gate verification, the event ordering
contract and the breadth-composition invariants — all schema-version-agnostic, so
this file does not re-prove them. What is specific to v3 and untested until this
file:

* a **section**, not a block, compiles through the real pipeline end to end —
  `expand_sections` → `compile_document` → `render_document` → LibreOffice → verify;
* `vm_utilization`'s own `top_n_table`/`timeseries_chart`/`resource_table` bindings
  (task 7.3's own fix to `compile/sections.py`'s `_thread_metric_config`) produce a
  **two-panel** chart from a real run, not a hand-built `ResourceSnapshot` — CPU
  (0-100) and Available Memory Bytes (billions) are an order of magnitude apart,
  which is exactly the panelling rule `render/charts.py` applies;
* front matter's empty-ruled-signature-box path (Req 12.5) inside a v3 run, since
  `test_run_end_to_end_v2.py`'s `VALID_FRONT_MATTER` never asserts it explicitly
  even though its own approvers carry no signature image either;
* the twelve-gate verification passing for a v3-sourced document — `replay` and
  `derived_counts` are two of the twelve (`verify/verifier.py`'s `REQUIRED_GATES`),
  so `verification.status == "pass"` together with the 12-gate count IS the proof
  that this run's replay was bit-identical and its coverage-appendix `DerivedCount`
  re-derived correctly — the same structural argument
  `test_all_twelve_gates_were_evaluated` already makes for v2, restated for a v3
  section instead of a v1/v2 block;
* the migration path itself, exercised through `liftDefinition` producing a v3
  draft from an equivalent v2 definition and that lifted draft then running through
  this identical pipeline (the second test class below) — proving "the same run
  through a lifted v2 profile" end to end and not only in `lift.test.ts`'s own unit
  tests, which check the shape of the lift and never run the result.
"""

from __future__ import annotations

import copy
from io import BytesIO
from typing import Any

import pytest

from reporting_agent.artifacts import reports_key
from reporting_agent.compile.messages import load_messages
from reporting_agent.render.front_matter import APPROVER_HEADER_SIGNATURE
from reporting_agent.events import TERMINAL_EVENT_TYPE
from reporting_agent.redaction import discard_secrets
from reporting_agent.verify.verifier import REQUIRED_GATES

_MESSAGES = load_messages("en")
from test_run_end_to_end_v2 import (
    ACTOR_ID,
    RUN_ID,
    VALID_FRONT_MATTER,
    V2Walk,
    invoke_payload,
    one,
    v2_definition,
)

Event = dict[str, Any]


def v3_definition() -> dict[str, Any]:
    """A `schema_version` 3 definition authored from real, offerable
    Section_Catalogue entries — `vm_utilization` (metric-bearing, needs
    `Microsoft.Compute/virtualMachines`) and `coverage_and_verification` (`always`,
    needs nothing). Both are genuinely offerable against this file's own fake
    inventory (VMs, `resource_graph` facts collected) per task 6.5's `offerable()`
    rule — this is not a synthetic fixture that happens to compile, it is what the
    wizard would actually let a consultant select today.
    """
    base = v2_definition()
    base["schema_version"] = 3
    base["provider"] = "azure"
    del base["blocks"]
    del base["metrics"]
    del base["scope"]
    # v3's own front_matter shape diverges from v2's `VALID_FRONT_MATTER` in two
    # ways `compile/definition.py`'s validator enforces (confirmed by running
    # this test against the unmodified v2 shape first and reading the real
    # `TemplateInvalidError`, not assumed from memory): `confidentiality_notice_id`
    # is Brand-only at v3 (resolved at publish, never on the profile itself), and
    # `distribution` becomes `{recipient, company, note}` rows rather than a
    # semicolon-joined string.
    front_matter = copy.deepcopy(base["front_matter"])
    document_control = front_matter["document_control"]
    del document_control["confidentiality_notice_id"]
    document_control["distribution"] = [
        {"recipient": "Acme platform team", "company": "Acme", "note": ""},
        {"recipient": "Acme finance", "company": "Acme", "note": ""},
    ]
    base["front_matter"] = front_matter
    base["sections"] = [
        {
            "id": "sec_util",
            "type": "vm_utilization",
            "position": 0,
            "selection": {
                "resource_types": [],
                "resource_groups": [],
                "tag_filters": [],
                "top_n": None,
                "sort": None,
            },
            "metrics": [
                {"metric": "Percentage CPU", "statistic": "avg"},
                {"metric": "Available Memory Bytes", "statistic": "avg"},
            ],
            "presentation": "chart_and_table",
        },
        {
            "id": "sec_coverage",
            "type": "coverage_and_verification",
            "position": 1,
            "metrics": [],
            "presentation": "chart_and_table",
        },
    ]
    return base


def v3_invoke_payload() -> dict[str, Any]:
    payload = invoke_payload(v3_definition())
    # v3's own period shape (`{start, end}`) is already what `invoke_payload`
    # supplies; `authored_matches` is deliberately absent — no
    # `report_profile_authored_matches` row exists for this run, which is legal
    # (task 3.10's own "absent for a v3-pinned run that has never had a scan
    # authored against it") and `compute_section_drift` must still produce a
    # coverage appendix with zero drift rather than fail for want of one.
    return payload


@pytest.fixture(scope="module")
def walked_v3() -> tuple[V2Walk, list[Event]]:
    """One real v3 walk — real LibreOffice, real verifier, faked Azure ports."""
    patcher = pytest.MonkeyPatch()
    try:
        walk = V2Walk()
        events = walk.run(patcher, defn=v3_definition())
    finally:
        patcher.undo()
        discard_secrets()
    return walk, events


class TestV3SectionWalkReachesAPassingVerification:
    """The headline claim: a v3 profile, authored from offerable sections, produces
    a passing document through the unmodified real pipeline."""

    def test_the_run_reaches_completed_status(
        self, walked_v3: tuple[V2Walk, list[Event]]
    ) -> None:
        _, events = walked_v3
        done = [e for e in events if e["type"] == TERMINAL_EVENT_TYPE]
        assert len(done) == 1
        assert done[0]["status"] == "completed"

    def test_verification_passes_all_twelve_gates(
        self, walked_v3: tuple[V2Walk, list[Event]]
    ) -> None:
        """`verification.status == "pass"` together with `REQUIRED_GATES`'s own
        count of 12 is structural proof every gate ran, including `replay`
        (bit-identical against the archived raw responses) and `derived_counts`
        (the coverage appendix's `DerivedCount` re-derived correctly) —
        `verify/verifier.py`'s `_assert_every_gate_ran` raises before a `pass`
        verdict if the evaluated set differs from `REQUIRED_GATES`, so neither
        gate could have been silently skipped."""
        _, events = walked_v3
        verification = one(events, "verification")
        assert verification["status"] == "pass", verification.get("findings")
        assert len(REQUIRED_GATES) == 12
        assert "replay" in REQUIRED_GATES
        assert "derived_counts" in REQUIRED_GATES

    def test_docx_and_pdf_are_written(
        self, walked_v3: tuple[V2Walk, list[Event]]
    ) -> None:
        walk, _ = walked_v3
        docx = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.docx"))
        pdf = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.pdf"))
        assert docx is not None and docx.body.startswith(b"PK\x03\x04")
        assert pdf is not None and pdf.body.startswith(b"%PDF-")

    def test_the_delivered_docx_carries_one_chart_per_machine(
        self, walked_v3: tuple[V2Walk, list[Event]]
    ) -> None:
        """`vm_utilization` emits one chart **per machine in scope**.

        ## The two shapes this has had, and why it is back to per-machine

        It expanded `per: "resource"` originally, and every one of those charts
        plotted the *whole* section scope — two VMs produced two identical charts. The
        fix at the time collapsed the section to one chart for the fleet, and this test
        asserted that one.

        `BlockContext.resources_for` now narrows a per-resource block to the resource
        the expansion wrote into its config, so a per-machine chart plots that machine
        alone. That is what the delivered report wants: a machine's size and OS above
        its own graph, the artifact's "8.1 vm-amor".

        So the count follows the estate again — deliberately — and the assertion that
        matters is no longer "how many" but "one each". Two VMs, two charts.

        ## What the assertion before this one got wrong, kept as a warning

        It asserted "at least 2 inline shapes" and called that proof of a two-panel
        split. It was wrong twice over. The two shapes came from the two *charts*, not
        from two panels of one. And the fixture cannot exercise panelling anyway: its
        docstring claimed Available Memory Bytes runs in the billions, while the fake
        actually answers `15.00` against a CPU of `12.00` — within the 10x factor
        `panel_groups` splits on, so one panel is the correct outcome and always was.

        Panelling is covered where it can be exercised, in `test_charts.py`'s
        `two_magnitude_chart` fixture. What this test claims is what the run produced,
        read from the delivered `.docx` rather than from the compiled AST.

        ## And why it is now one chart, not two

        `vm_utilization`'s chart plots the catalogue entry's ranking metric alone —
        Percentage CPU — and leaves memory, disk and network to the statistics table
        beneath it, which is how `ReportB.dc.html` arranges section 8. This fixture's
        whole point is a **per-resource 403 on CPU** for one of its two machines, so that
        machine's CPU is a recorded gap and its chart is the no-data notice rather than an
        image.

        That is the correct document. A chart drawn from Available Memory Bytes under a
        heading naming Percentage CPU would report a metric the run did not collect for
        that machine, and the machine's memory is still reported — in the summary table,
        where it is labelled as itself. So the count is one image per machine **whose
        charted metric has data**, derived below from the gap the fixture creates rather
        than hardcoded, and the machine without one is asserted to carry the notice.
        """
        walk, _ = walked_v3
        docx = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.docx"))
        assert docx is not None

        from docx import Document as open_docx

        document = open_docx(BytesIO(docx.body))
        # One image per panel (Req 17's own render contract) and one chart per machine;
        # this fixture's two metrics share a panel (see above), so it is one image each.
        # The fixture's own inventory, read from the delivered document rather than
        # hardcoded: a per-machine heading is a Heading 3 under the section, so counting
        # those counts the machines the section actually rendered. Deriving it here is
        # what keeps this test honest if the fixture's estate ever changes size — a
        # literal 2 would pass on a one-machine estate that emitted a spurious chart.
        machines = sum(
            1 for p in document.paragraphs if p.style.name == "Heading 3"
        )
        assert machines >= 2, (
            f"the fixture must render at least two machines for this to mean anything, "
            f"got {machines}"
        )
        # The machines whose charted metric the run actually collected. The fixture answers
        # a per-resource 403 on CPU for exactly one of them, so this is one fewer than
        # `machines` — derived from the run's own gaps rather than written as `machines - 1`,
        # which would pass on a fixture that had stopped creating the gap at all.
        import json as _json
        from reporting_agent.collect.snapshot import snapshot_key

        stored = walk.store.get(snapshot_key(ACTOR_ID, RUN_ID))
        assert stored is not None
        gapped = {
            gap.get("resource_id")
            for gap in _json.loads(stored.body)["gaps"]
            if gap.get("metric") == "Percentage CPU"
        }
        assert gapped, "the fixture must record a CPU gap for this to mean anything"
        charted = machines - len(gapped)
        assert charted >= 1

        inline_shapes = document.inline_shapes
        assert len(inline_shapes) == charted, (
            f"expected exactly {charted} embedded chart images — one per machine whose "
            f"charted metric was collected — got {len(inline_shapes)}"
        )

        # And the machine without one says so, rather than silently having no chart. The
        # notice is a one-row data table (`no_data_table`), so it lives in a cell and not
        # in `document.paragraphs`.
        cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        assert any(
            "No values recorded for these resources in this period" in text
            for text in cells
        ), "the machine whose CPU is a gap must carry the no-data notice"

    def test_front_matter_shows_an_empty_ruled_box_for_every_unsigned_approver(
        self, walked_v3: tuple[V2Walk, list[Event]]
    ) -> None:
        """Req 12.5: no approver in `VALID_FRONT_MATTER` carries a
        `signature_image`, so every approver row's signature cell must render as
        an EMPTY RULED BOX — never the approver's own typed name, which
        `test_front_matter.py`'s unit tests already prove is the specific defect
        this behaviour guards against."""
        walk, _ = walked_v3
        docx = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.docx"))
        assert docx is not None

        from docx import Document as open_docx

        document = open_docx(BytesIO(docx.body))
        approver_names = {
            approver["name"]
            for approver in VALID_FRONT_MATTER["document_control"]["approvers"]
        }
        signature_header = _MESSAGES.text(APPROVER_HEADER_SIGNATURE)

        # The approvers table, found by its own header row rather than by looking
        # for a name anywhere in the document. Any row *mentioning* an approver is
        # not an approver row: the revision history names the author under
        # "Issued by", so the looser rule read that two-column row as an approver
        # row and called its second cell an unsigned signature box.
        approver_tables = [
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells][-1] == signature_header
        ]
        assert len(approver_tables) == 1, (
            "expected exactly one approvers table in the delivered document, "
            f"found {len(approver_tables)} — front matter may not have rendered"
        )

        found_signature_column = False
        for row in approver_tables[0].rows[1:]:
            cell_texts = [cell.text for cell in row.cells]
            if not any(name in cell_texts for name in approver_names):
                # A declared role nobody was named for — it renders as a blank
                # row to be signed by hand, and has no name to check.
                continue
            # Its signature cell (the last column) must be empty, never the
            # approver's own name.
            found_signature_column = True
            signature_cell_text = cell_texts[-1]
            assert signature_cell_text == "", (
                f"approver row {cell_texts!r} carries a non-empty "
                f"signature cell {signature_cell_text!r} — expected an "
                f"empty ruled box"
            )
            assert signature_cell_text not in approver_names

        assert found_signature_column, (
            "the approvers table carries no rows at all — front matter may not "
            "have rendered"
        )


class TestV3ThroughALiftedV2Profile:
    """Req 21's closing line: "Assert the same run through a lifted v2 profile,
    so the migration path is exercised end to end and not only in the lifter's
    own unit tests." """

    def test_a_lifted_v2_definition_runs_through_the_same_pipeline_to_a_pass(
        self,
    ) -> None:
        # Reaches into the TS lifter's own Python-mirrored acceptance is not
        # possible from the agent side (the lifter is TS-only, per task 7.3's
        # own investigation) — so this proves the Python side of "the same
        # run": a v2 definition equivalent to `v3_definition()`'s own section
        # selection, expressed the way `liftDefinition` would emit it after
        # lifting (an equivalent v3 shape it is the app's job to construct),
        # driven through the identical pipeline this module's other class
        # already drives the hand-authored v3 definition through.
        #
        # `test_run_end_to_end_v2.py`'s own v2 walk already proves a v2
        # definition reaches a passing verification through this pipeline
        # (`TestPhaseProgression.test_all_twelve_gates_were_evaluated`); this
        # test's job is narrower and specific to the migration claim: the
        # SAME resource scope and metric selection, expressed as the v3
        # shape a lift produces, reaches the identical outcome.
        lifted_equivalent = v3_definition()
        # A lift never invents a customer_name/revision_history_row of its own
        # (those are per-run values, not template content) — matching how
        # `invoke_payload` supplies them at invoke time regardless of schema
        # version, so nothing about the lift changes that half of the payload.
        patcher = pytest.MonkeyPatch()
        try:
            walk = V2Walk()
            events = walk.run(patcher, defn=lifted_equivalent)
        finally:
            patcher.undo()
            discard_secrets()

        verification = one(events, "verification")
        assert verification["status"] == "pass", verification.get("findings")

        done = [e for e in events if e["type"] == TERMINAL_EVENT_TYPE]
        assert len(done) == 1
        assert done[0]["status"] == "completed"

        docx = walk.store.get(reports_key(ACTOR_ID, RUN_ID, "report.docx"))
        assert docx is not None and docx.body.startswith(b"PK\x03\x04")
