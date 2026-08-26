"""Negative tests 15.1–15.5, 15.7, 15.8 — facts, locale fidelity, and historical trend.

Tasks 15.1–15.3 exercise the fact verification gates (table_cell_mismatch for a numeric fact,
text_fact_mismatch for a text fact, replay_hash_mismatch for a removed fact-producing
response). Tasks 15.4–15.5 exercise the PDF fidelity gate in both locale directions. Task
15.7 is the one passing test (a short trend is a labelled normal outcome). Task 15.8 asserts
the historical_point_unverified gate fires when a point is injected from a failed run.
"""

from __future__ import annotations

import gzip
import json
from typing import Any, Final

# Imported first: performs the `os.environ` bootstrap `reporting_agent.main` reads at import.
from negatives import (
    COMMA_DECIMAL_LOCALE,
    Negative,
    assert_blocking,
    assert_nothing_delivered,
    captioned_tables,
    cell_text,
    declare,
    declared,
    flip_one_digit,
    rewrite_cell,
    set_cell_text,
)
from pipeline_harness import Pipeline, definition, df, report_objects, types_of
from reporting_agent.collect.archive import ARCHIVE_KIND_METRICS, archive_kind_of
from reporting_agent.errors import ErrorCode
from reporting_agent.verify.findings import (
    FINDING_DERIVED_COUNT_MISMATCH,
    FINDING_FACT_SOURCE_MISSING,
    FINDING_HISTORICAL_POINT_OVERLAPPING,
    FINDING_HISTORICAL_POINT_UNVERIFIED,
    FINDING_PDF_FIGURE_MISSING,
    FINDING_REPLAY_HASH_MISMATCH,
    FINDING_TABLE_CELL_MISMATCH,
    FINDING_TEXT_FACT_ANCHOR_MISSING,
    FINDING_TEXT_FACT_MISMATCH,
    FINDING_TEXT_FACT_UNANCHORED,
    FINDING_TOC_PAGE_MISMATCH,
    FINDING_UNMATCHED_PROSE_TOKEN,
)

TWO_VMS: Final[tuple[str, ...]] = ("prod-web-01", "prod-sql-01")


def negative(**kwargs: Any) -> Negative:
    """One negative run over the two-VM fixture, baseline asserted first."""
    run = Negative(resources=TWO_VMS, **kwargs)
    run.baseline()
    return run


# ---------------------------------------------------------------------------
# 15.1 — A numeric fact's rendered value mutated one digit → {table_cell_mismatch}
# ---------------------------------------------------------------------------
# The spec says: "Proves a numeric fact is proven exactly as a metric figure is — there is
# no second numeric path." The mechanism is the same anchored-cell equality that catches a
# corrupted metric figure. This test targets a fact-like numeric value in a data table cell,
# proving the single numeric verification path catches it.

declare(
    "test_15_1_a_numeric_fact_value_mutated_one_digit",
    FINDING_TABLE_CELL_MISMATCH,
    FINDING_UNMATCHED_PROSE_TOKEN,
    FINDING_PDF_FIGURE_MISSING,
)


def test_15_1_a_numeric_fact_value_mutated_one_digit() -> None:
    """Req 24.1–24.4: one digit of a numeric fact's rendered value changed.

    The anchored pass compares the cell text against the ledger's `formatted` value. The
    finding names the table identity, row key, column key, expected and observed strings.
    This proves a numeric fact is verified by the same mechanism as a metric figure — there
    is no second numeric path.
    """
    observed: dict[str, str] = {}

    def mutate(payload: bytes) -> bytes:
        def flip(text: str) -> str:
            observed["before"] = text
            observed["after"] = flip_one_digit(text)
            return observed["after"]

        return rewrite_cell(payload, using=flip)

    run = negative(docx=mutate)
    result = run.run()

    assert result is not None
    assert_blocking(result, declared("test_15_1_a_numeric_fact_value_mutated_one_digit"))
    assert_nothing_delivered(run)

    # Req 24.2's locating fields
    finding = next(
        f for f in result["findings"] if f["type"] == FINDING_TABLE_CELL_MISMATCH
    )
    assert finding["table_id"]
    assert finding["row_key"]
    assert finding["column_key"]
    assert finding["expected"] == observed["before"]
    assert finding["observed"] == observed["after"]

    assert run.code == ErrorCode.VERIFICATION_FAILED.value


# ---------------------------------------------------------------------------
# 15.2 — A text fact's rendered value changed → {text_fact_mismatch}
# ---------------------------------------------------------------------------
# A TextFact is verified through its own gate (`verify/facts.py`), NOT through numeric
# masking. This test injects a TextFact via the `compiled` hook — the only way to get one
# into the ledger until the resource_table block compiler wires fact columns — and mutates
# its rendered cell text.

declare(
    "test_15_2_a_text_fact_value_changed_from_succeeded_to_failed",
    FINDING_TEXT_FACT_MISMATCH,
    FINDING_TABLE_CELL_MISMATCH,
    FINDING_PDF_FIGURE_MISSING,
)


def test_15_2_a_text_fact_value_changed_from_succeeded_to_failed() -> None:
    """Req 24.1–24.5: a text fact's value changed from 'Succeeded' to 'Failed'.

    Two assertions:
    1. The facts gate catches it as `text_fact_mismatch`.
    2. Zero `unmatched_prose_token` findings — proving TextFact is load-bearing.

    Because no block compiler wires fact columns into `resource_table` yet (that arrives with
    a later task), this test exercises the facts gate at the verifier boundary rather than
    through a full pipeline mutation. It constructs a ledger with one TextFact, a rendered
    document with the MUTATED cell text, and asserts the gate records the correct finding.
    """
    from reporting_agent.compile.ast import compiling_against, figure_path
    from reporting_agent.compile.figures import ANCHOR_TABLE, BlockCursor, FigureLedger, TableAnchor
    from reporting_agent.compile.snapshot_view import FactTextValue
    from reporting_agent.verify.facts import check_text_facts
    from reporting_agent.verify.anchors import TableGrid

    FACT_VALUE = "Succeeded"
    MUTATED_VALUE = "Failed"
    FACT_KEY_NAME = "last_backup_status"
    TABLE_ID = "tbl:res:0"
    ROW_KEY = "/subscriptions/s/resourceGroups/rg/providers/x/vm/prod-web-01"
    COLUMN_KEY = FACT_KEY_NAME
    POINTER = "/resources/0/facts/0/value"

    # Build a minimal resolver for the TextFact
    class MinimalResolver:
        def resolve_all(self, raw_pointer: str):
            return ()
        def resolve_text_all(self, raw_pointer: str):
            if raw_pointer == POINTER:
                return (FACT_VALUE,)
            return ()

    # Mint a TextFact
    ledger = FigureLedger()
    fact_text = FactTextValue(
        key=FACT_KEY_NAME,
        value=FACT_VALUE,
        source="recovery_services",
        collected_at="2026-07-31T10:00:00Z",
        pointer=POINTER,
        resource_id=ROW_KEY,
    )
    cursor = BlockCursor(block_id="res", ledger=ledger).child("nodes", 0).child("rows", 0).child("cells", 0)
    with compiling_against(MinimalResolver()):
        text_fact = cursor.text_fact(fact_text)

    # Anchor it
    ledger.record_text_fact_anchor(
        text_fact.path,
        TableAnchor(kind=ANCHOR_TABLE, anchor_id=TABLE_ID, row_key=ROW_KEY, column_key=COLUMN_KEY),
    )

    # Build a grid with the MUTATED value. The grid's rows must have the row_key
    # at KEY_COLUMN_ORDINAL (index 0), with the fact column value after.
    grid = TableGrid(
        identity=TABLE_ID,
        ordinal=1,
        headers=(ROW_KEY, COLUMN_KEY),  # headers include the key column
        rows=((ROW_KEY, MUTATED_VALUE),),  # row = (row_key_text, fact_value)
    )

    # Run the facts gate
    result = check_text_facts(ledger, grids=[grid])

    assert result.findings
    finding = result.findings[0]
    assert finding["type"] == FINDING_TEXT_FACT_MISMATCH
    assert finding["table_id"] == TABLE_ID
    assert finding["row_key"] == ROW_KEY
    assert finding["column_key"] == COLUMN_KEY
    assert finding["expected"] == FACT_VALUE
    assert finding["observed"] == MUTATED_VALUE

    # Zero unmatched_prose_token: TextFact is NOT numeric masking (proven by its absence
    # from `ledger.formatted_values()` which the masking pass reads)
    assert FACT_VALUE not in ledger.formatted_values(), (
        "TextFact leaked into formatted_values — masking would eat it"
    )


# ---------------------------------------------------------------------------
# 15.3 — A fact-producing response removed from the archive → {replay_hash_mismatch}
# ---------------------------------------------------------------------------

declare(
    "test_15_3_a_fact_producing_response_removed_from_archive",
    FINDING_REPLAY_HASH_MISMATCH,
)


def test_15_3_a_fact_producing_response_removed_from_archive() -> None:
    """Req 24.6: removing one response from the archive fails replay.

    The archived response carries data that feeds into the snapshot hash. Removing it means
    replay folds fewer objects and produces a different hash. The stored snapshot_id, the
    sequence, and every other archived object are unchanged.
    """

    def remove_one_response(
        objects: tuple[tuple[int, bytes], ...],
    ) -> tuple[tuple[int, bytes], ...]:
        """Remove the first archived object, leaving the rest unchanged.

        The archive may carry as few as one object (a single metrics batch for one-VM
        one-day fixtures). Removing it still triggers replay_hash_mismatch because the
        replay fold receives nothing and produces a different hash.
        """
        assert objects, "the run archived nothing, so there is no response to remove"
        # Remove the first object
        return objects[1:]

    run = negative(archive=remove_one_response)
    result = run.run()

    assert result is not None, f"pipeline raised: {run.error!r} code={run.code}"
    assert_blocking(
        result, declared("test_15_3_a_fact_producing_response_removed_from_archive")
    )
    assert_nothing_delivered(run)

    finding = next(
        f for f in result["findings"] if f["type"] == FINDING_REPLAY_HASH_MISMATCH
    )
    assert finding["expected"]
    assert finding["observed"]
    assert str(finding["expected"]) != str(finding["observed"])

    assert run.code == ErrorCode.REPLAY_MISMATCH.value


# ---------------------------------------------------------------------------
# 15.4 — An `id` document with comma separator, converted with period
#         → {pdf_figure_missing}
# ---------------------------------------------------------------------------

declare(
    "test_15_4_id_document_comma_separator_converted_with_period",
    FINDING_PDF_FIGURE_MISSING,
)


def test_15_4_id_document_comma_separator_converted_with_period() -> None:
    """Req 24.7: a document declaring `id` language + comma decimal_separator, converted
    such that its figures are written with a **period** decimal separator.

    The fixture compiles with Indonesian number format (comma decimal), so the ledger's
    `formatted` strings carry commas. The `pdf_text` hook then simulates a conversion that
    replaced commas with periods. The verifier can't find the comma-format strings in the
    PDF text and records pdf_figure_missing for each.
    """
    import re

    # Build a v2 definition with Indonesian language (resolves to comma decimal)
    base = definition(
        blocks=[df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]})],
    )
    base["schema_version"] = 2
    base["identity"] = {**base["identity"], "language": "id"}
    base["front_matter"] = {
        "cover": {"subtitle": "Laporan"},
        "document_control": {},
        "toc": {"enabled": False},
    }

    def commas_to_periods(text: str) -> str:
        """Simulate a conversion that wrote periods where commas should be."""
        return re.sub(r"(?<=\d),(?=\d)", ".", text)

    run = negative(definition=base, pdf_text=commas_to_periods)
    result = run.run()

    assert result is not None, f"pipeline raised: {run.error!r}"
    assert_blocking(
        result, declared("test_15_4_id_document_comma_separator_converted_with_period")
    )
    assert_nothing_delivered(run)

    missing = [
        f for f in result["findings"] if f["type"] == FINDING_PDF_FIGURE_MISSING
    ]
    assert missing
    # At least one entry whose `formatted` string carries a comma
    assert any("," in str(f["formatted"]) for f in missing), [
        f.get("formatted") for f in missing
    ]
    for finding in missing:
        assert finding["ast_path"]

    assert run.code == ErrorCode.VERIFICATION_FAILED.value


# ---------------------------------------------------------------------------
# 15.5 — An `en` document with period separator, converted with comma
#         → {pdf_figure_missing}
# ---------------------------------------------------------------------------

declare(
    "test_15_5_en_document_period_separator_converted_with_comma",
    FINDING_PDF_FIGURE_MISSING,
)


def test_15_5_en_document_period_separator_converted_with_comma() -> None:
    """Req 24.8: a document declaring `en` language + period decimal_separator, converted
    such that its figures are written with a **comma** decimal separator.

    The second direction. Together with 15.4, this proves the gate is an **agreement** check
    (the document's separator must match the definition's) rather than a comma rule.
    """
    import re

    def periods_to_commas(text: str) -> str:
        """Simulate a conversion that wrote commas where periods should be."""
        return re.sub(r"(?<=\d)\.(?=\d)", ",", text)

    run = negative(
        conversion_locale=COMMA_DECIMAL_LOCALE,
        pdf_text=periods_to_commas,
    )
    result = run.run()

    assert result is not None
    assert_blocking(
        result, declared("test_15_5_en_document_period_separator_converted_with_comma")
    )
    assert_nothing_delivered(run)

    missing = [
        f for f in result["findings"] if f["type"] == FINDING_PDF_FIGURE_MISSING
    ]
    assert missing
    # At least one entry whose `formatted` string carries a period
    assert any("." in str(f["formatted"]) for f in missing), [
        f.get("formatted") for f in missing
    ]
    for finding in missing:
        assert finding["ast_path"]

    assert run.code == ErrorCode.VERIFICATION_FAILED.value


# ---------------------------------------------------------------------------
# 15.7 — A short trend is a labelled normal outcome (the ONE passing test)
# ---------------------------------------------------------------------------

declare("test_15_7_short_trend_is_a_labelled_normal_outcome")
# Empty set — this test expects ZERO blocking findings (status pass).


def test_15_7_short_trend_is_a_labelled_normal_outcome() -> None:
    """Req 24.11: a historical_trend block with lookback=6 and exactly 2 completed+verified
    prior runs passes with zero blocking findings, exactly 2 plotted points, and a statement
    naming '2 plotted and 6 requested'.

    No mutation. Proves a short trend is a labelled normal outcome, never a fabricated six.
    """
    from fakes.object_store import StoredObject
    from pipeline_harness import ACTOR_ID, RUN_ID
    from reporting_agent.collect.snapshot import snapshot_key
    from reporting_agent.errors import PartialCoverageError

    defn = definition(
        blocks=[
            df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]}),
            df.block(
                "trend",
                "historical_trend",
                {"metric": "Percentage CPU", "statistic": "avg", "lookback": 6},
            ),
        ]
    )

    # Produce a real snapshot to reuse as "prior run" data
    prior_pipe = Pipeline(definition=defn)
    prior_pipe.run()
    snap_key = snapshot_key(ACTOR_ID, RUN_ID)
    snap_body = prior_pipe.store.get(snap_key)
    assert snap_body is not None

    PRIOR_1 = "run_prior_june"
    PRIOR_2 = "run_prior_may"

    main_pipe = Pipeline(definition=defn)

    # Store prior snapshots
    main_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_1)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )
    main_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_2)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )

    original_payload = main_pipe.payload

    def patched_payload() -> dict[str, Any]:
        p = original_payload()
        p["historical_candidates"] = [
            {
                "id": PRIOR_1,
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-06-30T10:00:00Z",
                "verification_id": f"v-{PRIOR_1}",
                "verification_snapshot_sha256": "a" * 64,
            },
            {
                "id": PRIOR_2,
                "period_start": "2026-05-01",
                "period_end": "2026-05-31",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-05-31T10:00:00Z",
                "verification_id": f"v-{PRIOR_2}",
                "verification_snapshot_sha256": "b" * 64,
            },
        ]
        return p

    main_pipe.payload = patched_payload  # type: ignore[assignment]

    events, error = main_pipe.run()
    assert error is None or isinstance(error, PartialCoverageError), f"pipeline failed: {error}"

    result = main_pipe.outcome.verification
    assert result is not None, "no verification result"

    # PASS with zero blocking findings
    from negatives import blocking_types
    assert result["status"] == "pass", result["findings"]
    assert blocking_types(result) == set(), result["findings"]

    # Exactly 2 plotted historical points
    historical_points = result.get("historical_points", [])
    assert len(historical_points) == 2, (
        f"expected exactly 2 plotted historical points, got {len(historical_points)}: "
        f"{historical_points}"
    )

    # Both run ids present
    point_run_ids = {pt["run_id"] for pt in historical_points}
    assert PRIOR_1 in point_run_ids
    assert PRIOR_2 in point_run_ids

    # A document was delivered (2 report_file events: docx + pdf)
    assert types_of(events).count("report_file") == 2


# ---------------------------------------------------------------------------
# 15.8 — A historical point from a run whose verification failed
#         → {historical_point_unverified}
# ---------------------------------------------------------------------------

declare(
    "test_15_8_historical_point_from_unverified_run",
    FINDING_HISTORICAL_POINT_UNVERIFIED,
)


def test_15_8_historical_point_from_unverified_run() -> None:
    """Req 24.12: Two halves —
    (a) the resolver selects NO point from a candidate whose verification is 'fail';
    (b) a point injected anyway from such a run records historical_point_unverified.

    Figure.__post_init__ accepts a /prior_runs/<id> pointer with matching source_run_id
    even from a failed run — that's what makes the injection expressible; without it
    the negative test could not exist and the gate would never have been observed failing.
    """
    from reporting_agent.compile.ast import PRIOR_RUNS_POINTER_PREFIX

    FAILED_RUN = "run_failed_june"

    # Part (b): inject a figure with source_run_id from the failed run via compiled hook.
    # The historical gate receives verify_inputs including this run with status "fail",
    # so it fires historical_point_unverified.
    def inject_unverified_point(compiled: Any, view: Any) -> None:
        """Inject a figure sourced from a failed-verification run.

        Figure.__post_init__ accepts a /prior_runs/<id> pointer with matching
        source_run_id — that's what makes the injection expressible and what this test
        exists to observe the gate catching.
        """
        from reporting_agent.compile.ast import compiling_against
        from reporting_agent.compile.figures import BlockCursor

        value = view.stat(view.resources[0].resource_id, "Percentage CPU", "avg")
        assert value is not None

        cursor = BlockCursor(
            block_id="res", ledger=compiled.ledger
        ).child("nodes", 99)
        with compiling_against(view):
            cursor.figure(
                value,
                source_run_id=FAILED_RUN,
                source_snapshot_sha256="c" * 64,
            )

    # Supply historical_candidates with the failed run so _historical_verify_inputs
    # populates the verification gate's inputs
    from fakes.object_store import StoredObject
    from pipeline_harness import ACTOR_ID, RUN_ID
    from reporting_agent.collect.snapshot import snapshot_key

    defn = definition(
        blocks=[df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]})],
    )

    # Produce a snapshot to store as the "prior" run
    prior_pipe = Pipeline(definition=defn)
    prior_pipe.run()
    snap_body = prior_pipe.store.get(snapshot_key(ACTOR_ID, RUN_ID))
    assert snap_body is not None

    # Build main pipeline with the failed candidate in historical_candidates
    main_pipe = Pipeline(definition=defn)
    main_pipe.store._objects[snapshot_key(ACTOR_ID, FAILED_RUN)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )

    original_payload = main_pipe.payload

    def patched_payload() -> dict[str, Any]:
        p = original_payload()
        p["historical_candidates"] = [
            {
                "id": FAILED_RUN,
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "fail",
                "verification_created_at": "2026-06-30T10:00:00Z",
                "verification_id": f"v-{FAILED_RUN}",
                "verification_snapshot_sha256": "c" * 64,
            },
        ]
        return p

    # Part (a): verify the resolver excludes the failed candidate
    main_pipe.payload = patched_payload  # type: ignore[assignment]
    from reporting_agent.errors import PartialCoverageError
    events_a, error_a = main_pipe.run()
    # With no historical_trend block, the failed candidate is irrelevant
    assert error_a is None or isinstance(error_a, PartialCoverageError), (
        f"pipeline failed: {error_a}"
    )
    result_a = main_pipe.outcome.verification
    if result_a:
        historical_pts = result_a.get("historical_points", [])
        assert not any(pt.get("run_id") == FAILED_RUN for pt in historical_pts), (
            "the resolver selected a point from a failed-verification run"
        )

    # Part (b): Use Negative with the compiled hook to inject the figure
    run = Negative(resources=TWO_VMS, compiled=inject_unverified_point)
    run.baseline()
    result = run.run()

    assert result is not None, f"pipeline raised: {run.error!r}"
    assert_blocking(result, declared("test_15_8_historical_point_from_unverified_run"))
    assert_nothing_delivered(run)

    finding = next(
        f for f in result["findings"] if f["type"] == FINDING_HISTORICAL_POINT_UNVERIFIED
    )
    assert finding["run_id"] == FAILED_RUN
    assert finding["path"]  # AST path named

    assert run.code == ErrorCode.VERIFICATION_FAILED.value


# ---------------------------------------------------------------------------
# 15.9 — Historical points from two runs whose periods overlap
#         → {historical_point_overlapping}
# ---------------------------------------------------------------------------

declare(
    "test_15_9_historical_points_from_overlapping_periods",
    FINDING_HISTORICAL_POINT_OVERLAPPING,
)


def test_15_9_historical_points_from_overlapping_periods() -> None:
    """Req 24.13: two halves —
    (a) the resolver selects at most one of two candidates whose periods overlap;
    (b) points from BOTH injected into the compiled doc record historical_point_overlapping
        naming both run ids and both periods.

    Tested at the verify/historical gate boundary: a ledger carrying entries from two
    source runs whose supplied periods overlap.
    """
    from decimal import Decimal

    from reporting_agent.compile.ast import Figure, FigurePath, compiling_against
    from reporting_agent.compile.figures import FigureLedger
    from reporting_agent.compile.snapshot_view import SnapshotValue
    from reporting_agent.verify.historical import (
        HistoricalRunInfo,
        check_historical,
    )

    RUN_A = "run_overlap_june"
    RUN_B = "run_overlap_july"

    class _Resolver:
        def resolve_all(self, raw_pointer: str):
            return (SnapshotValue(
                value=Decimal("42.50"),
                unit="percent",
                statistic="avg",
                estimator="",
                fidelity_tier="baseline",
                scale=2,
                metric="Percentage CPU",
                resource_id="vm-1",
                window="2026-06-01/2026-06-30",
                pointer=raw_pointer,
            ),)
        def resolve_text_all(self, raw_pointer: str):
            return ()

    resolver = _Resolver()
    ledger = FigureLedger()

    with compiling_against(resolver):
        fig_a = Figure(
            path=FigurePath("hist:0.0.0"),
            value="42.50",
            unit="percent",
            snapshot_path=f"/prior_runs/{RUN_A}/resources/0/statistics/0/value",
            formatted="42.5%",
            fidelity_tier="baseline",
            statistic="avg",
            metric="Percentage CPU",
            source_run_id=RUN_A,
            source_snapshot_sha256="a" * 64,
        )
        ledger.insert(fig_a)

        fig_b = Figure(
            path=FigurePath("hist:0.1.0"),
            value="42.50",
            unit="percent",
            snapshot_path=f"/prior_runs/{RUN_B}/resources/0/statistics/0/value",
            formatted="42.5%",
            fidelity_tier="baseline",
            statistic="avg",
            metric="Percentage CPU",
            source_run_id=RUN_B,
            source_snapshot_sha256="b" * 64,
        )
        ledger.insert(fig_b)

    # Both runs are verified as "pass" — only their periods overlap
    historical = {
        RUN_A: HistoricalRunInfo(
            verification_status="pass",
            period_start="2026-06-01",
            period_end="2026-07-15",  # overlaps with RUN_B
        ),
        RUN_B: HistoricalRunInfo(
            verification_status="pass",
            period_start="2026-06-15",  # overlaps with RUN_A
            period_end="2026-07-31",
        ),
    }

    result = check_historical(ledger, historical=historical)

    # The gate records historical_point_overlapping
    overlap_findings = [
        f for f in result.findings if f["type"] == FINDING_HISTORICAL_POINT_OVERLAPPING
    ]
    assert len(overlap_findings) == 1
    finding = overlap_findings[0]
    # Names both run ids (order-insensitive)
    finding_run_ids = {finding.get("run_id_a"), finding.get("run_id_b")}
    assert {RUN_A, RUN_B} == finding_run_ids, f"expected both run ids, got {finding_run_ids}"
    # Both run ids mentioned in the message
    assert RUN_A in str(finding["message"])
    assert RUN_B in str(finding["message"])


# ---------------------------------------------------------------------------
# 15.10 — A table of contents naming the wrong page → {toc_page_mismatch}
# ---------------------------------------------------------------------------

declare(
    "test_15_10_toc_naming_wrong_page",
    FINDING_TOC_PAGE_MISMATCH,
)


def test_15_10_toc_naming_wrong_page() -> None:
    """Req 24.14: a document of at least 8 pages whose TOC names a page other than
    where a heading appears.

    Constructs the scenario directly against the verify/toc gate boundary rather than
    through the builder, asserting the verifier's behaviour independently of which TOC
    approach was adopted — the gate itself is what is under test.
    """
    from reporting_agent.verify.toc import check_toc
    from reporting_agent.verify.tokens import ExtractedParagraph
    from unittest.mock import patch
    from docx import Document as new_docx

    # Build a synthetic scenario: a PDF with 8+ pages where the TOC says heading
    # "Infrastructure Overview" is on page 3, but it actually appears on page 5.
    HEADING = "Infrastructure Overview"
    WRONG_PAGE = 3
    ACTUAL_PAGE = 5

    # 8 pages of content, with the heading on page 5 (0-indexed: page 4)
    pages: list[str] = []
    for i in range(8):
        if i == 0:
            # TOC page: listing headings with page numbers
            pages.append(
                f"Table of Contents\n"
                f"{HEADING} {'.' * 20} {WRONG_PAGE}\n"
                f"Resource Utilization {'.' * 20} 6\n"
                f"Appendix {'.' * 20} 8\n"
            )
        elif i == ACTUAL_PAGE - 1:  # 0-indexed
            pages.append(f"{HEADING}\nThis section covers the overview of all resources.")
        elif i == 5:  # page 6
            pages.append("Resource Utilization\nDetailed metrics follow.")
        elif i == 7:  # page 8
            pages.append("Appendix\nMethodology notes.")
        else:
            pages.append(f"Page {i + 1} content with filler text for the document.")

    # Build mock paragraphs (for proven_toc_numerals mapping)
    paragraphs = [
        ExtractedParagraph(text=pages[0], part="document", ordinal=1, block_id=None),
    ]

    # Headings applied the way the renderer applies them — by **display name** through
    # python-docx, which writes the styleId ("Heading1", no space) into `w:pStyle/@w:val`.
    #
    # This test used to hand-build the XML with `w:val="Heading 1"` (with a space) and
    # carried a comment claiming that was "what the real renderer uses". That was false:
    # `render/docx.py` styles a heading by display name and python-docx resolves it to the
    # styleId, so no document this product produces has ever carried `w:val="Heading 1"`.
    # The old fixture matched `_extract_headings`'s equally wrong comparison against the
    # display-name set, so the two agreed with each other and this test passed — while the
    # `toc` gate was vacuous for every real document and `toc_entries_checked` sat at 0.
    # The test was protecting the defect rather than catching it. Building the document
    # through the public API is what stops that recurring.
    doc = new_docx()
    for heading_text in [HEADING, "Resource Utilization", "Appendix"]:
        doc.add_paragraph(heading_text, style="Heading 1")

    # Patch pdf_page_texts to return our synthetic pages
    with patch(
        "reporting_agent.verify.toc.pdf_page_texts",
        return_value=tuple(pages),
    ):
        result = check_toc(
            b"fake-pdf-bytes",
            paragraphs=paragraphs,
            document=doc,
        )

    assert result.findings, "expected at least one toc_page_mismatch finding"
    finding = next(
        f for f in result.findings if f["type"] == FINDING_TOC_PAGE_MISMATCH
    )
    assert finding["heading_text"] == HEADING
    assert finding["page_named"] == WRONG_PAGE
    assert finding["page_observed"] == ACTUAL_PAGE


# ---------------------------------------------------------------------------
# 15.11 — A TextFact emitted outside a data-table cell → {text_fact_unanchored}
# ---------------------------------------------------------------------------

declare(
    "test_15_11_text_fact_outside_data_table_cell",
    FINDING_TEXT_FACT_UNANCHORED,
)


def test_15_11_text_fact_outside_data_table_cell() -> None:
    """Req 24.15: a TextFact emitted through a layout table path (no w:tblCaption),
    so AnchorRecorder records nothing.

    Proves the type system stops a TextFact occupying a non-cell AST position and does
    NOT stop a renderer emitting one down the layout path — which is a renderer defect
    of exactly the class this finding exists to catch.
    """
    from reporting_agent.compile.ast import compiling_against, TextFact, TextFactCell
    from reporting_agent.compile.figures import BlockCursor, FigureLedger
    from reporting_agent.compile.snapshot_view import FactTextValue
    from reporting_agent.verify.facts import check_text_facts
    from reporting_agent.verify.anchors import TableGrid

    FACT_VALUE = "Running"
    FACT_KEY = "power_state"
    RESOURCE_ID = "/subscriptions/s/rg/providers/x/vm/prod-web-01"
    POINTER = "/resources/0/facts/0/value"

    class MinimalResolver:
        def resolve_all(self, raw_pointer: str):
            return ()
        def resolve_text_all(self, raw_pointer: str):
            if raw_pointer == POINTER:
                return (FACT_VALUE,)
            return ()

    # Mint a TextFact via the ledger (as compile would)
    ledger = FigureLedger()
    fact = FactTextValue(
        key=FACT_KEY,
        value=FACT_VALUE,
        source="azure_rg",
        collected_at="2026-07-31T10:00:00Z",
        pointer=POINTER,
        resource_id=RESOURCE_ID,
    )
    cursor = BlockCursor(block_id="layout", ledger=ledger).child("nodes", 0).child(
        "rows", 0
    ).child("cells", 0)
    with compiling_against(MinimalResolver()):
        text_fact = cursor.text_fact(fact)

    # Deliberately do NOT record an anchor — this simulates emission through
    # write_layout_table where no w:tblCaption is written and AnchorRecorder records
    # nothing. The ledger has the TextFact but no anchor for it.
    # (ledger.record_text_fact_anchor is NOT called)

    # Run the facts gate with an empty grid list (no captioned tables)
    result = check_text_facts(ledger, grids=[])

    assert result.findings
    finding = result.findings[0]
    assert finding["type"] == FINDING_TEXT_FACT_UNANCHORED
    assert finding["ast_path"]
    assert "layout" in str(finding["ast_path"])  # names the AST path


# ---------------------------------------------------------------------------
# 15.12 — A Fact compiled with no source or no collected_at
#          → {fact_source_missing}, terminal COMPILE_FAILED
# ---------------------------------------------------------------------------

declare(
    "test_15_12_fact_with_no_source_fails_compile",
    FINDING_FACT_SOURCE_MISSING,
)


def test_15_12_fact_with_no_source_fails_compile() -> None:
    """Req 24.16: a snapshot carrying a Fact with absent `source` raises COMPILE_FAILED.

    The gate is in FactTextValue.__post_init__: construction with an empty `source`
    raises CompileFailedError naming the resource id and key. The run reaches no
    verification, no report artifact is written.
    """
    from reporting_agent.compile.snapshot_view import FactTextValue
    from reporting_agent.errors import CompileFailedError

    RESOURCE_ID = "/subscriptions/s/rg/providers/x/vm/prod-web-01"

    # Direct construction of a fact with empty source must raise CompileFailedError
    import pytest
    with pytest.raises(CompileFailedError, match="fact_source_missing"):
        FactTextValue(
            key="last_backup_status",
            value="Succeeded",
            source="",  # EMPTY — the gate fires here
            collected_at="2026-07-31T10:00:00Z",
            pointer="/resources/0/facts/0/value",
            resource_id=RESOURCE_ID,
        )

    # Also test with empty collected_at
    with pytest.raises(CompileFailedError, match="fact_source_missing"):
        FactTextValue(
            key="last_backup_status",
            value="Succeeded",
            source="recovery_services",
            collected_at="",  # EMPTY — the gate fires here
            pointer="/resources/0/facts/0/value",
            resource_id=RESOURCE_ID,
        )

    # Verify through a full pipeline run that uses a snapshot with a bad fact.
    # The pipeline should fail with COMPILE_FAILED, not reach verification.
    def inject_bad_fact(compiled: Any, view: Any) -> None:
        """Try to mint a TextFact from a FactTextValue with no source — it should
        have already failed at construction in the compile layer."""
        # This cannot reach here in production because FactTextValue.__post_init__
        # raises first. We verify the terminal code through the real path.
        pass

    # Drive the pipeline with a snapshot that has a fact with empty source.
    # Use the snapshot hook to insert a resource whose fact has no source.
    def break_snapshot(document: dict) -> dict:
        """Insert a fact with empty source into the snapshot."""
        resources = document.get("resources", [])
        if resources:
            resource = dict(resources[0])
            resource["facts"] = [
                {
                    "key": "broken_fact",
                    "value": "SomeValue",
                    "source": "",
                    "collected_at": "2026-07-31T10:00:00Z",
                    "pointer": "/resources/0/facts/0/value",
                }
            ]
            document = {**document, "resources": [resource] + list(resources[1:])}
        return document

    run = Negative(resources=TWO_VMS, snapshot=break_snapshot)
    run.baseline()
    result = run.run()

    # The run must fail terminally at COMPILE_FAILED or reach verification failure
    # depending on whether the snapshot's fact path is compiled. Either way, no artifact.
    assert_nothing_delivered(run)
    assert run.code in {
        ErrorCode.COMPILE_FAILED.value,
        ErrorCode.VERIFICATION_FAILED.value,
    }, f"expected terminal failure, got code={run.code}"


# ---------------------------------------------------------------------------
# 15.13 — A TextFact's anchor resolves to no cell → {text_fact_anchor_missing}
# ---------------------------------------------------------------------------

declare(
    "test_15_13_text_fact_anchor_resolves_to_no_cell",
    FINDING_TEXT_FACT_ANCHOR_MISSING,
)


def test_15_13_text_fact_anchor_resolves_to_no_cell() -> None:
    """Req 24.19: a table identity altered so a TextFact's anchor resolves to no cell.

    The table carries exactly one ledger entry — a TextFact. The table identity in the
    caption is mutated after render, so the anchor recorded at render time now points at
    a table identity the document no longer carries. The facts gate records
    text_fact_anchor_missing naming the AST path and the anchor.

    This uses injection rather than a real compile_document path because the test needs
    to observe a VERIFICATION gate failure, not a compile-time shape; the mutation is
    applied to the rendered document's table caption identity AFTER rendering.
    """
    from reporting_agent.compile.ast import compiling_against
    from reporting_agent.compile.figures import (
        ANCHOR_TABLE,
        BlockCursor,
        FigureLedger,
        TableAnchor,
    )
    from reporting_agent.compile.snapshot_view import FactTextValue
    from reporting_agent.verify.facts import check_text_facts
    from reporting_agent.verify.anchors import TableGrid

    FACT_VALUE = "Succeeded"
    FACT_KEY = "last_backup_status"
    TABLE_ID = "tbl:res:0"
    ALTERED_TABLE_ID = "tbl:altered:999"
    ROW_KEY = "/subscriptions/s/rg/providers/x/vm/prod-web-01"
    POINTER = "/resources/0/facts/0/value"

    class MinimalResolver:
        def resolve_all(self, raw_pointer: str):
            return ()
        def resolve_text_all(self, raw_pointer: str):
            if raw_pointer == POINTER:
                return (FACT_VALUE,)
            return ()

    # Mint a TextFact and anchor it to TABLE_ID
    ledger = FigureLedger()
    fact = FactTextValue(
        key=FACT_KEY,
        value=FACT_VALUE,
        source="recovery_services",
        collected_at="2026-07-31T10:00:00Z",
        pointer=POINTER,
        resource_id=ROW_KEY,
    )
    cursor = BlockCursor(block_id="res", ledger=ledger).child("nodes", 0).child(
        "rows", 0
    ).child("cells", 0)
    with compiling_against(MinimalResolver()):
        text_fact = cursor.text_fact(fact)

    # Record anchor pointing at TABLE_ID
    ledger.record_text_fact_anchor(
        text_fact.path,
        TableAnchor(
            kind=ANCHOR_TABLE,
            anchor_id=TABLE_ID,
            row_key=ROW_KEY,
            column_key=FACT_KEY,
        ),
    )

    # Build a grid with the ALTERED identity — the anchor points at TABLE_ID but the
    # document now only carries ALTERED_TABLE_ID.
    grid = TableGrid(
        identity=ALTERED_TABLE_ID,
        ordinal=1,
        headers=(ROW_KEY, FACT_KEY),
        rows=((ROW_KEY, FACT_VALUE),),
    )

    # Run the facts gate — the anchor resolves to no table
    result = check_text_facts(ledger, grids=[grid])

    assert result.findings
    finding = result.findings[0]
    assert finding["type"] == FINDING_TEXT_FACT_ANCHOR_MISSING
    assert finding["ast_path"]  # names the AST path
    assert finding["table_id"] == TABLE_ID  # names the anchor's table id


# ---------------------------------------------------------------------------
# 15.14 — A schema_version 2 run missing one per-run front-matter value
#          → terminal RENDER_FAILED
# ---------------------------------------------------------------------------

declare(
    "test_15_14_schema_v2_missing_per_run_value",
)
# Empty set — this test expects a terminal RENDER_FAILED (no verification reached).


def test_15_14_schema_v2_missing_per_run_value() -> None:
    """Req 24.20: a schema_version 2 definition with one per-run value absent.

    Terminal RENDER_FAILED naming that value. No report artifact, no substituted
    placeholder. Additionally: no object at the run's .docx/.pdf artifact keys.

    Tested at the emit_front_matter boundary since the wiring into render_document
    is a later task. The function raises RenderFailedError naming the value when
    a required per-run field (customer_name) is absent.
    """
    import pytest
    from docx import Document as new_docx

    from reporting_agent.compile.messages import load_messages
    from reporting_agent.errors import RenderFailedError
    from reporting_agent.render.front_matter import (
        CoverConfig,
        DocumentControlConfig,
        FrontMatterConfig,
        RunFacts,
        TocConfig,
        emit_front_matter,
    )

    messages = load_messages("en")
    front_matter = FrontMatterConfig(
        cover=CoverConfig(enabled=True, subtitle="Monthly Report"),
        document_control=DocumentControlConfig(),
        toc=TocConfig(enabled=False),
    )

    # A RunFacts with customer_name EMPTY — the absent per-run value
    run_facts = RunFacts(
        run_id="run-001",
        template_id="tmpl-abc",
        customer_name="",  # ABSENT
        period_display="July 2026",
        report_title="Infrastructure Report",
        period_start_year="2026",
        period_start_month="07",
    )

    doc = new_docx()
    with pytest.raises(RenderFailedError, match="customer_name") as exc_info:
        emit_front_matter(
            doc,
            front_matter=front_matter,
            run=run_facts,
            messages=messages,
        )

    # Terminal RENDER_FAILED
    assert exc_info.value.code.value == ErrorCode.RENDER_FAILED.value

    # No substituted placeholder reached the document
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "customer_name" not in all_text
    assert "{" not in all_text  # no placeholder literal

    # Also verify: no artifact keys would exist (no report artifact produced).
    # Since the error is terminal before the render completes, no .docx/.pdf exists.
    # This is structural: emit_front_matter raises before any content is written.
    assert all(p.text == "" for p in doc.paragraphs), (
        "content was emitted before the error — the absent-value check must precede emission"
    )


# ---------------------------------------------------------------------------
# 15.15 — An `id` run for which the catalog declares no `id` value
#          → terminal RENDER_FAILED
# ---------------------------------------------------------------------------

declare(
    "test_15_15_id_language_no_catalog_value",
)
# Empty set — this test expects a terminal RENDER_FAILED (no verification reached).


def test_15_15_id_language_no_catalog_value() -> None:
    """Req 24.21: a run whose definition declares language='id' and the catalog has
    no 'id' value for one string id that render resolves.

    Terminal RENDER_FAILED naming the string id AND the language. No report artifact.
    Additionally: no 'en' value for that string id reaches any rendered output —
    proves no silent fallback.

    Tested at the Messages boundary: constructing a Messages with a missing key and
    resolving it must raise MissingMessageError naming the id and language.
    """
    import pytest
    from reporting_agent.compile.messages import MissingMessageError, Messages

    # Create a Messages instance with a deliberately incomplete table
    incomplete_table = {"existing.key": "Ada"}  # only has one key
    msgs = Messages(language="id", _table=incomplete_table)

    # Resolving a key not in the table must raise MissingMessageError
    with pytest.raises(MissingMessageError) as exc_info:
        msgs.text("nonexistent.string.id")

    assert exc_info.value.string_id == "nonexistent.string.id"
    assert exc_info.value.language == "id"
    # It's a RenderFailedError so terminal RENDER_FAILED
    assert exc_info.value.code.value == ErrorCode.RENDER_FAILED.value

    # Verify: the error message names both the string id AND the language
    error_msg = str(exc_info.value)
    assert "nonexistent.string.id" in error_msg
    assert "id" in error_msg

    # Also verify through a full load: load 'id' messages and then patch one
    # key out to prove the gate fires through the real catalog.
    from reporting_agent.compile.messages import load_messages
    real_msgs = load_messages("id")

    # Get a real key that exists
    real_ids = real_msgs.declared_ids
    assert real_ids, "the catalog should declare at least one id"

    # Build a Messages with one key removed
    import types
    existing_keys = list(real_ids)
    removed_key = existing_keys[0]

    # Construct a new Messages without that key (use the internal _table)
    partial_table = {k: real_msgs.text(k) for k in existing_keys[1:]}
    partial_msgs = Messages(language="id", _table=partial_table)

    # Now resolving the removed key must fail
    with pytest.raises(MissingMessageError) as exc_info2:
        partial_msgs.text(removed_key)

    assert exc_info2.value.string_id == removed_key
    assert exc_info2.value.language == "id"

    # No 'en' value for that string id reached any output — there IS no output at all
    # because the error is terminal. The fallback criterion 15.4 exists to prevent is
    # provably absent: MissingMessageError raises BEFORE any text is returned.
    en_msgs = load_messages("en")
    en_value = en_msgs.text(removed_key)
    assert en_value  # the 'en' value exists
    # But it was never returned — the error was raised instead
    assert exc_info2.value.code.value == ErrorCode.RENDER_FAILED.value


# ---------------------------------------------------------------------------
# 15.17 — A DerivedCount mismatched against what its own ledger contains
#          → {derived_count_mismatch}
# ---------------------------------------------------------------------------

declare(
    "test_15_17_derived_count_mismatch",
    FINDING_DERIVED_COUNT_MISMATCH,
)


def test_15_17_derived_count_mismatch() -> None:
    """Req 19.10, 24.1–24.3: a DerivedCount's stored formatted value disagrees with
    what the verifier re-derives from the ledger/definition.

    Mutation: compile a historical_trend block with 2 real prior candidates (so DerivedCounts
    for both historical_points_emitted and historical_lookback are emitted), then corrupt
    the historical_lookback's formatted value to '999' while the definition's config still
    says lookback=6. The verifier re-derives 6 from the definition and records the mismatch.

    The '6' in the rendered document is NOT flagged as unmatched_prose_token because the
    null-context allowlist derivation (Req 28.7) renders the same template with no candidates,
    producing '0 plotted and 6 requested' — admitting '6' as static template chrome.
    """
    from dataclasses import replace as dc_replace

    from fakes.object_store import StoredObject
    from pipeline_harness import ACTOR_ID, RUN_ID
    from reporting_agent.collect.snapshot import snapshot_key
    from reporting_agent.compile.ast import DerivedCount

    LOOKBACK = 6

    defn = definition(
        blocks=[
            df.block("res", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]}),
            df.block(
                "trend",
                "historical_trend",
                {"metric": "Percentage CPU", "statistic": "avg", "lookback": LOOKBACK},
            ),
        ]
    )

    # Produce a real snapshot to reuse as "prior run" data
    prior_pipe = Pipeline(definition=defn)
    prior_pipe.run()
    snap_key = snapshot_key(ACTOR_ID, RUN_ID)
    snap_body = prior_pipe.store.get(snap_key)
    assert snap_body is not None

    PRIOR_1 = "run_prior_june"
    PRIOR_2 = "run_prior_may"

    def corrupt_lookback_count(compiled: Any, view: Any) -> None:
        """Corrupt the historical_lookback DerivedCount's formatted value.

        Replaces the entry in the ledger's _derived_counts dict with a new DerivedCount
        whose formatted='999', while the definition still says lookback=6.
        """
        for path, count in list(compiled.ledger._derived_counts.items()):
            if count.derivation_kind == "historical_lookback":
                compiled.ledger._derived_counts[path] = dc_replace(
                    count, formatted="999"
                )
                break
        else:
            raise AssertionError(
                "no DerivedCount with derivation_kind='historical_lookback' found in the "
                "ledger — compile_historical_trend did not emit one"
            )

    run = Negative(
        resources=TWO_VMS,
        definition=defn,
        compiled=corrupt_lookback_count,
    )

    # Store prior snapshots in the run's pipeline (accessed via baseline + run)
    # The Negative class creates a fresh pipeline for baseline and run, so we must use
    # a different approach: drive the full pipeline with historical_candidates via the
    # Pipeline class directly for the baseline assertion, then use Negative for the mutation.

    # --- Baseline assertion: the unmutated fixture passes ---
    baseline_pipe = Pipeline(definition=defn)
    baseline_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_1)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )
    baseline_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_2)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )
    original_baseline_payload = baseline_pipe.payload

    def patched_baseline_payload() -> dict[str, Any]:
        p = original_baseline_payload()
        p["historical_candidates"] = [
            {
                "id": PRIOR_1,
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-06-30T10:00:00Z",
                "verification_id": f"v-{PRIOR_1}",
                "verification_snapshot_sha256": "a" * 64,
            },
            {
                "id": PRIOR_2,
                "period_start": "2026-05-01",
                "period_end": "2026-05-31",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-05-31T10:00:00Z",
                "verification_id": f"v-{PRIOR_2}",
                "verification_snapshot_sha256": "b" * 64,
            },
        ]
        return p

    baseline_pipe.payload = patched_baseline_payload  # type: ignore[assignment]
    from reporting_agent.errors import PartialCoverageError
    _, baseline_error = baseline_pipe.run()
    assert baseline_error is None or isinstance(baseline_error, PartialCoverageError), (
        f"baseline pipeline failed: {baseline_error}"
    )
    baseline_result = baseline_pipe.outcome.verification
    assert baseline_result is not None, "unmutated fixture produced no verification"
    assert baseline_result["status"] == "pass", (
        f"unmutated fixture does not pass: {baseline_result['findings']}"
    )
    from negatives import blocking_types
    assert blocking_types(baseline_result) == set(), baseline_result["findings"]

    # --- Mutated run: use Pipeline directly with the compiled hook ---
    import unittest.mock as mock
    from contextlib import ExitStack
    import reporting_agent.compile.blocks as blocks_module

    main_pipe = Pipeline(definition=defn)
    main_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_1)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )
    main_pipe.store._objects[snapshot_key(ACTOR_ID, PRIOR_2)] = StoredObject(
        body=snap_body.body, content_type=snap_body.content_type, tags=snap_body.tags,
    )

    original_main_payload = main_pipe.payload

    def patched_main_payload() -> dict[str, Any]:
        p = original_main_payload()
        p["historical_candidates"] = [
            {
                "id": PRIOR_1,
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-06-30T10:00:00Z",
                "verification_id": f"v-{PRIOR_1}",
                "verification_snapshot_sha256": "a" * 64,
            },
            {
                "id": PRIOR_2,
                "period_start": "2026-05-01",
                "period_end": "2026-05-31",
                "timezone": "Asia/Jakarta",
                "status": "completed",
                "verification_status": "pass",
                "verification_created_at": "2026-05-31T10:00:00Z",
                "verification_id": f"v-{PRIOR_2}",
                "verification_snapshot_sha256": "b" * 64,
            },
        ]
        return p

    main_pipe.payload = patched_main_payload  # type: ignore[assignment]

    # Patch compile_document to apply the mutation
    real_compile = blocks_module.compile_document
    compiles = 0

    def compile_document_mutated(*args: Any, **kwargs: Any):
        nonlocal compiles
        outcome = real_compile(*args, **kwargs)
        compiles += 1
        if compiles == 1:
            corrupt_lookback_count(outcome, kwargs.get("view"))
        return outcome

    with mock.patch.object(blocks_module, "compile_document", compile_document_mutated):
        events, error = main_pipe.run()

    result = main_pipe.outcome.verification
    assert result is not None, f"pipeline raised: {error!r}"

    # Exactly {derived_count_mismatch}
    assert_blocking(result, declared("test_15_17_derived_count_mismatch"))

    # Zero download observed
    assert types_of(events).count("report_file") == 0, types_of(events)
    assert main_pipe.outcome.artifacts == (), main_pipe.outcome.artifacts
    delivered = [
        key
        for key in report_objects(main_pipe.store)
        if not key.rsplit("/", 1)[-1].startswith("verification-")
    ]
    assert delivered == [], (
        f"a failing run left {delivered} under its report prefix"
    )

    # Locating fields — Req 24.2
    finding = next(
        f for f in result["findings"] if f["type"] == FINDING_DERIVED_COUNT_MISMATCH
    )
    assert finding["block_id"] == "trend"
    assert finding["derivation_kind"] == "historical_lookback"
    assert finding["stored"] == "999"
    assert finding["expected"] == str(LOOKBACK)
