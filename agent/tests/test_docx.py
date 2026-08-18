"""The DOCX emitter and the anchor contracts (Req 20, Req 21).

Read against the emitted `.docx` **as Word stores it** — `w:t` nodes, `w:tbl` elements,
`w:tblPr/w:tblCaption` — rather than through `python-docx`'s `.paragraphs` / `.tables`
convenience properties. Those enumerate only direct children of the body, so a paragraph
inside a table cell is invisible to them, and a test that used them would pass while
asserting nothing about half the document. That is the same trap `verify/tokens.py` has to
avoid (Req 26.2), and it is worth avoiding here too so the tests and the verifier read the
document the same way.

## The figure-run contract is the one to break first

Req 20.3 is what the whole verifier rests on: every figure is **exactly one run** in the
theme's `Figure` character style, holding the `formatted` string in full and no other
character. If that slips — a unit appended outside the run, a figure split across two runs,
a stray space inside it — the Token_Extractor stops finding figures and every affected
number becomes an unmatched prose token, withholding a document that is entirely correct.
So it is asserted at every position the AST can place a figure, not once.
"""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import zipfile
from collections.abc import Mapping
from pathlib import Path
from typing import Final

import pytest
from docx.oxml.ns import qn

import definition_factory as df
import snapshot_factory as sf
from reporting_agent.compile.ast import (
    Column,
    Document,
    EmptyCell,
    LayoutColumn,
    LayoutRow,
    PageBreak,
    Paragraph,
    Row,
    Table,
    Text,
    TextCell,
    compiling_against,
    figure_path,
)
from reporting_agent.compile.blocks import compile_document
from reporting_agent.compile.blocks.base import (
    EMPTY_SCOPE_TEXT,
    NOTICE_COLUMN_HEADER,
    DesignSettings,
)
from reporting_agent.compile.figures import FigureLedger
from reporting_agent.compile.snapshot_view import build_snapshot_view
from reporting_agent.errors import ErrorCode, RenderFailedError
from reporting_agent.render import anchors as A
from reporting_agent.render import docx as D
from reporting_agent.render.themes import FIGURE_CHARACTER_STYLE

AGENT_ROOT: Final[Path] = Path(__file__).resolve().parent.parent
W: Final[str] = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# --------------------------------------------------------------------------- #
# Helpers — read the package, not the convenience API
# --------------------------------------------------------------------------- #


DEFAULT_DESIGN: Final[dict[str, object]] = {
    "preset": "editorial",
    "accent_color": "#1f6f78",
    "density": "normal",
    "table_style": "hairline",
    "number_format": {"decimal_places": 2, "group_thousands": True},
    "cover_page": True,
    "logo": None,
    "page_size": "A4",
}


def render(blocks, *, design: Mapping[str, object] | None = None, view=None, **kwargs):
    """Compile a definition and render it, returning `(compiled, outcome)`.

    `design` is a **definition** design object, not a `DesignSettings`, and it is handed to
    both halves. That matters: `table_style`, `preset` and `cover_page` are decided by the
    *compiler* — the AST already carries the resolved table style name, and a false
    `cover_page` means the compiler emits no cover at all. Handing a `DesignSettings` only
    to the renderer would silently test nothing, which is what the first version of this
    helper did.
    """
    resolved_view = build_snapshot_view(sf.two_vm_snapshot()) if view is None else view
    design_object = {**DEFAULT_DESIGN, **(design or {})}
    compiled = compile_document(
        df.definition(blocks, design=design_object, **kwargs), view=resolved_view
    )
    outcome = D.render_document(
        compiled.document,
        ledger=compiled.ledger,
        design=DesignSettings.from_plain(design_object),
    )
    return compiled, outcome


def row_block(block_id: str, columns: list[list[dict]]) -> dict:
    """A `row` block.

    Carries **no** `config` key: the definition schema declares a row's fields as `id`,
    `type` and `columns`, and the validator rejects an unrecognised `config` on one. Worth a
    helper rather than repeating the literal, because the omission looks like a mistake.
    """
    return {"id": block_id, "type": "row", "columns": columns}


def document_xml(payload: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def body_element(payload: bytes):
    from lxml import etree

    return etree.fromstring(document_xml(payload).encode("utf-8")).find(f"{W}body")


def all_text(payload: bytes) -> list[str]:
    """Every `w:t` value, at every depth — including inside table cells."""
    return [node.text or "" for node in body_element(payload).iter(f"{W}t")]


def runs_with_style(payload: bytes, style_id: str) -> list[str]:
    """The text of every run carrying `style_id` as its character style."""
    found: list[str] = []
    for run in body_element(payload).iter(f"{W}r"):
        properties = run.find(f"{W}rPr")
        if properties is None:
            continue
        style = properties.find(f"{W}rStyle")
        if style is not None and style.get(f"{W}val") == style_id:
            found.append("".join(node.text or "" for node in run.iter(f"{W}t")))
    return found


def tables(payload: bytes) -> list:
    """Every `w:tbl` at every depth, in document order."""
    return list(body_element(payload).iter(f"{W}tbl"))


def caption_of(table_element) -> str | None:
    properties = table_element.find(f"{W}tblPr")
    if properties is None:
        return None
    for caption in properties.findall(f"{W}tblCaption"):
        value = caption.get(f"{W}val")
        if value and value.strip():
            return value
    return None


def table_style_of(table_element) -> str | None:
    properties = table_element.find(f"{W}tblPr")
    if properties is None:
        return None
    style = properties.find(f"{W}tblStyle")
    return None if style is None else style.get(f"{W}val")


def row_texts(table_element) -> list[list[str]]:
    return [
        ["".join(node.text or "" for node in cell.iter(f"{W}t")) for cell in row.findall(f"{W}tc")]
        for row in table_element.findall(f"{W}tr")
    ]


# --------------------------------------------------------------------------- #
# Req 20.3 — every figure is one run in the Figure character style
# --------------------------------------------------------------------------- #


def test_every_ledger_figure_is_emitted_as_a_figure_styled_run() -> None:
    compiled, outcome = render(
        [
            df.block("h", "heading", {"text": "Utilization", "level": 1}),
            df.block("p", "rich_text", {"text": "Prose."}),
            df.block("t", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]}),
            df.block("k", "kpi_row", {"metrics": [df.CPU_AVG]}),
            df.block("v", "verification_record", {}),
        ]
    )
    emitted = runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE)
    expected = [figure.formatted for figure in compiled.ledger.entries.values()]

    assert len(emitted) == len(expected), (len(emitted), len(expected))
    assert sorted(emitted) == sorted(expected)
    assert outcome.figures_emitted == len(expected)


def test_a_figure_run_holds_the_formatted_string_and_no_other_character() -> None:
    """Req 20.3's "no other character".

    A trailing space or an appended unit inside the run would break the exact-equality
    comparison the verifier performs, and the report would be withheld for a correct number.
    """
    compiled, outcome = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG]})]
    )
    formatted = {figure.formatted for figure in compiled.ledger.entries.values()}
    for run_text in runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE):
        assert run_text in formatted, run_text
        assert run_text == run_text.strip(), f"{run_text!r} carries padding"


def test_a_figure_in_prose_is_its_own_run_beside_the_prose() -> None:
    """The split is what makes a figure locatable without re-parsing the sentence."""
    view = build_snapshot_view(sf.two_vm_snapshot())
    ledger = FigureLedger()
    with compiling_against(view):
        value = view.stat(view.resources[0].resource_id, sf.CPU, "avg")
        assert value is not None
        from reporting_agent.compile.figures import BlockCursor

        cursor = BlockCursor(block_id="prose", ledger=ledger)
        figure = cursor.child("nodes", 0).child("inlines", 1).figure(value)
        document = Document(
            blocks=(
                Paragraph(
                    path=figure_path("prose", 0),
                    style="Body Text",
                    inlines=(Text(path=figure_path("prose", 0, 0), text="CPU averaged "), figure),
                ),
            )
        )
        outcome = D.render_document(document, ledger=ledger, design=DesignSettings())

    assert runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE) == [figure.formatted]
    # And the prose is a separate, unstyled run.
    assert "CPU averaged " in all_text(outcome.docx_bytes)


def test_a_figure_inside_a_data_table_cell_is_figure_styled() -> None:
    """Req 20.3 names the data-table cell explicitly, and a cell is where most figures are."""
    compiled, outcome = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG]})]
    )
    styled = runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE)
    assert styled, "no figure-styled run inside the table"
    assert len(styled) == len(compiled.ledger)


def test_a_figure_on_the_cover_is_figure_styled() -> None:
    """Req 20.3 names the cover field. The cover carries no metric value (Req 16.13), but
    the assertion is about the *wrapping* applying wherever a figure appears."""
    compiled, outcome = render(
        [
            df.block("cover", "cover", {"subtitle": "Monthly review"}),
            df.block("v", "verification_record", {}),
        ],
        design={"cover_page": True},
    )
    styled = runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE)
    assert len(styled) == len(compiled.ledger)


def test_a_figure_carrying_no_formatted_string_is_refused() -> None:
    """Req 7.9 — the Formatter is the only path from a value to a display string, so a
    figure that never took it has nothing legitimate to emit."""
    view = build_snapshot_view(sf.two_vm_snapshot())
    ledger = FigureLedger()
    with compiling_against(view):
        value = view.stat(view.resources[0].resource_id, sf.CPU, "avg")
        assert value is not None
        from reporting_agent.compile.figures import BlockCursor

        cursor = BlockCursor(block_id="p", ledger=ledger)
        figure = cursor.child("nodes", 0).child("inlines", 0).figure(value)
        object.__setattr__(figure, "formatted", "")
        document = Document(
            blocks=(Paragraph(path=figure_path("p", 0), style="Body Text", inlines=(figure,)),)
        )
        with pytest.raises(RenderFailedError, match="no formatted string"):
            D.render_document(document, ledger=ledger, design=DesignSettings())


# --------------------------------------------------------------------------- #
# Req 20.12 — a numeric in a text position that is not a Figure
# --------------------------------------------------------------------------- #


def test_a_non_figure_in_an_inline_position_is_refused_naming_the_path() -> None:
    """The AST refuses this at construction, so this drives the renderer directly to prove
    its own refusal exists rather than relying on the type that normally prevents it."""

    class Sneaky:
        """Neither Text nor Figure, carrying a quantity."""

        text = "12.4"

    document = Document(
        blocks=(Paragraph(path=figure_path("p", 0), style="Body Text", inlines=()),)
    )
    object.__setattr__(document.blocks[0], "inlines", (Sneaky(),))

    with pytest.raises(RenderFailedError) as raised:
        D.render_document(document, ledger=FigureLedger(), design=DesignSettings())
    assert "Sneaky" in raised.value.message
    assert "p:0" in raised.value.message
    assert raised.value.code is ErrorCode.RENDER_FAILED


def test_an_unknown_block_type_is_refused_rather_than_skipped() -> None:
    """Emitting nothing would leave the document quietly missing a section, which is
    indistinguishable from a block that was never configured."""

    class Mystery:
        path = "mystery:0"

    document = Document(blocks=())
    object.__setattr__(document, "blocks", (Mystery(),))
    with pytest.raises(RenderFailedError, match="no emitter is declared"):
        D.render_document(document, ledger=FigureLedger(), design=DesignSettings())


# --------------------------------------------------------------------------- #
# Req 20.4, 20.5 — styles by name, and a missing style names every one
# --------------------------------------------------------------------------- #


def test_every_paragraph_and_table_carries_a_style_the_theme_declares() -> None:
    _, outcome = render(
        [
            df.block("h1", "heading", {"text": "One", "level": 1}),
            df.block("h4", "heading", {"text": "Four", "level": 4}),
            df.block("p", "rich_text", {"text": "Prose."}),
            df.block("t", "resource_table", {"columns": [df.CPU_AVG], "caption": "Cap"}),
        ]
    )
    xml = document_xml(outcome.docx_bytes)
    for style_id in ("Heading1", "Heading4", "BodyText", "TableHairline", "Caption"):
        assert f'w:val="{style_id}"' in xml, style_id


@pytest.mark.parametrize("table_style", ["hairline", "banded", "bordered"])
def test_the_declared_table_style_reaches_the_document(table_style: str) -> None:
    _, outcome = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG]})],
        design={"table_style": table_style},
    )
    expected = {"hairline": "TableHairline", "banded": "TableBanded", "bordered": "TableBordered"}
    assert table_style_of(tables(outcome.docx_bytes)[0]) == expected[table_style]


def test_a_style_the_theme_does_not_declare_is_terminal_render_failed() -> None:
    document = Document(
        blocks=(Paragraph(path=figure_path("p", 0), style="Heading 7", inlines=()),)
    )
    with pytest.raises(RenderFailedError) as raised:
        D.render_document(document, ledger=FigureLedger(), design=DesignSettings())
    assert "Heading 7" in raised.value.message
    assert "editorial" in raised.value.message
    assert raised.value.terminal is True


def test_assert_theme_declares_names_every_missing_style_not_the_first() -> None:
    """Req 20.5 — so one fix pass clears the render."""
    with pytest.raises(RenderFailedError) as raised:
        D.assert_theme_declares("editorial", ["Heading 7", "Nonexistent", "Body Text"])
    assert "Heading 7" in raised.value.message
    assert "Nonexistent" in raised.value.message
    assert "Body Text" not in raised.value.message


# --------------------------------------------------------------------------- #
# Req 21.1, 21.2, 21.10, 21.11 — the caption asymmetry
# --------------------------------------------------------------------------- #


def test_every_data_table_carries_exactly_one_caption_equal_to_its_identity() -> None:
    compiled, outcome = render(
        [
            df.block("t", "resource_table", {"columns": [df.CPU_AVG]}),
            df.block("v", "verification_record", {}),
        ]
    )
    emitted = tables(outcome.docx_bytes)
    assert len(emitted) == 2

    for element in emitted:
        properties = element.find(f"{W}tblPr")
        found = properties.findall(f"{W}tblCaption")
        assert len(found) == 1, f"{len(found)} captions on one table"
        assert found[0].get(f"{W}val") in outcome.table_identities

    assert set(outcome.table_identities) == set(compiled.ledger.table_identities())


def test_writing_a_caption_twice_leaves_exactly_one() -> None:
    """Idempotent by removal. Two captions is not a documented state: Word shows the first
    and the verifier's search would find both."""
    import docx as docx_module

    from reporting_agent.render.themes import load_theme

    document = load_theme("editorial")
    table = document.add_table(rows=1, cols=1)
    A.write_data_table_caption(table, "tbl:x:0")
    A.write_data_table_caption(table, "tbl:x:0")

    properties = table._tbl.tblPr
    assert len(properties.findall(qn("w:tblCaption"))) == 1
    assert A.read_table_caption(table) == "tbl:x:0"
    assert docx_module is not None  # keep the import meaningful to linters


def test_a_layout_table_carries_no_caption_and_no_header_row() -> None:
    """Req 21.2 — the asymmetry that lets the verifier exclude layout by construction."""
    _, outcome = render(
        [
            row_block("r", [
                    [df.block("a", "rich_text", {"text": "Left."})],
                    [df.block("b", "rich_text", {"text": "Right."})],
                ])
        ]
    )
    emitted = tables(outcome.docx_bytes)
    assert len(emitted) == 1
    assert caption_of(emitted[0]) is None
    assert table_style_of(emitted[0]) == "LayoutTable"

    look = emitted[0].find(f"{W}tblPr").find(f"{W}tblLook")
    if look is not None:
        assert look.get(f"{W}firstRow") == "0"

    # No w:tblHeader anywhere in it: that is the marker a data table's header row carries.
    assert not list(emitted[0].iter(f"{W}tblHeader"))


def test_a_data_table_nested_in_a_layout_cell_keeps_its_own_caption() -> None:
    """Req 21.10 — a data-bearing child of a `row` is checked while its container is
    skipped. The rule is about the table, never about where the table sits."""
    compiled, outcome = render(
        [
            row_block("r", [
                    [df.block("inner", "resource_table", {"columns": [df.CPU_AVG]})],
                    [df.block("b", "rich_text", {"text": "Right."})],
                ])
        ]
    )
    emitted = tables(outcome.docx_bytes)
    assert len(emitted) == 2

    captions = [caption_of(element) for element in emitted]
    assert captions.count(None) == 1, captions  # the layout container
    inner_caption = next(value for value in captions if value is not None)
    assert inner_caption.startswith("tbl:inner:")
    assert inner_caption in compiled.ledger.table_identities()


def test_a_data_table_carrying_zero_figures_still_registers_its_identity() -> None:
    """Req 21.11. Without this the verifier finds a captioned table it cannot resolve and
    reports `table_anchor_unexpected` for a table the compiler emitted correctly."""
    narrow = df.scope(tag_filters=[{"key": "env", "value": "nope"}])
    compiled, outcome = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG]}, scope_override=narrow)]
    )
    assert len(compiled.ledger) == 0, "the fixture must produce a figureless table"

    emitted = tables(outcome.docx_bytes)
    assert len(emitted) == 1
    identity = caption_of(emitted[0])
    assert identity is not None
    assert identity in compiled.ledger.table_identities()
    assert compiled.ledger.anchors() == {}

    # And the explicit row is present, with a real column header.
    texts = row_texts(emitted[0])
    assert texts[0] == [NOTICE_COLUMN_HEADER]
    assert [EMPTY_SCOPE_TEXT] in texts


def test_a_blank_caption_reads_as_absent() -> None:
    """Req 26.5's reasoning, enforced where the caption is read: a whitespace caption can
    smuggle a table neither into nor out of the data pass."""
    from reporting_agent.render.themes import load_theme

    document = load_theme("editorial")
    table = document.add_table(rows=1, cols=1)
    properties = table._tbl.tblPr
    properties.append(properties.makeelement(qn("w:tblCaption"), {qn("w:val"): "   "}))
    assert A.read_table_caption(table) is None


def test_an_empty_identity_is_refused() -> None:
    from reporting_agent.render.themes import load_theme

    document = load_theme("editorial")
    table = document.add_table(rows=1, cols=1)
    with pytest.raises(RenderFailedError, match="non-empty"):
        A.write_data_table_caption(table, "")


def test_an_over_long_identity_is_refused() -> None:
    from reporting_agent.render.themes import load_theme

    document = load_theme("editorial")
    table = document.add_table(rows=1, cols=1)
    with pytest.raises(RenderFailedError, match="at most 255"):
        A.write_data_table_caption(table, "tbl:" + "x" * 300)


# --------------------------------------------------------------------------- #
# Req 21.3, 21.8 — the anchor triple
# --------------------------------------------------------------------------- #


def test_every_figure_in_a_data_table_records_one_anchor_triple() -> None:
    compiled, _ = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX]})]
    )
    anchors = compiled.ledger.anchors()
    assert len(anchors) == len(compiled.ledger)

    triples = [
        (anchor.anchor_id, anchor.row_key, anchor.column_key) for anchor in anchors.values()
    ]
    assert len(set(triples)) == len(triples), "a triple addresses one cell"
    for anchor_id, row_key, column_key in triples:
        assert anchor_id.startswith("tbl:")
        assert row_key
        assert column_key


def test_the_recorded_row_key_is_the_key_columns_text_not_the_internal_row_key() -> None:
    """Req 21.5 defines the row key as what the verifier can read back out of the document.

    `Row.key` is a resource **id**; the first column shows a resource **name**. Recording
    the id would record something absent from the `.docx`.
    """
    compiled, outcome = render(
        [df.block("t", "resource_table", {"columns": [df.CPU_AVG]})]
    )
    recorded = {anchor.row_key for anchor in compiled.ledger.anchors().values()}
    assert recorded == {"prod-sql-01", "prod-web-01"}

    internal = {row.key for row in _first_table_node(compiled.document).rows}
    assert recorded != internal, "the two must differ, or this test proves nothing"

    # And what was recorded is genuinely in the document.
    emitted = all_text(outcome.docx_bytes)
    for key in recorded:
        assert key in emitted


def test_no_anchor_is_recorded_for_a_figure_outside_a_data_table() -> None:
    """Req 21.8 — not in a layout cell, a heading, a paragraph, a header or a footer."""
    view = build_snapshot_view(sf.two_vm_snapshot())
    ledger = FigureLedger()
    with compiling_against(view):
        from reporting_agent.compile.figures import BlockCursor

        value = view.stat(view.resources[0].resource_id, sf.CPU, "avg")
        assert value is not None
        cursor = BlockCursor(block_id="p", ledger=ledger)
        figure = cursor.child("nodes", 0).child("inlines", 0).figure(value)
        document = Document(
            blocks=(Paragraph(path=figure_path("p", 0), style="Body Text", inlines=(figure,)),)
        )
        outcome = D.render_document(document, ledger=ledger, design=DesignSettings())

    assert runs_with_style(outcome.docx_bytes, FIGURE_CHARACTER_STYLE) == [figure.formatted]
    assert ledger.anchors() == {}


def _first_table_node(document: Document) -> Table:
    for block in document.blocks:
        if isinstance(block, Table):
            return block
    raise AssertionError("no table in the document")


# --------------------------------------------------------------------------- #
# Req 21.4 — header text
# --------------------------------------------------------------------------- #


def test_the_header_row_is_the_tables_first_row_and_repeats_across_pages() -> None:
    _, outcome = render([df.block("t", "resource_table", {"columns": [df.CPU_AVG]})])
    element = tables(outcome.docx_bytes)[0]
    assert row_texts(element)[0] == ["Resource", "Percentage CPU (avg)"]
    # w:tblHeader on the first row, so a table spanning a break repeats the text the
    # verifier resolves columns by.
    assert list(element.findall(f"{W}tr")[0].iter(f"{W}tblHeader"))


@pytest.mark.parametrize(
    ("columns", "expected"),
    [
        ((Column(key="a", header=""),), "empty header"),
        ((Column(key="a", header="x" * 256),), "over 255"),
        ((Column(key="a", header="Same"), Column(key="b", header="Same")), "share the header"),
        ((), "declares no columns"),
    ],
)
def test_a_header_row_that_cannot_be_resolved_is_refused(columns, expected: str) -> None:
    with pytest.raises(RenderFailedError, match=expected):
        A.assert_header_row(columns, at="table 't:0'")


def test_every_header_problem_is_reported_not_the_first() -> None:
    columns = (
        Column(key="a", header=""),
        Column(key="b", header="y" * 300),
        Column(key="c", header="Dup"),
        Column(key="d", header="Dup"),
    )
    with pytest.raises(RenderFailedError) as raised:
        A.assert_header_row(columns, at="table 't:0'")
    message = raised.value.message
    assert "'a'" in message and "'b'" in message and "'d'" in message


def test_a_valid_header_row_returns_the_header_text_in_column_order() -> None:
    columns = (Column(key="a", header="Alpha"), Column(key="b", header="Beta"))
    assert A.assert_header_row(columns, at="t") == ("Alpha", "Beta")


# --------------------------------------------------------------------------- #
# Req 21.5 — row keys
# --------------------------------------------------------------------------- #


def test_a_row_with_no_readable_key_column_text_is_refused() -> None:
    row = Row(
        path=figure_path("t", 0, 0),
        key="k",
        cells=(EmptyCell(path=figure_path("t", 0, 0, 0)),),
    )
    with pytest.raises(RenderFailedError, match="no readable text in key column"):
        A.document_row_key(row, at="table 't:0' row 0")


def test_a_row_with_no_cells_is_refused() -> None:
    row = Row(path=figure_path("t", 0, 0), key="k", cells=())
    with pytest.raises(RenderFailedError, match="carries no key column"):
        A.document_row_key(row, at="table 't:0' row 0")


def test_a_whitespace_only_key_column_is_refused() -> None:
    row = Row(
        path=figure_path("t", 0, 0),
        key="k",
        cells=(TextCell(path=figure_path("t", 0, 0, 0), text="   "),),
    )
    with pytest.raises(RenderFailedError, match="no readable text"):
        A.document_row_key(row, at="table 't:0' row 0")


def test_repeated_row_keys_are_refused_naming_the_table() -> None:
    """Two rows the verifier cannot tell apart. Refused rather than disambiguated: a
    suffix invented here would put a string in the document that came from neither the
    snapshot nor the template."""
    node = Table(
        path=figure_path("t", 0),
        style="Table Hairline",
        columns=(Column(key="name", header="Resource"),),
        rows=(
            Row(
                path=figure_path("t", 0, 0),
                key="a",
                cells=(TextCell(path=figure_path("t", 0, 0, 0), text="vm-1"),),
            ),
            Row(
                path=figure_path("t", 0, 1),
                key="b",
                cells=(TextCell(path=figure_path("t", 0, 1, 0), text="vm-1"),),
            ),
        ),
    )
    recorder = A.AnchorRecorder(ledger=FigureLedger())
    with pytest.raises(RenderFailedError) as raised:
        recorder.row_keys_for(node)
    assert "repeated row keys" in raised.value.message
    assert "'vm-1'" in raised.value.message
    assert "t:0" in raised.value.message


# --------------------------------------------------------------------------- #
# Req 21.6, 21.9 — identity from the path alone, unique in the document
# --------------------------------------------------------------------------- #


def test_a_table_identity_is_derived_from_the_ast_path_alone() -> None:
    compiled, first = render([df.block("t", "resource_table", {"columns": [df.CPU_AVG]})])
    second = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings()
    )
    assert first.table_identities == second.table_identities
    assert first.table_identities == ("tbl:t:0",)


def test_a_duplicate_identity_is_refused_naming_both_paths() -> None:
    recorder = A.AnchorRecorder(ledger=FigureLedger())
    node = Table(
        path=figure_path("t", 0),
        style="Table Hairline",
        columns=(Column(key="a", header="A"),),
    )
    recorder.claim_identity(node, "tbl:t:0")
    other = Table(
        path=figure_path("u", 0),
        style="Table Hairline",
        columns=(Column(key="a", header="A"),),
    )
    with pytest.raises(RenderFailedError, match="emitted twice"):
        recorder.claim_identity(other, "tbl:t:0")


def test_identities_are_unique_across_a_document_using_many_tables() -> None:
    _, outcome = render(
        [
            df.block("t1", "resource_table", {"columns": [df.CPU_AVG]}),
            df.block("t2", "top_n_table", {"columns": [df.CPU_AVG], "order_by": df.CPU_AVG}),
            df.block("k", "kpi_row", {"metrics": [df.CPU_AVG]}),
            df.block("v", "verification_record", {}),
            df.block("g", "gaps_and_coverage", {}),
        ]
    )
    assert len(set(outcome.table_identities)) == len(outcome.table_identities)
    assert len(outcome.table_identities) >= 5


# --------------------------------------------------------------------------- #
# Req 20.6, 20.7 — layout and order
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize("column_count", [2, 3])
def test_a_row_becomes_one_layout_table_of_its_declared_column_count(
    column_count: int,
) -> None:
    columns = [
        [df.block(f"c{index}", "rich_text", {"text": f"Column {index}."})]
        for index in range(column_count)
    ]
    _, outcome = render([row_block("r", columns)])
    emitted = tables(outcome.docx_bytes)
    assert len(emitted) == 1
    assert len(emitted[0].findall(f"{W}tr")) == 1
    assert len(emitted[0].findall(f"{W}tr")[0].findall(f"{W}tc")) == column_count


def test_each_child_is_emitted_into_the_column_it_declares_in_declared_order() -> None:
    _, outcome = render(
        [
            row_block("r", [
                    [
                        df.block("a1", "rich_text", {"text": "Left first."}),
                        df.block("a2", "rich_text", {"text": "Left second."}),
                    ],
                    [df.block("b1", "rich_text", {"text": "Right only."})],
                ])
        ]
    )
    cells = tables(outcome.docx_bytes)[0].findall(f"{W}tr")[0].findall(f"{W}tc")
    left = [node.text for node in cells[0].iter(f"{W}t")]
    right = [node.text for node in cells[1].iter(f"{W}t")]
    assert left == ["Left first.", "Left second."]
    assert right == ["Right only."]


def test_a_layout_column_leaves_no_leading_blank_paragraph() -> None:
    """python-docx gives a new cell one empty paragraph; emitting appends after it, which
    would put a blank line above every column's first block."""
    _, outcome = render(
        [
            row_block("r", [
                    [df.block("a", "rich_text", {"text": "Left."})],
                    [df.block("b", "rich_text", {"text": "Right."})],
                ])
        ]
    )
    cells = tables(outcome.docx_bytes)[0].findall(f"{W}tr")[0].findall(f"{W}tc")
    for cell in cells:
        paragraphs = cell.findall(f"{W}p")
        assert len(paragraphs) == 1, [
            "".join(n.text or "" for n in p.iter(f"{W}t")) for p in paragraphs
        ]


def test_blocks_are_emitted_in_ast_order_with_no_reordering() -> None:
    """Req 20.7 — no reordering derived from type, content length or figure count."""
    _, outcome = render(
        [
            df.block("z", "rich_text", {"text": "First."}),
            df.block("t", "resource_table", {"columns": [df.CPU_AVG]}),
            df.block("a", "rich_text", {"text": "Third."}),
        ]
    )
    texts = all_text(outcome.docx_bytes)
    assert texts.index("First.") < texts.index("Resource")
    assert texts.index("Resource") < texts.index("Third.")


def test_a_layout_row_inside_a_layout_cell_is_refused() -> None:
    """Nesting is one level deep, so this cannot arise from a valid definition — and a
    nested layout table would hand the verifier a container it has no rule for."""
    inner = LayoutRow(
        path=figure_path("inner", 0),
        columns=(
            LayoutColumn(path=figure_path("inner", 0, 0)),
            LayoutColumn(path=figure_path("inner", 0, 1)),
        ),
    )
    outer = LayoutRow(
        path=figure_path("outer", 0),
        columns=(
            LayoutColumn(path=figure_path("outer", 0, 0), blocks=(inner,)),
            LayoutColumn(path=figure_path("outer", 0, 1)),
        ),
    )
    with pytest.raises(RenderFailedError, match="nesting is one level deep"):
        D.render_document(
            Document(blocks=(outer,)), ledger=FigureLedger(), design=DesignSettings()
        )


def test_a_page_break_is_emitted_as_a_break(  ) -> None:
    _, outcome = render(
        [
            df.block("a", "rich_text", {"text": "Before."}),
            df.block("pb", "page_break", {}),
            df.block("b", "rich_text", {"text": "After."}),
        ]
    )
    breaks = [
        node
        for node in body_element(outcome.docx_bytes).iter(f"{W}br")
        if node.get(f"{W}type") == "page"
    ]
    assert len(breaks) == 1


def test_a_page_break_inside_a_layout_cell_is_refused() -> None:
    """A break there applies to the cell rather than to the page."""
    row = LayoutRow(
        path=figure_path("r", 0),
        columns=(
            LayoutColumn(
                path=figure_path("r", 0, 0), blocks=(PageBreak(path=figure_path("pb", 0)),)
            ),
            LayoutColumn(path=figure_path("r", 0, 1)),
        ),
    )
    with pytest.raises(RenderFailedError, match="applies to the cell"):
        D.render_document(
            Document(blocks=(row,)), ledger=FigureLedger(), design=DesignSettings()
        )


# --------------------------------------------------------------------------- #
# Short rows and empty cells
# --------------------------------------------------------------------------- #


def test_a_short_row_is_padded_rather_than_refused() -> None:
    """A truncation row carries two cells in a wider table, deliberately: it is a statement
    about the table rather than another resource."""
    node = Table(
        path=figure_path("t", 0),
        style="Table Hairline",
        columns=(
            Column(key="a", header="A"),
            Column(key="b", header="B"),
            Column(key="c", header="C"),
        ),
        rows=(
            Row(
                path=figure_path("t", 0, 0),
                key="only",
                cells=(TextCell(path=figure_path("t", 0, 0, 0), text="just one"),),
            ),
        ),
    )
    outcome = D.render_document(
        Document(blocks=(node,)), ledger=FigureLedger(), design=DesignSettings()
    )
    grid = row_texts(tables(outcome.docx_bytes)[0])
    assert grid[0] == ["A", "B", "C"]
    assert grid[1] == ["just one", "", ""]


def test_a_row_with_more_cells_than_columns_is_refused() -> None:
    """A cell with no column cannot be addressed as (row key, column key)."""
    node = Table(
        path=figure_path("t", 0),
        style="Table Hairline",
        columns=(Column(key="a", header="A"),),
        rows=(
            Row(
                path=figure_path("t", 0, 0),
                key="k",
                cells=(
                    TextCell(path=figure_path("t", 0, 0, 0), text="one"),
                    TextCell(path=figure_path("t", 0, 0, 1), text="two"),
                ),
            ),
        ),
    )
    with pytest.raises(RenderFailedError, match="cannot be addressed"):
        D.render_document(
            Document(blocks=(node,)), ledger=FigureLedger(), design=DesignSettings()
        )


def test_an_empty_cell_emits_nothing_rather_than_a_zero() -> None:
    """The distinction is the product: a metric a resource does not emit is a recorded gap,
    and a zero would read as measured idleness."""
    node = Table(
        path=figure_path("t", 0),
        style="Table Hairline",
        columns=(Column(key="a", header="A"), Column(key="b", header="B")),
        rows=(
            Row(
                path=figure_path("t", 0, 0),
                key="k",
                cells=(
                    TextCell(path=figure_path("t", 0, 0, 0), text="vm-1"),
                    EmptyCell(path=figure_path("t", 0, 0, 1)),
                ),
            ),
        ),
    )
    outcome = D.render_document(
        Document(blocks=(node,)), ledger=FigureLedger(), design=DesignSettings()
    )
    assert row_texts(tables(outcome.docx_bytes)[0])[1] == ["vm-1", ""]
    assert "0" not in "".join(all_text(outcome.docx_bytes))


# --------------------------------------------------------------------------- #
# Req 20.8 — determinism
# --------------------------------------------------------------------------- #

_EVERYTHING = [
    df.block("cover", "cover", {"subtitle": "Monthly"}),
    df.block("h", "heading", {"text": "Utilization", "level": 2}),
    df.block("p", "rich_text", {"text": "Prose."}),
    df.block("t", "resource_table", {"columns": [df.CPU_AVG, df.CPU_MAX], "caption": "Cap"}),
    df.block("k", "kpi_row", {"metrics": [df.CPU_AVG]}),
    df.block("g", "gaps_and_coverage", {}),
    df.block("v", "verification_record", {}),
    df.block("m", "appendix_methodology", {}),
]


def test_two_emissions_of_one_ast_are_byte_identical() -> None:
    compiled, first = render(_EVERYTHING)
    second = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings()
    )
    assert first.docx_bytes == second.docx_bytes


def test_the_timestamps_are_a_fixed_sentinel_rather_than_the_clock() -> None:
    """The one thing a document writer reaches for that makes byte equality impossible."""
    _, outcome = render(_EVERYTHING)
    with zipfile.ZipFile(io.BytesIO(outcome.docx_bytes)) as archive:
        core = archive.read("docProps/core.xml").decode("utf-8")
    assert core.count(D.FIXED_TIMESTAMP) >= 2
    assert not re.search(r"20[0-9]{2}-[0-9]{2}-[0-9]{2}T", core.replace(D.FIXED_TIMESTAMP, ""))


def test_the_volatile_part_list_names_one_part_and_it_exists() -> None:
    """"The digests match apart from timestamps" is only meaningful if the exclusion is a
    fixed list rather than whatever happened to differ."""
    _, outcome = render(_EVERYTHING)
    with zipfile.ZipFile(io.BytesIO(outcome.docx_bytes)) as archive:
        names = set(archive.namelist())
    assert D.VOLATILE_PACKAGE_PARTS == ("docProps/core.xml",)
    for part in D.VOLATILE_PACKAGE_PARTS:
        assert part in names


def test_every_part_except_the_volatile_one_is_byte_identical_across_renders() -> None:
    """The stronger form: not just equal digests, equal parts, with the exclusion named."""
    compiled, first = render(_EVERYTHING)
    second = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings()
    )

    def parts(payload: bytes) -> dict[str, bytes]:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            return {
                name: archive.read(name)
                for name in archive.namelist()
                if name not in D.VOLATILE_PACKAGE_PARTS
            }

    assert parts(first.docx_bytes) == parts(second.docx_bytes)


# --------------------------------------------------------------------------- #
# Req 7.5, 20.11 — the cover flag and the single write
# --------------------------------------------------------------------------- #


def test_a_false_cover_flag_emits_no_cover_content_and_no_leading_blank_page() -> None:
    _, without = render(
        [
            df.block("cover", "cover", {"subtitle": "Monthly review"}),
            df.block("p", "rich_text", {"text": "Body."}),
        ],
        design={"cover_page": False},
    )
    texts = all_text(without.docx_bytes)
    assert "Monthly review" not in texts
    assert texts[0] == "Body.", texts
    breaks = [
        node
        for node in body_element(without.docx_bytes).iter(f"{W}br")
        if node.get(f"{W}type") == "page"
    ]
    assert breaks == []


def test_a_true_cover_flag_emits_the_cover(  ) -> None:
    _, with_cover = render(
        [
            df.block("cover", "cover", {"subtitle": "Monthly review"}),
            df.block("p", "rich_text", {"text": "Body."}),
        ],
        design={"cover_page": True},
    )
    assert "Monthly review" in all_text(with_cover.docx_bytes)


def test_the_renderer_returns_bytes_rather_than_writing_them() -> None:
    """Req 20.11 — the completed document is written as **one** artifact, and the cleanest
    guarantee is a renderer that cannot write at all."""
    _, outcome = render(_EVERYTHING)
    assert isinstance(outcome.docx_bytes, bytes)
    assert outcome.docx_bytes[:2] == b"PK"


def test_a_non_document_argument_is_refused() -> None:
    with pytest.raises(RenderFailedError, match="compiled Document"):
        D.render_document({}, ledger=FigureLedger(), design=DesignSettings())


# --------------------------------------------------------------------------- #
# Page size and preview
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(("page_size", "width"), [("A4", 11906), ("Letter", 12240)])
def test_the_declared_page_size_reaches_the_section(page_size: str, width: int) -> None:
    """The unit conversion this asserts is not cosmetic: assigning raw twips to a `Length`
    is read as EMU, which makes the page about 0.013 inch wide. Every paragraph then
    overflows it and the converted PDF is a stack of blank pages — a document that renders,
    converts and contains nothing."""
    _, outcome = render(
        [df.block("p", "rich_text", {"text": "Body."})],
        design={"page_size": page_size},
    )
    section = body_element(outcome.docx_bytes).find(f"{W}sectPr").find(f"{W}pgSz")
    assert int(section.get(f"{W}w")) == width


def test_an_undeclared_page_size_is_refused() -> None:
    """Driven directly rather than through a definition, because the **validator** rejects
    an undeclared page size first (`design.page_size must be one of: A4, Letter`).

    So this refusal is unreachable from a valid definition — which is the point of asserting
    it. It is the renderer's own floor, and it holds for a caller that builds a
    `DesignSettings` by hand: a page size nobody has a twips value for would otherwise fall
    through to a `KeyError` or, worse, to a default nobody chose.
    """
    compiled, _ = render([df.block("p", "rich_text", {"text": "Body."})])
    with pytest.raises(RenderFailedError, match="page size"):
        D.render_document(
            compiled.document,
            ledger=compiled.ledger,
            design=DesignSettings(page_size="Legal"),
        )


def header_texts(payload: bytes) -> list[str]:
    """Every `w:t` value in every header part of the package.

    A header part is what makes a notice **per-page**: Word repeats it on every page of the
    section it belongs to. Read as its own part rather than through `all_text`, which walks
    `word/document.xml`'s body and by construction cannot see a header at all.
    """
    from lxml import etree

    found: list[str] = []
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/header"):
                continue
            root = etree.fromstring(archive.read(name))
            found.extend(node.text or "" for node in root.iter(f"{W}t"))
    return found


def test_preview_mode_emits_the_notice_and_normal_mode_does_not() -> None:
    """The notice is per-page, which means it is in the header and not in the body.

    Both halves matter. In the header, Word repeats it on every page, so page seven of a
    forwarded preview says what it is as loudly as page one. Out of the body, it shifts no
    block — a notice that changed the pagination would be changing the thing the preview
    exists to show.
    """
    compiled, _ = render([df.block("p", "rich_text", {"text": "Body."})])
    preview = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings(), preview=True
    )
    plain = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings(), preview=False
    )

    assert D.PREVIEW_NOTICE_TEXT in header_texts(preview.docx_bytes)
    assert D.PREVIEW_NOTICE_TEXT not in all_text(preview.docx_bytes)
    assert D.PREVIEW_NOTICE_TEXT not in header_texts(plain.docx_bytes)
    assert D.PREVIEW_NOTICE_TEXT not in all_text(plain.docx_bytes)


def test_the_preview_notice_reaches_every_page_including_the_first() -> None:
    """A theme that set `different_first_page_header_footer` would otherwise make page one —
    the cover, the page most likely to be screenshotted — the one page with no notice.

    Asserted by writing **both** headers when the flag is set rather than by clearing the
    flag. Clearing it would reach page one too, and would do it by overriding a theme
    decision: a preset that suppresses its cover-page header would then preview a page 1
    layout the real render does not produce, which is the one thing a preview must not do.
    """
    from docx import Document as open_docx

    compiled, _ = render([df.block("p", "rich_text", {"text": "Body."})])
    preview = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings(), preview=True
    )

    document = open_docx(io.BytesIO(preview.docx_bytes))
    for section in document.sections:
        assert section.header.is_linked_to_previous is False
        assert D.PREVIEW_NOTICE_TEXT in [p.text for p in section.header.paragraphs]
        if section.different_first_page_header_footer:
            assert D.PREVIEW_NOTICE_TEXT in [
                p.text for p in section.first_page_header.paragraphs
            ]


def test_the_preview_does_not_override_the_themes_first_page_header_setting() -> None:
    """The flag a preview must not decide for the theme.

    Every preset ships `False` today, so this reads as a tautology and is not: it is the
    assertion that fails if somebody reintroduces `different_first_page_header_footer = False`
    to reach page one the easy way. The preview would then differ from the real render for any
    future preset that suppresses its cover-page header.
    """
    from docx import Document as open_docx

    compiled, _ = render([df.block("p", "rich_text", {"text": "Body."})])
    for preset in ("editorial", "corporate", "technical", "minimal"):
        design = DesignSettings(preset=preset)

        preview = D.render_document(
            compiled.document, ledger=compiled.ledger, design=design, preview=True
        )
        plain = D.render_document(
            compiled.document, ledger=compiled.ledger, design=design, preview=False
        )

        previewed = [
            section.different_first_page_header_footer
            for section in open_docx(io.BytesIO(preview.docx_bytes)).sections
        ]
        rendered = [
            section.different_first_page_header_footer
            for section in open_docx(io.BytesIO(plain.docx_bytes)).sections
        ]
        assert previewed == rendered, preset


@pytest.mark.parametrize("preset", ["editorial", "corporate", "technical", "minimal"])
def test_every_preset_renders_the_same_document(preset: str) -> None:
    compiled, _ = render(_EVERYTHING)
    outcome = D.render_document(
        compiled.document, ledger=compiled.ledger, design=DesignSettings(preset=preset)
    )
    assert outcome.figures_emitted == len(compiled.ledger)
    assert outcome.docx_bytes[:2] == b"PK"


# --------------------------------------------------------------------------- #
# One real conversion — the assertion a fake cannot make
# --------------------------------------------------------------------------- #

_SOFFICE: Final[str | None] = shutil.which("soffice")


@pytest.mark.skipif(_SOFFICE is None, reason="LibreOffice is not installed")
def test_every_ledger_figure_survives_into_the_converted_pdf(tmp_path: Path) -> None:
    """The contract Req 23's PDF pass will depend on, asserted through real LibreOffice.

    A `.docx` that is structurally perfect can still convert to a PDF in which a figure is
    unfindable — the two ways that happen are a case-transforming style (`w:caps` rewrites
    the glyphs and leaves `w:t` alone) and a page size in the wrong unit (every paragraph
    overflows a microscopic page and the PDF is blank). Both were real defects in this
    module's first version, and neither is visible from the `.docx` alone.

    So: render, convert, extract, and require every ledger `formatted` string to be a
    contiguous substring of the normalized text. Normalized the way `verify/tokens.py` will
    normalize it — every whitespace run collapsed to one space — because a conversion is
    free to break a line wherever it likes.
    """
    compiled, outcome = render(_EVERYTHING)
    source = tmp_path / "report.docx"
    source.write_bytes(outcome.docx_bytes)

    result = subprocess.run(
        [
            _SOFFICE,
            "--headless",
            "--norestore",
            f"-env:UserInstallation=file://{tmp_path / 'profile'}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(source),
        ],
        capture_output=True,
        text=True,
        env={"LANG": "C.UTF-8", "PATH": "/usr/bin:/bin", "HOME": str(tmp_path)},
        timeout=300,
        check=False,
    )
    produced = tmp_path / "report.pdf"
    assert result.returncode == 0, result.stderr
    assert produced.is_file() and produced.stat().st_size > 0

    from pypdf import PdfReader

    reader = PdfReader(str(produced))
    assert len(reader.pages) >= 1
    normalized = " ".join(
        " ".join(page.extract_text().split()) for page in reader.pages
    )

    assert normalized.strip(), (
        "the converted PDF carries no extractable text at all — the usual cause is a page "
        "size assigned in the wrong unit, which makes every page blank"
    )

    missing = [
        figure.formatted
        for figure in compiled.ledger.entries.values()
        if figure.formatted not in normalized
    ]
    assert missing == [], (
        f"{len(missing)} ledger figure(s) are not contiguous substrings of the converted "
        f"PDF text: {missing}. The PDF pass would report pdf_figure_missing for a document "
        f"that is entirely correct."
    )


@pytest.mark.skipif(_SOFFICE is None, reason="LibreOffice is not installed")
def test_a_figure_beside_a_case_transforming_style_survives_conversion(
    tmp_path: Path,
) -> None:
    """The narrow case the `Figure` style's `w:caps w:val="0"` exists for.

    `technical` sets `w:caps` on `Heading 4`, and a figure in a heading is ordinary. Digits
    have no case, so without the override this would fail only for units and estimator
    labels — intermittently, on some presets, for some metrics.
    """
    compiled, _ = render(
        [
            df.block("h", "heading", {"text": "Memory", "level": 4}),
            df.block("t", "resource_table", {"columns": [df.MEMORY_USED_PCT_AVG]}),
        ],
        design={"preset": "technical"},
    )
    outcome = D.render_document(
        compiled.document,
        ledger=compiled.ledger,
        design=DesignSettings(preset="technical"),
    )
    source = tmp_path / "caps.docx"
    source.write_bytes(outcome.docx_bytes)

    subprocess.run(
        [
            _SOFFICE,
            "--headless",
            "--norestore",
            f"-env:UserInstallation=file://{tmp_path / 'profile'}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(source),
        ],
        capture_output=True,
        text=True,
        env={"LANG": "C.UTF-8", "PATH": "/usr/bin:/bin", "HOME": str(tmp_path)},
        timeout=300,
        check=True,
    )

    from pypdf import PdfReader

    normalized = " ".join(
        " ".join(page.extract_text().split())
        for page in PdfReader(str(tmp_path / "caps.pdf")).pages
    )
    for figure in compiled.ledger.entries.values():
        assert figure.formatted in normalized, (figure.formatted, normalized)
