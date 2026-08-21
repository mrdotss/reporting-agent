"""The one table-of-contents measurement, used by the evaluation and by the proof test
(Req 14.2, 14.11).

## Why there is exactly one of these

Task 2.3 evaluates three candidate approaches and records a verdict for each; task 2.4's
proof test then asserts the adopted one is still correct on every run. If those two measured
through different code, a `correct` verdict would mean "candidate B satisfied the
evaluation's notion of correct" while the proof asserted something else — and the
disagreement would be invisible, because both would be green.

So this module owns the measurement and both callers import it. `agent/tests/
test_toc_evidence.py` additionally asserts that the fixture the evidence record names is the
fixture the proof test uses, compared by path **and** by content digest, which closes the
other half of the same hole.

## What is measured, and the one subtlety in it

A table of contents claims that a heading is on a page. Checking that claim needs two
numbers per heading:

* `named_pages` — the page number the **document itself prints** in its table of contents.
* `observed_pages` — the page the heading **actually landed on** after conversion.

`observed_pages` resolves a heading to the page carrying its **first rendered character**.
That rule is not a tie-break detail; it is what makes the measurement single-valued. A
heading long enough to wrap can straddle a page boundary, and "the page it is on" then has
two defensible answers. Taking the first character means a heading is on exactly one page,
which is also the page a reader turns to — a table of contents that pointed at the page
holding a heading's *second* line would send them one page late.

## Not a unit test's shape, deliberately

`measure` runs the real `compile → render/docx.py → render/pdf.py` path, which means it
spawns LibreOffice and takes seconds rather than milliseconds. That cost is the point:
pagination is decided by the converter, so a measurement that stubbed the conversion would
be measuring an assumption. There is no faster honest version of this.
"""

from __future__ import annotations

import hashlib
import io
import re
import subprocess
import tempfile
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Final

from docx import Document as open_docx
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

import definition_factory as df  # noqa: F401  (imported for callers' convenience)
from reporting_agent.compile.blocks import compile_document
from reporting_agent.compile.blocks.base import DesignSettings
from reporting_agent.compile.snapshot_view import build_snapshot_view
from reporting_agent.render import docx as D
from reporting_agent.render.pdf import (
    CONVERT_TIMEOUT_S,
    SOFFICE_BINARY,
    assert_lang_in_effect,
    convert_to_pdf,
    profile_path,
)
from reporting_agent.render.themes import TOC_ENTRY_STYLE
from reporting_agent.render.toc import (
    TOC_APPROACH_CONVERSION_MACRO,
    TOC_APPROACH_LIBREOFFICE_INDEX,
    TOC_APPROACH_NONE,
    TOC_APPROACH_TWO_PASS,
    TOC_APPROACHES,
)
from reporting_agent.verify.tokens import pdf_page_texts

__all__ = [
    "FIXTURE_DIR",
    "MACRO_LIBRARY_RELATIVE_PATH",
    "MACRO_NAME",
    "TOC_FIELD_INSTRUCTION",
    "TOC_LABEL",
    "TocApproachUnavailableError",
    "TocMeasurement",
    "fixture_digest",
    "load_fixture",
    "measure",
    "soffice_version",
]

FIXTURE_DIR: Final[Path] = Path(__file__).resolve().parent / "fixtures" / "toc"

TOC_LABEL: Final[str] = "Contents"
"""The heading above the table of contents.

Styled `Title` rather than `Heading 1`, and that is load-bearing rather than typographic:
`\\o "1-3"` collects `Heading 1` through `Heading 3`, so a `Heading 1` label would put
"Contents" into its own table of contents. It would also add a seventh heading to
`observed_pages` for a section the definition never declared, which would make the fixture's
`headings` count disagree with the six headings it actually declares.
"""

TOC_FIELD_INSTRUCTION: Final[str] = ' TOC \\o "1-3" \\h \\z \\u '
"""Candidate A's field instruction, spelled exactly as criterion 14.1 names it.

`\\o "1-3"` collects heading levels 1 to 3, `\\h` makes each entry a hyperlink, `\\z` hides
tab leaders and page numbers in Web layout, `\\u` uses the paragraph outline level.
"""

MACRO_NAME: Final[str] = "Standard.Module1.ConvertWithIndexes"
MACRO_LIBRARY_RELATIVE_PATH: Final[str] = "user/basic/Standard/Module1.xba"
"""Where candidate C's macro would have to live inside the LibreOffice profile.

Named as a constant because :func:`_convert_via_macro` **reads** it to decide whether the
candidate is available at all, and never writes it. Installing the macro would be a write
into the profile the image pre-warms at build time — which is exactly the "writable macro
library" criterion 14.3 makes a rejection.
"""

_PAGE_NUMBER_SUFFIX: Final[str] = r"[\s.·…\u2022\t]*(\d+)"
"""What separates a contents entry's text from its page number: a run of dot-leader
characters, a tab, plain whitespace, or nothing at all.

Not anchored to a line boundary, because the PDF text extractor returns an entire contents
page as **one line** — six entries and their leaders, joined. What keeps the pattern
unambiguous is *where* it is applied: only to the contents pages
(:func:`_contents_page_indices`), never to body prose."""


class TocApproachUnavailableError(RuntimeError):
    """A candidate could not be exercised at all in this environment.

    Distinct from producing a wrong measurement, and the distinction is the whole reason
    criterion 14.1's verdict set has three members rather than two: `unavailable` records
    that the mechanism needs a facility the image does not provide, which is a different fact
    from `incorrect` — the mechanism ran and got the page numbers wrong. Collapsing them
    would put "we could not try it" and "we tried it and it lies" under one word.

    Raised rather than returned so a measurement is never fabricated for a candidate that did
    not run; :func:`measure`'s caller catches this and records the verdict.
    """


def load_fixture(name: str = "long_report") -> tuple[dict, dict]:
    """The committed `(definition, snapshot)` pair.

    Returned as plain data read from disk rather than rebuilt from the factories, because the
    evidence record names this file and its digest: a fixture rebuilt at measurement time
    could drift from the one a recorded verdict was measured against, and the verdict would
    then describe a document nobody has.
    """
    import json

    definition = json.loads(
        (FIXTURE_DIR / f"{name}.definition.json").read_text(encoding="utf-8")
    )
    snapshot = json.loads(
        (FIXTURE_DIR / f"{name}.snapshot.json").read_text(encoding="utf-8")
    )
    return definition, snapshot


def fixture_digest(name: str = "long_report") -> str:
    """A digest over both fixture files, so the evidence record can pin what was measured.

    Over the **bytes on disk**, in a fixed order, rather than over the parsed objects: the
    parsed form would hash the same after a reformat that changed what a future reader sees,
    and the point of pinning is that the recorded verdict describes an exact input.
    """
    digest = hashlib.sha256()
    for suffix in ("definition", "snapshot"):
        digest.update((FIXTURE_DIR / f"{name}.{suffix}.json").read_bytes())
    return digest.hexdigest()


@dataclass(frozen=True, slots=True)
class TocMeasurement:
    """One rendered-and-converted measurement of one approach.

    `named_pages` and `observed_pages` are both `{heading text: page number}`, one-based.
    `named_pages` is **empty** for an approach that ships no table of contents — which is not
    a failure but the recorded outcome of :data:`TOC_APPROACH_NONE`, and is why the proof test
    branches on the adopted approach rather than asserting equality unconditionally.
    """

    docx_bytes: bytes
    pdf_bytes: bytes
    pdf_sha256: str
    observed_pages: Mapping[str, int]
    named_pages: Mapping[str, int]

    @property
    def pages(self) -> int:
        """How many pages the converted PDF has."""
        return self._page_count

    _page_count: int = 0

    @property
    def headings(self) -> int:
        return len(self.observed_pages)

    @property
    def distinct_heading_pages(self) -> int:
        """How many **different** pages carry a heading (Req 14.11).

        The fixture's whole purpose is that this is greater than one: a table of contents
        every entry of which points at page 1 would satisfy "named equals observed" while
        proving nothing about pagination.
        """
        return len(set(self.observed_pages.values()))


def _heading_texts(definition: Mapping[str, object]) -> tuple[str, ...]:
    """The headings the definition declares, in document order.

    Read from the **definition** rather than found in the rendered text, so a heading the
    renderer dropped entirely is a missing observation rather than a heading nobody looked
    for. That direction matters: scanning the PDF for things that look like headings would
    silently pass on a document that emitted none.
    """
    texts: list[str] = []
    for block in definition.get("blocks") or []:
        if not isinstance(block, Mapping) or block.get("type") != "heading":
            continue
        config = block.get("config")
        text = config.get("text") if isinstance(config, Mapping) else None
        if isinstance(text, str) and text.strip():
            texts.append(text)
    return tuple(texts)


def _contents_page_indices(pages: Sequence[str]) -> frozenset[int]:
    """The zero-based indices of the pages holding the table of contents itself.

    Identified by the page **starting** with :data:`TOC_LABEL`, which is exact for the section
    :func:`_prepend_toc` emits — the label is its first paragraph and a page break closes it,
    so no body content shares the page. A substring test would also match a body page whose
    prose happened to contain the word.

    Both halves of the measurement are bounded by this one rule, in opposite directions, and
    that symmetry is the point:

    * `observed_pages` **skips** these pages, because a contents entry names a heading and is
      not the heading. Without this, every heading is observed on the contents page — its text
      appears there first — and `distinct_heading_pages` collapses to 1 while the numbers
      printed beside them stay right, which reads as a *correct* table of contents for a
      document with no pagination.
    * `named_pages` searches **only** these pages, because outside them "a heading followed by
      a number" is ordinary prose rather than an entry.

    For :data:`TOC_APPROACH_NONE` the set is empty, so the baseline measurement is untouched.
    """
    return frozenset(
        index
        for index, page in enumerate(pages)
        if page.lstrip().startswith(TOC_LABEL)
    )


def _first_character_page(pages: Sequence[str], text: str) -> int | None:
    """The one-based page carrying `text`'s first rendered character, or `None`.

    Scans pages in order and returns the first that contains the heading, which **is** the
    first-character rule for any heading that fits on a page and the correct reading for one
    that does not: a heading straddling a boundary appears in full only on neither page, so
    the search falls back to progressively shorter prefixes, and the earliest page carrying
    any prefix is the page its first character is on.
    """
    contents = _contents_page_indices(pages)

    for index, page in enumerate(pages, start=1):
        if index - 1 in contents:
            continue
        if text in page:
            return index

    # Straddling a boundary: find the longest prefix that appears, earliest page first.
    for length in range(len(text) - 1, 0, -1):
        prefix = text[:length].strip()
        if not prefix:
            break
        for index, page in enumerate(pages, start=1):
            if index - 1 in contents:
                continue
            if prefix and page.endswith(prefix):
                return index
    return None


def _observed_pages(pages: Sequence[str], headings: Sequence[str]) -> dict[str, int]:
    observed: dict[str, int] = {}
    for heading in headings:
        page = _first_character_page(pages, heading)
        if page is not None:
            observed[heading] = page
    return observed


def _named_pages(pages: Sequence[str], headings: Sequence[str]) -> dict[str, int]:
    """The page number the **document itself prints** for each heading, read back out of the
    converted PDF.

    Read from the PDF rather than returned from the dict the emitter wrote, for both candidates
    that print numbers. That direction is what makes the measurement a measurement: a number
    the emitter intended but the converter dropped, truncated or re-laid-out would otherwise be
    reported as named correctly, and candidate B — whose numbers *come from* a previous
    measurement — would then be checking its own arithmetic rather than the document.

    A contents entry is a declared heading followed by a leader — a run of dots, a tab, or
    plain space — and then digits. Searched **only on the contents pages**
    (:func:`_contents_page_indices`), which is what makes that pattern unambiguous: in body
    prose "a heading followed by a number" is an ordinary sentence, and anchoring the pattern
    to a line boundary instead would not help, because the extractor returns a whole contents
    page as **one line** with every entry joined by its dot leader.
    """
    contents = _contents_page_indices(pages)
    named: dict[str, int] = {}
    for heading in headings:
        pattern = re.compile(re.escape(heading) + _PAGE_NUMBER_SUFFIX)
        for index in sorted(contents):
            match = pattern.search(pages[index])
            if match is not None:
                named[heading] = int(match.group(1))
                break
    return named


# --------------------------------------------------------------------------- #
# The candidates' emission paths
# --------------------------------------------------------------------------- #


def _add_field_paragraph(document: object, instruction: str) -> object:
    """One paragraph holding a Word field with **no cached result** (candidate A).

    `begin` → `instrText` → `end`, with no `separate` and therefore no cached result run
    between them. That shape is deliberate and is what criterion 14.1 asks for: a field
    carrying a cached result would let the PDF display a page number that was computed when
    the `.docx` was written rather than by the converter, which is a stale number that *looks*
    resolved. `w:dirty="true"` marks the field as needing an update, which is the only signal
    a consumer gets that the empty result is intentional.
    """
    paragraph = document.add_paragraph(style=TOC_ENTRY_STYLE)  # type: ignore[attr-defined]
    element = paragraph._p

    def append_run(child: object) -> None:
        run = element.makeelement(qn("w:r"), {})
        run.append(child)
        element.append(run)

    begin = element.makeelement(qn("w:fldChar"), {})
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    append_run(begin)

    instruction_element = element.makeelement(qn("w:instrText"), {})
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = instruction
    append_run(instruction_element)

    end = element.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    append_run(end)
    return paragraph


def _prepend_toc(
    docx_bytes: bytes,
    *,
    field: bool = False,
    entries: Sequence[tuple[str, int | None]] = (),
) -> bytes:
    """`docx_bytes` with a table-of-contents section inserted **before the first block**.

    Either a field (candidate A and C) or literal entry paragraphs (candidate B), never both.
    Built by reopening the rendered bytes rather than by changing `render/docx.py`, because
    task 2.3 is an *evaluation*: two of the three candidates are going to be rejected, and
    putting three emission paths into the production renderer to delete two of them would
    leave the renderer carrying the shape of the experiment. The adopted one moves into
    `render/docx.py` in a later task, alone.

    Paragraphs are appended and then moved to the front, because `python-docx` has no
    insert-at-index: `add_paragraph` only appends. The move preserves their relative order, so
    the section reads label → entries → page break.
    """
    document = open_docx(io.BytesIO(docx_bytes))
    body = document.element.body
    first = next(iter(body.iterchildren()))

    added: list[object] = [document.add_paragraph(TOC_LABEL, style="Title")._p]

    if field:
        added.append(_add_field_paragraph(document, TOC_FIELD_INSTRUCTION)._p)
    for text, page in entries:
        paragraph = document.add_paragraph(style=TOC_ENTRY_STYLE)
        paragraph.add_run(text)
        # The tab is emitted whether or not a number follows it, so the entry occupies the
        # same single line in both of candidate B's passes. That is what "at full size with
        # no numbers" means: pass 1 must lay out to exactly the height pass 2 will, or the
        # numbers pass 1 measured describe a document with a different pagination.
        paragraph.add_run().add_tab()
        if page is not None:
            paragraph.add_run(str(page))
        added.append(paragraph._p)

    page_break = document.add_paragraph(style="Normal")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    added.append(page_break._p)

    for element in added:
        body.remove(element)
        first.addprevious(element)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def soffice_version() -> str:
    """The `soffice --version` string, for the evidence record (criterion 14.1).

    Recorded because pagination is decided by this binary: a verdict measured under one
    LibreOffice is not evidence about another, and the record has to say which one it is.
    """
    try:
        completed = subprocess.run(
            [SOFFICE_BINARY, "--version"],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
    except (OSError, subprocess.SubprocessError) as error:
        return f"unavailable: {type(error).__name__}"
    return (completed.stdout or completed.stderr).strip().splitlines()[0]


def _convert_via_macro(docx_bytes: bytes) -> bytes:
    """Candidate C: convert by driving `updateIndexes()` through `soffice`'s scripting URL.

    Refuses **before** invoking anything when the pre-warmed profile carries no Basic module
    for the macro to live in, because the alternative is to install one — and installing it is
    a write into the profile the image warms at build time, which criterion 14.3 names as a
    rejection in as many words. This function therefore only ever *reads* the profile.
    """
    assert_lang_in_effect()
    profile = profile_path()
    module = profile / MACRO_LIBRARY_RELATIVE_PATH

    # The messages below name the profile-relative path and never the absolute one. This
    # string is committed into `evidence/toc/evaluation.json`, and an absolute path would pin
    # the record to the machine that ran the evaluation rather than to the facility that was
    # missing — which is the durable fact a future reader needs.
    if not module.is_file():
        raise TocApproachUnavailableError(
            f"the LibreOffice profile carries no Basic module at "
            f"{MACRO_LIBRARY_RELATIVE_PATH}, so {MACRO_NAME!r} cannot be addressed. "
            f"Installing it would be a write into the profile the image pre-warms at build "
            f"time, which is the writable macro library criterion 14.3 rejects"
        )
    if MACRO_NAME.rsplit(".", 1)[-1] not in module.read_text(
        encoding="utf-8", errors="replace"
    ):
        raise TocApproachUnavailableError(
            f"the profile's {MACRO_LIBRARY_RELATIVE_PATH} exists but declares no "
            f"{MACRO_NAME!r}: LibreOffice warms an **empty** Standard library, so the macro "
            f"would have to be written into the profile the image builds. That is the "
            f"writable macro library criterion 14.3 rejects, and the alternative — a second "
            f"`soffice` invocation to install it — contends on the one pre-warmed profile, "
            f"which criterion 14.3 also rejects"
        )

    with tempfile.TemporaryDirectory(prefix="rpt-toc-macro-") as scratch:
        workspace = Path(scratch)
        source = workspace / "report.docx"
        source.write_bytes(docx_bytes)
        target = workspace / "report.pdf"
        command = [
            SOFFICE_BINARY,
            "--headless",
            "--norestore",
            f"-env:UserInstallation=file://{profile}",
            f'macro:///{MACRO_NAME}("{source}","{target}")',
        ]
        try:
            subprocess.run(
                command, capture_output=True, timeout=CONVERT_TIMEOUT_S, check=False
            )
        except subprocess.TimeoutExpired as error:
            raise TocApproachUnavailableError(
                f"the macro invocation did not return within {CONVERT_TIMEOUT_S}s; a "
                f"`macro:///` URL keeps the office process alive as a service unless the "
                f"macro itself terminates it, and a conversion that cannot be bounded is "
                f"not one this pipeline can run"
            ) from error
        if not target.is_file() or not target.stat().st_size:
            raise TocApproachUnavailableError(
                "the macro invocation produced no PDF and reported no error, which is the "
                "worst available failure shape: a silent no-op indistinguishable from a "
                "successful conversion"
            )
        return target.read_bytes()


def _base_docx(
    definition: Mapping[str, object], snapshot: Mapping[str, object]
) -> bytes:
    """The document as the production path renders it — no table of contents anywhere."""
    view = build_snapshot_view(dict(snapshot))
    compiled = compile_document(dict(definition), view=view)
    design = DesignSettings.from_plain(definition.get("design"))
    outcome = D.render_document(
        compiled.document, ledger=compiled.ledger, design=design
    )
    return outcome.docx_bytes


def _pages_of(pdf_bytes: bytes) -> tuple[str, ...]:
    with tempfile.TemporaryDirectory() as directory:
        path = Path(directory) / "measured.pdf"
        path.write_bytes(pdf_bytes)
        return pdf_page_texts(path)


async def measure(
    definition: Mapping[str, object],
    snapshot: Mapping[str, object],
    *,
    approach: str,
) -> TocMeasurement:
    """Compile, render, convert and measure, once.

    `approach` selects how the table of contents is produced. It is validated against
    :data:`TOC_APPROACHES` rather than accepted as free text, so a typo in an evaluation run
    fails instead of silently measuring :data:`TOC_APPROACH_NONE` and recording a verdict for
    a candidate that was never exercised.

    `async` because task 2.3's evaluation and task 2.5's two-pass candidate both need to
    serialize conversions — LibreOffice contends on the one pre-warmed profile in the image —
    and an `async` signature is what lets a caller await them in a defined order rather than
    discovering the contention as an intermittent profile-in-use failure.

    Raises :class:`TocApproachUnavailableError` when a candidate cannot be exercised in this
    environment at all, so a verdict is never recorded for a measurement that was fabricated.

    ## What `named_pages` means per candidate, and why B needs no extra field

    * :data:`TOC_APPROACH_NONE` — `{}`. No section is emitted, so the document names nothing.
      Empty rather than a mirror of `observed_pages`; mirroring would make task 2.4's equality
      assertion true by construction for every approach, the one that emits nothing included.
    * :data:`TOC_APPROACH_LIBREOFFICE_INDEX` and :data:`TOC_APPROACH_CONVERSION_MACRO` — the
      numbers **read back out of the converted PDF**, which is what the field resolved to.
    * :data:`TOC_APPROACH_TWO_PASS` — also read back out of the PDF, and that PDF is **pass
      2's**, whose printed numbers are **pass 1's observations**. So `named_pages` is pass 1's
      pagination and `observed_pages` is pass 2's, which makes `named == observed` exactly
      criterion 14.3's fixed-point test — "no heading's observed page differs between its two
      passes" — with no second observation field to keep in step. The two readings of
      `named_pages` coincide here rather than competing.
    """
    if approach not in TOC_APPROACHES:
        raise ValueError(
            f"{approach!r} is not one of the declared TOC approaches "
            f"{list(TOC_APPROACHES)}; a typo must fail rather than silently measure 'none'"
        )

    headings = _heading_texts(definition)
    base = _base_docx(definition, snapshot)

    if approach == TOC_APPROACH_NONE:
        docx_bytes = base
        pdf_bytes = convert_to_pdf(docx_bytes).pdf_bytes
    elif approach == TOC_APPROACH_LIBREOFFICE_INDEX:
        docx_bytes = _prepend_toc(base, field=True)
        pdf_bytes = convert_to_pdf(docx_bytes).pdf_bytes
    elif approach == TOC_APPROACH_CONVERSION_MACRO:
        docx_bytes = _prepend_toc(base, field=True)
        pdf_bytes = _convert_via_macro(docx_bytes)
    elif approach == TOC_APPROACH_TWO_PASS:
        # Pass 1: the section at full size with no numbers, only to be measured. Its bytes
        # are held here and never returned — criterion 14.2 makes pass 2's the artifacts.
        first_docx = _prepend_toc(
            base, entries=tuple((heading, None) for heading in headings)
        )
        first_pages = _pages_of(convert_to_pdf(first_docx).pdf_bytes)
        measured = _observed_pages(first_pages, headings)

        # Pass 2: the same section carrying pass 1's numbers as literal text. Serialized
        # after pass 1 by awaiting nothing in between — `convert_to_pdf` holds the process
        # lock, and the two calls are sequential statements, so they cannot contend.
        docx_bytes = _prepend_toc(
            base,
            entries=tuple((heading, measured.get(heading)) for heading in headings),
        )
        pdf_bytes = convert_to_pdf(docx_bytes).pdf_bytes
    else:  # pragma: no cover - unreachable while TOC_APPROACHES has four members
        # Exhaustive by construction rather than by a trailing `else` that silently absorbs a
        # fifth approach into the two-pass path.
        raise NotImplementedError(
            f"{approach!r} is a declared approach with no emission path here; measuring it "
            f"through another candidate's path would record a verdict for something that "
            f"was never rendered"
        )

    pages = _pages_of(pdf_bytes)

    return TocMeasurement(
        docx_bytes=docx_bytes,
        pdf_bytes=pdf_bytes,
        pdf_sha256=hashlib.sha256(pdf_bytes).hexdigest(),
        observed_pages=_observed_pages(pages, headings),
        named_pages=(
            {} if approach == TOC_APPROACH_NONE else _named_pages(pages, headings)
        ),
        _page_count=len(pages),
    )
