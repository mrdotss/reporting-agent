"""Emit the front matter: cover, document control and table of contents, in that order,
before every content block (Req 13.4, 13.5, 13.6, 13.8, 13.9, 13.15, 13.16, 15.3).

## The three clauses that bind this module

(a) Where the cover-page flag is FALSE, emit no cover content AND NO LEADING BLANK PAGE,
    retain the cover configuration in the definition, and emit the document control page
    and table of contents UNCHANGED — disabling the cover does not disable the front matter.

(b) A role with no supplied signature image gets an EMPTY RULED SIGNATURE BOX at the height
    the theme declares, and emphatically NOT that role's typed name: a typed name in a
    signature position presents an approval nobody gave.

(c) A per-run value that is absent is RENDER_FAILED NAMING THAT VALUE, with no report
    artifact and NO SUBSTITUTED PLACEHOLDER in its position — a cover carrying invented
    copy is a document that cannot be signed.

## `document_number` resolves identically on every render of one run

Two renders of one run resolve one identical number. Two runs of one template and one
resolved period resolve the SAME number (distinguished by the revision history row),
because a re-run of one period is a revision of one document rather than a second document.

## Every fixed string is resolved by id through `messages`

This module contains no literal user-facing copy. Every heading, label, and column header
is a string id resolved through the `Messages` instance the caller supplies, pinned to the
run's language.

## TOC emission

Table-of-contents entries are emitted at levels 1-3 by :func:`_emit_toc` as pass 1 of
the two-pass approach: heading text + tab, no page number. Pass 2 is handled by
``render/toc.py::apply_toc_page_numbers`` which operates on the serialized docx bytes
after conversion to measure actual positions and fill in page numbers.

## Call site in render/docx.py

``render_document`` calls ``emit_front_matter`` after ``_apply_page_size`` and the preview
notice, before the first content block. ``front_matter=None`` (v1 definitions, thumbnails)
skips the call entirely.
"""

from __future__ import annotations

import io
import re
from collections.abc import Sequence
from dataclasses import dataclass, field
from typing import Final

from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

from reporting_agent.compile.definition import (
    APPROVER_ROLES,
)
from reporting_agent.compile.messages import Messages
from reporting_agent.errors import RenderFailedError

__all__ = [
    "front_matter_sections",
    "ContentsEntry",
    "FrontMatterSection",
    "FrontMatterBackground",
    "FrontMatterPairs",
    "FrontMatterPageBreak",
    "FrontMatterLogo",
    "FrontMatterNote",
    "FrontMatterHeading",
    "FrontMatterGrid",
    "FrontMatterContents",
    "APPROVER_HEADER_COMPANY",
    "APPROVER_HEADER_NAME",
    "APPROVER_HEADER_ROLE",
    "APPROVER_HEADER_SIGNATURE",
    "APPROVER_ROLE_LABEL_IDS",
    "DOC_CONTROL_CONFIDENTIALITY",
    "DOC_CONTROL_DISTRIBUTION",
    "DOC_CONTROL_DOCUMENT_NAME",
    "REVISION_ISSUE_DATE_LABEL",
    "REVISION_PAGES_CHANGED_LABEL",
    "REVISION_ISSUE_DATE_LABEL",
    "REVISION_VERSION_LABEL",
    "COVER_VERIFICATION_NOTE",
    "COVER_PERIOD_LABEL",
    "COVER_CUSTOMER_LABEL",
    "DOC_CONTROL_TITLE_ROW",
    "DOC_CONTROL_CUSTOMER_NAME",
    "DOC_CONTROL_DOCUMENT_TITLE",
    "DOC_CONTROL_DOCUMENT_NUMBER",
    "DOC_CONTROL_REVISION_HISTORY",
    "DOC_CONTROL_TITLE",
    "SIGNATURE_BOX_HEIGHT_TWIPS",
    "ApproverEntry",
    "FrontMatterConfig",
    "RunFacts",
    "document_number",
    "emit_front_matter",
]


# ---------------------------------------------------------------------------
# String ids resolved through the message catalog (Req 15.3)
# ---------------------------------------------------------------------------

DOC_CONTROL_TITLE: Final[str] = "doc.front_matter.document_control"
DOC_CONTROL_DOCUMENT_NAME: Final[str] = "doc.front_matter.document_name"
DOC_CONTROL_DOCUMENT_TITLE: Final[str] = "doc.front_matter.document_title"
DOC_CONTROL_CUSTOMER_NAME: Final[str] = "doc.front_matter.customer_name"
DOC_CONTROL_TITLE_ROW: Final[str] = "doc.front_matter.title"
COVER_CUSTOMER_LABEL: Final[str] = "doc.front_matter.customer"
COVER_PERIOD_LABEL: Final[str] = "doc.front_matter.period"
COVER_VERIFICATION_NOTE: Final[str] = "doc.front_matter.verification_note"
REVISION_VERSION_LABEL: Final[str] = "doc.front_matter.revision_version"
REVISION_ISSUE_DATE_LABEL: Final[str] = "doc.front_matter.revision_issue_date"
REVISION_PAGES_CHANGED_LABEL: Final[str] = "doc.front_matter.revision_pages_changed"
DOC_CONTROL_DOCUMENT_NUMBER: Final[str] = "doc.front_matter.document_number"
DOC_CONTROL_CONFIDENTIALITY: Final[str] = "doc.front_matter.confidentiality"
DOC_CONTROL_DISTRIBUTION: Final[str] = "doc.front_matter.distribution"
DISTRIBUTION_HEADER_RECIPIENT: Final[str] = "doc.front_matter.distribution_recipient"
DISTRIBUTION_HEADER_COMPANY: Final[str] = "doc.front_matter.distribution_company"
DISTRIBUTION_HEADER_NOTE: Final[str] = "doc.front_matter.distribution_note"
DOC_CONTROL_REVISION_HISTORY: Final[str] = "doc.front_matter.revision_history"
REVISION_LABEL: Final[str] = "doc.front_matter.revision"
REVISION_AUTHOR_LABEL: Final[str] = "doc.front_matter.revision_author"
REVISION_NOTE_LABEL: Final[str] = "doc.front_matter.revision_note"

APPROVER_HEADER_ROLE: Final[str] = "doc.front_matter.approver_role"
APPROVER_HEADER_COMPANY: Final[str] = "doc.front_matter.approver_company"
APPROVER_HEADER_NAME: Final[str] = "doc.front_matter.approver_name"
APPROVER_HEADER_SIGNATURE: Final[str] = "doc.front_matter.approver_signature"

APPROVER_ROLE_LABEL_IDS: Final[dict[str, str]] = {
    "author": "doc.front_matter.role.author",
    "reviewer": "doc.front_matter.role.reviewer",
    "approver": "doc.front_matter.role.approver",
    "recipient": "doc.front_matter.role.recipient",
}
"""Requirement 12.3 — the four roles are relabelled **positionally**
(`Author` / `Quality Control` / `Reviewed By` / `Customer`), through the
message catalogue rather than by renaming a stored id, a fixture, or a
mirror region. `_emit_approvers_table` resolves the role column through
this map instead of writing the raw role id string into the cell — a role
id is an internal symbol (Requirement 1's own rule for this rename), not
copy a reader should see."""


# ---------------------------------------------------------------------------
# Theme constants
# ---------------------------------------------------------------------------

from reporting_agent.render.themes import SIGNATURE_ROW_HEIGHT_TWIPS

SIGNATURE_BOX_HEIGHT_TWIPS: Final[int] = SIGNATURE_ROW_HEIGHT_TWIPS
"""The height of an empty ruled signature box, read from the theme rather than restated.

The task's wording is "an empty ruled signature box at the height the theme declares", so
the theme is the single declaration and this is an alias for reading it by the name the
front matter thinks in. It was briefly a second literal `907` here; two constants that must
agree, with nothing asserting they do, is the drift this codebase avoids everywhere else.

The box is exactly this tall. A supplied signature image would grow it
(``w:hRule="atLeast"``), but no image is emitted when none is supplied — only an empty cell
at this declared height, because a typed name in a signature position presents an approval
nobody gave."""


# ---------------------------------------------------------------------------
# Placeholder grammar (closed set from task 7.2 — read, not reinvented)
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE: Final[re.Pattern[str]] = re.compile(r"\{[^{}]*\}")


# ---------------------------------------------------------------------------
# Configuration and per-run facts
# ---------------------------------------------------------------------------


@dataclass(frozen=True, slots=True)
class ApproverEntry:
    """One role in the approvers list."""

    role: str
    name: str = ""
    title: str = ""
    company: str = ""
    signature_image: bytes | None = None
    """The decoded signature image, if one was supplied for this role.

    Bytes, not a key: Requirement 13.5's `signature_key` is resolved to bytes
    **server-side in the app** and passed inline in the invoke payload — the
    runtime holds no session and must not fetch content from the app back —
    so by the time a definition reaches this renderer, there is no key left
    to resolve, only the image or its absence. `None` is the ordinary,
    expected value for a role with no supplied signature; it is what
    produces the empty ruled box (Req 13.1), not an error condition."""


@dataclass(frozen=True, slots=True)
class CoverConfig:
    """The cover subsection of `front_matter`."""

    enabled: bool = True
    logo: str | None = None
    logo_key: str | None = None
    """The stored object the logo's bytes live at, resolved from :attr:`logo` by the app.

    `logo` is the URL a profile author typed and is what the wizard shows back to them.
    Fetching it at render time would mean the runtime issuing a request to a
    user-supplied address from inside the VPC — an endpoint chosen by the person whose
    report it is, which is the shape of every SSRF. The app resolves it once when a
    version is saved, validates the bytes as an image the way it already validates an
    uploaded signature, and writes the key here."""

    logo_image: bytes | None = None
    """The bytes at :attr:`logo_key`, where this run could read them."""

    background: str | None = None
    background_key: str | None = None
    """The full-bleed image the cover's text sits on, and the stored object its bytes were
    fetched into.

    A different picture from the logo with a different job — the logo is a mark a few
    centimetres wide, this covers the page — so it carries its own URL, its own key and
    its own byte ceiling. A cover may have either, both or neither."""

    background_image: bytes | None = None
    """The bytes at :attr:`background_key`, where this run could read them."""
    contact_block: str | None = None
    subtitle: str | None = None


@dataclass(frozen=True, slots=True)
class DistributionRow:
    """One row of the v3 distribution list (Req 12.6)."""

    recipient: str = ""
    company: str = ""
    note: str = ""


@dataclass(frozen=True, slots=True)
class DocumentControlConfig:
    """The document_control subsection of `front_matter`.

    `distribution` carries the **v1/v2** free-text block; `distribution_rows` carries the
    **v3** ordered `{recipient, company, note}` rows (Req 12.6). Two fields rather than one
    union because the two render differently — a paragraph and a table — and a renderer
    that had to ask which shape it held on every use would ask in more places than the one
    that matters.

    A definition populates at most one of them. Both empty is the ordinary case for a
    profile that names no distribution at all.
    """

    document_name: str | None = None
    document_number_pattern: str | None = None
    confidentiality_notice_id: str | None = None

    confidentiality_notice: str | None = None
    """The notice itself, as prose, resolved from the Brand when the version was saved.

    Prose and not a message id, because the notice names the consultancy that owns the
    document — "...is owned by PT. Helios Informatika Nusantara..." — and no catalogue
    entry can carry a per-tenant name: a catalogue holds fixed copy every tenant shares.
    Requirement 12.7 makes it Brand-owned and not editable per profile, and the Brand
    glossary entry lists "confidentiality-notice text" among what a Brand carries.

    :attr:`confidentiality_notice_id` stays for the catalogue-sourced notice it was always
    meant to name. Where both are present the prose wins, because it is the one a person
    wrote for this account."""

    distribution: str | None = None
    distribution_rows: tuple[DistributionRow, ...] = ()
    approvers: tuple[ApproverEntry, ...] = ()


@dataclass(frozen=True, slots=True)
class TocConfig:
    """The toc subsection of `front_matter`."""

    enabled: bool = True
    max_level: int = 3


@dataclass(frozen=True, slots=True)
class FrontMatterConfig:
    """The full `front_matter` section, resolved from a validated definition.

    This is a data container only — no I/O, no network. The definition validator has
    already constrained every field.
    """

    cover: CoverConfig = field(default_factory=CoverConfig)
    document_control: DocumentControlConfig = field(
        default_factory=DocumentControlConfig
    )
    toc: TocConfig = field(default_factory=TocConfig)


@dataclass(frozen=True, slots=True)
class RevisionHistoryRow:
    """One row of the revision history table, supplied per-run at enqueue (Req 13.7)."""

    revision: str
    note: str = ""
    author: str = ""


@dataclass(frozen=True, slots=True)
class RunFacts:
    """Per-run values that the front matter needs (Req 13.7).

    Template-once values live in `FrontMatterConfig`; per-run values live here. This
    separation is the schema's: template config is stored on the definition and per-run
    values arrive at enqueue.
    """

    run_id: str
    template_id: str
    customer_name: str
    period_display: str
    report_title: str
    revision_history: RevisionHistoryRow | None = None
    period_start_year: str = ""
    period_start_month: str = ""

    issued_on: str = ""
    """The date this report was published, already formatted for the document.

    Read from the **snapshot's** `collected_at`, not from a clock. A clock would print a
    different date every time the document was re-rendered, and `verify/allowlist.py`
    derives the document's numeric chrome from a fresh null-context render — so a
    re-verification would compare a stored document carrying one date against an
    allowlist carrying another, and the original's digits would survive masking as
    `unmatched_prose_token`. The snapshot is stored, replayed, and says exactly when the
    data behind this document was collected."""


# ---------------------------------------------------------------------------
# `document_number` — the closed placeholder grammar (Req 13.8, 13.16)
# ---------------------------------------------------------------------------



# ---------------------------------------------------------------------------
# The neutral description — what the front matter *is*, before either emitter
# ---------------------------------------------------------------------------


@dataclass(frozen=True, slots=True)
class FrontMatterHeading:
    """A line of its own, in a named paragraph style."""

    text: str
    style: str


@dataclass(frozen=True, slots=True)
class FrontMatterPairs:
    """A label/value block — the cover's facts, the naming block, the revision row."""

    rows: tuple[tuple[str, str], ...]
    table_style: str
    paragraph_style: str


@dataclass(frozen=True, slots=True)
class FrontMatterGrid:
    """A headed table — the approvers, the distribution list.

    `paragraph_style` is `None` for a table whose cells take the table style's own
    formatting rather than a paragraph style, which is what the approvers table does and
    the distribution table does not. Carried rather than inferred, because the difference
    is visible in the delivered document and neither emitter should have to guess it.

    `signature_column` names a column whose cells are left **empty** whatever the row
    holds — Req 13.6's ruled box, which must never be filled with the approver's typed
    name — and `images` supplies the bytes to place there where a signature was given.
    """

    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]
    table_style: str
    paragraph_style: str | None = None
    signature_column: int | None = None
    images: tuple[bytes | None, ...] = ()
    row_height_twips: int | None = None


@dataclass(frozen=True, slots=True)
class FrontMatterNote:
    """A paragraph of prose — the free-text distribution block, the confidentiality
    notice."""

    text: str
    style: str


@dataclass(frozen=True, slots=True)
class FrontMatterBackground:
    """The cover's full-bleed image, behind everything else on the page.

    Emitted as its own section kind rather than as a field on the logo, because the two
    are independent: a cover may carry either, both or neither, and a background with no
    logo must still reach both emitters. `test_front_matter_sections.py` asserts every
    declared kind is produced and that both emitters render it, which is what keeps a kind
    added for one from going missing from the other.
    """

    image: bytes
    """The bytes themselves. A background with none is not emitted at all — there is no
    space to reserve for a picture that covers the page."""


@dataclass(frozen=True, slots=True)
class FrontMatterLogo:
    """Reserved space at the top of the cover, where a logo goes.

    **Space, not a picture.** `front_matter.cover.logo` holds a key — a URL on the
    profiles written so far — and nothing in this pipeline fetches a remote image at
    render time. Drawing one would mean a network call from inside the VPC that no other
    stage makes and that fails closed on a private subnet, and a cover that sometimes has
    a logo and sometimes does not is worse than one that reliably has room for it.

    So this reserves the block the artifact reserves — `ReportA.dc.html` opens its cover
    with a 104x34 box above the eyebrow — and stays empty. When the logo arrives as bytes,
    the way an approver's `signature_image` already does, it fills this box and nothing
    else about the cover moves.

    Emitted only where a logo is configured: a profile that names none gets no gap where
    one would have been.
    """

    height_pt: float
    width_pt: float
    space_after_pt: float = 0.0
    """The gap beneath it. See :data:`LOGO_SPACE_AFTER_PT`."""

    space_before_pt: float = 0.0
    """The gap above it. See :data:`COVER_SPACE_BEFORE_PT`."""

    image: bytes | None = None
    """The logo itself, where this run could read it.

    `None` reserves the space and draws nothing — a profile that names a logo the run
    could not fetch lays out exactly as it will once the fetch succeeds, rather than
    shifting every element on the cover the first time it works."""


@dataclass(frozen=True, slots=True)
class ContentsEntry:
    """One line of the contents: a heading's text, its level, its section number and the
    id of the heading it points at.

    `number` and `anchor` are printed by the styled reading copy and by nothing else —
    see `render/toc.py::section_numbers` for why a section number in the `.docx` would
    fail the prose gate on a correct report."""

    text: str
    level: int
    number: str
    anchor: str


@dataclass(frozen=True, slots=True)
class FrontMatterContents:
    """The table of contents: a heading and one entry per document heading.

    Carries **no page numbers**. The `.docx` gets them from
    `render/toc.py::apply_toc_page_numbers`, which measures the converted PDF; a CSS
    print stylesheet generates them at pagination time. Neither is something this
    description can know, and a guessed page number is a promise the document breaks.
    """

    label: str
    entries: tuple[ContentsEntry, ...]
    entry_style: str
    label_style: str = "Title"
    """`Title` so the contents heading does not appear in its own contents."""


@dataclass(frozen=True, slots=True)
class FrontMatterPageBreak:
    """Where the front matter starts a new page."""


FrontMatterSection = (
    FrontMatterHeading
    | FrontMatterPairs
    | FrontMatterGrid
    | FrontMatterNote
    | FrontMatterLogo
    | FrontMatterBackground
    | FrontMatterContents
    | FrontMatterPageBreak
)
"""One piece of front matter, described without saying how it is drawn.

## Why this exists

`emit_front_matter` wrote straight into a `DocxDocument`, so the cover, the document
control page and the contents existed **only** as Word calls. The HTML emitter walks the
block AST, and the front matter is not in the AST — it is fixed rather than composed and
accepts no block — so `document.html` began at the first section heading, with no cover,
no approvers and no contents at all.

Giving the second emitter its own front-matter renderer would put the *order and content*
of the front matter in two places, which is exactly what Req 24.1 forbids for blocks and
for the same reason: two statements of one layout is one layout and one latent bug. So the
order and content live here, once, and each emitter chooses elements for what it is given.

The style names ride along rather than being resolved per emitter. That follows the body:
`Table.style` on the AST is a Word style name and `render/html.py` emits it as
`data-style`, so a stylesheet can key off the same vocabulary the theme uses.
"""


LOGO_HEIGHT_PT: Final[float] = 34.0
LOGO_WIDTH_PT: Final[float] = 104.0

COVER_SPACE_BEFORE_PT: Final[float] = 150.0
"""The space above the cover's first element.

`ReportA.dc.html` centres the whole cover block in the page rather than starting it at the
top margin — the logo sits about a third of the way down, with the title and the fact list
below it and the confidentiality line at the foot. Emitted as space **before** the first
element rather than as a vertical alignment property, because that is the one expression
both emitters have: Word's `w:vAlign` centres a whole section and would take the footer
line with it, and the reading copy's page box has no height for `justify-content` to
distribute until it is paginated.

150pt is a little under a fifth of A4's 842pt, which puts a cover carrying a logo, an
eyebrow, a title, a period and four fact rows on the page's optical centre rather than its
arithmetic one — the artifact's own arrangement, where the block sits slightly high."""

LOGO_SPACE_AFTER_PT: Final[float] = 44.0
"""The gap between the logo and the eyebrow beneath it.

`ReportA.dc.html` sets 56px under a 34px logo on a 640px-wide page standing in for A4 —
a gap two-thirds again the logo's own height, which is what stops the mark reading as a
heading for the title. Emitted as paragraph space rather than as an empty paragraph so
the two emitters can express it identically and neither has a blank line to lose."""
"""The block `ReportA.dc.html` reserves for the logo, in points.

Its own 104x34 CSS pixels read as points here because the artifact is drawn at 96dpi on an
A4 page and this stylesheet lays one out at 72 — close enough that the cover's proportions
carry, and a reserved block is not a measurement anybody checks."""


RUN_PLACEHOLDER_LENGTH: Final[int] = 8
"""How much of the run id `{run}` contributes to a document number.

A run id is a 36-character UUID, and substituting it whole produced
`FRM.SOP.019-2026-08.91286d3a-bb51-4c1c-88af-af9069dcf321` on the cover of a signed
document — a number nobody can read out, quote in an email, or write on a printout, which
is the entire job of a document number. The artifact this front matter is styled from
prints `FRM.SOP.0XX-26-816.01`.

Eight hex characters, because `{run}` is what distinguishes two runs of one template and
one period (see below) and it has to keep doing that: four billion values within a single
template-and-period pair is not a collision anyone will meet. It is a **prefix**, not a
hash, so the number still reads back to its run by inspection.
"""


def document_number(pattern: str, *, run: RunFacts) -> str:
    """Apply the document-number pattern to one run.

    The placeholder grammar is the closed set declared in `compile/definition.py` (task 7.2):
    `{template}`, `{year}`, `{month}`, `{run}`.

    Two renders of one run resolve one identical number. Two runs of one template and one
    resolved period resolve the SAME number — only `{run}` distinguishes them, and the
    definition validator already rejects a pattern naming no varying placeholder — so a
    re-run of one period is a revision of one document rather than a second document
    (Req 13.16).

    A placeholder whose substitution value is absent raises `RENDER_FAILED` naming the
    placeholder (clause c).
    """
    substitutions: dict[str, str] = {
        "{template}": run.template_id,
        "{year}": run.period_start_year,
        "{month}": run.period_start_month,
        "{run}": run.run_id[:RUN_PLACEHOLDER_LENGTH],
    }

    # Validate that every required substitution value is non-empty.
    for placeholder, value in substitutions.items():
        if placeholder in pattern and not value:
            raise RenderFailedError(
                f"the document-number pattern names {placeholder} but the per-run value "
                f"for it is absent; a cover carrying invented copy is a document that "
                f"cannot be signed (Req 13.15)"
            )

    # Apply all placeholders from the closed set.
    result = pattern
    for placeholder, value in substitutions.items():
        result = result.replace(placeholder, value)

    return result


# ---------------------------------------------------------------------------
# emit_front_matter — the main entry point
# ---------------------------------------------------------------------------



def front_matter_sections(
    *,
    front_matter: FrontMatterConfig,
    run: RunFacts,
    messages: Messages,
    heading_entries: tuple[tuple[str, int], ...] = (),
    include_toc: bool = True,
) -> tuple[FrontMatterSection, ...]:
    """The front matter, in order, described without saying how it is drawn.

    Cover, then document control, then the table of contents (Req 13.4). Not composable,
    not reorderable, no block accepted inside it.

    Where the cover-page flag is false, no cover content and no leading blank page — and
    the document control page and the contents unchanged, because disabling the cover does
    not disable the front matter (Req 13.9).

    Pure over its arguments, so both emitters see the same front matter and a test can read
    it without rendering anything.
    """
    from reporting_agent.compile.blocks.base import LAYOUT_TABLE_STYLE
    from reporting_agent.render.themes import (
        COVER_META_STYLE,
        COVER_TITLE_STYLE,
        DOCUMENT_CONTROL_STYLE,
        SIGNATURE_TABLE_STYLE,
        TOC_ENTRY_STYLE,
    )
    from reporting_agent.render.toc import TOC_LABEL_ID

    _require_run_value(run.customer_name, "customer_name")
    _require_run_value(run.period_display, "period_display")
    _require_run_value(run.report_title, "report_title")
    _require_run_value(run.run_id, "run_id")

    pattern = front_matter.document_control.document_number_pattern
    doc_number = document_number(pattern, run=run) if pattern else None
    control = front_matter.document_control

    sections: list[FrontMatterSection] = []

    # --- cover (Req 13.4, 13.9) ----------------------------------------------
    if front_matter.cover.enabled:
        # Behind everything, first, so both emitters meet it before the content it sits
        # under — the docx anchors it to the first paragraph of the page, and the print
        # stylesheet paints it as the first page's own background.
        if front_matter.cover.background_image:
            sections.append(
                FrontMatterBackground(image=front_matter.cover.background_image)
            )
        # The logo's space, above everything, as `ReportA.dc.html` opens its cover.
        if front_matter.cover.logo:
            sections.append(
                FrontMatterLogo(
                    height_pt=LOGO_HEIGHT_PT,
                    width_pt=LOGO_WIDTH_PT,
                    space_before_pt=COVER_SPACE_BEFORE_PT,
                    space_after_pt=LOGO_SPACE_AFTER_PT,
                    image=front_matter.cover.logo_image,
                )
            )
        # The eyebrow, above the title: the document's own name in small muted caps, the
        # way `ReportA.dc.html` opens its cover with "Preventive Maintenance Report" over
        # "Marketing Riset Azure Usage Report". The two are different things — one names
        # the kind of document, the other names this one — and the cover said only the
        # second, so a reader could not tell a maintenance report from a usage report
        # without reading the contents.
        if control.document_name:
            sections.append(
                FrontMatterHeading(control.document_name, COVER_META_STYLE)
            )
        sections.append(FrontMatterHeading(run.report_title, COVER_TITLE_STYLE))
        # The period, large, under the title — `ReportA`'s "July 2026". It also appears
        # in the pairs below as a precise range with its timezone; this is the same fact
        # at reading distance, which is what a cover is for.
        sections.append(FrontMatterHeading(run.period_display, COVER_META_STYLE))
        if front_matter.cover.subtitle:
            sections.append(FrontMatterHeading(front_matter.cover.subtitle, COVER_META_STYLE))

        # No customer / period / document-number table here any more.
        #
        # It restated, in a bordered grid a third of the way down the page, three facts the
        # cover had already said above it — the customer is in the title, the period is set
        # large under it — and the third, the document number, belongs to the document
        # control page, which states it and is where a reader looks for it. Req 13.8 asks
        # that the number resolve to the same string wherever it appears; it does not ask
        # that it appear twice.
        #
        # `contact_block` goes with it. A prepared-by line is a document-control fact, and
        # it is on that page already.
        if front_matter.cover.contact_block:
            cover_rows = [
                (
                    messages.text("doc.front_matter.prepared_by"),
                    front_matter.cover.contact_block,
                )
            ]
            sections.append(
                FrontMatterPairs(tuple(cover_rows), LAYOUT_TABLE_STYLE, COVER_META_STYLE)
            )
        sections.append(FrontMatterPageBreak())

    # --- document control (Req 13.5, 13.6) -----------------------------------
    sections.append(
        FrontMatterHeading(messages.text(DOC_CONTROL_TITLE), DOCUMENT_CONTROL_STYLE)
    )

    # `ReportA.dc.html`'s own four rows, under their own subheading: who it is for, what
    # it is called, what kind of document it is, and its number. The previous three
    # (document name, number, prepared for) named the document twice and the reader once.
    sections.append(
        FrontMatterHeading(
            messages.text(DOC_CONTROL_DOCUMENT_TITLE), DOCUMENT_CONTROL_STYLE
        )
    )
    naming: list[tuple[str, str]] = [
        (messages.text(DOC_CONTROL_CUSTOMER_NAME), run.customer_name),
        (messages.text(DOC_CONTROL_TITLE_ROW), run.report_title),
    ]
    if control.document_name:
        naming.append((messages.text(DOC_CONTROL_DOCUMENT_NAME), control.document_name))
    if doc_number:
        naming.append((messages.text(DOC_CONTROL_DOCUMENT_NUMBER), doc_number))
    sections.append(
        FrontMatterPairs(tuple(naming), LAYOUT_TABLE_STYLE, DOCUMENT_CONTROL_STYLE)
    )

    # --- approvers (Req 13.6) ------------------------------------------------
    approver_rows: list[tuple[str, ...]] = []
    approver_images: list[bytes | None] = []
    for role in APPROVER_ROLES:
        entry = _find_approver(control.approvers, role)
        role_label_id = APPROVER_ROLE_LABEL_IDS.get(role)
        approver_rows.append(
            (
                messages.text(role_label_id) if role_label_id else role,
                # The COMPANY column shows the approver's company; `title` remains the
                # fallback so a profile that put its company there keeps rendering as it did.
                (entry.company or entry.title) if entry else "",
                entry.name if entry else "",
                "",
            )
        )
        approver_images.append(entry.signature_image if entry is not None else None)

    sections.append(
        FrontMatterGrid(
            headers=(
                # Blank, as `ReportA.dc.html` leaves it. The column holds Author, Quality
                # Control, Reviewed By and Customer — a reader does not need to be told
                # those are roles, and "Role" over them is a caption for the obvious that
                # the artifact spends no width on.
                "",
                messages.text(APPROVER_HEADER_COMPANY),
                messages.text(APPROVER_HEADER_NAME),
                messages.text(APPROVER_HEADER_SIGNATURE),
            ),
            rows=tuple(approver_rows),
            table_style=SIGNATURE_TABLE_STYLE,
            signature_column=3,
            images=tuple(approver_images),
            row_height_twips=SIGNATURE_BOX_HEIGHT_TWIPS,
        )
    )

    # --- revision history ----------------------------------------------------
    if run.revision_history:
        row = run.revision_history
        sections.append(
            FrontMatterHeading(
                messages.text(DOC_CONTROL_REVISION_HISTORY), DOCUMENT_CONTROL_STYLE
            )
        )
        # A table with a header row, as `ReportA.dc.html` draws it — `Version | Issue
        # Date | Pages Changed | Notes`, one row per revision. It was a label/value pairs
        # block, which reads as a property of the document rather than as its history:
        # a revision history with one row still has to look like a list somebody will add
        # a second row to.
        #
        # `Published date` carries the run's own publication date and `Pages Changed` an
        # em dash. Neither is invented: the date is the snapshot's `collected_at` — see
        # `RunFacts.issued_on`, which explains why it is not a clock — and nothing in the
        # run record knows which pages a re-run changed, so an em dash fills that column
        # exactly as the artifact does for a first issue.
        #
        # The author is deliberately **not** a column here. It was, carried under the
        # `Issue Date` header, which is a name in a column asking for a date; the person
        # is named in the approvers table above, where the document already identifies who
        # wrote it.
        sections.append(
            FrontMatterGrid(
                headers=(
                    messages.text(REVISION_VERSION_LABEL),
                    messages.text(REVISION_ISSUE_DATE_LABEL),
                    messages.text(REVISION_PAGES_CHANGED_LABEL),
                    messages.text(REVISION_NOTE_LABEL),
                ),
                rows=((row.revision, run.issued_on or "\u2014", "\u2014", row.note or ""),),
                table_style=LAYOUT_TABLE_STYLE,
                paragraph_style=DOCUMENT_CONTROL_STYLE,
            )
        )

    # --- distribution --------------------------------------------------------
    if control.distribution_rows or control.distribution:
        sections.append(
            FrontMatterHeading(messages.text(DOC_CONTROL_DISTRIBUTION), DOCUMENT_CONTROL_STYLE)
        )
        if control.distribution_rows:
            sections.append(
                FrontMatterGrid(
                    headers=(
                        messages.text(DISTRIBUTION_HEADER_RECIPIENT),
                        messages.text(DISTRIBUTION_HEADER_COMPANY),
                        messages.text(DISTRIBUTION_HEADER_NOTE),
                    ),
                    rows=tuple(
                        (row.recipient, row.company, row.note)
                        for row in control.distribution_rows
                    ),
                    table_style=LAYOUT_TABLE_STYLE,
                    paragraph_style=DOCUMENT_CONTROL_STYLE,
                )
            )
        else:
            # v1/v2 — one free-text block, unchanged.
            sections.append(
                FrontMatterNote(str(control.distribution), DOCUMENT_CONTROL_STYLE)
            )

    # --- confidentiality notice ----------------------------------------------
    # Under its own heading, as `ReportA.dc.html` prints it: the notice is a distinct
    # section of the document control page, not a trailing paragraph of the distribution
    # list. Nothing is emitted at all when no notice is configured — an empty heading
    # would be a promise of text that is not there.
    notice = _confidentiality_notice(control, messages)
    if notice is not None:
        sections.append(
            FrontMatterHeading(
                messages.text(DOC_CONTROL_CONFIDENTIALITY), DOCUMENT_CONTROL_STYLE
            )
        )
        sections.append(FrontMatterNote(notice, DOCUMENT_CONTROL_STYLE))

    sections.append(FrontMatterPageBreak())

    # --- table of contents (Req 14.3, 14.5, 14.11) ---------------------------
    if include_toc and front_matter.toc.enabled:
        sections.append(
            FrontMatterContents(
                label=messages.text(TOC_LABEL_ID),
                entries=_contents_entries(heading_entries),
                entry_style=TOC_ENTRY_STYLE,
            )
        )
        sections.append(FrontMatterPageBreak())

    return tuple(sections)


def emit_front_matter(
    document: DocxDocument,
    *,
    front_matter: FrontMatterConfig,
    run: RunFacts,
    messages: Messages,
    cursor: object = None,
    ledger: object = None,
) -> None:
    """Render the front matter into a Word document.

    **What** the front matter is and **in what order** is `front_matter_sections`; this
    chooses Word elements for it. The split is the same one the body already has between
    the block compiler and `render/docx.py`, and it exists for the same reason: the styled
    PDF renders the same sections through `render/html.py`, and two statements of one
    layout is one layout and one latent bug.

    A per-run value that is absent is `RENDER_FAILED` naming that value, with no report
    artifact and no substituted placeholder in its position (Req 13.15) — raised by the
    builder, before anything is written.
    """
    from reporting_agent.render.toc import should_emit_toc, toc_entries_from_document

    sections = front_matter_sections(
        front_matter=front_matter,
        run=run,
        messages=messages,
        heading_entries=toc_entries_from_document(cursor) if cursor is not None else (),
        include_toc=should_emit_toc(),
    )
    for section in sections:
        _emit_section(document, section)


def _emit_section(document: DocxDocument, section: FrontMatterSection) -> None:
    """One described section, as Word elements."""
    if isinstance(section, FrontMatterHeading):
        document.add_paragraph(section.text, style=section.style)

    elif isinstance(section, FrontMatterPairs):
        _emit_label_value_table(
            document,
            section.rows,
            style_name=section.table_style,
            paragraph_style=section.paragraph_style,
        )

    elif isinstance(section, FrontMatterGrid):
        _emit_grid(document, section)

    elif isinstance(section, FrontMatterNote):
        document.add_paragraph(section.text, style=section.style)

    elif isinstance(section, FrontMatterBackground):
        _place_cover_background(document, section.image)

    elif isinstance(section, FrontMatterLogo):
        # The logo, or the space it will occupy. No border and no placeholder text: an
        # empty dashed box labelled LOGO is a mock-up's device, and this is a document
        # somebody signs.
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(section.space_before_pt)
        paragraph.paragraph_format.space_after = Pt(section.space_after_pt)
        if section.image is None:
            paragraph.paragraph_format.line_spacing = Pt(section.height_pt)
        else:
            # Height alone, so the image keeps its own aspect ratio inside the block the
            # cover reserved for it — constraining both would stretch somebody's logo.
            paragraph.add_run().add_picture(
                io.BytesIO(section.image), height=Pt(section.height_pt)
            )

    elif isinstance(section, FrontMatterContents):
        # Text plus a tab, whether or not a number follows, so pass 1 lays out to exactly
        # the height pass 2 will — `render/toc.py::apply_toc_page_numbers` measures the
        # converted PDF and re-emits with the numbers it found.
        document.add_paragraph(section.label, style=section.label_style)
        for entry in section.entries:
            paragraph = document.add_paragraph(style=section.entry_style)
            # The text alone. The section number is deliberately not printed here: see
            # `render/toc.py::section_numbers`.
            paragraph.add_run(entry.text)
            paragraph.add_run().add_tab()

    elif isinstance(section, FrontMatterPageBreak):
        _add_page_break(document)

    else:  # pragma: no cover - the union is closed and every member is handled above
        raise RenderFailedError(f"unhandled front matter section {type(section).__name__}")


def _emit_grid(document: DocxDocument, section: FrontMatterGrid) -> None:
    """A headed table, with the signature column's rules where it has one.

    Req 13.6 clause (b): the signature cell is emitted **empty unconditionally** and never
    the approver's typed name — set before any image placement, so an exception raised while
    placing an image cannot leave a typed name behind it.
    """
    table = document.add_table(rows=1 + len(section.rows), cols=len(section.headers))
    table.style = section.table_style

    header_cells = table.rows[0].cells
    for index, header in enumerate(section.headers):
        if section.paragraph_style is None:
            header_cells[index].text = header
        else:
            _set_cell_text(header_cells[index], header, section.paragraph_style)

    for row_index, row in enumerate(section.rows):
        cells = table.rows[row_index + 1].cells
        for index, value in enumerate(row):
            text = "" if index == section.signature_column else value
            if section.paragraph_style is None:
                cells[index].text = text
            else:
                _set_cell_text(cells[index], text, section.paragraph_style)

        # Clause (c): a supplied signature is scaled to fit the theme's declared row
        # height without changing it, so a signed row and an unsigned row occupy the same
        # space and pagination never depends on who signed.
        if section.signature_column is not None and row_index < len(section.images):
            image = section.images[row_index]
            if image is not None:
                _place_signature_image(cells[section.signature_column], image)

        if section.row_height_twips is not None:
            # After either path, so `w:hRule="atLeast"` plus a scaled-to-fit image is what
            # makes signed and unsigned rows the same size, rather than the image growing
            # the row past it.
            _set_row_height(table.rows[row_index + 1], section.row_height_twips)


def _require_run_value(value: str, name: str) -> None:
    """Clause (c): a per-run value that is absent is RENDER_FAILED naming it."""
    if not value or not value.strip():
        raise RenderFailedError(
            f"the per-run value {name!r} is absent; a cover carrying invented copy is "
            f"a document that cannot be signed. No report artifact is produced and no "
            f"substituted placeholder is emitted in its position (Req 13.15)."
        )


def _emit_cover(
    document: DocxDocument,
    *,
    front_matter: FrontMatterConfig,
    run: RunFacts,
    messages: Messages,
    doc_number: str | None,
) -> None:
    """Emit the cover page: report title, then a labelled block of the run's facts.

    Uses theme styles ``Cover Title`` and ``Cover Meta``, and ``Layout Table`` for the
    facts block. Ends with a page break so document control starts on a fresh page.

    ## Why the facts are a table and not a run of paragraphs

    They were five bare `Cover Meta` paragraphs — customer, period, document number — one
    under the next with nothing saying which was which. A reader who does not already know
    the document cannot tell the customer's name from the report's subtitle, because
    nothing distinguishes them.

    A two-column layout table gives each value its label. ``Layout Table`` is the theme's
    own borderless style, already used by `render/anchors.py` for exactly this — structure
    with no rules drawn — so the cover gains alignment without gaining a grid.
    """
    from reporting_agent.compile.blocks.base import LAYOUT_TABLE_STYLE
    from reporting_agent.render.themes import COVER_META_STYLE, COVER_TITLE_STYLE

    document.add_paragraph(run.report_title, style=COVER_TITLE_STYLE)

    # The subtitle sits directly under the title, where it reads as part of it, rather
    # than as another labelled fact.
    if front_matter.cover.subtitle:
        document.add_paragraph(front_matter.cover.subtitle, style=COVER_META_STYLE)

    # Req 13.8 — the document number is the same string here and on the document
    # control page; both read it from `doc_number`.
    rows: list[tuple[str, str]] = [
        (messages.text("doc.front_matter.prepared_for"), run.customer_name),
        (messages.text("doc.front_matter.reporting_period"), run.period_display),
    ]
    if doc_number:
        rows.append((messages.text(DOC_CONTROL_DOCUMENT_NUMBER), doc_number))
    if front_matter.cover.contact_block:
        rows.append(
            (
                messages.text("doc.front_matter.prepared_by"),
                front_matter.cover.contact_block,
            )
        )

    _emit_label_value_table(
        document, rows, style_name=LAYOUT_TABLE_STYLE, paragraph_style=COVER_META_STYLE
    )

    _add_page_break(document)


def _emit_label_value_table(
    document: DocxDocument,
    rows: Sequence[tuple[str, str]],
    *,
    style_name: str,
    paragraph_style: str,
) -> None:
    """A two-column label/value block, in the given table and paragraph styles.

    Emits nothing for an empty `rows` — a table with no rows is a table a reader has to
    wonder about. Every cell's paragraph carries `paragraph_style` rather than the
    document default, so the block inherits the theme's front-matter type rather than the
    body face.
    """
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.style = style_name
    for index, (label, value) in enumerate(rows):
        cells = table.rows[index].cells
        _set_cell_text(cells[0], label, paragraph_style)
        _set_cell_text(cells[1], value, paragraph_style)


def _set_cell_text(cell: object, text: str, paragraph_style: str) -> None:
    """Write `text` into `cell`'s first paragraph in `paragraph_style`.

    `python-docx` gives a new cell exactly one empty paragraph; this writes into it rather
    than adding a second, so a cell never carries a blank line above its content.
    """
    paragraph = cell.paragraphs[0]  # type: ignore[attr-defined]
    paragraph.style = paragraph_style
    paragraph.add_run(text)


def _emit_document_control(
    document: DocxDocument,
    *,
    front_matter: FrontMatterConfig,
    run: RunFacts,
    messages: Messages,
    doc_number: str | None,
) -> None:
    """Emit the document control page (Req 13.5, 13.6).

    Carries: the page title, a labelled block naming the document, the approvers table,
    the revision history, the distribution list and the confidentiality notice.

    The approvers table emits one row per role. Where no signature image is supplied,
    emit an EMPTY RULED SIGNATURE BOX — NOT the typed name (clause b).

    Every fixed string is resolved by id through ``messages`` (Req 15.3).

    ## Labelled blocks, not "Label: value" sentences

    The naming block and the revision history were emitted as `Document Control`
    paragraphs holding `f"{label}: {value}"`. That put the page's whole structure inside
    string concatenation: nothing aligned, every label re-read as prose, and a document
    control page indistinguishable from a paragraph of running text. They are tables now,
    in the theme's own borderless ``Layout Table``.
    """
    from reporting_agent.compile.blocks.base import LAYOUT_TABLE_STYLE
    from reporting_agent.render.themes import DOCUMENT_CONTROL_STYLE

    document.add_paragraph(messages.text(DOC_CONTROL_TITLE), style=DOCUMENT_CONTROL_STYLE)

    # --- what this document is ------------------------------------------------
    naming: list[tuple[str, str]] = []
    if front_matter.document_control.document_name:
        naming.append(
            (
                messages.text(DOC_CONTROL_DOCUMENT_NAME),
                front_matter.document_control.document_name,
            )
        )
    # Req 13.8 — identical to the cover's.
    if doc_number:
        naming.append((messages.text(DOC_CONTROL_DOCUMENT_NUMBER), doc_number))
    naming.append(
        (messages.text("doc.front_matter.prepared_for"), run.customer_name)
    )
    _emit_label_value_table(
        document,
        naming,
        style_name=LAYOUT_TABLE_STYLE,
        paragraph_style=DOCUMENT_CONTROL_STYLE,
    )

    # --- approvers (Req 13.6) -------------------------------------------------
    _emit_approvers_table(document, front_matter=front_matter, messages=messages)

    # --- revision history -----------------------------------------------------
    if run.revision_history:
        document.add_paragraph(
            messages.text(DOC_CONTROL_REVISION_HISTORY), style=DOCUMENT_CONTROL_STYLE
        )
        row = run.revision_history
        # One labelled row per field the run actually carried, in the same label/value
        # shape as the naming block above it.
        #
        # Previously one row of `(revision, note + " (author)")`, which put a bare "11" in
        # the label column — a number in the position every other row of this block uses
        # for the name of the thing — and, on a run whose note is empty, a value cell
        # reading " (Mayer Reflino Sitorus)": a parenthetical qualifying nothing. Since
        # schema_version 3 the note is *always* empty, because the revision is derived
        # from the count of prior runs for the period rather than typed, so that was the
        # shape every current run rendered.
        #
        # Concatenation is what made the empty note visible, so there is none: a field the
        # run did not carry contributes no row instead of a fragment of one. A pinned v1
        # or v2 run whose note was typed by hand still renders it, which is what keeps
        # those replays reading as they did.
        revision_rows = [(messages.text(REVISION_LABEL), row.revision)]
        if row.author:
            revision_rows.append((messages.text(REVISION_AUTHOR_LABEL), row.author))
        if row.note:
            revision_rows.append((messages.text(REVISION_NOTE_LABEL), row.note))
        _emit_label_value_table(
            document,
            revision_rows,
            style_name=LAYOUT_TABLE_STYLE,
            paragraph_style=DOCUMENT_CONTROL_STYLE,
        )

    # --- distribution ---------------------------------------------------------
    _emit_distribution(
        document, front_matter=front_matter, messages=messages,
        paragraph_style=DOCUMENT_CONTROL_STYLE,
    )

    # --- confidentiality notice ----------------------------------------------
    notice = _confidentiality_notice(front_matter.document_control, messages)
    if notice is not None:
        # A paragraph in the document-control style, exactly as `_emit_distribution`
        # emits its own heading: the two are the same kind of thing on the same page, and
        # a second heading mechanism here would be a second thing to style.
        document.add_paragraph(
            messages.text(DOC_CONTROL_CONFIDENTIALITY), style=DOCUMENT_CONTROL_STYLE
        )
        document.add_paragraph(notice, style=DOCUMENT_CONTROL_STYLE)

    _add_page_break(document)




def _contents_entries(
    heading_entries: Sequence[tuple[str, int]],
) -> tuple[ContentsEntry, ...]:
    """Number and anchor a heading sequence, in the order the document presents it."""
    from reporting_agent.render.toc import heading_anchor, section_numbers

    numbers = section_numbers([level for _text, level in heading_entries])
    return tuple(
        ContentsEntry(
            text=text,
            level=level,
            number=number,
            anchor=heading_anchor(ordinal),
        )
        for ordinal, ((text, level), number) in enumerate(
            zip(heading_entries, numbers, strict=True), start=1
        )
    )



def _place_cover_background(document: DocxDocument, image: bytes) -> None:
    """Anchor `image` behind the cover, filling the page.

    python-docx emits a picture as `<wp:inline>`, which flows in the text and pushes
    everything below it down the page — the opposite of a background. Word's own model for
    "behind the text" is `<wp:anchor behindDoc="1">` positioned relative to the **page**,
    which python-docx does not expose, so the inline element it builds is rewritten into
    an anchor here. The graphic itself, and the relationship to the image part, are
    python-docx's; only the wrapper changes.

    Sized to the section's full page rather than to its text width, because a background
    that stops at the margins is a picture with a white frame. Anchored to a paragraph in
    the body rather than to a header: a header picture repeats on every page of the
    section, and this belongs to the cover alone.
    """
    section = document.sections[0]
    width = section.page_width
    height = section.page_height
    if width is None or height is None:  # pragma: no cover - a theme always sets both
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image), width=width, height=height)

    drawing = run._r.find(qn("w:drawing"))
    inline = drawing[0]
    graphic = inline.find(qn("a:graphic"))
    doc_pr = inline.find(qn("wp:docPr"))

    anchor = parse_xml(
        f'''<wp:anchor {nsdecls("wp", "a", "r")}
              behindDoc="1" distT="0" distB="0" distL="0" distR="0"
              simplePos="0" locked="0" layoutInCell="1" allowOverlap="1"
              relativeHeight="0">
             <wp:simplePos x="0" y="0"/>
             <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
             <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
             <wp:extent cx="{int(width)}" cy="{int(height)}"/>
             <wp:effectExtent l="0" t="0" r="0" b="0"/>
             <wp:wrapNone/>
           </wp:anchor>'''
    )
    if doc_pr is not None:
        anchor.append(doc_pr)
    if graphic is not None:
        anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def _confidentiality_notice(
    control: DocumentControlConfig, messages: Messages
) -> str | None:
    """The notice's text, or `None` where none is configured.

    One function for both emitters, so the `.docx` and the reading copy cannot disagree
    about which of the two fields wins — the prose a person wrote for this account, over
    a catalogue id naming copy every tenant shares.
    """
    if control.confidentiality_notice:
        return control.confidentiality_notice
    if control.confidentiality_notice_id:
        return messages.text(control.confidentiality_notice_id)
    return None


def _emit_distribution(
    document: DocxDocument,
    *,
    front_matter: FrontMatterConfig,
    messages: Messages,
    paragraph_style: str,
) -> None:
    """Emit the distribution list, in whichever of its two shapes the definition carries.

    ## The v3 rows were being rendered as their own Python repr

    `distribution` is a free-text block at schema_version 1 and 2, and **ordered
    `{recipient, company, note}` rows at 3** (Req 12.6). The emitter knew only the string,
    and the pipeline coerced whatever arrived with `str(...)` — so a v3 profile's
    distribution reached the delivered document as `[{'recipient': ...}]`, the list's repr,
    in a signed report. The rows get a table; the string keeps the paragraph it always had.
    """
    from reporting_agent.compile.blocks.base import LAYOUT_TABLE_STYLE

    control = front_matter.document_control
    if not control.distribution_rows and not control.distribution:
        return

    document.add_paragraph(messages.text(DOC_CONTROL_DISTRIBUTION), style=paragraph_style)

    if not control.distribution_rows:
        # v1/v2 — one free-text block, unchanged.
        document.add_paragraph(str(control.distribution), style=paragraph_style)
        return

    table = document.add_table(rows=1 + len(control.distribution_rows), cols=3)
    table.style = LAYOUT_TABLE_STYLE
    header = table.rows[0].cells
    _set_cell_text(header[0], messages.text(DISTRIBUTION_HEADER_RECIPIENT), paragraph_style)
    _set_cell_text(header[1], messages.text(DISTRIBUTION_HEADER_COMPANY), paragraph_style)
    _set_cell_text(header[2], messages.text(DISTRIBUTION_HEADER_NOTE), paragraph_style)

    for index, row in enumerate(control.distribution_rows):
        cells = table.rows[index + 1].cells
        _set_cell_text(cells[0], row.recipient, paragraph_style)
        _set_cell_text(cells[1], row.company, paragraph_style)
        _set_cell_text(cells[2], row.note, paragraph_style)


def _emit_approvers_table(
    document: DocxDocument,
    *,
    front_matter: FrontMatterConfig,
    messages: Messages,
) -> None:
    """Emit the approvers table: one row per role (Req 13.6).

    Clause (b): a role with no supplied signature image gets an EMPTY RULED SIGNATURE BOX
    at the height the theme declares, NOT that role's typed name.

    The table uses the ``Table Signature`` style from ``render/themes.py``.
    """
    from reporting_agent.render.themes import SIGNATURE_TABLE_STYLE

    # Headers resolved by id
    header_role = messages.text(APPROVER_HEADER_ROLE)
    header_company = messages.text(APPROVER_HEADER_COMPANY)
    header_name = messages.text(APPROVER_HEADER_NAME)
    header_signature = messages.text(APPROVER_HEADER_SIGNATURE)

    # Create table: header row + one row per declared role
    num_roles = len(APPROVER_ROLES)
    table = document.add_table(rows=1 + num_roles, cols=4)
    table.style = SIGNATURE_TABLE_STYLE

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = header_role
    hdr_cells[1].text = header_company
    hdr_cells[2].text = header_name
    hdr_cells[3].text = header_signature

    # One row per role
    for i, role in enumerate(APPROVER_ROLES):
        row_cells = table.rows[i + 1].cells
        entry = _find_approver(front_matter.document_control.approvers, role)

        # Req 12.3 — the role column shows the positional label
        # (Author / Quality Control / Reviewed By / Customer), resolved
        # through the message catalog. The stored role id is an internal
        # symbol; a reader never sees it.
        role_label_id = APPROVER_ROLE_LABEL_IDS.get(role)
        row_cells[0].text = messages.text(role_label_id) if role_label_id else role
        # The COMPANY column shows the approver's **company**. It showed `title` — the
        # person's job title — under a header reading "Company", while the `company` the
        # app collects was dropped by the pipeline and never reached this renderer at
        # all. `title` remains the fallback so a profile that put its company in that
        # field before this fix keeps rendering exactly as it did.
        row_cells[1].text = (entry.company or entry.title) if entry else ""
        row_cells[2].text = entry.name if entry else ""

        # Clause (b): the signature cell text is EMPTY unconditionally, and
        # NEVER the typed name — set before any image placement, so an
        # exception raised while placing an image cannot leave a stale
        # typed name behind it.
        row_cells[3].text = ""

        # Clause (c): where a signature image was supplied, place it inside
        # the signature cell, scaled to fit the theme's declared row height
        # without changing that height — a signed row and an unsigned row
        # occupy the same space, so pagination never depends on who signed.
        if entry is not None and entry.signature_image is not None:
            _place_signature_image(row_cells[3], entry.signature_image)

        # Set row height to the theme's declared signature box height —
        # AFTER either path, so `w:hRule="atLeast"` plus a scaled-to-fit
        # image is what makes the signed and unsigned rows the same size,
        # rather than the image growing the row past it.
        _set_row_height(table.rows[i + 1], SIGNATURE_BOX_HEIGHT_TWIPS)


def _place_signature_image(cell: object, image_bytes: bytes) -> None:
    """Place a signature image inside a table cell, height-constrained to fit
    within the theme's declared signature row height (Req 13.3).

    Height-constrained rather than width-constrained: the box this cell sits
    in is fixed-**height** (`SIGNATURE_BOX_HEIGHT_TWIPS`, `w:hRule="atLeast"`),
    so the image's height is what must not exceed it — an image scaled only
    by width could still be taller than the box and grow the row, which is
    exactly the coupling Req 13.3 exists to rule out. `python-docx` preserves
    aspect ratio automatically when only one of `width`/`height` is given.
    """
    from docx.shared import Twips

    paragraph = cell.paragraphs[0]  # type: ignore[attr-defined]
    run = paragraph.add_run()
    run.add_picture(
        io.BytesIO(image_bytes),
        height=Twips(SIGNATURE_BOX_HEIGHT_TWIPS),
    )


def _set_row_height(row: object, height_twips: int) -> None:
    """Set a table row's height to exactly ``height_twips`` with hRule=atLeast.

    This creates the empty ruled signature box — the row is at least this tall, so an
    empty signature cell renders as a box of the declared height rather than collapsing.
    """
    tr = row._tr  # type: ignore[attr-defined]
    tr_pr = tr.get_or_add_trPr()
    tr_height = tr_pr.makeelement(
        qn("w:trHeight"),
        {qn("w:val"): str(height_twips), qn("w:hRule"): "atLeast"},
    )
    tr_pr.append(tr_height)


def _add_page_break(document: DocxDocument) -> None:
    """Add a page break after the current content."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def _find_approver(
    approvers: tuple[ApproverEntry, ...], role: str
) -> ApproverEntry | None:
    """Find the approver entry for a given role, or None."""
    for entry in approvers:
        if entry.role == role:
            return entry
    return None



def _emit_toc(
    document: DocxDocument,
    *,
    heading_entries: tuple[tuple[str, int], ...],
    messages: Messages,
) -> None:
    """Emit the table of contents section (Req 14.3, 14.5, 14.11).

    Emits a TOC section heading followed by one entry per heading block at levels 1-3
    in document order. Each entry carries the heading's text and a tab stop; the page
    number position is left empty for pass 1 of the two-pass approach. Pass 2
    (``render.toc.apply_toc_page_numbers``) fills it after measuring actual page
    positions from the converted PDF.

    Followed by a page break so content starts on a fresh page.
    """
    from reporting_agent.render.themes import TOC_ENTRY_STYLE
    from reporting_agent.render.toc import TOC_LABEL_ID

    # Section heading — styled Title so it doesn't appear in its own TOC.
    toc_label = messages.text(TOC_LABEL_ID)
    document.add_paragraph(toc_label, style="Title")

    # One entry per heading, at levels 1-3. The text + tab is emitted whether or not
    # a number follows, so pass 1 lays out to exactly the height pass 2 will.
    for heading_text, _level in heading_entries:
        paragraph = document.add_paragraph(style=TOC_ENTRY_STYLE)
        paragraph.add_run(heading_text)
        paragraph.add_run().add_tab()

    # Page break after the TOC section.
    _add_page_break(document)
