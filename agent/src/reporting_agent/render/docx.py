"""The DOCX emitter: one walk of one AST, against one theme.

Reads the compiled AST as its **only** source of content (Req 20.1) and the ledger's
`formatted` strings as its only source of numerals (Req 20.10). It composes no number,
rounds nothing, re-scales nothing and re-units nothing — `compile/format.py` is the only
place a value becomes a string, and this module's job is to put that string in the document
without touching it.

No `docxtpl`, no placeholder substitution, no user-supplied `.docx` (Req 20.2). There is
nothing to substitute into: the theme carries styles and zero content, and the document is
built from the tree.

## Every figure is exactly one run in the `Figure` character style

Req 20.3, and it is the contract the whole verifier rests on. At every position the AST
places a figure — prose, heading, data cell, cover field, chart companion cell — the emitter
writes one run, in that style, holding the `formatted` string in full and **no other
character**. No leading space, no trailing unit, no punctuation. That is what lets the
Token_Extractor find figures without re-parsing prose, and it is why a figure adjacent to
punctuation gets its own run and the punctuation gets another.

## Determinism

Two emissions of one AST against one theme produce identical bytes once
`docProps/core.xml`'s timestamps are excluded (Req 20.8). Nothing here reads the clock, the
locale, the hostname, the environment or a directory listing. `python-docx` writes `created`
and `modified` into the core properties from `datetime.now()` unless told otherwise, so
:data:`FIXED_TIMESTAMP` is set explicitly and the byte-equality test excludes that part by
name rather than hoping it does not matter.

## One write, at the end

The completed document is serialized once, after every block is emitted (Req 20.11). There
is no incremental flush and no partially emitted artifact: :func:`render_document` returns
bytes, and the caller writes them. A renderer that streamed to S3 would be able to leave a
truncated `.docx` behind that looks like a report.
"""

from __future__ import annotations

import io
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass, field
from typing import Final

from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Twips
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell
from docx.text.paragraph import Paragraph as DocxParagraph

from reporting_agent.compile.ast import (
    Chart,
    Document,
    EmptyCell,
    Figure,
    FigureCell,
    LayoutColumn,
    LayoutRow,
    PageBreak,
    Paragraph,
    Row,
    Table,
    Text,
    TextCell,
    TextFact,
    TextFactCell,
)
from reporting_agent.compile.blocks.base import (
    CAPTION_STYLE,
    LAYOUT_TABLE_STYLE,
    NOTICE_STYLE,
    DesignSettings,
)
from reporting_agent.compile.figures import ANCHOR_CHART, ANCHOR_TABLE, FigureLedger
from reporting_agent.compile.messages import Messages
from reporting_agent.errors import RenderFailedError
from reporting_agent.render.anchors import (
    AnchorRecorder,
    assert_header_row,
    record_cell_anchor,
    write_data_table_caption,
    write_layout_table,
)
from reporting_agent.render.charts import (
    CHART_ALT_TEXT_PREFIX,
    SIDECAR_SUFFIX,
    render_chart,
)
from reporting_agent.render.tablefit import allocate, column_demands, header_demands
from reporting_agent.render.themes import (
    FIGURE_CHARACTER_STYLE,
    PREVIEW_NOTICE_STYLE,
    load_theme,
    missing_styles,
)

__all__ = [
    "FIXED_TIMESTAMP",
    "PREVIEW_NOTICE_ID",
    "VOLATILE_PACKAGE_PARTS",
    "RenderOutcome",
    "render_document",
]

FIXED_TIMESTAMP: Final[str] = "2026-01-01T00:00:00Z"
"""The sentinel written into `docProps/core.xml`'s created and modified fields.

A sentinel rather than the real time because Req 20.8 requires two emissions of one AST to
be byte-identical, and a timestamp is the one thing a document writer reaches for that makes
that impossible. The run's own timing is recorded on `report_runs`, which is where an
auditor looks for it — a `.docx` property is not an audit trail.
"""

VOLATILE_PACKAGE_PARTS: Final[tuple[str, ...]] = ("docProps/core.xml",)
"""The parts the determinism test excludes, named explicitly.

One entry, and it is excluded because `python-docx` writes a `dcterms:modified` value this
module cannot fully control through the public API. Naming the part is the point: "the
digests match apart from timestamps" is only a meaningful claim if the exclusion is a fixed
list rather than whatever happened to differ.
"""

PREVIEW_NOTICE_ID: Final[str] = "doc.preview.notice"
"""Resolved through the message catalog so the notice appears in the run's pinned language.
Emitted only in preview mode, in the theme's `PreviewNotice` style, so a
"Render real preview" artifact still says what it is after it leaves the app."""

_PAGE_SIZES: Final[Mapping[str, tuple[int, int]]] = {
    # (width, height) in **twips** — twentieths of a point — which is the unit OOXML's
    # `w:pgSz` uses and the unit the theme documents are written in.
    #
    # They are converted with `Twips()` before being assigned, and that conversion is not
    # optional: `Section.page_width` is a `Length` measured in EMU, and assigning a raw
    # twips integer is read as EMU. 11906 EMU is about 0.013 inch, so the page becomes
    # microscopic, every paragraph overflows it, and LibreOffice emits a stack of blank
    # pages — a document that renders, converts, and contains nothing.
    "A4": (11906, 16838),
    "Letter": (12240, 15840),
}


@dataclass(slots=True)
class RenderOutcome:
    """What one render produced.

    `advisories` carries the non-fatal findings Req 7.6 describes — today only an
    unresolvable logo. They are returned rather than raised because the render **succeeds**:
    a cover without a logo is a complete report, and failing it would withhold a document
    over a decoration.
    """

    docx_bytes: bytes
    table_identities: tuple[str, ...]
    advisories: tuple[str, ...] = ()
    figures_emitted: int = 0
    text_facts_emitted: int = 0
    """Counted **separately** from `figures_emitted` (Req 6.15). A text fact is a checked
    value and not a figure: it enters neither `figure_count` nor the unused-figure warning,
    so a single total would make an unrendered figure and an unrendered fact
    indistinguishable in the one number an operator reads first."""
    chart_hashes: Mapping[str, str] = field(default_factory=dict)
    chart_sidecars: Mapping[str, bytes] = field(default_factory=dict)
    """The chart data hash per chart identity, and the sidecar bytes to write beside each
    embedded image (Req 22.3). Returned rather than written for the same reason the document
    is: this module cannot write, so it cannot leave a partial set behind."""


def _apply_column_widths(table: DocxTable, node: Table, *, text_width: int) -> None:
    """Divide `text_width` (in twips) between the columns in proportion to their demand.

    Word divides a table's width equally by default, which is fine until one column holds
    something much longer than its neighbours: a companion table with four CPU columns
    (`0.20%`) and one memory column (`3,187,970,789.00 bytes`) gave all five the same
    width, and LibreOffice wrapped the memory values across two lines inside their cells.
    Every character was still on the page, but `verify/pdf.py` searches for the ledger
    string contiguously and a line break falls in the middle of it — run ef01a404 came back
    with 30 `pdf_figure_missing` findings, one per day of July, all in that one column, on a
    table with room to spare in its other four.

    Fixed layout rather than autofit, and the width is set on **every cell** as well as on
    the grid column: Word honours the grid, LibreOffice honours the cells, and a document
    this product renders has to convert correctly under LibreOffice specifically.

    Sizing alone cannot save a table whose columns cannot *all* fit — see
    `tablefit.fits_page`, which is what keeps one from being built.
    """
    allocation = allocate(column_demands(node), header_demands(node))
    total = sum(allocation)
    if not total:
        return

    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)

    widths = [Twips(int(text_width * share / total)) for share in allocation]
    for ordinal, width in enumerate(widths):
        table.columns[ordinal].width = width
    for row in table.rows:
        for ordinal, cell in enumerate(row.cells):
            if ordinal < len(widths):
                cell.width = widths[ordinal]


@dataclass(slots=True)
class _Emitter:
    """One render's mutable state. Not reused across renders."""

    document: DocxDocument
    ledger: FigureLedger
    design: DesignSettings
    messages: Messages
    recorder: AnchorRecorder
    declared_styles: frozenset[str]
    advisories: list[str] = field(default_factory=list)
    figures_emitted: int = 0
    text_facts_emitted: int = 0
    chart_hashes: dict[str, str] = field(default_factory=dict)
    chart_sidecars: dict[str, bytes] = field(default_factory=dict)

    @property
    def text_width(self) -> int:
        """The width available to a table, in twips: the page less its two margins.

        Read off the section rather than assumed from the page size, because the theme
        sets the margins and `_apply_page_size` sets the page — a table sized against a
        constant would overflow the moment either changed.
        """
        section = self.document.sections[0]
        return int(section.page_width - section.left_margin - section.right_margin)

    # --- styles ---------------------------------------------------------------

    def style(self, name: str, *, at: str) -> str:
        """Assert the theme declares `name`, then return it.

        Checked against the set read once from the theme rather than by catching `KeyError`
        from `python-docx`, so the failure message can name the AST node instead of the
        style lookup that happened to blow up.
        """
        if name not in self.declared_styles:
            raise RenderFailedError(
                f"{at} references the style {name!r}, which theme "
                f"{self.design.preset!r} does not declare"
            )
        return name

    # --- inline content -------------------------------------------------------

    def write_inlines(
        self, paragraph: DocxParagraph, inlines: Iterable[object], *, at: str
    ) -> None:
        """Write a paragraph's inline content, one run per inline.

        A `Text` becomes one run in the paragraph's own style; a `Figure` becomes one run in
        the `Figure` character style holding `formatted` and nothing else. The split is what
        makes a figure locatable: `"CPU averaged "` and `"12.4%"` are two runs, so the
        extractor finds the second without parsing the first.
        """
        for ordinal, inline in enumerate(inlines):
            where = f"{at} inline {ordinal}"
            if isinstance(inline, Figure):
                self.write_figure_run(paragraph, inline)
            elif isinstance(inline, Text):
                paragraph.add_run(inline.text)
            else:
                # Unreachable through the AST's own validation, which admits only Text and
                # Figure in an inline position. Reaching it means a node type was added to
                # the union without being taught here, and Req 20.12 makes that a refusal
                # rather than a silently dropped value.
                raise RenderFailedError(
                    f"{where} is {type(inline).__name__}; an inline position admits only "
                    f"Text or Figure, and a quantity reaching a text position without "
                    f"being a Figure has no provenance"
                )

    def write_figure_run(self, paragraph: DocxParagraph, figure: Figure) -> None:
        """One run, the `Figure` character style, `formatted` in full and nothing else.

        Req 20.3 and Req 20.10 together: the string comes from the ledger entry verbatim,
        and the run holds no other character. The emitter never appends a unit or a
        separator — `compile/format.py` already put the unit inside `formatted`, and a
        second one here would break the exact-equality comparison the verifier performs.
        """
        if not figure.formatted:
            raise RenderFailedError(
                f"figure {figure.path!r} carries no formatted string; the Formatter is the "
                f"only path from a value to a display string and this figure never took it"
            )
        run = paragraph.add_run(figure.formatted)
        run.style = self.document.styles[FIGURE_CHARACTER_STYLE]
        self.figures_emitted += 1

    def write_text_fact_run(self, paragraph: DocxParagraph, fact: TextFact) -> None:
        """One run, in exactly one paragraph, the `Figure` character style, `formatted` and
        nothing else (Req 6.4, 6.9).

        **The same character style a figure takes**, and that is deliberate rather than
        convenient: what the style marks is *this text is a checked value*, which is exactly
        as true of `Succeeded` as of `12.5%`. It is also what lets the token extractor find a
        fact without re-parsing prose — a fact's value carries no digit in the common case, so
        nothing about the text itself would identify it.

        **Exactly one run** is the load-bearing part. `verify/facts.py` concatenates a cell's
        runs in document order **with no character inserted between them** and compares the
        result to `formatted` character for character. Two runs would still concatenate
        correctly today, but every additional run is a place a later edit could insert a
        space, and the comparison admits none — so the invariant is stated here, where the
        runs are created, and asserted by the paragraph check below.

        The paragraph must be empty when this is called. A cell's first paragraph is fresh
        from `python-docx`, so a non-empty one means something already wrote into this cell,
        and the concatenation would then include text that is not the fact.
        """
        if not fact.formatted:
            raise RenderFailedError(
                f"text fact {fact.path!r} carries no formatted string; "
                f"`compile/format.py::format_text_fact` is the only path to one and this "
                f"entry never took it"
            )
        if paragraph.runs:
            raise RenderFailedError(
                f"text fact {fact.path!r} is emitted into a paragraph that already carries "
                f"{len(paragraph.runs)} run(s); the facts pass concatenates a cell's runs "
                f"with no separator and compares the result to `formatted`, so any earlier "
                f"text in this cell would be read as part of the fact"
            )

        run = paragraph.add_run(fact.formatted)
        run.style = self.document.styles[FIGURE_CHARACTER_STYLE]
        self.text_facts_emitted += 1

    # --- blocks ---------------------------------------------------------------

    def emit_block(self, node: object, *, container: DocxCell | None = None) -> None:
        """Emit one block-level node, in AST order, with no reordering (Req 20.7)."""
        if isinstance(node, Paragraph):
            self.emit_paragraph(node, container=container)
        elif isinstance(node, Table):
            self.emit_table(node, container=container)
        elif isinstance(node, Chart):
            self.emit_chart(node, container=container)
        elif isinstance(node, LayoutRow):
            self.emit_layout_row(node, container=container)
        elif isinstance(node, PageBreak):
            self.emit_page_break(node, container=container)
        else:
            raise RenderFailedError(
                f"cannot emit {type(node).__name__} at "
                f"{getattr(node, 'path', '<unknown path>')!r}: no emitter is declared for "
                f"that node type, and emitting nothing would leave the document quietly "
                f"missing a section"
            )

    def emit_paragraph(self, node: Paragraph, *, container: DocxCell | None) -> None:
        style = self.style(node.style, at=f"paragraph {node.path!r}")
        paragraph = self._new_paragraph(container, style)
        self.write_inlines(paragraph, node.inlines, at=f"paragraph {node.path!r}")

    def emit_page_break(self, node: PageBreak, *, container: DocxCell | None) -> None:
        if container is not None:
            # A page break inside a layout cell would break the column, not the page. The
            # compiler cannot emit one there (a `row`'s children are validated block types
            # and `page_break` is not among them), so this is a contract violation rather
            # than a case to handle.
            raise RenderFailedError(
                f"page break {node.path!r} sits inside a layout cell; a break there "
                f"applies to the cell rather than to the page"
            )
        paragraph = self.document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def emit_chart(self, node: Chart, *, container: DocxCell | None) -> None:
        """One image, then its companion table, with nothing between them (Req 22.2).

        Both carry the same `cht:<path>` identity — the image in its alternative text, the
        table in its `w:tblCaption` — so the verifier pairs them by identity rather than by
        proximity. The adjacency is still enforced here because a reader relies on it: the
        table explains the picture above it.

        The companion table goes through :meth:`emit_table`, so it takes the same caption
        writer, the same header assertion and the same anchor recorder as any other data
        table. A second table emitter would be a second chance to get the caption wrong.

        When the chart carries a ``period_label``, it is emitted as a caption paragraph
        between the image and the table (Req 17.12). The string is identical to the one
        rendered inside the chart image, so the document and the chart cannot disagree
        about the period plotted.
        """
        at = f"chart {node.path!r}"
        artifacts = render_chart(node, table_style=self.design.table_style_name, messages=self.messages)

        picture_paragraph = self._new_paragraph(container, self.style(CAPTION_STYLE, at=at))
        run = picture_paragraph.add_run()
        run.add_picture(io.BytesIO(artifacts.image_png), width=Inches(_CHART_WIDTH_INCHES))
        _set_picture_alt_text(picture_paragraph, artifacts.identity)

        # Req 17.12 — present the period_label identically to how the chart image renders it.
        if node.period_label:
            period_para = self._new_paragraph(container, self.style(CAPTION_STYLE, at=at))
            period_para.add_run(node.period_label)

        self.chart_hashes[artifacts.identity] = artifacts.data_hash
        self.chart_sidecars[f"{artifacts.identity}{SIDECAR_SUFFIX}"] = artifacts.sidecar_json

        self.emit_table(
            artifacts.table,
            container=container,
            identity=artifacts.identity,
            anchor_kind=ANCHOR_CHART,
        )

    # --- tables ---------------------------------------------------------------

    def emit_table(
        self,
        node: Table,
        *,
        container: DocxCell | None,
        identity: str | None = None,
        anchor_kind: str = ANCHOR_TABLE,
    ) -> None:
        """One data table: a header row, then one row per data row, then its caption.

        `identity` overrides the node's own `tbl:<path>` and exists for exactly one caller: a
        chart's companion table, which must carry the **chart's** `cht:<path>` identity so the
        verifier can pair the table with the image by identity rather than by proximity
        (Req 22.2). Two different identities on the two halves of one chart would leave every
        chart unpairable — and because both strings are derived from the same AST path, the
        mismatch would look correct in isolation.

        The caption paragraph comes **after** the table and carries the theme's `Caption`
        style. The table's *identity* is a different thing entirely and lives in
        `w:tblPr/w:tblCaption` — Word calls both "caption", which is a naming collision
        worth being careful about: the reader sees the paragraph, the verifier reads the
        element.
        """
        at = f"table {node.path!r}"
        style = self.style(node.style, at=at)
        headers = assert_header_row(node.columns, at=at)
        resolved_identity = node.anchor_id if identity is None else identity

        self.recorder.claim_identity(node, resolved_identity)
        row_keys = self.recorder.row_keys_for(node)

        table = self._new_table(container, rows=0, cols=len(node.columns), style=style)
        write_data_table_caption(table, resolved_identity)

        # Req 21.4 — the header row is the table's first row, and its text is what the
        # verifier resolves a column by.
        header_cells = table.add_row().cells
        for ordinal, header in enumerate(headers):
            self._set_cell_text(header_cells[ordinal], header, style=None)
        self._mark_header_row(table)

        for row_ordinal, row in enumerate(node.rows):
            self.emit_row(
                table,
                node,
                row,
                identity=resolved_identity,
                row_key=row_keys[row_ordinal],
                at=f"{at} row {row_ordinal}",
                anchor_kind=anchor_kind,
            )

        # After every row, because the demand is measured from the cells the table
        # actually carries rather than from its declaration.
        _apply_column_widths(table, node, text_width=self.text_width)

        if node.caption:
            caption = self._new_paragraph(container, self.style(CAPTION_STYLE, at=at))
            caption.add_run(node.caption)

        # The one instant every fact in this table was collected at, when they agree —
        # what used to be an `observed_at` column beside every value. Its own paragraph
        # rather than a run appended to the caption: the caption is the template author's
        # string and this is the compiler's, and a theme that wants to set them
        # differently needs two paragraphs to do it with.
        if node.note:
            note = self._new_paragraph(container, self.style(CAPTION_STYLE, at=at))
            note.add_run(node.note)

    def emit_row(
        self,
        table: DocxTable,
        node: Table,
        row: Row,
        *,
        identity: str,
        row_key: str,
        at: str,
        anchor_kind: str = ANCHOR_TABLE,
    ) -> None:
        """One data row.

        Short rows are **padded** rather than refused. A truncation row carries two cells in
        a five-column table (`compile/blocks/base.py`'s `omitted_row`), which is deliberate
        on the compiler's side — the row is a statement about the table, not another
        resource — so the renderer completes the grid. Padding with genuinely empty cells
        rather than with repeated text keeps the verifier's `(row_key, column_key)`
        resolution honest: an empty cell resolves to no figure, which is what it holds.
        """
        if len(row.cells) > len(node.columns):
            raise RenderFailedError(
                f"{at} carries {len(row.cells)} cells but the table declares "
                f"{len(node.columns)} columns; a cell with no column cannot be addressed "
                f"as (row key, column key)"
            )

        cells = table.add_row().cells
        notice_style = self._notice_style_for(node, row)

        for column_ordinal, column in enumerate(node.columns):
            docx_cell = cells[column_ordinal]
            if column_ordinal >= len(row.cells):
                continue  # padded: no content, no anchor

            cell = row.cells[column_ordinal]
            where = f"{at} cell {column_ordinal}"

            if isinstance(cell, FigureCell | TextFactCell):
                paragraph = docx_cell.paragraphs[0]
                if isinstance(cell, FigureCell):
                    self.write_figure_run(paragraph, cell.figure)
                else:
                    self.write_text_fact_run(paragraph, cell.fact)
                # Req 21.3, 6.9 — the anchor triple, completed here because only the
                # renderer knows the emitted grid, and built in **one** place for both
                # kinds so a change to its shape cannot reach one and miss the other.
                record_cell_anchor(
                    self.ledger,
                    cell,
                    anchor_kind=anchor_kind,
                    anchor_id=identity,
                    row_key=row_key,
                    # The **header text**, not `column.key`, for exactly the reason
                    # `document_row_key` records emitted text rather than `Row.key`: Req 27.1
                    # resolves a column by the header the document carries, and a chart's
                    # companion table has `key="value"` under the header `"Value"`. Recording
                    # the key would record a string the verifier cannot find.
                    column_key=column.header,
                )
            elif isinstance(cell, TextCell):
                self._set_cell_text(docx_cell, cell.text, style=notice_style)
            elif isinstance(cell, EmptyCell):
                # Deliberately nothing. Distinct from "0", and the distinction is the
                # product: a metric a resource does not emit is a recorded gap, and a zero
                # would read as measured idleness.
                continue
            else:
                raise RenderFailedError(
                    f"{where} is {type(cell).__name__}; a cell admits only FigureCell, "
                    f"TextFactCell, TextCell or EmptyCell"
                )

    def _notice_style_for(self, node: Table, row: Row) -> str | None:
        """`Notice` for an explicit no-data row, otherwise the table's own cell style.

        The empty-scope and no-gaps rows read as **information**, not as an error —
        `design-system.md` is explicit that they belong in mist neutrals rather than in
        `--destructive`. Identified by the row's compiler-assigned key rather than by
        matching its text, so a wording change does not silently restyle it.
        """
        if row.key in _NOTICE_ROW_KEYS:
            return self.style(NOTICE_STYLE, at=f"table {node.path!r} row {row.key!r}")
        return None

    # --- layout ---------------------------------------------------------------

    def emit_layout_row(self, node: LayoutRow, *, container: DocxCell | None) -> None:
        """A `row` block: one borderless layout table, one cell per declared column.

        Nesting is one level deep by the definition schema, so a layout row inside a layout
        cell cannot arise from a valid definition; it is refused rather than handled, because
        a nested layout table would give the verifier a container it has no rule for.
        """
        at = f"layout row {node.path!r}"
        if container is not None:
            raise RenderFailedError(
                f"{at} sits inside a layout cell; nesting is one level deep and a row "
                f"holds no row"
            )

        style = self.style(LAYOUT_TABLE_STYLE, at=at)
        table = self._new_table(None, rows=1, cols=len(node.columns), style=style)

        # Req 21.2 — no caption, no header row, no row key. The verifier's table pass
        # enumerates captioned tables, so this table is excluded by construction.
        write_layout_table(table)

        cells = table.rows[0].cells
        for ordinal, column in enumerate(node.columns):
            if not isinstance(column, LayoutColumn):  # pragma: no cover - AST validates
                raise RenderFailedError(f"{at} column {ordinal} is not a LayoutColumn")
            cell = cells[ordinal]
            # python-docx gives a new cell one empty paragraph. Emitting into the cell
            # appends after it, which would leave a blank line above every column's first
            # block, so the placeholder is removed once the column has content.
            placeholder = cell.paragraphs[0]
            for child in column.blocks:
                self.emit_block(child, container=cell)
            if column.blocks:
                self._remove_paragraph(placeholder)

    # --- python-docx plumbing -------------------------------------------------

    def _new_paragraph(self, container: DocxCell | None, style: str) -> DocxParagraph:
        if container is None:
            return self.document.add_paragraph(style=style)
        paragraph = container.add_paragraph(style=style)
        return paragraph

    def _new_table(
        self, container: DocxCell | None, *, rows: int, cols: int, style: str
    ) -> DocxTable:
        if container is None:
            table = self.document.add_table(rows=rows, cols=cols)
        else:
            table = container.add_table(rows=rows, cols=cols)
        table.style = self.document.styles[style]
        return table

    def _set_cell_text(
        self, cell: DocxCell, text: str, *, style: str | None
    ) -> None:
        """Write text into a cell's first paragraph, in one run.

        Uses the existing paragraph rather than `cell.text = ...`, because that setter
        replaces the paragraph and drops the style with it.
        """
        paragraph = cell.paragraphs[0]
        if style is not None:
            paragraph.style = self.document.styles[style]
        paragraph.add_run(text)

    def _mark_header_row(self, table: DocxTable) -> None:
        """Mark the first row as a repeating header (`w:trPr/w:tblHeader`).

        So a table spanning a page break repeats its header, which matters because the
        header text is what a reader *and* the verifier resolve columns by. Structural, not
        formatting: it duplicates no style the theme declares.
        """
        row = table.rows[0]
        properties = row._tr.get_or_add_trPr()
        properties.append(properties.makeelement(qn("w:tblHeader"), {}))

    @staticmethod
    def _remove_paragraph(paragraph: DocxParagraph) -> None:
        element = paragraph._element
        element.getparent().remove(element)


_CHART_WIDTH_INCHES: Final[float] = 6.0
"""The embedded image's display width, matching `chartstyle.CHART_SIZE_INCHES`'s width so the
PNG is placed at its natural scale and no resampling blurs the direct labels."""


def _set_picture_alt_text(paragraph: DocxParagraph, identity: str) -> None:
    """Write `identity` into the inline picture's alternative text (Req 22.2).

    Into both `descr` and `name` on the drawing's `wp:docPr`: Word surfaces `descr` as the
    Alt Text description, and a consumer reading `name` instead still finds the identity. The
    verifier pairs an image with its table by this string, so it has to be present verbatim
    rather than as part of a sentence.
    """
    for doc_pr in paragraph._p.iter(qn("wp:docPr")):
        doc_pr.set("descr", f"{CHART_ALT_TEXT_PREFIX}{identity}")
        doc_pr.set("name", identity)
        return
    raise RenderFailedError(
        f"the chart image for {identity!r} carries no drawing properties, so its identity "
        f"cannot be recorded in its alternative text and the verifier could not pair it "
        f"with its companion table"
    )


_NOTICE_ROW_KEYS: Final[frozenset[str]] = frozenset({"empty-scope", "no-gaps"})
"""The compiler's row keys for the two explicit no-data rows.

Matched by key rather than by text so a wording change in `EMPTY_SCOPE_TEXT` does not
silently stop the row being styled as a notice."""


def render_document(
    compiled: object,
    *,
    ledger: FigureLedger,
    design: DesignSettings,
    messages: Messages,
    preview: bool = False,
    front_matter: object | None = None,
    run: object | None = None,
) -> RenderOutcome:
    """Emit `compiled` as `.docx` bytes, once (Req 20.1, 20.11).

    `compiled` is a :class:`~reporting_agent.compile.ast.Document`. `ledger` is the ledger
    that render produced — the same object, not a reconstruction, because the anchors this
    function records are written onto its entries.

    `messages` is **required** rather than optional, because this function passes it to the
    chart renderer and the preview notice. A default that silently fell back to English is
    the exact defect criterion 15.11 closes.

    Returns bytes rather than writing them: Req 20.11 requires the completed document to be
    written as **one** artifact object, and the cleanest way to guarantee that is for the
    renderer to be unable to write at all.
    """
    if not isinstance(compiled, Document):
        raise RenderFailedError(
            f"render_document takes a compiled Document, got {type(compiled).__name__}"
        )

    document = load_theme(design.preset)
    _apply_page_size(document, design)

    declared = frozenset(style.name for style in document.styles if style.name)
    emitter = _Emitter(
        document=document,
        ledger=ledger,
        design=design,
        messages=messages,
        recorder=AnchorRecorder(ledger=ledger),
        declared_styles=declared,
    )

    if preview:
        _emit_preview_notice(document, emitter, messages=messages)

    # --- front matter (cover, document control, TOC) — Req 13.4 -------------------
    # Emitted after _apply_page_size and preview notice, before the first content block.
    # `front_matter=None` means no front matter (v1 definitions, thumbnails).
    if front_matter is not None and run is not None:
        from reporting_agent.render.front_matter import (
            FrontMatterConfig,
            RunFacts,
            emit_front_matter,
        )

        if isinstance(front_matter, FrontMatterConfig) and isinstance(run, RunFacts):
            emit_front_matter(
                document,
                front_matter=front_matter,
                run=run,
                messages=messages,
                cursor=compiled,
                ledger=ledger,
            )

    for ordinal, block in enumerate(compiled.blocks):
        try:
            emitter.emit_block(block)
        except RenderFailedError:
            raise
        except Exception as error:
            raise RenderFailedError(
                f"emitting block {ordinal} at "
                f"{getattr(block, 'path', '<unknown>')!r} failed: "
                f"{type(error).__name__}: {error}"
            ) from error

    _fix_timestamps(document)

    # One write, after every block. A partially emitted document is never bytes anybody
    # could mistake for a report.
    buffer = io.BytesIO()
    document.save(buffer)

    return RenderOutcome(
        docx_bytes=buffer.getvalue(),
        table_identities=emitter.recorder.identities(),
        advisories=tuple(emitter.advisories),
        figures_emitted=emitter.figures_emitted,
        text_facts_emitted=emitter.text_facts_emitted,
        chart_hashes=dict(emitter.chart_hashes),
        chart_sidecars=dict(emitter.chart_sidecars),
    )


def _apply_page_size(document: DocxDocument, design: DesignSettings) -> None:
    """Set the section's page size from the template's `design.page_size`.

    The theme ships A4; a template asking for Letter gets it here rather than by shipping
    eight theme documents.
    """
    size = _PAGE_SIZES.get(design.page_size)
    if size is None:
        raise RenderFailedError(
            f"page size {design.page_size!r} is not one of {sorted(_PAGE_SIZES)}"
        )
    width, height = size
    for section in document.sections:
        section.page_width = Twips(width)
        section.page_height = Twips(height)


def _emit_preview_notice(document: DocxDocument, emitter: _Emitter, *, messages: Messages) -> None:
    """Put the preview notice in every section's header, so it lands on **every page**.

    A leading body paragraph would say it once, on page one, and a preview PDF is a thing
    people excerpt, screenshot and forward. Page seven of a preview has to be as obviously
    a preview as page one, because the failure this notice exists to prevent is somebody
    acting on an unverified figure — and nobody scrolls back to page one to check.

    A header rather than a body paragraph for a second reason: the preview exists to show
    what the layout will do, and a paragraph inserted at the top of the body shifts every
    block below it. The notice would change the pagination it is meant to preview.

    `is_linked_to_previous = False` on each section so the header is that section's own.

    **`different_first_page_header_footer` is left exactly as the theme set it**, and the
    notice is written into `first_page_header` as well. Forcing the flag to `False` would put
    the notice on page one, which is the goal — but by overriding a theme decision, and a
    preset that suppresses its own cover-page header would then render a preview whose page 1
    layout differs from the real render. That is precisely the class of problem the pagination
    argument above is about: a preview must not change the layout it exists to preview. All
    four themes currently carry `False` with empty single-paragraph headers, so writing both
    is a no-op today and stays correct for a theme that does not.
    """
    style = emitter.style(PREVIEW_NOTICE_STYLE, at="the preview notice")
    for section in document.sections:
        headers = [section.header]
        if section.different_first_page_header_footer:
            headers.append(section.first_page_header)
        for header in headers:
            header.is_linked_to_previous = False
            paragraph = (
                header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            )
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            paragraph.style = style
            paragraph.add_run(messages.text(PREVIEW_NOTICE_ID))


def _fix_timestamps(document: DocxDocument) -> None:
    """Pin the core properties' timestamps to the sentinel (Req 20.8).

    `python-docx` writes `dcterms:created` and `dcterms:modified` from the wall clock on
    save unless they already carry a value. Setting them through the XML rather than through
    `core_properties` because the latter takes a `datetime` and re-serializes it, which
    reintroduces a formatting decision this module does not need to own.
    """
    core = document.core_properties._element
    for field_name in ("created", "modified"):
        for element in core.findall(qn(f"dcterms:{field_name}")):
            element.text = FIXED_TIMESTAMP
            element.set(qn("xsi:type"), "dcterms:W3CDTF")


def assert_theme_declares(preset: str, names: Sequence[str]) -> None:
    """Assert `preset`'s theme declares every name in `names` (Req 20.5).

    Reports **every** missing style rather than the first, so one fix pass clears the
    render. `load_theme` already asserts the full declared union; this exists for a caller
    holding a specific set — a compiled document's actual style references.
    """
    document = load_theme(preset)
    absent = missing_styles(document, required=tuple(names))
    if absent:
        raise RenderFailedError(
            f"theme {preset}.docx is missing {len(absent)} style(s) the compiled document "
            f"references: {', '.join(absent)}"
        )
