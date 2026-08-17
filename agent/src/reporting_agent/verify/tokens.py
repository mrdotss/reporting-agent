"""Reading the document the way Word stores it.

**`document.paragraphs` and `document.tables` are the wrong readers, and the reason is
a silent total failure** (Req 26.2). Both collections enumerate only the *direct
children of the body element*. A paragraph inside a table cell, inside a nested table,
inside a text box (`w:txbxContent`) or inside a content control (`w:sdt`) is not a
direct child of the body — and neither is a table nested in a cell. This product emits
a `row` block as a borderless layout table, so a chart's companion table lives nested
one level down as a matter of course. A verifier reading through those collections
extracts **nothing** from it, finds no unmatched token, records no finding, and
**passes the document**. There is no error, no empty result to notice, nothing that
distinguishes the outcome from a genuinely clean document. So the reader here is
`body.iter(qn("w:p"))`, and `agent/tests/test_boundaries.py` fails the suite if any
module under `verify/` so much as names `.paragraphs` or `.tables` on a document.

**Tokenize the concatenated paragraph, never a run** (Req 26.3, 28.9). One formatted
number is routinely stored as several consecutive `w:r`/`w:t` pairs — Word splits runs
on spell-check state, revision marks and `rPr` changes alone — so `1,234.56` commonly
arrives as `1,`, `234.` and `56`. Per-run tokenization yields three fragments matching
no ledger value and three spurious survivors, turning a correct document into three
blocking findings.

**Join with no inserted character** (Req 26.8). Adjacent `w:t` nodes concatenate
directly; every space a `w:t` carries is preserved; `w:tab` and `w:br` each become one
space; leading and trailing whitespace is stripped; nothing else is altered. Inserting
a space between runs breaks a figure into two tokens exactly as per-run tokenization
does — the join and the tokenizer have to agree, and the way they agree is that the
join adds nothing.

**A blank caption counts as absent** (Req 26.5), so a present-but-whitespace
`w:tblCaption` can smuggle a layout table neither into the data pass nor a data table
out of it.

**An unreadable document is a proven failure, never an empty token set** (Req 26.10).
A `.docx` that will not open, or one carrying no body element, raises
:class:`VerificationFailedError` here rather than returning nothing — because
returning nothing is indistinguishable from a document with nothing wrong, and every
later pass would agree that a document it could not read contained no defects.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Final, NamedTuple

from docx.oxml.ns import qn

from reporting_agent.errors import VerificationFailedError

__all__ = [
    "PART_BODY",
    "PART_FOOTER",
    "PART_HEADER",
    "ExtractedParagraph",
    "ExtractedTable",
    "NumericToken",
    "data_tables",
    "normalize_pdf_text",
    "numeric_tokens",
    "open_document",
    "paragraph_text",
    "paragraph_texts",
    "read_pdf_text",
    "table_caption",
]

PART_BODY: Final[str] = "body"
PART_HEADER: Final[str] = "header"
PART_FOOTER: Final[str] = "footer"

_W_P: Final[str] = qn("w:p")
_W_TBL: Final[str] = qn("w:tbl")
_W_T: Final[str] = qn("w:t")
_W_TAB: Final[str] = qn("w:tab")
_W_BR: Final[str] = qn("w:br")
_W_TBLPR: Final[str] = qn("w:tblPr")
_W_TBLCAPTION: Final[str] = qn("w:tblCaption")
_W_VAL: Final[str] = qn("w:val")

_DIGIT: Final[re.Pattern[str]] = re.compile(r"\d")


class ExtractedParagraph(NamedTuple):
    """One paragraph's concatenated text with where it came from.

    `block_id` is the caption identity of the innermost captioned data table the
    paragraph sits inside, or `None` when it belongs to no data table. It is the only
    block attribution the document actually carries — the renderer writes
    `w:tblCaption` on data tables and nothing per-paragraph — which is why a finding
    for a paragraph outside a table locates itself by `region` and `ordinal` instead.
    """

    text: str
    part: str
    ordinal: int
    block_id: str | None


class ExtractedTable(NamedTuple):
    """One `w:tbl` carrying a non-blank caption, with its document ordinal."""

    identity: str
    ordinal: int
    element: object


class NumericToken(NamedTuple):
    """One maximal whitespace-delimited substring carrying at least one digit."""

    text: str
    part: str
    ordinal: int
    block_id: str | None
    offset: int


def open_document(path: str | Path) -> object:
    """Open a `.docx`, refusing an unreadable one loudly.

    Every failure mode collapses to one terminal outcome (Req 26.10): a file that is
    not a valid package, a package that is not a Word document, and a document
    carrying no body element all raise rather than yielding an empty extraction that
    every later pass would read as "nothing wrong here".
    """
    from docx import Document  # imported here so a bad install fails at call, not import

    try:
        document = Document(str(path))
    # Broad by intent: a corrupt zip, a non-Word package and a truncated file are
    # three different exceptions and one outcome — the document cannot be read, so it
    # cannot be verified. Narrowing this would let an unanticipated exception type
    # escape as a crash instead of a proven verification failure.
    except Exception as exc:
        raise VerificationFailedError(
            f"the rendered document at {path} could not be opened: "
            f"{type(exc).__name__}"
        ) from exc

    body = getattr(document.element, "body", None)
    if body is None:
        raise VerificationFailedError(
            f"the rendered document at {path} carries no body element"
        )
    return document


def _caption_of(table_element: object) -> str | None:
    """The `w:tblCaption` value of a `w:tbl`, or `None` when absent or blank."""
    properties = table_element.find(_W_TBLPR)  # type: ignore[attr-defined]
    if properties is None:
        return None
    for caption in properties.findall(_W_TBLCAPTION):
        value = caption.get(_W_VAL)
        if value is not None and value.strip():
            return value
    return None


def table_caption(table_element: object) -> str | None:
    """Public form of :func:`_caption_of`, for callers holding a raw `w:tbl`."""
    return _caption_of(table_element)


def paragraph_text(paragraph_element: object) -> str:
    """Concatenate one `w:p` the way Word stores it.

    Descends into the paragraph at every depth so a `w:t` inside a hyperlink, a
    revision wrapper or a content control is not skipped; inserts nothing between
    adjacent `w:t` nodes; maps `w:tab` and `w:br` to one space each; strips the ends.
    """
    pieces: list[str] = []
    for node in paragraph_element.iter():  # type: ignore[attr-defined]
        tag = node.tag
        if tag == _W_T:
            if node.text:
                pieces.append(node.text)
        elif tag in (_W_TAB, _W_BR):
            pieces.append(" ")
    return "".join(pieces).strip()


def _enclosing_caption(paragraph_element: object, captions: dict[object, str]) -> str | None:
    """The innermost captioned table a paragraph sits inside, if any."""
    node = paragraph_element.getparent()  # type: ignore[attr-defined]
    while node is not None:
        if node.tag == _W_TBL:
            identity = captions.get(node)
            if identity is not None:
                return identity
        node = node.getparent()
    return None


def _extract_part(root: object, part: str, start: int, captions: dict[object, str]) -> tuple[list[ExtractedParagraph], int]:
    paragraphs: list[ExtractedParagraph] = []
    ordinal = start
    for element in root.iter(_W_P):  # type: ignore[attr-defined]
        ordinal += 1
        paragraphs.append(
            ExtractedParagraph(
                text=paragraph_text(element),
                part=part,
                ordinal=ordinal,
                block_id=_enclosing_caption(element, captions),
            )
        )
    return paragraphs, ordinal


def _caption_index(document: object) -> dict[object, str]:
    body = document.element.body  # type: ignore[attr-defined]
    index: dict[object, str] = {}
    for table_element in body.iter(_W_TBL):
        identity = _caption_of(table_element)
        if identity is not None:
            index[table_element] = identity
    return index


def paragraph_texts(document: object) -> tuple[ExtractedParagraph, ...]:
    """Every paragraph of the body plus every header and footer part (Req 26.1, 26.6).

    Ordinals are 1-based and continuous across the whole extraction in body → headers →
    footers order, so a finding's `paragraph_ordinal` identifies one paragraph of one
    document rather than one paragraph of whichever part it happened to come from.
    """
    captions = _caption_index(document)
    body = document.element.body  # type: ignore[attr-defined]
    paragraphs, ordinal = _extract_part(body, PART_BODY, 0, captions)

    # Headers and footers are separate parts: they are not reachable from the body
    # element at any depth, so an extraction that stops at the body silently ignores
    # every figure a cover block or a running footer placed there.
    for section in document.sections:  # type: ignore[attr-defined]
        for attribute, part in (
            ("header", PART_HEADER),
            ("first_page_header", PART_HEADER),
            ("even_page_header", PART_HEADER),
            ("footer", PART_FOOTER),
            ("first_page_footer", PART_FOOTER),
            ("even_page_footer", PART_FOOTER),
        ):
            container = getattr(section, attribute, None)
            if container is None:
                continue
            # `_element` rather than the public `.paragraphs` / `.tables`, and this is
            # the one place a private attribute is the correct call: those two
            # collections carry the same direct-children-only blind spot on a header
            # that they carry on the body, so using them here would reintroduce
            # exactly the silent failure this module exists to prevent. python-docx
            # exposes no public accessor for a header part's raw element.
            root = getattr(container, "_element", None)
            if root is None:
                continue
            extracted, ordinal = _extract_part(root, part, ordinal, captions)
            paragraphs.extend(extracted)

    return tuple(paragraphs)


def data_tables(document: object) -> tuple[ExtractedTable, ...]:
    """Every `w:tbl` carrying a non-blank caption, in document order (Req 26.4, 26.5).

    Layout tables — the borderless tables a `row` block emits — carry no caption and
    are therefore excluded **by construction** rather than by guessing from borders or
    cell count. The ordinal counts captioned tables only, so it is stable against a
    layout table being added or removed around them.
    """
    body = document.element.body  # type: ignore[attr-defined]
    tables: list[ExtractedTable] = []
    ordinal = 0
    for element in body.iter(_W_TBL):
        identity = _caption_of(element)
        if identity is None:
            continue
        ordinal += 1
        tables.append(ExtractedTable(identity=identity, ordinal=ordinal, element=element))
    return tuple(tables)


def numeric_tokens(paragraph: ExtractedParagraph) -> tuple[NumericToken, ...]:
    """Each maximal whitespace-delimited substring carrying at least one digit.

    A paragraph boundary terminates a token (Req 26.7): tokens are never joined across
    paragraphs, so a figure ending one paragraph and a figure starting the next do not
    fuse into one unmatched survivor.

    `offset` is the token's start index within the concatenated paragraph text, which
    is what lets `masking.py` overwrite a matched span in place and keep every later
    stage's offsets stable.
    """
    tokens: list[NumericToken] = []
    for match in re.finditer(r"\S+", paragraph.text):
        candidate = match.group()
        if _DIGIT.search(candidate) is None:
            continue
        tokens.append(
            NumericToken(
                text=candidate,
                part=paragraph.part,
                ordinal=paragraph.ordinal,
                block_id=paragraph.block_id,
                offset=match.start(),
            )
        )
    return tuple(tokens)


# --- PDF ------------------------------------------------------------------------------

_WHITESPACE_RUN: Final[re.Pattern[str]] = re.compile(r"\s+")


def normalize_pdf_text(pages: list[str]) -> str:
    """One contiguous string from a PDF's pages (Req 33.5).

    Join pages in ascending order with a single space, collapse every whitespace run to
    one space, trim. The point is that a `formatted` string the conversion split across
    text-show operators, across a line break or across a page boundary is still one
    contiguous substring of the result — so the fidelity gate can ask "is this string
    present" without assuming any correspondence between one operator and one figure,
    because there is none.
    """
    return _WHITESPACE_RUN.sub(" ", " ".join(pages)).strip()


def read_pdf_text(path: str | Path) -> tuple[str, int]:
    """The normalized text of a PDF and the page count it was read from.

    An unreadable PDF is the same proven failure an unreadable `.docx` is: the
    fidelity gate must not conclude that every ledger string is absent because the
    file would not parse.
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    # Broad for the same reason as `open_document`: every way a PDF fails to parse is
    # the same outcome, and the gate must not conclude that every ledger string is
    # absent because the file would not open.
    except Exception as exc:
        raise VerificationFailedError(
            f"the converted PDF at {path} could not be read: {type(exc).__name__}"
        ) from exc
    return normalize_pdf_text(pages), len(pages)
